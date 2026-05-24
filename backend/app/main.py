from fastapi import FastAPI, HTTPException, UploadFile, File, Form, Header, BackgroundTasks, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel
import uuid, os, re, io, sqlite3, shutil, random, hashlib, json, time, collections
import urllib.request, urllib.parse, urllib.error, ssl, xml.etree.ElementTree as ET
try:
    import certifi as _certifi
    _SSL_CTX = ssl.create_default_context(cafile=_certifi.where())
except ImportError:
    _SSL_CTX = ssl.create_default_context()
from datetime import datetime, timedelta, timezone
from typing import Optional
from dotenv import load_dotenv
load_dotenv()

app = FastAPI()
_SERVER_START_TIME = time.time()
_ALLOWED_ORIGINS = [o.strip() for o in os.environ.get("ALLOWED_ORIGINS", "*").split(",")]
_use_credentials = "*" not in _ALLOWED_ORIGINS
app.add_middleware(CORSMiddleware, allow_origins=_ALLOWED_ORIGINS, allow_credentials=_use_credentials, allow_methods=["*"], allow_headers=["*"])

DB_PATH = os.environ.get("DB_PATH", "/app/bversity.db")
SUBMISSIONS_DIR = os.environ.get("SUBMISSIONS_DIR", "/app/submissions")
IMAGE_CONFIG_PATH = os.environ.get("IMAGE_CONFIG_PATH", "/app/image_config.json")
os.makedirs(SUBMISSIONS_DIR, exist_ok=True)

DEFAULT_IMAGE_CONFIG = {
  "careers": {
    "bioinformatics_scientist":       {"url": "https://images.unsplash.com/photo-1532187863486-abf9dbad1b69?w=900&q=80", "label": "Bioinformatics Scientist"},
    "genomics_data_analyst":          {"url": "https://images.unsplash.com/photo-1614935151651-0bea6508db6b?w=900&q=80", "label": "Genomics Data Analyst"},
    "drug_discovery_scientist":       {"url": "https://images.unsplash.com/photo-1576086213369-97a306d36557?w=900&q=80", "label": "Drug Discovery Scientist"},
    "clinical_research_associate":    {"url": "https://images.unsplash.com/photo-1581093806997-124204d9fa9d?w=900&q=80", "label": "Clinical Research Associate"},
    "regulatory_affairs_associate":   {"url": "https://images.unsplash.com/photo-1450101499163-c8848c66ca85?w=900&q=80", "label": "Regulatory Affairs Associate"},
    "computational_biologist":        {"url": "https://images.unsplash.com/photo-1518770660439-4636190af475?w=900&q=80", "label": "Computational Biologist"},
    "pharmacovigilance_scientist":    {"url": "https://images.unsplash.com/photo-1587854692152-cbe660dbde88?w=900&q=80", "label": "Pharmacovigilance Scientist"},
    "medical_science_liaison":        {"url": "https://images.unsplash.com/photo-1559757148-5c350d0d3c56?w=900&q=80", "label": "Medical Science Liaison"},
    "biomarker_scientist":            {"url": "https://images.unsplash.com/photo-1559757175-0eb30cd8c063?w=900&q=80", "label": "Biomarker Scientist"},
    "clinical_data_manager":          {"url": "https://images.unsplash.com/photo-1551288049-bebda4e38f71?w=900&q=80", "label": "Clinical Data Manager"},
    "biotech_bd_associate":           {"url": "https://images.unsplash.com/photo-1521737711867-e3b97375f902?w=900&q=80", "label": "Biotech BD Associate"},
    "market_access_analyst":          {"url": "https://images.unsplash.com/photo-1611974789855-9c2a0a7236a3?w=900&q=80", "label": "Market Access Analyst"},
    "medical_affairs_associate":      {"url": "https://images.unsplash.com/photo-1582719508461-905c673771fd?w=900&q=80", "label": "Medical Affairs Associate"},
    "genomics_commercial_specialist": {"url": "https://images.unsplash.com/photo-1568219557405-376e23e4f7cf?w=900&q=80", "label": "Genomics Commercial Specialist"},
    "biotech_product_manager":        {"url": "https://images.unsplash.com/photo-1542744094-3a31f272c490?w=900&q=80", "label": "Biotech Product Manager"},
    "life_sciences_consultant":       {"url": "https://images.unsplash.com/photo-1556761175-4b46a572b786?w=900&q=80", "label": "Life Sciences Consultant"},
    "biotech_venture_analyst":        {"url": "https://images.unsplash.com/photo-1559136555-9303baea8ebd?w=900&q=80", "label": "Biotech Venture Analyst"},
    "licensing_partnerships":         {"url": "https://images.unsplash.com/photo-1507003211169-0a1dd7228f2d?w=900&q=80", "label": "Licensing & Partnerships"},
    "ai_drug_discovery":              {"url": "https://images.unsplash.com/photo-1677442135703-1787eea5ce01?w=900&q=80", "label": "AI Drug Discovery"},
    "precision_medicine_specialist":  {"url": "https://images.unsplash.com/photo-1576086476234-1103be98f096?w=900&q=80", "label": "Precision Medicine Specialist"},
    "biotech_founder":                {"url": "https://images.unsplash.com/photo-1519389950473-47ba0277781c?w=900&q=80", "label": "Biotech Founder"},
  },
  "clusters": {
    "Science & Technical":  {"url": "https://images.unsplash.com/photo-1579165466741-7f35e4755660?w=1200&q=80", "label": "Science & Technical cluster"},
    "Business & Commercial": {"url": "https://images.unsplash.com/photo-1454165804606-c3d57bc86b40?w=1200&q=80", "label": "Business & Commercial cluster"},
    "Emerging & Hybrid":    {"url": "https://images.unsplash.com/photo-1507003211169-0a1dd7228f2d?w=1200&q=80", "label": "Emerging & Hybrid cluster"},
  },
  "degrees": {
    "msc": {"url": "https://images.unsplash.com/photo-1576086213369-97a306d36557?w=800&h=320&fit=crop&q=80", "label": "M.Sc Biotechnology"},
    "bsc": {"url": "https://images.unsplash.com/photo-1530026405186-ed1f139313f8?w=800&h=320&fit=crop&q=80", "label": "B.Sc (Hons) Bioengineering"},
  },
}

def load_image_config():
    if os.path.exists(IMAGE_CONFIG_PATH):
        with open(IMAGE_CONFIG_PATH) as f:
            return json.load(f)
    return DEFAULT_IMAGE_CONFIG

def save_image_config(config):
    with open(IMAGE_CONFIG_PATH, "w") as f:
        json.dump(config, f, indent=2)

# ── Rate limiting ──────────────────────────────────────────────────────────────
RATE_LIMIT_MAX = int(os.environ.get("CHAT_RATE_LIMIT", "40"))   # messages per window
RATE_LIMIT_WINDOW = 3600                                          # 1 hour in seconds
_rate_buckets: dict = collections.defaultdict(collections.deque)

_anthropic_status = {"ok": None, "error": None, "checked_at": None}

def check_rate_limit(student_id: str):
    now = time.time()
    bucket = _rate_buckets[student_id]
    cutoff = now - RATE_LIMIT_WINDOW
    while bucket and bucket[0] < cutoff:
        bucket.popleft()
    if len(bucket) >= RATE_LIMIT_MAX:
        raise HTTPException(status_code=429, detail=f"Rate limit reached  -  max {RATE_LIMIT_MAX} messages per hour.")
    bucket.append(now)


def get_db():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn


def init_db():
    conn = get_db()
    conn.executescript("""
        CREATE TABLE IF NOT EXISTS students (
            id TEXT PRIMARY KEY, name TEXT NOT NULL, email TEXT UNIQUE NOT NULL, created_at TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS messages (
            id INTEGER PRIMARY KEY AUTOINCREMENT, student_id TEXT NOT NULL, subject_id TEXT NOT NULL,
            role TEXT NOT NULL, content TEXT NOT NULL, created_at TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS concept_progress (
            id INTEGER PRIMARY KEY AUTOINCREMENT, student_id TEXT NOT NULL, subject_id TEXT NOT NULL,
            concept_id TEXT NOT NULL, first_covered_at TEXT NOT NULL, mastered_at TEXT,
            UNIQUE(student_id, subject_id, concept_id)
        );
        CREATE TABLE IF NOT EXISTS materials (
            id TEXT PRIMARY KEY, subject_id TEXT NOT NULL, filename TEXT NOT NULL,
            chunk_count INTEGER DEFAULT 0, uploaded_at TEXT NOT NULL
        );
        CREATE VIRTUAL TABLE IF NOT EXISTS doc_chunks_fts USING fts5(
            material_id UNINDEXED, subject_id UNINDEXED, content, tokenize = 'porter ascii'
        );
        CREATE TABLE IF NOT EXISTS student_profile (
            student_id TEXT PRIMARY KEY, career_id TEXT, career_goal_raw TEXT, updated_at TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS capstone_submissions (
            id TEXT PRIMARY KEY, student_id TEXT NOT NULL, subject_id TEXT NOT NULL,
            filename TEXT NOT NULL, filepath TEXT NOT NULL, submitted_at TEXT NOT NULL,
            score INTEGER, feedback TEXT, marked_at TEXT,
            UNIQUE(student_id, subject_id)
        );
        CREATE TABLE IF NOT EXISTS approved_emails (
            email TEXT PRIMARY KEY,
            added_at TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS verification_codes (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            email TEXT NOT NULL,
            code TEXT NOT NULL,
            expires_at TEXT NOT NULL,
            used INTEGER DEFAULT 0,
            created_at TEXT DEFAULT (datetime('now'))
        );
    """)
    conn.executescript("""
        CREATE TABLE IF NOT EXISTS career_changes (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            student_id TEXT NOT NULL,
            from_career_id TEXT,
            to_career_id TEXT,
            reason TEXT NOT NULL,
            notes TEXT,
            changed_at TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS subject_status (
            student_id TEXT NOT NULL,
            subject_id TEXT NOT NULL,
            status TEXT NOT NULL DEFAULT 'locked',
            unlocked_at TEXT,
            paused_at TEXT,
            completed_at TEXT,
            pause_reason TEXT,
            pause_notes TEXT,
            PRIMARY KEY (student_id, subject_id)
        );
        CREATE TABLE IF NOT EXISTS module_quizzes (
            student_id  TEXT NOT NULL,
            subject_id  TEXT NOT NULL,
            module_id   TEXT NOT NULL,
            passed      INTEGER DEFAULT 0,
            completed_at TEXT,
            PRIMARY KEY (student_id, subject_id, module_id)
        );
        CREATE TABLE IF NOT EXISTS concept_resources (
            id           TEXT PRIMARY KEY,
            subject_id   TEXT NOT NULL,
            concept_id   TEXT NOT NULL,
            url          TEXT NOT NULL,
            title        TEXT NOT NULL,
            resource_type TEXT NOT NULL DEFAULT 'article',
            description  TEXT,
            added_at     TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS notes (
            id TEXT PRIMARY KEY,
            student_id TEXT NOT NULL,
            subject_id TEXT,
            content TEXT NOT NULL,
            created_at TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS study_plan (
            student_id TEXT NOT NULL,
            day_number INTEGER NOT NULL,
            subject_id TEXT NOT NULL,
            concept_id TEXT NOT NULL,
            target_date TEXT NOT NULL,
            PRIMARY KEY (student_id, day_number, concept_id)
        );
        CREATE TABLE IF NOT EXISTS concept_videos (
            subject_id TEXT NOT NULL,
            concept_id TEXT NOT NULL,
            drive_url  TEXT NOT NULL,
            title      TEXT,
            added_at   TEXT NOT NULL,
            PRIMARY KEY (subject_id, concept_id)
        );
        CREATE TABLE IF NOT EXISTS subject_completions (
            student_id    TEXT NOT NULL,
            subject_id    TEXT NOT NULL,
            completed_at  TEXT NOT NULL,
            credential_id TEXT NOT NULL,
            PRIMARY KEY (student_id, subject_id)
        );
        CREATE TABLE IF NOT EXISTS platform_feedback (
            id          TEXT PRIMARY KEY,
            student_id  TEXT NOT NULL,
            q1          TEXT,
            q2          TEXT,
            q3          TEXT,
            rating      INTEGER NOT NULL,
            comment     TEXT,
            submitted_at TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS module_quiz_questions (
            subject_id      TEXT NOT NULL,
            module_id       TEXT NOT NULL,
            questions_json  TEXT NOT NULL,
            generated_at    TEXT NOT NULL,
            PRIMARY KEY (subject_id, module_id)
        );
        CREATE TABLE IF NOT EXISTS access_requests (
            id           TEXT PRIMARY KEY,
            name         TEXT NOT NULL,
            email        TEXT NOT NULL UNIQUE,
            phone        TEXT,
            university   TEXT,
            year_of_study TEXT,
            country      TEXT,
            reason       TEXT,
            status       TEXT NOT NULL DEFAULT 'pending',
            submitted_at TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS session_summaries (
            id            INTEGER PRIMARY KEY AUTOINCREMENT,
            student_id    TEXT NOT NULL,
            subject_id    TEXT NOT NULL,
            summary       TEXT NOT NULL,
            message_count INTEGER DEFAULT 0,
            created_at    TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS concept_feedback (
            id          INTEGER PRIMARY KEY AUTOINCREMENT,
            student_id  TEXT NOT NULL,
            subject_id  TEXT NOT NULL,
            concept_title TEXT NOT NULL,
            value       TEXT NOT NULL,
            created_at  TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS message_feedback (
            id          INTEGER PRIMARY KEY AUTOINCREMENT,
            student_id  TEXT NOT NULL,
            subject_id  TEXT NOT NULL,
            message_idx INTEGER NOT NULL,
            value       TEXT NOT NULL,
            created_at  TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS concept_notes (
            subject_id  TEXT NOT NULL,
            concept_id  TEXT NOT NULL,
            notes       TEXT NOT NULL DEFAULT '',
            updated_at  TEXT NOT NULL,
            PRIMARY KEY (subject_id, concept_id)
        );
        CREATE TABLE IF NOT EXISTS saved_concepts (
            id          TEXT PRIMARY KEY,
            student_id  TEXT NOT NULL,
            subject_id  TEXT NOT NULL,
            title       TEXT NOT NULL,
            card_data   TEXT NOT NULL,
            saved_at    TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS lab_progress (
            student_id  TEXT NOT NULL,
            project_id  TEXT NOT NULL,
            subject_id  TEXT NOT NULL,
            steps_done  TEXT NOT NULL DEFAULT '[]',
            status      TEXT NOT NULL DEFAULT 'in_progress',
            submission  TEXT,
            ai_feedback TEXT,
            started_at  TEXT NOT NULL,
            submitted_at TEXT,
            PRIMARY KEY (student_id, project_id)
        );
        CREATE TABLE IF NOT EXISTS job_listings (
            id          INTEGER PRIMARY KEY AUTOINCREMENT,
            cert_id     TEXT NOT NULL,
            title       TEXT NOT NULL,
            company     TEXT,
            location    TEXT,
            salary_min  REAL,
            salary_max  REAL,
            url         TEXT,
            description TEXT,
            posted_date TEXT,
            fetched_at  TEXT NOT NULL
        );
    """)
    for col in [
        "ALTER TABLE capstone_submissions ADD COLUMN ai_score INTEGER",
        "ALTER TABLE capstone_submissions ADD COLUMN ai_feedback TEXT",
        "ALTER TABLE capstone_submissions ADD COLUMN ai_graded_at TEXT",
        "ALTER TABLE concept_progress ADD COLUMN mastered_at TEXT",
        "ALTER TABLE student_profile ADD COLUMN college TEXT",
        "ALTER TABLE student_profile ADD COLUMN year_of_study TEXT",
        "ALTER TABLE student_profile ADD COLUMN aspirations TEXT",
        "ALTER TABLE student_profile ADD COLUMN motivation TEXT",
        "ALTER TABLE student_profile ADD COLUMN tutor_note TEXT",
        "ALTER TABLE student_profile ADD COLUMN onboarded_at TEXT",
        "ALTER TABLE student_profile ADD COLUMN avatar_color TEXT",
        "ALTER TABLE student_profile ADD COLUMN linkedin_url TEXT",
        "ALTER TABLE student_profile ADD COLUMN github_url TEXT",
        "ALTER TABLE student_profile ADD COLUMN bio TEXT",
        "ALTER TABLE student_profile ADD COLUMN city TEXT",
        "ALTER TABLE student_profile ADD COLUMN state TEXT",
        "ALTER TABLE student_profile ADD COLUMN show_on_map INTEGER DEFAULT 1",
        "ALTER TABLE student_profile ADD COLUMN avatar_num INTEGER",
        "ALTER TABLE student_profile ADD COLUMN is_placed INTEGER DEFAULT 0",
        "ALTER TABLE student_profile ADD COLUMN streak_count INTEGER DEFAULT 0",
        "ALTER TABLE student_profile ADD COLUMN streak_last_date TEXT",
        "ALTER TABLE student_profile ADD COLUMN last_active_at TEXT",
        "ALTER TABLE student_profile ADD COLUMN nudge_sent_at TEXT",
        "ALTER TABLE student_profile ADD COLUMN weekly_report_sent_at TEXT",
        "ALTER TABLE approved_emails ADD COLUMN product TEXT NOT NULL DEFAULT 'career_pathways'",
        "ALTER TABLE approved_emails ADD COLUMN expires_at TEXT",
        "ALTER TABLE approved_emails ADD COLUMN access_type TEXT NOT NULL DEFAULT 'manual'",
        "ALTER TABLE access_requests ADD COLUMN product TEXT NOT NULL DEFAULT 'career_pathways'",
        "ALTER TABLE student_profile ADD COLUMN learner_archetype TEXT",
        "ALTER TABLE subscriptions ADD COLUMN subscription_end TEXT",
        "ALTER TABLE subscriptions ADD COLUMN started_at TEXT",
        "ALTER TABLE subscriptions ADD COLUMN warning_2d_sent TEXT",
        "ALTER TABLE subscriptions ADD COLUMN warning_1d_sent TEXT",
    ]:
        try:
            conn.execute(col)
        except Exception:
            pass

    conn.executescript("""
        CREATE TABLE IF NOT EXISTS platform_config (
            product     TEXT PRIMARY KEY,
            mode        TEXT NOT NULL DEFAULT 'self_serve',
            trial_days  INTEGER NOT NULL DEFAULT 5,
            updated_at  TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS subscriptions (
            id                      TEXT PRIMARY KEY,
            student_id              TEXT NOT NULL,
            product                 TEXT NOT NULL,
            status                  TEXT NOT NULL DEFAULT 'trial',
            trial_start             TEXT,
            trial_end               TEXT,
            country_code            TEXT,
            payment_method          TEXT,
            gateway_subscription_id TEXT,
            gateway_customer_id     TEXT,
            amount_cents            INTEGER,
            currency                TEXT,
            created_at              TEXT NOT NULL,
            updated_at              TEXT NOT NULL,
            UNIQUE(student_id, product)
        );
        CREATE TABLE IF NOT EXISTS payment_events (
            id                  TEXT PRIMARY KEY,
            student_id          TEXT NOT NULL,
            product             TEXT NOT NULL,
            event_type          TEXT NOT NULL DEFAULT 'payment',
            amount_cents        INTEGER NOT NULL DEFAULT 0,
            currency            TEXT NOT NULL DEFAULT 'usd',
            gateway             TEXT NOT NULL,
            gateway_payment_id  TEXT,
            created_at          TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS site_settings (
            key   TEXT PRIMARY KEY,
            value TEXT NOT NULL,
            updated_at TEXT NOT NULL
        );
        INSERT OR IGNORE INTO site_settings (key, value, updated_at) VALUES (
            'hero_video_url',
            'https://videos.pexels.com/video-files/18069830/18069830-hd_1920_1080_24fps.mp4',
            datetime('now')
        );
        CREATE TABLE IF NOT EXISTS subject_videos (
            id           INTEGER PRIMARY KEY AUTOINCREMENT,
            subject_id   TEXT NOT NULL,
            title        TEXT NOT NULL,
            drive_url    TEXT NOT NULL,
            order_index  INTEGER NOT NULL,
            added_at     TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS archetype_scores (
            student_id          TEXT PRIMARY KEY,
            question_depth      REAL DEFAULT 0,
            exam_focus          REAL DEFAULT 0,
            anxiety_markers     REAL DEFAULT 0,
            topic_jumping       REAL DEFAULT 0,
            challenge_seeking   REAL DEFAULT 0,
            practical_anchoring REAL DEFAULT 0,
            mastery_push        REAL DEFAULT 0,
            sessions_scored     INTEGER DEFAULT 0,
            archetype           TEXT,
            archetype_reasoning TEXT,
            updated_at          TEXT
        );
        CREATE TABLE IF NOT EXISTS concept_summaries (
            concept_id   TEXT NOT NULL,
            subject_id   TEXT NOT NULL,
            what         TEXT NOT NULL,
            why          TEXT NOT NULL,
            key_points   TEXT NOT NULL,
            real_world   TEXT,
            interview_q  TEXT,
            generated_at TEXT NOT NULL,
            PRIMARY KEY (subject_id, concept_id)
        );
        CREATE TABLE IF NOT EXISTS industry_news (
            id           INTEGER PRIMARY KEY AUTOINCREMENT,
            title        TEXT NOT NULL,
            url          TEXT NOT NULL,
            source       TEXT NOT NULL,
            published_at TEXT,
            fetched_at   TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS industry_newsletter (
            id           INTEGER PRIMARY KEY AUTOINCREMENT,
            title        TEXT NOT NULL,
            content      TEXT,
            url          TEXT,
            published_at TEXT NOT NULL,
            created_at   TEXT NOT NULL
        );
    """)
    conn.execute("""
        INSERT OR IGNORE INTO platform_config (product, mode, trial_days, updated_at) VALUES
            ('career_pathways', 'self_serve', 5, datetime('now')),
            ('certifications',  'self_serve', 5, datetime('now'))
    """)
    conn.commit()
    conn.close()


_PLACED_ALUMNI = [
    {"id": "BVA001", "name": "Dristi Mohta",         "email": "dristi.mohta@bversity.alumni",      "city": "Raipur",        "state": "Chhattisgarh",   "linkedin_url": "https://www.linkedin.com/in/dristi-mohta-b62275187/", "avatar_color": "#00A896"},
    {"id": "BVA002", "name": "Arpita Ganguly",        "email": "arpita.ganguly@bversity.alumni",    "city": "Kolkata",       "state": "West Bengal",    "linkedin_url": None, "avatar_color": "#7B61FF"},
    {"id": "BVA003", "name": "Ananaya Jain",          "email": "ananaya.jain@bversity.alumni",      "city": "Ghaziabad",     "state": "Uttar Pradesh",  "linkedin_url": None, "avatar_color": "#FF6B6B"},
    {"id": "BVA004", "name": "Dheeraj Babu",          "email": "dheeraj.babu@bversity.alumni",      "city": "Bangalore",     "state": "Karnataka",      "linkedin_url": None, "avatar_color": "#00A896"},
    {"id": "BVA005", "name": "Ekant Kannam",          "email": "ekant.kannam@bversity.alumni",      "city": "Bijapur",       "state": "Chhattisgarh",   "linkedin_url": None, "avatar_color": "#F5A623"},
    {"id": "BVA006", "name": "Megha",                 "email": "megha@bversity.alumni",             "city": "Chennai",       "state": "Tamil Nadu",     "linkedin_url": None, "avatar_color": "#FF6B6B"},
    {"id": "BVA007", "name": "Khushi Tyagi",          "email": "khushi.tyagi@bversity.alumni",      "city": "Delhi",         "state": "Delhi",          "linkedin_url": None, "avatar_color": "#7B61FF"},
    {"id": "BVA008", "name": "Disha Jain",            "email": "disha.jain@bversity.alumni",        "city": "Bangalore",     "state": "Karnataka",      "linkedin_url": None, "avatar_color": "#00A896"},
    {"id": "BVA009", "name": "Arsha Adak",            "email": "arsha.adak@bversity.alumni",        "city": "Howrah",        "state": "West Bengal",    "linkedin_url": None, "avatar_color": "#F5A623"},
    {"id": "BVA010", "name": "Kayalvizhi",            "email": "kayalvizhi@bversity.alumni",        "city": "Perambalur",    "state": "Tamil Nadu",     "linkedin_url": None, "avatar_color": "#FF6B6B"},
    {"id": "BVA011", "name": "Risha Reddy",           "email": "risha.reddy@bversity.alumni",       "city": "Hyderabad",     "state": "Telangana",      "linkedin_url": None, "avatar_color": "#00A896"},
    {"id": "BVA012", "name": "Suchismita Shaw",       "email": "suchismita.shaw@bversity.alumni",   "city": "Kolkata",       "state": "West Bengal",    "linkedin_url": None, "avatar_color": "#7B61FF"},
    {"id": "BVA013", "name": "Priya Banthia",         "email": "priya.banthia@bversity.alumni",     "city": "Kolkata",       "state": "West Bengal",    "linkedin_url": None, "avatar_color": "#FF6B6B"},
    {"id": "BVA014", "name": "Abel George",           "email": "abel.george@bversity.alumni",       "city": "Nedumkandam",   "state": "Kerala",         "linkedin_url": None, "avatar_color": "#00A896"},
    {"id": "BVA015", "name": "Salam Akash Singh",     "email": "salam.akash@bversity.alumni",       "city": "Pune",          "state": "Maharashtra",    "linkedin_url": None, "avatar_color": "#F5A623"},
    {"id": "BVA016", "name": "Laxmi C M",             "email": "laxmi.cm@bversity.alumni",          "city": "Thrissur",      "state": "Kerala",         "linkedin_url": None, "avatar_color": "#7B61FF"},
    {"id": "BVA017", "name": "Reshma B",              "email": "reshma.b@bversity.alumni",          "city": "Madurai",       "state": "Tamil Nadu",     "linkedin_url": None, "avatar_color": "#FF6B6B"},
    {"id": "BVA018", "name": "Sowfika Dharshini",     "email": "sowfika.dharshini@bversity.alumni", "city": "Coimbatore",    "state": "Tamil Nadu",     "linkedin_url": None, "avatar_color": "#00A896"},
    {"id": "BVA019", "name": "Narendra Varma",        "email": "narendra.varma@bversity.alumni",    "city": "Tirupati",      "state": "Andhra Pradesh", "linkedin_url": None, "avatar_color": "#F5A623"},
    {"id": "BVA020", "name": "Sri Vaishnavi Dabberu", "email": "sri.vaishnavi@bversity.alumni",     "city": "Secunderabad",  "state": "Telangana",      "linkedin_url": None, "avatar_color": "#7B61FF"},
    {"id": "BVA021", "name": "Krishna Verma",         "email": "krishna.verma@bversity.alumni",     "city": "Ahmedabad",     "state": "Gujarat",        "linkedin_url": None, "avatar_color": "#00A896"},
    {"id": "BVA022", "name": "Amisha N",              "email": "amisha.n@bversity.alumni",          "city": "Davanagere",    "state": "Karnataka",      "linkedin_url": None, "avatar_color": "#FF6B6B"},
    {"id": "BVA023", "name": "Brindha",               "email": "brindha@bversity.alumni",           "city": "Nagercoil",     "state": "Tamil Nadu",     "linkedin_url": None, "avatar_color": "#F5A623"},
]

def seed_placed_alumni():
    conn = get_db()
    now = datetime.utcnow().isoformat()
    for a in _PLACED_ALUMNI:
        conn.execute(
            "INSERT OR IGNORE INTO students (id, name, email, created_at) VALUES (?, ?, ?, ?)",
            (a["id"], a["name"], a["email"], now),
        )
        conn.execute(
            """INSERT INTO student_profile (student_id, career_id, updated_at, city, state, show_on_map, is_placed, linkedin_url, avatar_color)
               VALUES (?, NULL, ?, ?, ?, 1, 1, ?, ?)
               ON CONFLICT(student_id) DO UPDATE SET
                 city = excluded.city, state = excluded.state,
                 show_on_map = 1, is_placed = 1,
                 linkedin_url = COALESCE(excluded.linkedin_url, student_profile.linkedin_url),
                 avatar_color = excluded.avatar_color,
                 updated_at = excluded.updated_at""",
            (a["id"], now, a["city"], a["state"], a["linkedin_url"], a["avatar_color"]),
        )
    conn.commit()
    conn.close()


init_db()
seed_placed_alumni()

# ── Industry News ─────────────────────────────────────────────────────────────

_NEWS_FEEDS = [
    {"name": "STAT News",            "url": "https://www.statnews.com/feed/"},
    {"name": "FierceBiotech",        "url": "https://www.fiercebiotech.com/rss/xml"},
    {"name": "Endpoints News",       "url": "https://endpts.com/feed/"},
    {"name": "BioPharma Dive",       "url": "https://www.biopharmadive.com/feeds/news/"},
    {"name": "The Scientist",        "url": "https://www.the-scientist.com/rss"},
    {"name": "FDA Press Releases",   "url": "https://www.fda.gov/about-fda/contact-fda/stay-informed/rss-feeds/press-announcements/rss.xml"},
    {"name": "GEN Biotech News",     "url": "https://www.genengnews.com/feed/"},
    {"name": "CenterWatch",          "url": "https://www.centerwatch.com/feed/"},
    {"name": "BioSpectrum India",    "url": "https://www.biospectrumindia.com/rss.xml"},
    {"name": "Express Pharma India", "url": "https://www.expresspharma.in/feed/"},
    {"name": "ET Pharma & Biotech",  "url": "https://economictimes.indiatimes.com/industry/healthcare/biotech/pharmaceuticals/rssfeeds/13353464.cms"},
]

_news_last_fetched: Optional[datetime] = None
_news_fetch_lock = False


def _parse_rss_date(date_str: Optional[str]) -> Optional[str]:
    if not date_str:
        return None
    fmts = [
        "%a, %d %b %Y %H:%M:%S %z",
        "%a, %d %b %Y %H:%M:%S GMT",
        "%Y-%m-%dT%H:%M:%S%z",
        "%Y-%m-%dT%H:%M:%SZ",
    ]
    for fmt in fmts:
        try:
            return datetime.strptime(date_str.strip(), fmt).isoformat()
        except Exception:
            pass
    return date_str


def _parse_feed_bytes(raw: bytes, source: str, max_items: int = 4):
    items = []
    try:
        root = ET.fromstring(raw)
    except ET.ParseError:
        return items
    ATOM = "http://www.w3.org/2005/Atom"

    channel = root.find("channel")
    if channel is not None:
        for item in list(channel.findall("item"))[:max_items]:
            title = (item.findtext("title") or "").strip()
            url   = (item.findtext("link") or "").strip()
            pub   = item.findtext("pubDate")
            if title and url:
                items.append({"title": title, "url": url, "source": source, "published_at": _parse_rss_date(pub)})
        return items

    def _at(t): return f"{{{ATOM}}}{t}"
    for entry in list(root.findall(_at("entry")))[:max_items]:
        t_el  = entry.find(_at("title"))
        l_el  = entry.find(_at("link"))
        p_el  = entry.find(_at("published")) or entry.find(_at("updated"))
        title = (t_el.text or "").strip() if t_el is not None else ""
        url   = l_el.get("href", "") if l_el is not None else ""
        pub   = p_el.text if p_el is not None else None
        if title and url:
            items.append({"title": title, "url": url, "source": source, "published_at": _parse_rss_date(pub)})
    return items


def _fetch_news_feeds():
    global _news_last_fetched, _news_fetch_lock
    if _news_fetch_lock:
        return
    _news_fetch_lock = True
    try:
        items = []
        for feed in _NEWS_FEEDS:
            try:
                req = urllib.request.Request(feed["url"], headers={"User-Agent": "Mozilla/5.0 (compatible; Bversity/1.0)"})
                with urllib.request.urlopen(req, context=_SSL_CTX, timeout=12) as resp:
                    raw = resp.read()
                items.extend(_parse_feed_bytes(raw, feed["name"]))
            except Exception as e:
                print(f"[news] {feed['name']} fetch error: {e}")
        if items:
            now = datetime.utcnow().isoformat()
            conn = get_db()
            conn.execute("DELETE FROM industry_news")
            for it in items[:30]:
                conn.execute(
                    "INSERT INTO industry_news (title, url, source, published_at, fetched_at) VALUES (?,?,?,?,?)",
                    (it["title"], it["url"], it["source"], it.get("published_at"), now),
                )
            conn.commit()
            conn.close()
            _news_last_fetched = datetime.utcnow()
            print(f"[news] fetched {len(items)} articles")
    except Exception as e:
        print(f"[news] fetch_all error: {e}")
    finally:
        _news_fetch_lock = False


# ── Curriculum ────────────────────────────────────────────────────────────────

CURRICULUM = {
    "bioinformatics": [
        {"id": "central_dogma_a",      "name": "DNA Structure & Replication",                  "desc": "Double helix, base pairing rules, nucleotides, semi-conservative replication, DNA polymerase fidelity and error rates"},
        {"id": "central_dogma_b",      "name": "Transcription & RNA Processing",               "desc": "RNA polymerase II, promoters, enhancers, pre-mRNA splicing, 5' capping, poly-A tail  -  how genes become transcripts"},
        {"id": "central_dogma_c",      "name": "Translation & Post-Translational Modification","desc": "Ribosomes, codons, tRNA, the genetic code, protein folding, PTMs (phosphorylation, glycosylation, ubiquitination)"},
        {"id": "seq_formats_a",        "name": "FASTA & FASTQ Formats",                        "desc": "FASTA sequence records, FASTQ quality scores, Phred encoding, per-base quality interpretation, multi-record files"},
        {"id": "seq_formats_b",        "name": "SAM/BAM & Alignment Formats",                  "desc": "SAM header sections, mandatory fields, CIGAR string decoding, FLAG values, sorting and indexing with SAMtools"},
        {"id": "seq_formats_c",        "name": "VCF/BCF & Variant Representation",             "desc": "VCF record structure, CHROM/POS/REF/ALT fields, INFO and FORMAT annotations, multi-sample VCFs, BCF compression"},
        {"id": "pairwise_alignment_a", "name": "Global Alignment: Needleman-Wunsch",           "desc": "Dynamic programming matrix, traceback, affine gap penalties, when to use global vs local alignment"},
        {"id": "pairwise_alignment_b", "name": "Local Alignment: Smith-Waterman",              "desc": "Local DP formulation, zero-floor rule, optimal local matches, computational complexity, hardware acceleration"},
        {"id": "pairwise_alignment_c", "name": "Scoring Matrices & Gap Penalties",             "desc": "PAM and BLOSUM matrices, how they were derived, choosing the right matrix and gap penalty for your alignment problem"},
        {"id": "blast_search_a",       "name": "BLAST Algorithm Internals",                    "desc": "Word seeding, HSP extension, two-hit method, E-value calculation, bit scores, statistical significance"},
        {"id": "blast_search_b",       "name": "BLAST Variants & Use Cases",                   "desc": "BLASTn, BLASTp, BLASTx, tBLASTn, tBLASTx, PSI-BLAST, DELTA-BLAST  -  when and why to use each"},
        {"id": "blast_search_c",       "name": "Interpreting BLAST Output",                    "desc": "Score, E-value, percent identity, query/subject coverage, low-complexity masking, filtering noise from results"},
        {"id": "msa_phylogenetics_a",  "name": "Multiple Sequence Alignment Methods",          "desc": "Progressive (ClustalW) vs iterative (MUSCLE, MAFFT) alignment, scoring MSAs, gap treatment, anchor columns"},
        {"id": "msa_phylogenetics_b",  "name": "Distance & Parsimony Tree Methods",            "desc": "Genetic distance models (Jukes-Cantor, K2P), UPGMA, neighbor-joining, maximum parsimony  -  concepts and limits"},
        {"id": "msa_phylogenetics_c",  "name": "Maximum Likelihood & Bayesian Trees",          "desc": "Substitution models, IQ-TREE, BEAST, bootstrapping, posterior probability, interpreting clade support values"},
        {"id": "bio_databases_a",      "name": "NCBI Database Suite",                          "desc": "GenBank, RefSeq, dbSNP, SRA, ClinVar, dbVar  -  database relationships, Entrez API, EDirect command-line access"},
        {"id": "bio_databases_b",      "name": "Protein & Structure Databases",                "desc": "UniProt/Swiss-Prot vs TrEMBL, PDB structure records, Pfam domains, InterPro, programmatic REST access"},
        {"id": "bio_databases_c",      "name": "Pathway & Ontology Databases",                 "desc": "KEGG pathways and BRITE hierarchies, Reactome, Gene Ontology (MF/BP/CC), using GO terms for enrichment analysis"},
        {"id": "ngs_qc_a",             "name": "Illumina Sequencing Technology",               "desc": "Flow cell clusters, bridge amplification, sequencing-by-synthesis, paired-end reads, index multiplexing, read length trade-offs"},
        {"id": "ngs_qc_b",             "name": "Long-Read Sequencing Platforms",               "desc": "PacBio HiFi (CCS) vs CLR, Oxford Nanopore (pore chemistry, basecalling), accuracy vs read length, appropriate use cases"},
        {"id": "ngs_qc_c",             "name": "NGS Quality Control & Trimming",               "desc": "FastQC report interpretation, per-base quality decay, adapter contamination, Trimmomatic/Cutadapt parameters, MultiQC"},
        {"id": "ngs_alignment_a",      "name": "Short-Read Alignment",                         "desc": "BWA-MEM algorithm, Bowtie2 modes, reference genome indexing, duplicate marking (Picard), alignment metrics"},
        {"id": "ngs_alignment_b",      "name": "RNA-seq Splice-Aware Alignment",               "desc": "STAR two-pass alignment, HISAT2, splice junctions, multi-mapping reads, alignment to transcriptome vs genome"},
        {"id": "ngs_alignment_c",      "name": "Variant Calling with GATK",                    "desc": "BQSR, HaplotypeCaller, joint genotyping with GenomicsDBImport, gVCF format, hard filtering vs VQSR"},
        {"id": "rnaseq_expression_a",  "name": "RNA-seq Library Prep & Design",                "desc": "Strand-specific libraries, rRNA depletion vs polyA selection, sequencing depth, replication strategy, batch effects"},
        {"id": "rnaseq_expression_b",  "name": "Quantification & Normalisation",               "desc": "featureCounts, Salmon pseudoalignment, TPM vs RPKM vs raw counts, DESeq2 size factors, edgeR TMM normalisation"},
        {"id": "rnaseq_expression_c",  "name": "Differential Expression Downstream Analysis",  "desc": "Volcano plots, MA plots, FDR correction, GSEA, fgsea, GO and KEGG enrichment, heatmaps, pathway visualisation"},
        {"id": "protein_structure_a",  "name": "Protein Structural Hierarchy",                 "desc": "Primary sequence, secondary elements (α-helix, β-sheet), tertiary fold, quaternary assemblies, domains and motifs"},
        {"id": "protein_structure_b",  "name": "Experimental Structure Determination",         "desc": "X-ray crystallography, NMR spectroscopy, cryo-EM  -  principles, resolution, PDB deposition, viewing in PyMOL/ChimeraX"},
        {"id": "protein_structure_c",  "name": "AlphaFold & Computational Structure Prediction","desc": "How AlphaFold2 works, pLDDT and PAE confidence scores, AlphaFold Multimer, ESMFold, structure-function applications"},
        {"id": "bio_programming_a",    "name": "Python & Biopython for Bioinformatics",        "desc": "SeqIO, pairwise2, Entrez API access, BLAST wrapper, pandas for tabular data, writing reusable bioinformatics scripts"},
        {"id": "bio_programming_b",    "name": "R & Bioconductor",                             "desc": "Biostrings, GenomicRanges, DESeq2, ggplot2, ComplexHeatmap  -  R ecosystem for genomics and visualisation"},
        {"id": "bio_programming_c",    "name": "Workflow Management & Reproducibility",        "desc": "Snakemake rules and DAGs, Nextflow channels, conda environments, containers (Docker/Singularity), reproducible pipelines"},
        {"id": "ml_bioinformatics_a",  "name": "Feature Engineering for Biological Data",      "desc": "Molecular descriptors, k-mer frequency, position weight matrices, one-hot encoding, handling high-dimensional omics features"},
        {"id": "ml_bioinformatics_b",  "name": "Classical ML on Omics Data",                   "desc": "Random forests, SVMs, gradient boosting for biomarker classification; clustering (k-means, hierarchical) for omics"},
        {"id": "ml_bioinformatics_c",  "name": "Deep Learning for Biological Sequences",       "desc": "CNNs for motif discovery, RNNs/LSTMs for sequence modelling, attention mechanisms, intro to protein language models"},
    ],
    "genomics": [
        {"id": "genome_structure_a",   "name": "Chromosome & Chromatin Organisation",         "desc": "Nucleosomes, chromatin compaction, topologically associating domains (TADs), lamina-associated domains, heterochromatin vs euchromatin"},
        {"id": "genome_structure_b",   "name": "Genes, Regulatory Elements & Non-Coding DNA", "desc": "Exon-intron structure, promoters, enhancers, silencers, insulators, ENCODE annotation, proportion of functional non-coding sequence"},
        {"id": "genome_structure_c",   "name": "Repetitive Elements & the Dark Genome",       "desc": "SINEs, LINEs, transposons, centromeric satellites, segmental duplications  -  what makes 45% of the genome hard to sequence"},
        {"id": "sequencing_tech_a",    "name": "Sanger Sequencing",                           "desc": "Chain termination chemistry, capillary electrophoresis, read length, clinical sequencing (Sanger confirmation), limitations"},
        {"id": "sequencing_tech_b",    "name": "Illumina Short-Read Sequencing",              "desc": "Paired-end library prep, flow cell clusters, sequencing-by-synthesis chemistry, error profiles, throughput vs cost"},
        {"id": "sequencing_tech_c",    "name": "Long-Read Sequencing: PacBio & Nanopore",     "desc": "PacBio HiFi (CCS) accuracy, Nanopore direct sequencing, basecalling models, resolving structural variants and repeats"},
        {"id": "genome_assembly_a",    "name": "De Novo Assembly Algorithms",                 "desc": "Overlap-layout-consensus vs de Bruijn graph assembly, k-mer selection, contig formation, assembly quality metrics (N50)"},
        {"id": "genome_assembly_b",    "name": "Scaffolding & Gap Filling",                   "desc": "Hi-C chromatin proximity, optical mapping (Bionano), reference-guided scaffolding, telomere-to-telomere assemblies"},
        {"id": "genome_assembly_c",    "name": "Genome Annotation",                           "desc": "Ab initio gene prediction (Augustus, GeneMark), evidence-based annotation (MAKER), RepeatMasker, functional annotation"},
        {"id": "variant_types_a",      "name": "SNPs & Small Indels",                         "desc": "SNP formation, MAF, synonymous vs non-synonymous, frameshift indels, variant nomenclature (HGVS), population frequencies"},
        {"id": "variant_types_b",      "name": "Copy Number Variants",                        "desc": "CNV detection methods (array CGH, WGS read depth), segmental duplications, gene dosage effects, clinical CNV examples"},
        {"id": "variant_types_c",      "name": "Structural Variants",                         "desc": "Inversions, translocations, complex rearrangements, mobile element insertions, long-read and linked-read SV calling"},
        {"id": "variant_calling_a",    "name": "GATK Best Practices Pipeline",                "desc": "Preprocessing (marking duplicates, BQSR), HaplotypeCaller, joint genotyping with GenomicsDBImport, gVCF workflow"},
        {"id": "variant_calling_b",    "name": "Variant Filtering & Annotation",              "desc": "VQSR vs hard filters, VEP/ANNOVAR consequence prediction, gnomAD frequencies, conservation scores, splicing predictors"},
        {"id": "variant_calling_c",    "name": "Clinical Variant Classification",             "desc": "ACMG/AMP 2015 criteria, ClinVar submissions, VUS interpretation, functional evidence categories, variant curation workflow"},
        {"id": "population_genetics_a","name": "Allele Frequencies & Population Dynamics",    "desc": "MAF, Hardy-Weinberg equilibrium testing, genetic drift, founder effects, bottlenecks, selection signatures"},
        {"id": "population_genetics_b","name": "Linkage Disequilibrium & Haplotypes",         "desc": "D' and r² measures, LD decay with distance, haplotype blocks, tagSNPs, HapMap and 1000 Genomes reference panels"},
        {"id": "population_genetics_c","name": "Population Stratification",                   "desc": "Principal component analysis, ADMIXTURE/STRUCTURE, ancestry inference, confounding in GWAS, genomic control"},
        {"id": "gwas_a",               "name": "GWAS Study Design & QC",                      "desc": "Cohort selection, genotyping arrays, imputation with TOPMed/HRC, QC steps (call rate, HWE, relatedness, ancestry)"},
        {"id": "gwas_b",               "name": "GWAS Statistical Analysis",                   "desc": "Logistic/linear regression, Bonferroni and FDR correction, Manhattan plots, Q-Q plots, lambda inflation factor"},
        {"id": "gwas_c",               "name": "Post-GWAS Analysis",                          "desc": "Fine-mapping (FINEMAP, SuSiE), colocalization, Mendelian randomization, polygenic risk scores (PRSice, LDpred)"},
        {"id": "transcriptomics_a",    "name": "Bulk RNA-seq Design & Processing",            "desc": "Experimental design, library prep choices, alignment (STAR), quantification (Salmon/featureCounts), QC metrics"},
        {"id": "transcriptomics_b",    "name": "Differential Expression Analysis",            "desc": "DESeq2 negative binomial model, edgeR, limma-voom, FDR correction, biological replication requirements"},
        {"id": "transcriptomics_c",    "name": "Functional Interpretation of RNA-seq",        "desc": "GO enrichment, GSEA, pathway analysis (KEGG, Reactome), cell type deconvolution (CIBERSORT, CellChat)"},
        {"id": "single_cell_a",        "name": "scRNA-seq Technologies",                      "desc": "10x Chromium droplet microfluidics, cell barcodes, UMIs, 3' vs 5' capture, SMART-seq, cell doublet rates"},
        {"id": "single_cell_b",        "name": "scRNA-seq Analysis Workflow",                 "desc": "Seurat/Scanpy pipeline  -  filtering, normalisation, HVGs, PCA, UMAP, clustering, marker genes, cell type annotation"},
        {"id": "single_cell_c",        "name": "Advanced Single-Cell Methods",                "desc": "RNA velocity (scVelo), pseudotime (Monocle, PAGA), spatial transcriptomics (Visium, MERFISH), multiome ATAC+RNA"},
        {"id": "epigenomics_a",        "name": "DNA Methylation & Bisulfite Sequencing",      "desc": "CpG methylation biology, WGBS and RRBS protocols, bismark alignment, DMR analysis, methylation arrays (EPIC)"},
        {"id": "epigenomics_b",        "name": "Histone Modifications & ChIP-seq",            "desc": "ChIP-seq library prep, peak calling with MACS2, H3K27ac enhancers, H3K4me3 promoters, differential ChIP"},
        {"id": "epigenomics_c",        "name": "Open Chromatin: ATAC-seq",                    "desc": "ATAC-seq protocol, Tn5 tagmentation, nucleosome-free regions, footprinting, linking regulatory elements to genes"},
        {"id": "metagenomics_a",       "name": "16S rRNA Amplicon Sequencing",                "desc": "Hypervariable regions (V3-V4), OTU vs ASV (DADA2), alpha and beta diversity, QIIME2 workflow, phyloseq"},
        {"id": "metagenomics_b",       "name": "Shotgun Metagenomics",                        "desc": "Whole metagenome sequencing, MetaPhlAn4 taxonomy, HUMAnN3 functional profiling, MAG binning (MetaBAT, GTDB-Tk)"},
        {"id": "metagenomics_c",       "name": "Microbiome Analysis & Disease Links",         "desc": "Differential abundance (LEfSe, ANCOM), confounding variables, gut-disease associations (IBD, diabetes, cancer)"},
        {"id": "precision_medicine_a", "name": "Pharmacogenomics",                            "desc": "CYP2D6, CYP2C19, TPMT, DPYD  -  pharmacogenomic variants, CPIC guidelines, clinical PGx testing and implementation"},
        {"id": "precision_medicine_b", "name": "Polygenic Risk Scores",                       "desc": "PRS construction (LDpred2, PRSice), external validation, ancestry transferability, clinical deployment ethics"},
        {"id": "precision_medicine_c", "name": "Precision Oncology & Companion Diagnostics",  "desc": "Tumour mutation burden, MSI, HER2/BRCA testing, NGS panels (Foundation One), CDx co-development, biomarker-led trials"},
    ],
    "drug_discovery": [
        {"id": "pipeline_overview_a",  "name": "Target to IND: Discovery Stages",             "desc": "Target → hit → lead → development candidate; timelines (12-15 years), costs ($2B+), attrition rates at each stage"},
        {"id": "pipeline_overview_b",  "name": "Clinical Phases & Regulatory Milestones",     "desc": "Phase I/II/III/IV overview, IND filing, NDA/BLA submission, probability of success by phase and therapeutic area"},
        {"id": "pipeline_overview_c",  "name": "Portfolio Management & Value Creation",       "desc": "Risk-adjusted NPV, pipeline diversification, go/no-go decision frameworks, partnering vs in-house development strategy"},
        {"id": "target_id_a",          "name": "Disease Biology for Target Identification",   "desc": "Genetics (GWAS, Mendelian disease), proteomics, pathway analysis, phenotypic screening as target-agnostic approaches"},
        {"id": "target_id_b",          "name": "Target Classes & Druggability",               "desc": "GPCRs, kinases, ion channels, nuclear receptors, PPIs, RNA targets  -  binding pocket properties, druggability scoring tools"},
        {"id": "target_id_c",          "name": "Target Validation Strategies",                "desc": "RNAi/CRISPR knockdown, patient genetics as human genetic validation, mouse models, biomarker evidence, clinical proof-of-concept"},
        {"id": "hit_discovery_a",      "name": "Compound Libraries & Screening Collections", "desc": "Diversity sets, fragment libraries (FBDD), DNA-encoded libraries, natural products, virtual screening as alternative to physical HTS"},
        {"id": "hit_discovery_b",      "name": "High-Throughput Screening",                   "desc": "Assay formats (biochemical, cell-based), miniaturisation (384/1536-well), Z'-factor, hit rate expectations, false positive sources"},
        {"id": "hit_discovery_c",      "name": "Hit Characterisation & Triage",               "desc": "Dose-response confirmation, selectivity counter-screens, mechanism of action, physicochemical properties, hit-to-lead criteria"},
        {"id": "lead_optimization_a",  "name": "Structure-Activity Relationships",            "desc": "SAR principles, analogue synthesis, matched molecular pair analysis, scaffold hopping, bioisosteric replacements"},
        {"id": "lead_optimization_b",  "name": "ADMET-Guided Medicinal Chemistry",            "desc": "Lipinski Ro5, multiparameter optimisation, metabolic soft spots, P-gp efflux, hERG liability in lead optimisation"},
        {"id": "lead_optimization_c",  "name": "Selectivity & Polypharmacology",              "desc": "Kinome selectivity profiling (Eurofins), off-target panels, designed polypharmacology, selectivity windows and safety margins"},
        {"id": "admet_a",              "name": "Absorption & Bioavailability",                "desc": "Solubility assays, Caco-2 permeability, PAMPA, first-pass hepatic metabolism, oral bioavailability calculation and prediction"},
        {"id": "admet_b",              "name": "Distribution & Metabolism",                   "desc": "Plasma protein binding, volume of distribution, blood-brain barrier penetration, CYP reaction phenotyping, metabolite ID"},
        {"id": "admet_c",              "name": "Excretion, Toxicity & In Silico ADMET",       "desc": "Renal clearance, biliary excretion, hERG patch-clamp, DILI prediction, in silico tools (pkCSM, SwissADME, Derek Nexus)"},
        {"id": "pk_pd_a",              "name": "Core PK Parameters",                          "desc": "Cmax, AUC, t½, clearance (CL), volume of distribution (Vd)  -  one and two-compartment models, non-compartmental analysis"},
        {"id": "pk_pd_b",              "name": "Pharmacodynamic Models",                      "desc": "Emax model, EC50, Hill coefficient, direct vs indirect response models, hysteresis, PD biomarkers in drug development"},
        {"id": "pk_pd_c",              "name": "PK/PD Integration & Dose Selection",          "desc": "Exposure-response relationships, PK/PD-driven dose selection, translational PK/PD from animal to human, PKPD modelling tools"},
        {"id": "preclinical_a",        "name": "In Vitro Safety Assessment",                  "desc": "Genotoxicity (Ames test, micronucleus), hERG inhibition assay, in vitro hepatotoxicity, reactive metabolite trapping"},
        {"id": "preclinical_b",        "name": "In Vivo Toxicology Studies",                  "desc": "GLP repeat-dose tox, MTD, NOAEL, LOAEL, species selection rationale, pathology, clinical observations, recovery groups"},
        {"id": "preclinical_c",        "name": "IND-Enabling Package",                        "desc": "Safety pharmacology (core battery), reproductive toxicology, impurity qualification, manufacturing readiness, IND content requirements"},
        {"id": "biologics_a",          "name": "Monoclonal Antibody Structure",               "desc": "IgG subclasses (IgG1-4), Fab/Fc regions, CDR loops, Fc effector functions (ADCC, CDC, FcRn half-life extension)"},
        {"id": "biologics_b",          "name": "mAb Engineering & Advanced Formats",          "desc": "Humanisation (CDR grafting), affinity maturation (phage display), ADCs (linker/payload), bispecifics (CrossMAb, BiTE), nanobodies"},
        {"id": "biologics_c",          "name": "Biosimilars",                                 "desc": "Analytical similarity (fingerprint-like), clinical extrapolation rationale, regulatory pathway (351(k) FDA, Article 10 EMA), interchangeability"},
        {"id": "cgt_dd_a",             "name": "Viral Vectors in Drug Discovery Programs",   "desc": "AAV serotype selection for in vivo gene therapy, lentiviral transduction for ex vivo programs, vector engineering for specificity"},
        {"id": "cgt_dd_b",             "name": "CAR-T in Drug Discovery",                    "desc": "Target antigen selection (safety, tumour expression), CAR construct design decisions, autologous vs allogeneic manufacturing trade-offs"},
        {"id": "cgt_dd_c",             "name": "CRISPR in Drug Discovery",                   "desc": "Genome-wide loss-of-function screens, target validation with CRISPR KO, base editing for disease modelling, in vivo CRISPR therapies"},
        {"id": "comp_drug_disc_a",     "name": "Molecular Docking & Virtual Screening",      "desc": "Receptor preparation, grid generation, scoring functions, virtual screening workflow, enrichment metrics, docking pitfalls"},
        {"id": "comp_drug_disc_b",     "name": "QSAR & ML for Property Prediction",          "desc": "Molecular descriptor selection, QSAR model building, applicability domain, uncertainty quantification, free energy perturbation"},
        {"id": "comp_drug_disc_c",     "name": "Generative AI for Molecules",                "desc": "CVAE, junction tree VAE, diffusion-based molecular design (DiffSBDD), multi-parameter optimisation, wet-lab validation loop"},
        {"id": "biomarkers_dd_a",      "name": "Biomarker Types & Validation",               "desc": "Predictive, prognostic, PD, safety, and susceptibility biomarkers; fit-for-purpose analytical validation (FDA guidance)"},
        {"id": "biomarkers_dd_b",      "name": "Companion Diagnostic Development",           "desc": "CDx co-development strategy, IVD analytical validation (pre-analytic, analytic, clinical), regulatory submission alongside NDA/BLA"},
        {"id": "biomarkers_dd_c",      "name": "Patient Stratification & Biomarker-Led Trials","desc": "Biomarker-enriched designs, basket and umbrella trials, master protocols, adaptive enrichment, genomic prescreening logistics"},
        {"id": "cmc_formulation_a",    "name": "API Synthesis & Process Chemistry",          "desc": "Chemical synthetic routes, process chemistry optimisation, polymorphism and salt screening, impurity control strategy"},
        {"id": "cmc_formulation_b",    "name": "Drug Product Formulation & Stability",       "desc": "Dosage form selection, excipient roles, ICH Q1 stability studies, forced degradation, packaging and storage requirements"},
        {"id": "cmc_formulation_c",    "name": "CMC Regulatory Requirements & Scale-Up",    "desc": "ICH Q8 (pharmaceutical development), Q9 (risk management), Q10 (PQS), QbD design space, tech transfer to CMO/CDMO"},
    ],
    "clinical_trials": [
        {"id": "trial_basics_a",       "name": "Why Clinical Trials Exist",                   "desc": "Evidence hierarchy (RCT vs observational), regulatory mandate for efficacy and safety, historical context (thalidomide, Kefauver-Harris)"},
        {"id": "trial_basics_b",       "name": "The Clinical Development Roadmap",            "desc": "IND filing → Phase I → II → III → NDA/BLA → Phase IV; typical timelines, costs, and probability of success at each stage"},
        {"id": "trial_basics_c",       "name": "Key Stakeholders in Clinical Trials",         "desc": "Sponsor, CRO, investigative site, IRB/IEC, DSMB, FDA/EMA  -  roles, responsibilities, and contractual relationships"},
        {"id": "phase1_a",             "name": "Phase I Objectives & First-in-Human Ethics",  "desc": "Safety, tolerability, and PK as primary goals; healthy volunteer vs patient studies; MRSD determination (NOAEL/HED/MABEL)"},
        {"id": "phase1_b",             "name": "Dose Escalation Designs",                     "desc": "3+3 rule-based, accelerated titration, mTPI, CRM, BOIN  -  Bayesian vs rule-based trade-offs, DLT definition, MTD and RP2D"},
        {"id": "phase1_c",             "name": "PK/PD Integration in Phase I",                "desc": "PK sampling strategy, PK/PD modelling for dose selection, biomarker integration (PD endpoints), food effect and DDI studies"},
        {"id": "phase2_3_a",           "name": "Phase II: Proof of Concept & Dose Selection", "desc": "Phase IIa vs IIb, PoC objectives, signal-seeking designs, dose-response, go/no-go criteria, seamless Phase II/III"},
        {"id": "phase2_3_b",           "name": "Phase III: Confirmatory Trial Design",        "desc": "Randomisation methods (stratified, minimisation), double-blind designs, parallel vs crossover, control arm rationale"},
        {"id": "phase2_3_c",           "name": "Endpoint Selection Strategy",                 "desc": "Primary, secondary, PRO endpoints; FDA endpoint qualification, surrogate vs clinical endpoints, patient-relevant outcomes"},
        {"id": "trial_design_stats_a", "name": "Hypothesis Testing in Clinical Trials",       "desc": "Null vs alternative hypothesis, Type I error (α, two-sided vs one-sided), Type II error (β), power  -  regulatory conventions"},
        {"id": "trial_design_stats_b", "name": "Sample Size Calculation",                     "desc": "Effect size assumptions, variance estimates, dropout/withdrawal allowance, power simulations, sample size inflation strategies"},
        {"id": "trial_design_stats_c", "name": "Adaptive & Bayesian Designs",                 "desc": "Interim analyses, futility stopping (O'Brien-Fleming), sample size re-estimation, alpha spending (Lan-DeMets), platform trials"},
        {"id": "regulatory_bodies_a",  "name": "FDA Structure & Meeting Types",               "desc": "CDER/CBER/CDRH divisions, review divisions by therapeutic area, Type A/B/C meeting requests, pre-IND, end-of-Phase II meetings"},
        {"id": "regulatory_bodies_b",  "name": "EMA, PMDA & Global Agencies",                "desc": "EMA CHMP procedure, centralised vs decentralised authorisation, PMDA consultations, ANVISA, NMPA  -  regional differences"},
        {"id": "regulatory_bodies_c",  "name": "ICH Guidelines for Clinical Development",     "desc": "ICH E6 GCP, E8 general considerations, E9 statistical principles, E10 control group choice  -  how these shape protocol design"},
        {"id": "submissions_a",        "name": "IND Application",                             "desc": "IND content (pharmacology, tox, CMC, clinical protocol), investigator IND vs commercial IND, annual reports, protocol amendments"},
        {"id": "submissions_b",        "name": "NDA & BLA Submissions",                       "desc": "CTD format (Modules 1-5), NDA vs BLA distinction, 505(b)(1) vs 505(b)(2), PDUFA user fees, standard vs priority review timelines"},
        {"id": "submissions_c",        "name": "Post-Submission Interactions",                "desc": "Discipline review letters, information requests, advisory committee prep, complete response letters (CRL), resubmission strategies"},
        {"id": "gcp_ethics_a",         "name": "GCP Principles & Responsibilities",           "desc": "ICH E6(R2) framework, sponsor quality systems, investigator obligations, monitor (CRA) roles, essential document management"},
        {"id": "gcp_ethics_b",         "name": "Informed Consent Process",                   "desc": "8 required elements (21 CFR 50.25), process (not just form), vulnerable populations (prisoners, minors), re-consent triggers"},
        {"id": "gcp_ethics_c",         "name": "IRB/IEC & Research Ethics",                  "desc": "IRB composition and review types (expedited, full board), Declaration of Helsinki principles, global harmonisation challenges"},
        {"id": "special_pathways_a",   "name": "FDA Expedited Programs",                      "desc": "Breakthrough Therapy Designation (intensive FDA guidance), Fast Track (rolling review), Priority Review (6-month PDUFA goal)"},
        {"id": "special_pathways_b",   "name": "Accelerated Approval",                        "desc": "Surrogate and intermediate endpoints, post-approval confirmatory trial requirements, FDORA reforms, withdrawal precedents"},
        {"id": "special_pathways_c",   "name": "Global Expedited Pathways",                   "desc": "EMA PRIME scheme, conditional marketing authorisation, exceptional circumstances, SAKIGAKE in Japan, Health Canada Priority Review"},
        {"id": "phase4_pv_a",          "name": "Phase IV Commitments & Post-Approval Studies","desc": "REMS requirements, PMC vs PMR commitments, label expansion studies, paediatric investigation plans (EU), PREA/BPCA (US)"},
        {"id": "phase4_pv_b",          "name": "Pharmacovigilance & Adverse Event Reporting", "desc": "Spontaneous reporting systems (FAERS, EudraVigilance), MedWatch, CIOMS I/II forms, expedited vs periodic reporting"},
        {"id": "phase4_pv_c",          "name": "Signal Detection & Risk Management",          "desc": "PRR, ROR disproportionality methods, PBRER/PSUR aggregate reports, EU risk minimisation measures, REMS effectiveness assessments"},
        {"id": "dct_adaptive_a",       "name": "Decentralised Trial Technologies",            "desc": "eConsent, wearables for digital endpoints, home nursing, telehealth visits, direct-to-patient drug shipment, data integrity challenges"},
        {"id": "dct_adaptive_b",       "name": "Real-World Evidence in Drug Development",    "desc": "RWD sources (EHR, claims, registries), FDA RWE framework, study design for RWE, limitations vs RCT, Sentinal system"},
        {"id": "dct_adaptive_c",       "name": "Platform & Master Protocol Designs",         "desc": "Basket, umbrella, and platform trials, RECOVERY/REMAP-CAP as examples, shared control arm, biomarker-defined subpopulations"},
        {"id": "labeling_postmarket_a","name": "Prescribing Information Structure",           "desc": "Highlights, full PI sections (indication, dosage, warnings, clinical studies), labelling regulations (21 CFR 201), SPL format"},
        {"id": "labeling_postmarket_b","name": "Label Negotiation & FDA Review",              "desc": "Draft label submission, FDA proposed changes, sponsor response, disputed sections, label as commercial and clinical asset"},
        {"id": "labeling_postmarket_c","name": "Lifecycle Management & Exclusivity",         "desc": "New indication filings, line extensions, paediatric exclusivity, patent cliff planning, reference product strategies"},
        {"id": "patient_recruitment_a","name": "Site Selection & Feasibility",               "desc": "Enrolment feasibility assessment, site selection criteria, country and region strategy, site performance metrics and risk"},
        {"id": "patient_recruitment_b","name": "Patient Recruitment Strategies",             "desc": "Disease registries, social media campaigns, digital outreach, patient advocacy partnerships, electronic health record screening"},
        {"id": "patient_recruitment_c","name": "Retention, Adherence & Protocol Deviations", "desc": "Dropout prediction, retention strategies (patient stipends, convenience), protocol deviation classification and impact on data integrity"},
    ],
    "genai_ml": [
        {"id": "ml_foundations_a",     "name": "Supervised & Unsupervised Learning",          "desc": "Regression, classification, clustering, dimensionality reduction  -  core paradigms, loss functions, gradient descent optimisation"},
        {"id": "ml_foundations_b",     "name": "Model Evaluation & Validation",               "desc": "Train/val/test splits, k-fold cross-validation, ROC/AUC, precision-recall, calibration curves for biological datasets"},
        {"id": "ml_foundations_c",     "name": "Overfitting, Regularisation & Hyperparameters","desc": "Bias-variance trade-off, L1/L2 regularisation, dropout, early stopping, hyperparameter search (grid, random, Bayesian)"},
        {"id": "bio_feature_eng_a",    "name": "Molecular Descriptors & Fingerprints",        "desc": "Morgan/ECFP fingerprints, RDKit descriptors, pharmacophore features  -  encoding small molecules for ML models"},
        {"id": "bio_feature_eng_b",    "name": "Sequence Encoding Strategies",                "desc": "One-hot encoding, k-mer frequency, PSS matrices, learned embeddings  -  representing DNA, RNA, and protein sequences"},
        {"id": "bio_feature_eng_c",    "name": "Graph & Multi-Modal Representations",         "desc": "Molecular graphs, protein contact maps, knowledge graphs, combining sequence + structure + omics in multi-modal models"},
        {"id": "classical_ml_lifesci_a","name": "Ensemble Methods for Life Science Data",    "desc": "Random forests, XGBoost, LightGBM applied to ADMET prediction, gene expression classifiers, biomarker selection"},
        {"id": "classical_ml_lifesci_b","name": "SVMs & Logistic Regression in Biology",     "desc": "Kernel trick for molecular data, regularised logistic regression for omics, handling class imbalance in clinical datasets"},
        {"id": "classical_ml_lifesci_c","name": "Survival Analysis & Patient Stratification","desc": "Kaplan-Meier curves, Cox proportional hazards, time-varying covariates, clustering patients for precision medicine"},
        {"id": "deep_learning_fund_a", "name": "Neural Network Fundamentals",                 "desc": "Layers, activation functions, backpropagation, SGD/Adam, batch normalisation  -  building and training simple networks"},
        {"id": "deep_learning_fund_b", "name": "CNNs & RNNs for Biological Data",            "desc": "CNNs for sequence motif discovery and microscopy images, RNNs/LSTMs for sequential biological data, vanishing gradients"},
        {"id": "deep_learning_fund_c", "name": "Attention Mechanisms & Transformers",        "desc": "Self-attention, multi-head attention, positional encoding, the transformer architecture  -  foundation of modern sequence AI"},
        {"id": "protein_lang_models_a","name": "Protein Language Model Architecture",        "desc": "BERT-style masked language models for amino acids, tokenisation, pre-training objectives, transfer learning for protein tasks"},
        {"id": "protein_lang_models_b","name": "ESM, ProtTrans & ProteinBERT",              "desc": "ESM-2 embeddings, ProtTrans models, zero-shot mutation effect prediction, using PLM representations in downstream tasks"},
        {"id": "protein_lang_models_c","name": "AlphaFold2 & AF3: How They Work",           "desc": "MSA-based Evoformer, structure module, pLDDT/PAE confidence, AlphaFold3 diffusion head, practical limitations"},
        {"id": "gnn_drug_disc_a",      "name": "Graph Theory for Molecules",                 "desc": "Atoms as nodes, bonds as edges, atom/bond features, molecular graph construction  -  why graphs suit chemistry better than SMILES"},
        {"id": "gnn_drug_disc_b",      "name": "GNN Architectures",                         "desc": "Message passing neural networks (MPNN), SchNet, DimeNet, AttentiveFP  -  how 3D geometry improves molecular property prediction"},
        {"id": "gnn_drug_disc_c",      "name": "GNNs in Drug Discovery Applications",       "desc": "Molecular property prediction (ADMET, activity), drug-target interaction prediction, reaction outcome prediction, retrosynthesis"},
        {"id": "gen_molecules_a",      "name": "VAEs & GANs for Molecular Design",          "desc": "Variational autoencoders for latent space molecular optimisation, junction tree VAE, GAN training instability and mode collapse"},
        {"id": "gen_molecules_b",      "name": "Diffusion Models for Molecules",            "desc": "Score-based and DDPM diffusion, 3D molecular generation (EDM, GeoDiff), structure-based drug design (DiffSBDD, DiffDock)"},
        {"id": "gen_molecules_c",      "name": "Multi-Parameter Optimisation in Generative Chemistry","desc": "Pareto optimisation for potency+selectivity+ADMET, REINFORCE and RL for molecular generation, closed-loop design-make-test"},
        {"id": "ai_genomics_omics_a",  "name": "Variant Effect Prediction",                 "desc": "Deep learning for splicing (SpliceAI), regulatory effects (Enformer), missense pathogenicity (AlphaMissense), zero-shot with PLMs"},
        {"id": "ai_genomics_omics_b",  "name": "Gene Expression & Single-Cell Foundation Models","desc": "scGPT, Geneformer, Universal Cell Embeddings  -  pre-training on cell atlases, downstream tasks, zero-shot cell type transfer"},
        {"id": "ai_genomics_omics_c",  "name": "Multi-Omics Integration",                   "desc": "MOFA+, CITE-seq analysis, multi-modal VAEs, graph-based integration, finding regulatory programs across omics layers"},
        {"id": "ai_clinical_rwe_a",    "name": "Clinical NLP & EHR Mining",                 "desc": "Named entity recognition for clinical text, ICD code prediction, BERT-based clinical models (BioBERT, ClinicalBERT, Med-PaLM)"},
        {"id": "ai_clinical_rwe_b",    "name": "AI for Patient Recruitment & Digital Biomarkers","desc": "Trial eligibility matching (TriNetX, Veeva Vault), wearable data modelling, passive digital biomarkers, FDA DDT program"},
        {"id": "ai_clinical_rwe_c",    "name": "AI-Driven Real-World Evidence",             "desc": "Observational study design with AI, propensity scoring, confounding control, synthetic control arms, FDA RWE pilots"},
        {"id": "responsible_ai_a",     "name": "Bias & Fairness in Medical AI",             "desc": "Training data demographic gaps, performance disparities across subgroups, fairness metrics, bias auditing frameworks"},
        {"id": "responsible_ai_b",     "name": "Explainability in Healthcare AI",           "desc": "SHAP values, LIME, attention visualisation, model cards, clinician trust, explainability vs performance trade-offs"},
        {"id": "responsible_ai_c",     "name": "FDA AI/ML Regulatory Framework",            "desc": "Predetermined change control plan (PCCP), SaMD risk classification, continuous learning, IEC 62304, validation standards"},
        {"id": "ai_dd_platforms_a",    "name": "Recursion & Phenomics-Based AI",            "desc": "Morphological profiling (Cell Painting), Recursion OS, phenomics-guided target ID, dataset scale and model architecture"},
        {"id": "ai_dd_platforms_b",    "name": "Insilico Medicine & Schrödinger",           "desc": "Insilico PandaOmics + Chemistry42 generative platform, INS018_055 IND milestone; Schrödinger FEP+, WaterMap, Glide workflow"},
        {"id": "ai_dd_platforms_c",    "name": "Isomorphic Labs & Next-Gen Platforms",      "desc": "AlphaFold3 in drug design, protein-ligand structure prediction, Isomorphic's Eli Lilly/Novartis deals, next-generation AI pipeline companies"},
        {"id": "ml_pipelines_a",       "name": "MLOps for Life Sciences",                   "desc": "Experiment tracking (MLflow, W&B), data versioning (DVC), model registry, CI/CD for ML, reproducibility requirements"},
        {"id": "ml_pipelines_b",       "name": "Cloud Infrastructure & Data Governance",    "desc": "AWS/GCP/Azure life sciences architectures, HIPAA/GDPR in ML systems, data access controls, federated learning concepts"},
        {"id": "ml_pipelines_c",       "name": "Production ML in Regulated Environments",  "desc": "FDA SaMD guidance, audit trails, model validation documentation, IQ/OQ/PQ for ML systems, change control post-deployment"},
    ],
    "biotech_business": [
        {"id": "biotech_biz_model_a",  "name": "Biotech Company Structures & Revenue Models", "desc": "Pipeline vs platform vs service company models, royalty and milestone revenue, the fully integrated vs virtual biotech spectrum"},
        {"id": "biotech_biz_model_b",  "name": "Startup vs Big Pharma Economics",            "desc": "Burn rate, runway, headcount efficiency, cost per IND, how venture-backed biotechs differ from large pharma in resource allocation"},
        {"id": "biotech_biz_model_c",  "name": "The Biotech Ecosystem",                      "desc": "VC firms (a16z bio, OrbiMed, Atlas, Flagship), academic spin-outs, biotech hubs (Kendall Sq, Mission Bay, Stevenage), CRO/CDMO ecosystem"},
        {"id": "biotech_financing_a",  "name": "Seed & Series A Financing",                  "desc": "What seed-stage investors look for (founder/team, target validation), typical seed terms, pre-money valuation, Series A readiness"},
        {"id": "biotech_financing_b",  "name": "Series B/C, Crossover & Late-Stage Financing","desc": "Later-stage milestone packages, crossover investor role, bridge rounds, venture debt (Hercules, Oxford Finance), dilution management"},
        {"id": "biotech_financing_c",  "name": "IPO & Public Market Lifecycle",              "desc": "S-1 anatomy, NASDAQ/NYSE listing, lock-up periods, institutional vs retail investor base, post-IPO stock performance patterns"},
        {"id": "ls_valuation_a",       "name": "Risk-Adjusted NPV (rNPV)",                   "desc": "Probability of success by phase, discount rate selection (10-15% for biotech), cash flow modelling, terminal value approaches"},
        {"id": "ls_valuation_b",       "name": "Comparable Transactions & Market Multiples", "desc": "Precedent deal database (EvaluatePharma, BioCentury), upfront vs milestone vs royalty benchmarks, peak sales multiples"},
        {"id": "ls_valuation_c",       "name": "Communicating Pipeline Value to Investors",  "desc": "Investor presentation anatomy, pipeline table with catalysts, data read-out timelines, translating science into financial narrative"},
        {"id": "bd_licensing_a",       "name": "Deal Structures in Biotech",                 "desc": "Option deals, co-development/co-commercialisation, full license, collaboration agreements  -  economics of each structure"},
        {"id": "bd_licensing_b",       "name": "Term Sheets & Deal Economics",               "desc": "Upfront payment, milestones (development, regulatory, commercial), royalty tiers, sublicensing rights, diligence obligations"},
        {"id": "bd_licensing_c",       "name": "Due Diligence in BD Deals",                 "desc": "IP due diligence, clinical data package review, CMC/manufacturing diligence, commercial diligence, typical timeline and red flags"},
        {"id": "market_strategy_a",    "name": "Patient Population Sizing",                  "desc": "Epidemiology → diagnosed population → treatment-eligible → accessible market; bottom-up vs top-down market sizing"},
        {"id": "market_strategy_b",    "name": "Competitive Landscape Analysis",             "desc": "Approved products, pipeline competitors, differentiation (efficacy, safety, convenience), SWOT, indication priority"},
        {"id": "market_strategy_c",    "name": "Launch Strategy & Go-to-Market",             "desc": "Geography sequencing (US-first vs EU), indication prioritisation, KOL mapping, pre-launch activities, formulary access strategy"},
        {"id": "market_access_heor_a", "name": "The Global Payer Landscape",                 "desc": "US commercial payers, Medicare/Medicaid, EU national payers, formulary tiers, P&T committee dynamics, prior authorisation"},
        {"id": "market_access_heor_b", "name": "Health Economics & HTA",                    "desc": "QALY, ICER thresholds, cost-effectiveness models, NICE/G-BA/HAS/AIFA appraisal processes, HTA-aligned trial design"},
        {"id": "market_access_heor_c", "name": "Value-Based Contracts & Pricing Strategy",  "desc": "Outcomes-based contracts, annuity payment models for gene therapy, reference pricing, parallel trade, net price negotiations"},
        {"id": "ip_biotech_a",         "name": "Patent Types in Life Sciences",              "desc": "Composition of matter, method of use, formulation, process patents  -  scope, strength, and litigation value of each type"},
        {"id": "ip_biotech_b",         "name": "Patent Strategy & Lifecycle",               "desc": "Filing timing relative to publication, continuation and CIP strategy, Orange Book listing, patent term extension, patent cliffs"},
        {"id": "ip_biotech_c",         "name": "FTO, Licensing & Trade Secrets",            "desc": "Freedom-to-operate analysis, in-licensing technology IP, trade secrets as alternative to patents, IP in partnership agreements"},
        {"id": "mfg_ops_a",            "name": "GMP & Quality Systems",                     "desc": "FDA 21 CFR Part 211, EMA GMP Annex 1, GMP audit preparation, quality agreements, change control in manufacturing"},
        {"id": "mfg_ops_b",            "name": "CMO & CDMO Strategy",                       "desc": "CMO vs CDMO distinction, selection criteria, QAA scope, tech transfer packages, dual-sourcing for supply security"},
        {"id": "mfg_ops_c",            "name": "Supply Chain & Cold Chain Logistics",       "desc": "Cold chain requirements for biologics, serialisation, track-and-trace, demand forecasting, supply chain risk management"},
        {"id": "reg_strategy_biz_a",   "name": "Regulatory Risk in Financial Models",       "desc": "How probability of approval by pathway affects rNPV, regulatory risk events (CRL), modelling regulatory timelines in valuations"},
        {"id": "reg_strategy_biz_b",   "name": "Approval Pathways as Business Value Drivers","desc": "Breakthrough designation premium, label breadth vs narrow indication trade-offs, regulatory strategy shaping competitive position"},
        {"id": "reg_strategy_biz_c",   "name": "Post-Approval Commercial Lifecycle",        "desc": "REMS commercial impact, label update strategies, line extension economics, loss of exclusivity (LOE) planning and defensive moves"},
        {"id": "people_culture_a",     "name": "Biotech Org Design & Cross-Functional Teams","desc": "Programme team structures, matrix vs functional org, cross-functional team dynamics, governance frameworks (steering committees)"},
        {"id": "people_culture_b",     "name": "Recruiting & Retaining Scientists",         "desc": "Academic-to-industry transition, equity compensation (options vs RSUs), culture building in high-uncertainty environments"},
        {"id": "people_culture_c",     "name": "Board Governance & Investor Relations",     "desc": "Board composition (independent directors, investors, management), fiduciary duties, investor communication cadence, proxy advisory firms"},
        {"id": "digital_health_biz_a", "name": "Digital Therapeutics & SaMD",              "desc": "DTx definition and reimbursement challenges (Pear Therapeutics lessons), FDA De Novo/510(k) for SaMD, prescription vs OTC DTx"},
        {"id": "digital_health_biz_b", "name": "Health Data as a Business Asset",          "desc": "Real-world data platforms, data partnerships, privacy regulations (HIPAA, GDPR), data monetisation models, synthetic data"},
        {"id": "digital_health_biz_c", "name": "AI Company Business Models in Life Sciences","desc": "Software-as-a-service vs milestone deals, partnership economics with pharma, build-vs-buy decisions, recurring revenue vs one-time"},
        {"id": "building_biotech_a",   "name": "Spinning Out from Academia",               "desc": "IP assignment agreements, university TTO relationships, founder equity, conflict-of-interest management, lab-to-company transition"},
        {"id": "building_biotech_b",   "name": "Founding Team & Seed Round",               "desc": "Ideal founding team composition (scientist + operator), equity splits, SAFE vs priced round, what seed investors evaluate"},
        {"id": "building_biotech_c",   "name": "Series A Readiness",                       "desc": "Data package for Series A (what's sufficient), narrative construction, lead investor identification, milestones that drive value inflection"},
    ],
    "cell_gene_therapy": [
        {"id": "cgt_foundations_a",    "name": "What Are Cell & Gene Therapies?",             "desc": "How CGTs differ from small molecules and biologics, types (gene addition, silencing, editing, cell therapy), approved products landscape"},
        {"id": "cgt_foundations_b",    "name": "Approved CGT Products Deep Dive",             "desc": "Zolgensma (AAV9-SMN1), Luxturna (RPE65), Casgevy (CRISPR-SCD/β-thal), Kymriah/Yescarta (CAR-T)  -  mechanisms and clinical results"},
        {"id": "cgt_foundations_c",    "name": "ATMP Regulatory Category",                    "desc": "EU ATMP classification (somatic CT, gene therapy, TEP), CBER vs CDER jurisdiction, RMAT and PRIME designations"},
        {"id": "viral_vectors_a",      "name": "AAV Biology & Serotype Selection",            "desc": "Capsid structure, receptor interactions, tissue tropism by serotype (AAV9 CNS/muscle, AAV8 liver, AAV5 eye), packaging capacity (~4.7 kb)"},
        {"id": "viral_vectors_b",      "name": "Lentiviral & Retroviral Vectors",             "desc": "LV integration mechanism, SIN design, insertional mutagenesis history, γ-retrovirus vs lentivirus safety profile, ex vivo use cases"},
        {"id": "viral_vectors_c",      "name": "Adenovirus & Next-Gen Vectors",               "desc": "Adenovirus immunogenicity, transient expression, vaccine applications (ChAdOx1), helper-dependent AdV, AAV capsid engineering approaches"},
        {"id": "nonviral_delivery_a",  "name": "Lipid Nanoparticles for Gene Delivery",      "desc": "LNP ionisable lipid design, endosomal escape mechanism, organ tropism (MC3 liver vs 9A1P1 lung), LNP stability and immunogenicity"},
        {"id": "nonviral_delivery_b",  "name": "Electroporation & Physical Delivery",        "desc": "Electroporation for ex vivo cell transfection (MaxCyte), hydrodynamic injection, sonoporation, nucleofection for T cells and HSCs"},
        {"id": "nonviral_delivery_c",  "name": "Polymer Nanoparticles & Emerging Delivery",  "desc": "PEI, PLGA, lipid-polymer hybrids, GalNAc conjugates for hepatocyte targeting, VLPs, limitations vs viral vectors"},
        {"id": "crispr_mechanisms_a",  "name": "Cas9 Mechanism & Guide RNA Design",          "desc": "SpCas9 RuvC/HNH domains, NGG PAM recognition, spacer design rules, guide RNA secondary structure, efficacy prediction tools"},
        {"id": "crispr_mechanisms_b",  "name": "Off-Target Effects & Minimisation",          "desc": "Off-target prediction (CRISPOR, Cas-OFFinder), unbiased detection (GUIDE-seq, CIRCLE-seq), high-fidelity Cas9 variants (eSpCas9, HiFi)"},
        {"id": "crispr_mechanisms_c",  "name": "Variant Cas Systems",                        "desc": "Cas12a (staggered cuts, T-rich PAM), Cas13 (RNA targeting), CasRx, nickases (D10A)  -  clinical and research applications"},
        {"id": "base_prime_editing_a", "name": "Cytosine & Adenine Base Editors",            "desc": "CBE mechanism (cytidine deaminase + Cas9-D10A), C→T editing, ABE mechanism (adenosine deaminase), A→G conversions, activity window"},
        {"id": "base_prime_editing_b", "name": "Prime Editing",                              "desc": "pegRNA design (spacer + RT template + PBS), PE2/PE3/PE3b systems, small insertions/deletions/all 12 base changes, efficiency vs CBE/ABE"},
        {"id": "base_prime_editing_c", "name": "Clinical Applications of Precision Editing", "desc": "BE4max for SCD (Beam Therapeutics), ABE8e for TTR amyloidosis, prime editing for PRNP  -  current IND filings and clinical status"},
        {"id": "cart_therapy_a",       "name": "CAR Construct Architecture",                 "desc": "scFv antigen-binding domain, hinge and transmembrane regions, CD28 vs 4-1BB costimulatory signalling, CD3ζ, armoured CARs"},
        {"id": "cart_therapy_b",       "name": "T Cell Manufacturing Process",               "desc": "Leukapheresis, T cell activation (CD3/CD28 beads), lentiviral transduction, expansion (G-Rex), harvest, QC, cryopreservation"},
        {"id": "cart_therapy_c",       "name": "Autologous vs Allogeneic CAR-T",             "desc": "Autologous limitations (manufacturing time, cost, vein-to-vein), allogeneic off-the-shelf approaches, TALEN/CRISPR for TCR/HLA knockout"},
        {"id": "other_cell_therapies_a","name": "NK Cell & iPSC-Derived Therapies",         "desc": "NK cell advantages (no GvHD, off-the-shelf), FT596 iPSC-derived NK, CD16 engineering, NK cell activation and persistence challenges"},
        {"id": "other_cell_therapies_b","name": "TCR-T Cell Therapy",                        "desc": "Endogenous TCR knockdown, HLA restriction, neoantigen vs shared antigen targets, Adaptimmune (SPEAR T cells), afami-cel"},
        {"id": "other_cell_therapies_c","name": "HSC Gene Correction & TIL Therapy",        "desc": "HSC mobilisation/collection, lentiviral correction (SCD, ADA-SCID, CGD), TIL therapy for solid tumours (Lifileucel), manufacturing scale"},
        {"id": "invivo_exvivo_a",      "name": "Liver & Eye In Vivo Gene Therapy",           "desc": "AAV8/9 hepatic tropism (Hemgenix, BioMarin valoctocogene), subretinal vs intravitreal delivery, dose, immunosuppression needs"},
        {"id": "invivo_exvivo_b",      "name": "CNS, Muscle & Systemic In Vivo Delivery",   "desc": "Intrathecal/intracerebroventricular AAV, Zolgensma systemic AAV9, Sarepta DMD programs, dose-immunogenicity trade-offs"},
        {"id": "invivo_exvivo_c",      "name": "Ex Vivo HSC Editing Workflow",              "desc": "G-CSF mobilisation, apheresis, electroporation-based editing, myeloablative conditioning, infusion, engraftment monitoring"},
        {"id": "cgt_immunogenicity_a", "name": "Pre-Existing AAV Immunity",                 "desc": "NAb prevalence by serotype and geography, patient screening assays, NAb exclusion criteria, strategies for re-dosing or switching"},
        {"id": "cgt_immunogenicity_b", "name": "Insertional Mutagenesis & Genotoxicity",    "desc": "Integration site analysis (INSPIIRED), γ-retrovirus safety incidents (X-SCID), LV safety improvements, monitoring for clonal expansion"},
        {"id": "cgt_immunogenicity_c", "name": "Immune Tolerance Strategies",               "desc": "Transient immunosuppression (prednisolone, rituximab, bortezomib), capsid engineering to evade immune detection, empty capsid depletion"},
        {"id": "cgt_manufacturing_a",  "name": "AAV Vector Production",                     "desc": "Triple transfection HEK293 vs baculovirus/Sf9 system, upstream optimisation, downstream purification (CsCl vs affinity), titer and purity specs"},
        {"id": "cgt_manufacturing_b",  "name": "Cell Therapy Vein-to-Vein Workflow",        "desc": "Apheresis scheduling, closed-system leukapheresis, GMP transduction and expansion, cryopreservation, cold chain logistics, vein-to-vein time"},
        {"id": "cgt_manufacturing_c",  "name": "Critical Quality Attributes for CGTs",      "desc": "AAV: capsid full/empty ratio, residual DNA/protein; Cell therapy: viability, potency assay (cytotoxicity), identity (flow), transduction efficiency"},
        {"id": "cgt_regulatory_a",     "name": "FDA CBER & RMAT Designation",               "desc": "RMAT criteria and benefits (intensive FDA interaction), pre-BLA meetings, 15-year long-term follow-up requirement, CBER guidance documents"},
        {"id": "cgt_regulatory_b",     "name": "EMA CAT & ATMP Regulation",                "desc": "CAT scientific review, PRIME scheme for ATMPs, hospital exemption, conditional marketing authorisation, EMA compassionate use"},
        {"id": "cgt_regulatory_c",     "name": "CMC Challenges for CGTs",                  "desc": "Comparability without traditional analytical similarity, adventitious agent testing, lot release assays, potency challenges, container closure"},
        {"id": "cgt_clinical_a",       "name": "Phase I/II Design for CGTs",               "desc": "Modified 3+3 and mTPI for gene therapy, patient selection (no prior gene therapy), long-term follow-up cohorts, endpoint selection"},
        {"id": "cgt_clinical_b",       "name": "Long-Term Follow-Up Requirements",         "desc": "FDA 15-year LTFU requirement, patient registry design, late-onset adverse events (insertional mutagenesis, immune reactions), data challenges"},
        {"id": "cgt_clinical_c",       "name": "CGT Economics & Access",                   "desc": "One-time curative pricing rationale, outcomes-based annuity payments, payer challenges, Zolgensma at $2.1M, global access equity"},
    ],
    "protein_engineering": [
        {"id": "protein_eng_found_a",  "name": "Structure Determines Function",               "desc": "How active sites, binding interfaces, allosteric pockets, and flexible loops define what a protein does  -  and can be engineered to do"},
        {"id": "protein_eng_found_b",  "name": "Sequence-Structure-Function Relationships",  "desc": "How single mutations propagate structurally, epistasis in protein fitness landscapes, deep mutational scanning as empirical mapping"},
        {"id": "protein_eng_found_c",  "name": "Engineering Objectives & Trade-Offs",        "desc": "Stability vs activity trade-off, thermostabilisation strategies, expression yield as engineering target, half-life extension considerations"},
        {"id": "directed_evolution_a", "name": "Diversity Generation Methods",               "desc": "Error-prone PCR (epPCR), DNA shuffling, StEP, OmniChange, chemical mutagenesis  -  how to create large, random and semi-random libraries"},
        {"id": "directed_evolution_b", "name": "Display & Selection Technologies",           "desc": "Phage display, yeast surface display, ribosome display, mRNA display  -  genotype-phenotype linkage strategies and selection pressure"},
        {"id": "directed_evolution_c", "name": "SELEX & In Vitro Evolution",                 "desc": "Systematic evolution of ligands by exponential enrichment, aptamer selection, continuous directed evolution (PACE), machine learning-guided DE"},
        {"id": "rational_design_a",    "name": "Structure-Based Mutagenesis",                "desc": "Reading electron density maps for active site engineering, FoldX and Rosetta ΔΔG prediction, thermostabilising mutations from structure"},
        {"id": "rational_design_b",    "name": "Computational Alanine Scanning",             "desc": "Hot spot identification, Robetta alanine scanning, energy decomposition, experimental validation of predicted hot spots"},
        {"id": "rational_design_c",    "name": "Semi-Rational Design Strategies",            "desc": "Combining structural insight with focused library screening (ISM, CAST), consensus sequence design, ancestral sequence reconstruction"},
        {"id": "alphafold_practical_a","name": "How AlphaFold2 Works",                      "desc": "MSA depth requirements, Evoformer architecture, triangle updates, structure module, recycling iterations  -  understanding the model"},
        {"id": "alphafold_practical_b","name": "Interpreting AlphaFold Confidence Scores",  "desc": "pLDDT per-residue confidence, PAE for domain orientations, predicted TM-score, reliable vs uncertain regions in engineering use"},
        {"id": "alphafold_practical_c","name": "AlphaFold Limitations for Protein Engineering","desc": "Conformational states problem, ligand-bound vs apo structures, novel backbone accuracy, MSA scarcity effects, AF3 improvements"},
        {"id": "rfdiffusion_a",        "name": "Diffusion Models for Protein Backbone Design","desc": "Forward noise process, reverse denoising, RFdiffusion architecture, conditioning on motifs/constraints, symmetric protein design"},
        {"id": "rfdiffusion_b",        "name": "RFdiffusion Applications",                   "desc": "De novo binder design, enzyme active site scaffolding, cyclic peptide design, protein-protein interface design, experimental success rates"},
        {"id": "rfdiffusion_c",        "name": "From In Silico to In Vitro: Validation Pipeline","desc": "Rosetta energy filtering, ProteinMPNN sequence design, AlphaFold2 structure prediction filter, expression screening, wet-lab hit rate"},
        {"id": "proteinmpnn_a",        "name": "Inverse Folding Concept",                   "desc": "Designing amino acid sequences that fold into a given backbone  -  why this is hard, how graph neural networks solve it, training data"},
        {"id": "proteinmpnn_b",        "name": "ProteinMPNN Architecture",                  "desc": "Graph encoding of backbone geometry, edge features (Cα-Cα distances, dihedral angles), autoregressive decoding, tied design for multimers"},
        {"id": "proteinmpnn_c",        "name": "ProteinMPNN + RFdiffusion Design Pipelines","desc": "Practical workflow: RFdiffusion backbone → ProteinMPNN sequences → AF2 filter → expression → biophysical characterisation"},
        {"id": "antibody_engineering_a","name": "mAb Structure & CDR Engineering",          "desc": "VH/VL domain organisation, CDR-H3 length distribution, germline selection impact, structural basis of antigen recognition"},
        {"id": "antibody_engineering_b","name": "Humanisation & Affinity Maturation",       "desc": "CDR grafting, vernier position back-mutations, phage display affinity maturation, yeast display for de novo discovery"},
        {"id": "antibody_engineering_c","name": "Bispecific Formats & Fc Engineering",      "desc": "CrossMAb, DART, BiTE, DuoBody, IgG-like bispecifics; Fc engineering for ADCC/CDC modulation, FcRn-extended half-life, YTE/LS variants"},
        {"id": "enzyme_engineering_a", "name": "Thermostability & Expression Engineering",  "desc": "Consensus mutagenesis, disulfide introduction, proline substitution, salt bridges  -  systematic approaches to thermostabilisation"},
        {"id": "enzyme_engineering_b", "name": "Activity & Selectivity Engineering",        "desc": "Active site mutagenesis for substrate scope, cofactor engineering, substrate tunnel design, enantioselectivity improvement"},
        {"id": "enzyme_engineering_c", "name": "Industrial Biocatalysis",                   "desc": "Directed evolution for process conditions (organic solvents, extreme pH), Codexis and Arzeda approaches, cascade enzyme design"},
        {"id": "ppi_design_a",         "name": "PPI Hot Spot Analysis",                     "desc": "Experimental (alanine scanning) and computational (Rosetta, FoldX) hot spot mapping, structural biology of binding interfaces"},
        {"id": "ppi_design_b",         "name": "Designing PPI Inhibitors",                  "desc": "α-helix mimetics, stapled peptides, constrained peptides, macrocycles  -  strategies to disrupt hot spot interactions therapeutically"},
        {"id": "ppi_design_c",         "name": "Miniproteins & Peptide Binders",            "desc": "Lumazine synthase binders, WORM scaffold, RFdiffusion-designed miniprotein binders, peptide therapeutics, oral delivery challenges"},
        {"id": "developability_a",     "name": "Aggregation & Colloidal Stability",         "desc": "Hydrophobic patch analysis, AC-SINS, DLS, SEC-MALS, thermal denaturation  -  predicting and reducing aggregation propensity"},
        {"id": "developability_b",     "name": "Chemical Degradation & Sequence Liabilities","desc": "Deamidation (NG, NS), oxidation (Met, Trp), isomerisation (DG), fragmentation (DP, DS)  -  identification and mitigation strategies"},
        {"id": "developability_c",     "name": "Immunogenicity Prediction & Mitigation",   "desc": "In silico T-cell epitope prediction (EpiMatrix, iTope), MHC-II binding assessment, immunogenicity risk scoring, deimmunisation"},
        {"id": "therapeutic_formats_a","name": "Half-Life Extension Strategies",           "desc": "Fc fusion proteins, PEGylation, albumin binding domains, XTEN fusion, FcRn engineering (YTE, LS, Halozyme ENHANZE)  -  pros and cons"},
        {"id": "therapeutic_formats_b","name": "ADC Design Principles",                    "desc": "Antibody selection, linker chemistry (cleavable vs non-cleavable), payload selection (MMAE, DM1, PBD), DAR, site-specific conjugation"},
        {"id": "therapeutic_formats_c","name": "Multispecific & Novel Protein Formats",    "desc": "Trispecifics, nanobodies (single-domain VHH), DARPins, affibodies, Centyrins, monobodies  -  when novel formats beat conventional mAbs"},
        {"id": "protein_validation_a", "name": "Biophysical Characterisation Methods",     "desc": "SPR (KD, kon, koff), ITC (thermodynamics), DSF/nanoDSF (Tm), MST (KD in solution)  -  matching method to question"},
        {"id": "protein_validation_b", "name": "Structural Validation Techniques",         "desc": "X-ray co-crystallography, cryo-EM for complexes, HDX-MS for epitope mapping, NMR for conformational dynamics, SAXS"},
        {"id": "protein_validation_c", "name": "Cell & In Vivo Efficacy Assays",           "desc": "Target engagement assays (NanoBRET, CETSA), cell-based potency, MOA confirmation, PK/PD in mouse models, translational decision-making"},
    ],
    "rna_therapeutics": [
        {"id": "rna_bio_found_a",      "name": "RNA Classes & Their Functions",               "desc": "mRNA, tRNA, rRNA, lncRNA, miRNA, siRNA, circRNA, piRNA  -  biogenesis, subcellular localisation, and functional roles"},
        {"id": "rna_bio_found_b",      "name": "RNA Secondary & Tertiary Structure",          "desc": "Stem-loops, hairpins, pseudoknots, G-quadruplexes, riboswitches, SHAPE probing  -  RNA folds as drug targets"},
        {"id": "rna_bio_found_c",      "name": "Why RNA Expands the Druggable Space",         "desc": "Undruggable protein targets, RAS transcripts, pre-mRNA splicing intervention, transient expression control, quantitative target modulation"},
        {"id": "mrna_design_a",        "name": "5' Cap Structures & Translation Initiation",  "desc": "Cap-0 vs cap-1 methylation, CleanCap AG co-transcriptional capping, cap analogue effect on translation efficiency and innate immune evasion"},
        {"id": "mrna_design_b",        "name": "UTR Engineering & ORF Optimisation",          "desc": "5'/3' UTR stability elements, Kozak context, secondary structure near start codon, codon optimisation algorithms (CAI, tAI, Codon Tools)"},
        {"id": "mrna_design_c",        "name": "Poly-A Tail & mRNA Stability",               "desc": "Enzymatic vs encoded poly-A, optimal tail length (120-150 A), 3' UTR regulatory elements, circular mRNA as alternative for stability"},
        {"id": "rna_delivery_lnp_a",   "name": "LNP Component Roles",                        "desc": "Ionisable lipid (pKa <6.5 for endosomal escape), DSPC helper lipid, cholesterol for membrane fluidity, PEG-lipid for stability and stealth"},
        {"id": "rna_delivery_lnp_b",   "name": "LNP Formulation & Characterisation",         "desc": "Microfluidic mixing, particle size (PDI), encapsulation efficiency (RiboGreen), cryogenic storage, analytical methods for lot release"},
        {"id": "rna_delivery_lnp_c",   "name": "Organ Tropism Engineering",                  "desc": "MC3/DLin-MC3 liver tropism via ApoE, 9A1P1 lung targeting, spleen-targeting LNPs for immune cells, selective organ targeting (SORT)"},
        {"id": "sirna_rnai_a",         "name": "RNAi Mechanism & RISC Loading",              "desc": "Dicer processing, RISC assembly (AGO2), guide strand thermodynamic asymmetry rules, seed region off-targets, mismatch tolerance"},
        {"id": "sirna_rnai_b",         "name": "Chemical Modifications for Stability",       "desc": "2'-OMe and 2'-F for nuclease resistance, phosphorothioate backbone, end caps, GalNAc-siRNA conjugates for hepatocyte targeting"},
        {"id": "sirna_rnai_c",         "name": "Approved siRNA Drugs",                       "desc": "Onpattro (LNP-siRNA, TTR), Givlaari (GalNAc-siRNA, AHP), Inclisiran (PCSK9, twice-yearly dosing), Vutrisiran  -  mechanism and clinical data"},
        {"id": "aso_therapeutics_a",   "name": "RNase H Gapmers",                           "desc": "DNA gap flanked by 2'-modified wings (2'-MOE, LNA, cEt), RNase H recruitment mechanism, hepatic accumulation, naked ASO delivery"},
        {"id": "aso_therapeutics_b",   "name": "Steric-Block & Splice-Switching ASOs",      "desc": "Exon skipping for DMD (eteplirsen, golodirsen), SMN2 exon inclusion (nusinersen), exon exclusion for progeria (lonafarnib comparison)"},
        {"id": "aso_therapeutics_c",   "name": "CNS Delivery & Clinical Success Stories",   "desc": "Intrathecal delivery (nusinersen, tofersen for SOD1-ALS), CSF distribution, repeat dosing by LP, Huntingtin-targeting ASOs in trials"},
        {"id": "mirna_therapeutics_a", "name": "miRNA Biogenesis & Function",               "desc": "Drosha/DGCR8 pri-miRNA processing, Dicer pre-miRNA cleavage, RISC-mediated translational repression, one miRNA hundreds of targets"},
        {"id": "mirna_therapeutics_b", "name": "miRNA Mimics & AntagomiRs",                 "desc": "miRNA replacement for downregulated tumour suppressors, antimiRs for upregulated oncomiRs, miravirsen (miR-122), cobomarsen (miR-155)"},
        {"id": "mirna_therapeutics_c", "name": "Pleiotropic Targeting Challenges",          "desc": "Off-target effects of miRNA modulation, biomarker development for miRNA therapeutics, delivery bottleneck for mimics, clinical development status"},
        {"id": "mrna_vaccines_a",      "name": "Vaccine Antigen Design",                    "desc": "Spike protein optimisation (2P mutations, furin cleavage abolition), antigen display on nanoparticles, signal peptide selection"},
        {"id": "mrna_vaccines_b",      "name": "Immune Response to mRNA Vaccines",         "desc": "Innate sensing balance, TLR-mediated adjuvanticity, T follicular helper cells, germinal centre reaction, durability of antibody response"},
        {"id": "mrna_vaccines_c",      "name": "BNT162b2 & mRNA-1273 Case Studies",        "desc": "Key design decisions, clinical trial outcomes, variant adaptations (bivalent boosters), manufacturing at scale, global deployment lessons"},
        {"id": "chemical_mods_a",      "name": "Pseudouridine: The Karikó & Weissman Discovery","desc": "How Ψ substitution suppresses TLR7/8 activation, effect on translation efficiency, the Nobel Prize-winning insight and its development"},
        {"id": "chemical_mods_b",      "name": "N1-Methylpseudouridine (m1Ψ)",             "desc": "Superior immunosuppression vs Ψ, enhanced ribosome engagement, why BNT162b2 and mRNA-1273 use m1Ψ, impact on protein yield"},
        {"id": "chemical_mods_c",      "name": "Backbone Modifications: 2'-OMe, 2'-F, PS", "desc": "2'-OMe and 2'-F stability, nuclease resistance, effect on RISC loading, phosphorothioate backbone in ASOs vs RNA, modification interplay"},
        {"id": "circular_sarna_a",     "name": "Circular RNA Therapeutics",                "desc": "Back-splicing mechanism, PIE strategy for production, IRES-driven translation, exonuclease resistance, Orna Therapeutics, Laronde Obi platform"},
        {"id": "circular_sarna_b",     "name": "Self-Amplifying RNA",                      "desc": "Alphavirus replicon design, nsP1-4 replicase, sub-genomic promoter for antigen, saRNA dose-sparing, LUNAR LNP delivery, clinical data"},
        {"id": "circular_sarna_c",     "name": "Next-Generation mRNA Platforms",           "desc": "trans-amplifying RNA, circular-saRNA hybrids, in vivo base editing mRNA, mRNA-encoded gene editors, Laronde eRNA technology"},
        {"id": "rna_immunostim_a",     "name": "Innate Immune Sensing of RNA",             "desc": "TLR3 (dsRNA), TLR7/8 (ssRNA), RIG-I (5'-triphosphate RNA), MDA5 (long dsRNA)  -  how different RNA features trigger immune activation"},
        {"id": "rna_immunostim_b",     "name": "Suppressing Immunogenicity for Therapeutics","desc": "Modified nucleosides, sequence engineering (AU-rich element removal), dsRNA removal by HPLC, innate sensing and IFN response control"},
        {"id": "rna_immunostim_c",     "name": "Leveraging Immunostimulation for Vaccines","desc": "Adjuvant-free mRNA vaccines, self-adjuvanting saRNA, innate immune activation as an asset for immunisation, mucosal mRNA delivery"},
        {"id": "rna_clinical_dev_a",   "name": "Phase I Design for RNA Therapeutics",      "desc": "Dose escalation in oligonucleotide trials, tissue accumulation and hepatotoxicity monitoring, renal accumulation of ASOs, PK sampling"},
        {"id": "rna_clinical_dev_b",   "name": "Extrahepatic Delivery Challenges",         "desc": "CNS, lung, muscle, tumour  -  delivery bottlenecks beyond liver, conjugate strategies (GalNAc, folate, antibody), progress in each tissue"},
        {"id": "rna_clinical_dev_c",   "name": "Regulatory Precedents in RNA Drugs",       "desc": "Alnylam NDA (Onpattro), Ionis NDA (nusinersen), Moderna BLA (mRNA-1273)  -  CMC, clinical, and regulatory novelty each had to resolve"},
        {"id": "rna_platforms_a",      "name": "Moderna: Platform Breadth & Strategy",     "desc": "Investigational mRNA pipeline (personalised cancer vaccines, CMV, HIV), manufacturing platform, LNP technology, revenue diversification"},
        {"id": "rna_platforms_b",      "name": "Alnylam & Ionis: GalNAc & ASO Platforms", "desc": "Alnylam GalNAc-siRNA franchise (ATTR, PH1, hATTR), Ionis CNS-focused ASO pipeline, business models and delivery differentiation"},
        {"id": "rna_platforms_c",      "name": "Emerging RNA Platform Companies",          "desc": "Arrowhead (AREG), Silence Therapeutics (mRNAi GOLD), Laronde (eRNA), Orna Therapeutics (oRNA), Wave Life Sciences  -  differentiation"},
    ],
    "biomanufacturing": [
        {"id": "biomanuf_intro_a",     "name": "The Biologics Manufacturing Landscape",       "desc": "Biologics product types (mAbs, enzymes, vaccines, CGTs), market size, global manufacturing capacity, the biologics supply chain"},
        {"id": "biomanuf_intro_b",     "name": "GMP as Regulatory Framework",                "desc": "GMP vs GLP vs GCP distinctions, 21 CFR Part 211 (FDA) and EU GMP Annex 1/2, the quality system model, inspection consequences"},
        {"id": "biomanuf_intro_c",     "name": "Manufacturing Process as Part of the Drug",  "desc": "Process and product are inseparable for biologics, why process changes require comparability, CQA concept introduction"},
        {"id": "cell_line_dev_a",      "name": "Host Cell Selection",                        "desc": "CHO (glycosylation, productivity), HEK293 (transient expression, viral vectors), E. coli (no glycosylation, IBs), yeast (P. pastoris, secretion)"},
        {"id": "cell_line_dev_b",      "name": "Stable Cell Line Generation",               "desc": "Stable transfection (linearised plasmid, PiggyBac), selection (MTX amplification, GS system), single-cell cloning by FACS or limiting dilution"},
        {"id": "cell_line_dev_c",      "name": "Clone Screening & Cell Bank Establishment", "desc": "High-throughput mini-bioreactor screening, product quality analytics (CIEX, SEC), MCB/WCB establishment, ICH Q5B characterisation"},
        {"id": "upstream_biopro_a",    "name": "Bioreactor Modes & Selection",              "desc": "Batch vs fed-batch vs perfusion (TFF, ATF)  -  volumetric productivity, product quality trade-offs, single-use vs stainless steel"},
        {"id": "upstream_biopro_b",    "name": "Critical Process Parameters",               "desc": "Dissolved oxygen (kLa), pH, temperature, agitation, dissolved CO2  -  setpoints, control strategies, sensitivity to process conditions"},
        {"id": "upstream_biopro_c",    "name": "Metabolic Monitoring & Scale-Up",           "desc": "Online metabolite monitoring (glucose, lactate, amino acids), metabolic shift, off-gas analysis, geometric similarity rules for scale-up"},
        {"id": "media_feed_a",         "name": "Chemically Defined Media Development",      "desc": "CDM components, supplier qualification, hydrolysate vs CDM trade-offs, DOE-driven media development, proprietary media platforms"},
        {"id": "media_feed_b",         "name": "Fed-Batch Feeding Strategies",              "desc": "Glucose-controlled feeding (exponential, constant, feedback), amino acid bolus vs continuous feeds, cell-specific consumption rates"},
        {"id": "media_feed_c",         "name": "Metabolic Profiling & Quality Trade-Offs",  "desc": "Lactate accumulation and pH effects, ammonia toxicity, galactose-shifted feeding for glycosylation control, titer vs product quality optimisation"},
        {"id": "downstream_proc_a",    "name": "Protein A Affinity Capture",               "desc": "Protein A resin selection (MabSelect, Eshmuno A), loading, wash, low-pH elution, CIP with NaOH, resin lifetime qualification"},
        {"id": "downstream_proc_b",    "name": "Polishing Chromatography Steps",           "desc": "AEX flow-through for HCPs and DNA, CEX for charge variants, HIC for aggregates and variants  -  bind-and-elute vs flow-through modes"},
        {"id": "downstream_proc_c",    "name": "Viral Clearance & Final Filtration",       "desc": "Low-pH viral inactivation (≥60 min, pH 3.5), nanofiltration (Planova 20N), UF/DF for buffer exchange and concentration, sterile filtration"},
        {"id": "analytical_qc_a",      "name": "Critical Quality Attributes for mAbs",     "desc": "Glycosylation (G0F, G1F, afucosylation), aggregation (HMWS), charge variants (acidic/basic species), potency  -  ICH Q6B specs"},
        {"id": "analytical_qc_b",      "name": "Analytical Method Development & Validation","desc": "CIEX for charge variants, SEC for aggregation, CEX-MS for intact mass, potency bioassay design, ICH Q2(R1) validation parameters"},
        {"id": "analytical_qc_c",      "name": "Comparability Testing Principles",         "desc": "When comparability is needed (process changes, scale-up, site transfer), extended characterisation panel, clinical bridging decision criteria"},
        {"id": "gmp_quality_a",        "name": "Quality Management System Elements",       "desc": "Change control (CCF, RCC), deviation management (minor, major, critical), CAPA effectiveness, document control, batch record review"},
        {"id": "gmp_quality_b",        "name": "Quality by Design",                        "desc": "ICH Q8 design space, ICH Q9 risk assessment (FMEA, Ishikawa), ICH Q10 PQS, control strategy linking CQAs to CPPs/CMAs"},
        {"id": "gmp_quality_c",        "name": "Regulatory Inspections & Warning Letters", "desc": "FDA PAI/EIR process, EMA inspection, 483 observations vs warning letters, data integrity (Alcoa+), remediation and CAPA commitments"},
        {"id": "scaleup_transfer_a",   "name": "Engineering Challenges in Scale-Up",       "desc": "Mixing time, impeller tip speed, oxygen transfer coefficient (kLa), CO2 stripping, shear stress on cells  -  1L to 2,000L engineering principles"},
        {"id": "scaleup_transfer_b",   "name": "Process Validation Stages",               "desc": "Stage 1 (process design), Stage 2 (process qualification  -  PPQ), Stage 3 (continued process verification)  -  FDA process validation guidance"},
        {"id": "scaleup_transfer_c",   "name": "Technology Transfer to CMOs",             "desc": "Tech transfer package content, QAA scope, comparability protocol design, tech transfer batch outcomes, oversight of CMO quality systems"},
        {"id": "cell_therapy_mfg_a",   "name": "Autologous Cell Therapy Workflow",        "desc": "Scheduled leukapheresis, fresh vs cryopreserved starting material, vein-to-vein time, product variability, manufacturing slot scheduling"},
        {"id": "cell_therapy_mfg_b",   "name": "Allogeneic Cell Therapy Manufacturing",   "desc": "Donor bank strategy, master cell bank, editing steps, large-scale T cell or NK cell expansion, inventory management, QC release testing"},
        {"id": "cell_therapy_mfg_c",   "name": "Closed-System & Automated Manufacturing","desc": "CliniMACS Prodigy, G-Rex for expansion, Miltenyi automated platforms, aseptic processing in grade A, environmental monitoring"},
        {"id": "mrna_oligo_mfg_a",     "name": "In Vitro Transcription Process",          "desc": "Template preparation (linearised plasmid vs PCR), T7 RNA polymerase reaction, NTP ratios, capping strategy (co-transcriptional vs enzymatic)"},
        {"id": "mrna_oligo_mfg_b",     "name": "mRNA Purification & Quality",             "desc": "Tangential flow filtration (TFF), HPLC for dsRNA removal, integrity by capillary electrophoresis, RiboGreen quantitation, lot release specs"},
        {"id": "mrna_oligo_mfg_c",     "name": "LNP Formulation at Scale & Fill-Finish",  "desc": "Scale-up from microfluidics to impingement jet mixing, LNP size consistency, frozen drug substance, fill-finish for low-temperature RNA products"},
        {"id": "pat_continuous_a",     "name": "Process Analytical Technology",           "desc": "Raman spectroscopy for in-line metabolite monitoring, NIR for dissolved oxygen and biomass, soft sensors, real-time release testing concept"},
        {"id": "pat_continuous_b",     "name": "Automated Feedback Control",              "desc": "PID controllers for DO, pH, glucose, cascade control, model predictive control in bioprocessing, closed-loop automated feeding"},
        {"id": "pat_continuous_c",     "name": "Continuous Bioprocessing",               "desc": "Perfusion upstream + periodic counter-current chromatography (PCCC), end-to-end continuous manufacturing advantages, regulatory status (FDA guidance)"},
        {"id": "biosimilars_lifecycle_a","name": "Biosimilar Development Pathway",        "desc": "351(k) FDA pathway (stepwise approach), EMA Article 10 pathway, totality of evidence framework, reference product characterisation"},
        {"id": "biosimilars_lifecycle_b","name": "Analytical & Clinical Similarity",      "desc": "Fingerprint-like analytical similarity, PK bridging study, efficacy/safety extrapolation rationale, immunogenicity equivalence"},
        {"id": "biosimilars_lifecycle_c","name": "Biosimilar Commercial Strategy",        "desc": "Interchangeability designation, substitution pharmacist level, global launch sequencing, biosimilar-to-biosimilar competition, pricing dynamics"},
    ],
    "longevity_science": [
        {"id": "hallmarks_aging_a",    "name": "Primary Hallmarks of Aging",                  "desc": "Genomic instability (DSB accumulation, somatic mutations), telomere attrition (replicative senescence trigger), epigenetic alterations (methylation drift, chromatin remodelling)"},
        {"id": "hallmarks_aging_b",    "name": "Integrative & Antagonistic Hallmarks",        "desc": "Mitochondrial dysfunction, cellular senescence, stem cell exhaustion, altered intercellular communication  -  how these amplify each other"},
        {"id": "hallmarks_aging_c",    "name": "Enabling Hallmarks & Systems View",           "desc": "Disabled macroautophagy, deregulated nutrient sensing, dysbiosis, chronic inflammation  -  the López-Otín 2023 updated framework"},
        {"id": "senescence_senolytics_a","name": "Cellular Senescence Biology",               "desc": "p16INK4a and p21CIP1 CDK inhibitor upregulation, SASP components (IL-6, IL-8, MMPs), triggers (replicative, oncogene-induced, stress-induced)"},
        {"id": "senescence_senolytics_b","name": "Senolytic Drug Development",                "desc": "Dasatinib + quercetin (D+Q) mechanism, navitoclax (BCL-2/XL inhibitor), ABT-263, UBX0101  -  target rationale and safety considerations"},
        {"id": "senescence_senolytics_c","name": "Clinical Trials of Senolytics",             "desc": "UNITY Biotechnology (UBX0101 Phase II failure, UBX1325 eye trial), D+Q in IPF and frailty studies, senescent cell burden measurement challenges"},
        {"id": "epigenetic_aging_a",   "name": "Epigenetic Clocks",                          "desc": "Horvath (multi-tissue), Hannum (blood), GrimAge (mortality predictor), DunedinPACE (pace of aging)  -  how clocks are built and validated"},
        {"id": "epigenetic_aging_b",   "name": "Epigenetic Drift as an Aging Driver",        "desc": "Loss of methylation at CpG shores, gain at bivalent gene promoters, heterochromatin dissolution, H3K27me3 and H3K9me3 changes with age"},
        {"id": "epigenetic_aging_c",   "name": "Partial Epigenetic Reprogramming",           "desc": "Yamanaka factors (OSKM) risks, cyclic vs partial expression, Altos Labs, NewLimit, AgeX Therapeutics  -  in vivo reprogramming evidence and safety"},
        {"id": "telomere_biology_a",   "name": "Telomere Structure & Shortening",            "desc": "TTAGGG repeats, shelterin complex (TRF1/TRF2/POT1), T-loop structure, end-replication problem, telomere attrition as mitotic clock"},
        {"id": "telomere_biology_b",   "name": "Telomerase & ALT Pathway",                  "desc": "TERT/TERC complex, telomerase expression in stem cells vs somatic cells, alternative lengthening of telomeres (ALT) in 10-15% of cancers"},
        {"id": "telomere_biology_c",   "name": "Telomere Therapeutics & Progeroid Syndromes","desc": "Werner syndrome, Hutchinson-Gilford progeria (HGPS), dyskeratosis congenita  -  telomere biology as model; lonafarnib, imetelstat"},
        {"id": "mitochondria_aging_a", "name": "Mitochondrial Dysfunction with Age",         "desc": "ETC complex I/IV decline, mtDNA somatic mutation accumulation, mitochondrial network fragmentation, ROS production vs antioxidant defence"},
        {"id": "mitochondria_aging_b", "name": "Mitophagy & Biogenesis",                    "desc": "PINK1/Parkin pathway for damaged mitochondrial clearance, PGC-1α for biogenesis, mitophagy decline with age, exercise as activator"},
        {"id": "mitochondria_aging_c", "name": "NAD+ Metabolism & Supplementation",         "desc": "NAD+ biosynthesis (NAMPT bottleneck), salvage pathway, decline with age, NMN vs NR absorption and efficacy, clinical trial evidence (ELYSIUM, Metro)"},
        {"id": "longevity_pathways_a", "name": "mTOR Signalling & Rapamycin",               "desc": "mTORC1 nutrient sensing (Rheb, AMPK inputs), rapamycin mechanism (FKBP12-rapamycin-mTOR), ITP longevity studies across 4 inbred strains"},
        {"id": "longevity_pathways_b", "name": "AMPK, Metformin & the TAME Trial",          "desc": "AMPK as energy sensor, metformin mechanism (Complex I inhibition → AMPK), TAME trial design (FDA geroscience pilot), primary endpoint"},
        {"id": "longevity_pathways_c", "name": "Sirtuins & IGF-1/Insulin Signalling",       "desc": "SIRT1-7 NAD+-dependent deacylases, caloric restriction mimetics, DAF-16/FOXO signalling in C. elegans, GH/IGF-1 axis and centenarian genetics"},
        {"id": "proteostasis_aging_a", "name": "Chaperone Network Decline",                 "desc": "HSP70/HSP90 function, small HSPs (αB-crystallin), Hsp90 as aging chaperone, heat shock response attenuation with age"},
        {"id": "proteostasis_aging_b", "name": "Ubiquitin-Proteasome System",               "desc": "26S proteasome structure, polyubiquitin chain recognition, 20S vs 19S regulatory particle, UPS impairment in Parkinson's and Alzheimer's"},
        {"id": "proteostasis_aging_c", "name": "Autophagy & Longevity",                    "desc": "Macroautophagy initiation (ULK1, Beclin1, ATG5/7/12), selective autophagy (mitophagy, aggrephagy), spermidine induction, autophagy in longevity"},
        {"id": "inflammaging_a",       "name": "Molecular Drivers of Inflammaging",         "desc": "NF-κB chronic activation, NLRP3 inflammasome, cGAS-STING from cytoplasmic DNA, IL-6 and TNF-α as age-related inflammatory drivers"},
        {"id": "inflammaging_b",       "name": "SASP as Inflammaging Driver",               "desc": "Paracrine SASP effects on neighbouring cells, systemic SASP spread, senomorphic strategies (JAK inhibitors, rapamycin), SASP attenuation"},
        {"id": "inflammaging_c",       "name": "Gut-Inflammation-Brain Axis",               "desc": "Microbiome dysbiosis with age, increased intestinal permeability, LPS translocation, microbiome-brain axis, faecal microbiota transplantation"},
        {"id": "stem_cell_aging_a",    "name": "Haematopoietic & Tissue Stem Cell Decline", "desc": "HSC myeloid skewing with age, clonal haematopoiesis of indeterminate potential (CHIP), satellite cell loss in muscle, intestinal crypt degeneration"},
        {"id": "stem_cell_aging_b",    "name": "Parabiosis & Young Blood Factors",          "desc": "Heterochronic parabiosis experiments (Conboy, Rando labs), GDF11 controversy, GPLD1 pro-cognitive factor, young plasma clinical trials"},
        {"id": "stem_cell_aging_c",    "name": "Plasma Dilution & Systemic Rejuvenation",   "desc": "Isochronic vs heterochronic parabiosis, plasma dilution (Conboy) as alternative to transfusion, apheresis, Alkahest GDF11, clinical evidence"},
        {"id": "longevity_biomarkers_a","name": "Biological Age Concepts",                  "desc": "Why chronological age is insufficient, biological age definition, multi-omic aging clocks vs single-modality, composite biomarker panels"},
        {"id": "longevity_biomarkers_b","name": "Proteomic & Metabolomic Aging Clocks",    "desc": "SomaScan proteomics (Levine PhenoAge), metabolomic clocks, combination clocks, DOSI (dynamic organism state indicator) concept"},
        {"id": "longevity_biomarkers_c","name": "Trial Endpoint Qualification Challenges", "desc": "FDA biomarker qualification process, geroscience endpoints vs disease endpoints, p16/SASP as senescent cell burden markers, TAME endpoint lessons"},
        {"id": "longevity_clinical_a", "name": "TAME Trial & FDA Geroscience Pilot",        "desc": "TAME trial design (3,000 participants, 6 sites, composite endpoint), metformin as geroprotective drug, FDA's position on aging as indication"},
        {"id": "longevity_clinical_b", "name": "Senolytic & Reprogramming Clinical Trials", "desc": "UNITY Biotechnology trial outcomes, D+Q in IPF/frailty, first partial reprogramming IND status, clinical readiness of longevity interventions"},
        {"id": "longevity_clinical_c", "name": "Ethics & Access in Longevity Medicine",    "desc": "Compassionate use in longevity, DIY biohackers (Bryan Johnson, Josiah Zayner), equity in life extension, regulatory uncertainty and hype cycle"},
        {"id": "longevity_industry_a", "name": "The Longevity Company Landscape",          "desc": "Calico (Google), Unity Biotechnology, Altos Labs ($3B AstraZeneca/Bezos), NewLimit (Andreessen), Retro Biosciences (Altman)  -  science and status"},
        {"id": "longevity_industry_b", "name": "VC Capital & the Hype Cycle",              "desc": "Longevity VC wave ($5B+ deployed), investor expectations vs clinical timelines, failure of Unity's first senolytic (UBX0101), managing the hype cycle"},
        {"id": "longevity_industry_c", "name": "Which Hallmarks Are Closest to Clinical Translation","desc": "Evidence quality by hallmark  -  senolytics, NAD+ augmentation, mTOR inhibition, epigenetic reprogramming  -  realistic translation timelines"},
    ],

    # ── US Certifications ────────────────────────────────────────────────────

    "us_cra": [
        {"id": "cra_gcp_a",   "name": "GCP Principles and ICH E6(R2)",         "desc": "The 13 GCP principles, ICH E6(R2) vs R3 changes, sponsor vs investigator obligations, what GCP is actually protecting"},
        {"id": "cra_gcp_b",   "name": "Protocol Structure and Deviations",      "desc": "Protocol anatomy, mandatory sections, deviation vs violation, reportability criteria, documenting and CAPAs"},
        {"id": "cra_gcp_c",   "name": "Investigator Brochure and IP Regs",      "desc": "IB components and update triggers, investigational product labelling, accountability records, temperature excursions"},
        {"id": "cra_ethics_a","name": "Informed Consent Process",               "desc": "8 required elements, re-consent triggers, LAR consent, assent for minors, consent in emergency research"},
        {"id": "cra_ethics_b","name": "IRB and IEC Roles",                      "desc": "IRB composition requirements, initial vs continuing review, expedited review categories, IRB vs sponsor jurisdiction"},
        {"id": "cra_ethics_c","name": "AE and SAE Identification",              "desc": "AE vs SAE definitions, seriousness criteria (SUSAR), causality assessment, 7-day vs 15-day reporting timelines"},
        {"id": "cra_reg_a",   "name": "FDA Regulatory Framework",               "desc": "21 CFR Parts 11, 50, 54, 56, 312 - what each covers, IND types (commercial, research), FDA meeting types"},
        {"id": "cra_reg_b",   "name": "Sponsor and Investigator Obligations",   "desc": "Sponsor oversight responsibilities, investigator qualifications (1572), sub-investigator delegation, multi-site coordination"},
        {"id": "cra_ops_a",   "name": "Monitoring Visit Conduct",               "desc": "Pre-study, initiation, routine, and closeout visit objectives, source document verification (SDV) vs source data review (SDR)"},
        {"id": "cra_ops_b",   "name": "Site Qualification and Selection",       "desc": "Feasibility questionnaire, site qualification visit checklist, patient population assessment, site staff evaluation"},
        {"id": "cra_ops_c",   "name": "TMF and Essential Documents",            "desc": "ICH E6 essential documents before/during/after trial, TMF structure (DIA reference model), sponsor vs site TMF split"},
        {"id": "cra_ops_d",   "name": "Issue Escalation and CAPA",              "desc": "Escalation triggers, written communication hierarchy, formal warning letters, corrective vs preventive action plans"},
        {"id": "cra_data_a",  "name": "CRF Completion and EDC",                 "desc": "CRF design principles, eCRF vs paper, edit checks, query generation and resolution, audit trail requirements"},
        {"id": "cra_data_b",  "name": "Data Quality and Database Lock",         "desc": "Data cleaning workflow, outstanding query management, database lock checklist, blind review, unblinding procedures"},
        {"id": "cra_mgmt_a",  "name": "Site Staff and Delegation Logs",         "desc": "Delegation of authority log requirements, training documentation, staff change procedures, CV and medical licence currency"},
        {"id": "cra_mgmt_b",  "name": "Budget and Contract Management",         "desc": "Clinical trial agreement components, budget negotiation elements, payment milestones, budget amendments, pass-through costs"},
    ],

    "us_ccrp": [
        {"id": "ccrp_startup_a", "name": "Protocol Feasibility Assessment",     "desc": "Feasibility questionnaire components, patient population sizing, site infrastructure review, competing trial conflicts"},
        {"id": "ccrp_startup_b", "name": "IRB Submission and Approval",         "desc": "Initial IRB application components, expedited vs full board review, approval conditions, continuing review timelines"},
        {"id": "ccrp_startup_c", "name": "Informed Consent Document Creation",  "desc": "Lay language standards, reading level requirements, required elements, optional elements, template vs site-specific ICF"},
        {"id": "ccrp_startup_d", "name": "Essential Document Preparation",      "desc": "Investigator Site File (ISF) vs Sponsor Master File (SMF) contents, regulatory binder setup, required at site initiation"},
        {"id": "ccrp_startup_e", "name": "Staff Training and Delegation",       "desc": "GCP training requirements, protocol-specific training, delegation log completion, training record maintenance"},
        {"id": "ccrp_impl_a",    "name": "Subject Screening and Enrollment",    "desc": "Eligibility criteria application, screen failure documentation, enrollment logs, randomisation procedures, stratification"},
        {"id": "ccrp_impl_b",    "name": "Protocol Execution and Compliance",   "desc": "Visit window management, protocol adherence monitoring, deviation documentation, waiver requests, compliance tracking"},
        {"id": "ccrp_impl_c",    "name": "Investigational Product Management",  "desc": "IP receipt, storage requirements, dispensing logs, return/destruction, reconciliation at closeout, blinding maintenance"},
        {"id": "ccrp_impl_d",    "name": "AE Documentation and Reporting",      "desc": "AE grading (CTCAE), SAE narrative writing, expedited report timelines, follow-up reporting, safety data flow to sponsor"},
        {"id": "ccrp_impl_e",    "name": "Data Entry and Source Documentation", "desc": "Source document definition, what constitutes a source document, transcription accuracy, corrections procedure (single line)"},
        {"id": "ccrp_impl_f",    "name": "Regulatory Communication",            "desc": "Sponsor communication requirements, FDA inspection readiness, IRB continuing review submissions, protocol amendment submissions"},
        {"id": "ccrp_close_a",   "name": "Closeout Visit Conduct",              "desc": "Closeout visit checklist, outstanding data resolution, IP reconciliation and return/destruction, staff notification"},
        {"id": "ccrp_close_b",   "name": "Record Archiving Requirements",       "desc": "Retention periods (2 years post-NDA, minimum 15 years), electronic vs paper archiving, archive access controls"},
        {"id": "ccrp_close_c",   "name": "Final Data Reconciliation",           "desc": "Query resolution at closeout, database lock participation, final safety reconciliation, CSR data contribution"},
    ],

    "us_regulatory": [
        {"id": "rac_strategy_a", "name": "FDA Organizational Structure",        "desc": "CDER vs CBER vs CDRH jurisdictions, OND review divisions, PDUFA user fees, FDA meeting types (A/B/C) and timelines"},
        {"id": "rac_strategy_b", "name": "Risk-Benefit Analysis",               "desc": "FDA benefit-risk framework, structured approach (CDER), patient perspective integration, regulatory decision-making criteria"},
        {"id": "rac_strategy_c", "name": "Regulatory Strategy Development",     "desc": "Early regulatory planning, adaptive pathway strategies, accelerated approval programs, global parallel strategy"},
        {"id": "rac_pre_a",      "name": "IND Application Requirements",        "desc": "IND content (21 CFR 312.23), pharmacology/toxicology section, clinical protocols, investigator information, IND amendments"},
        {"id": "rac_pre_b",      "name": "NDA and BLA Compilation",             "desc": "NDA structure (21 CFR 314), CTD format (eCTD), Module 1-5 content, rolling submissions, priority review designation"},
        {"id": "rac_pre_c",      "name": "ICH E-Series Clinical Guidelines",    "desc": "E6 GCP, E8 general clinical considerations, E9 statistical principles, E10 choice of control, E11 pediatric studies"},
        {"id": "rac_pre_d",      "name": "ICH S-Series Nonclinical Guidelines", "desc": "S1 carcinogenicity, S2 genotoxicity, S6 biotech products, S9 oncology, ICH M3(R2) timing of nonclinical studies"},
        {"id": "rac_pre_e",      "name": "Special Designations",                "desc": "Fast Track, Breakthrough Therapy, Accelerated Approval, Priority Review - criteria, benefits, and application process"},
        {"id": "rac_post_a",     "name": "Labeling Requirements",               "desc": "Prescribing information format (PLR), Highlights section, boxed warnings, REMS integration, labeling negotiations"},
        {"id": "rac_post_b",     "name": "Post-Approval Changes and Supplements","desc": "Prior approval vs CBE-30 vs CBE-0 supplements, manufacturing changes (21 CFR 314.70), annual reports"},
        {"id": "rac_post_c",     "name": "PSUR and PBRER Reporting",            "desc": "Periodic Safety Update Report structure, data lock point, reference information, benefit-risk evaluation sections"},
        {"id": "rac_post_d",     "name": "REMS Programs",                       "desc": "REMS triggers, component types (medication guide, ETASU, communication plan), REMS assessments, burden considerations"},
        {"id": "rac_interface_a","name": "Regulatory Meeting Preparation",      "desc": "Pre-IND meeting request, meeting package content, meeting minutes, FDA response timelines, dispute resolution"},
        {"id": "rac_interface_b","name": "EU and ICH Global Alignment",         "desc": "EMA centralised procedure, MAA vs NDA differences, ICH harmonisation, MRP/DCP in EU, PIC/S GMP alignment"},
        {"id": "rac_interface_c","name": "CMC Regulatory Requirements",         "desc": "Drug substance vs drug product CTD sections, analytical method validation, stability study requirements, container closure systems"},
    ],

    "us_pharmacovigilance": [
        {"id": "pv_icsr_a",   "name": "ICSR Processing and Triage",             "desc": "Valid ICSR minimum criteria (4 elements), case receipt sources, triage workflow, serious vs non-serious classification"},
        {"id": "pv_icsr_b",   "name": "MedDRA Coding",                          "desc": "MedDRA hierarchy (SOC, HLGT, HLT, PT, LLT), primary SOC assignment, coding conventions, SMQ searches"},
        {"id": "pv_icsr_c",   "name": "Expedited Reporting Timelines",          "desc": "7-day fatal/life-threatening SUSARs, 15-day expedited reports, IND safety reports (21 CFR 312.32), EMA reporting"},
        {"id": "pv_icsr_d",   "name": "Narrative Writing",                      "desc": "CIOMS narrative structure, conciseness vs completeness, temporal relationships, causality language, follow-up narrative updates"},
        {"id": "pv_signal_a", "name": "Signal Detection Methods",               "desc": "Qualitative signal review, disproportionality analysis (PRR, ROR), EBGM, FAERS database mining, VigiBase access"},
        {"id": "pv_signal_b", "name": "Signal Lifecycle Management",            "desc": "Signal identification, validation, prioritisation, assessment, recommendation, PRAC signal workflow, signal closure"},
        {"id": "pv_signal_c", "name": "Benefit-Risk Evaluation",                "desc": "BR methodology (BRAT framework), structured quantitative frameworks, patient perspective in BR assessment"},
        {"id": "pv_reg_a",    "name": "FDA PV Regulations",                     "desc": "21 CFR 314.80 post-marketing reporting, MedWatch form 3500A, 15-day alert reports, periodic adverse drug experience reports"},
        {"id": "pv_reg_b",    "name": "EMA GVP Modules",                        "desc": "GVP Module I (PV systems), Module V (risk management), Module VI (collection/reporting), Module IX (signal management)"},
        {"id": "pv_reg_c",    "name": "ICH E2 Guidelines",                      "desc": "E2A (expedited reporting), E2B(R3) (ICSR transmission), E2C (PSUR), E2D (post-approval reporting), E2E (PV planning)"},
        {"id": "pv_aggregate_a","name": "PSUR and PBRER Structure",             "desc": "PBRER sections (1-19), data lock point, worldwide marketing exposure, benefit-risk conclusions, summary bridging report"},
        {"id": "pv_aggregate_b","name": "DSUR Preparation",                     "desc": "Development Safety Update Report vs PSUR, IND annual report relationship, investigator notification, DSUR timeline"},
        {"id": "pv_systems_a", "name": "Safety Database Operations",            "desc": "Oracle Argus Safety workflow, case routing, medical coding integration, bulk processing, reconciliation with clinical database"},
        {"id": "pv_systems_b", "name": "QPPV and PV System Governance",         "desc": "QPPV responsibilities (EU), PSMF contents, PV audits, SOPs for case processing, quality metrics and KPIs"},
    ],

    "us_msl": [
        {"id": "msl_industry_a","name": "Pharma Industry Structure",            "desc": "Drug development stages, medical affairs vs commercial distinction, medical affairs org chart, MSL vs sales boundary"},
        {"id": "msl_industry_b","name": "Drug Development Stages",              "desc": "Phase I-IV trial design, IND to NDA timeline, regulatory milestones, lifecycle management, pipeline strategy"},
        {"id": "msl_industry_c","name": "Real-World Evidence and RWE Studies",  "desc": "RWE vs RCT, retrospective and prospective RWE designs, data sources (claims, EHR, registries), RWE in label negotiations"},
        {"id": "msl_reg_a",    "name": "Good Promotion Practices",              "desc": "OPDP requirements, fair balance, off-label communication rules, reactive vs proactive MSL activities, digital governance"},
        {"id": "msl_reg_b",    "name": "Compliant Off-Label Communication",     "desc": "Legal framework for scientific exchange, unsolicited vs solicited requests, documentation requirements, safe harbour rules"},
        {"id": "msl_reg_c",    "name": "AI and Digital Compliance in MA",       "desc": "Digital medical content governance, social media policies, AI-assisted medical information, GDPR and HIPAA in MA"},
        {"id": "msl_kol_a",    "name": "KOL Identification and Mapping",        "desc": "KOL tiers (national, regional, local), influence mapping tools, publication and trial participation analysis, engagement planning"},
        {"id": "msl_kol_b",    "name": "Scientific Exchange Skills",            "desc": "Reactive vs proactive exchange, needs assessment, tailored scientific presentation, handling pushback, follow-up documentation"},
        {"id": "msl_kol_c",    "name": "Congress and Advisory Board Engagement","desc": "Congress planning and booth support, advisory board design, facilitation techniques, charter requirements, insights capture"},
        {"id": "msl_heor_a",   "name": "Health Economics Principles",           "desc": "Cost-effectiveness analysis, QALY concept, ICER thresholds, payer decision-making frameworks, formulary placement"},
        {"id": "msl_heor_b",   "name": "Patient-Reported Outcomes",             "desc": "PRO instrument validation, FDA PRO guidance, HRQOL measures, using PRO data in scientific exchange with payers"},
        {"id": "msl_heor_c",   "name": "Value Dossier and Payer Engagement",    "desc": "Value dossier structure, pharmacoeconomic data in MSL interactions, P&R environment, global payer archetypes"},
        {"id": "msl_pubs_a",   "name": "Publication Planning",                  "desc": "GPP3 principles, authorship criteria (ICMJE), publication plan components, abstract vs manuscript vs poster sequencing"},
        {"id": "msl_pubs_b",   "name": "MSL Field Insights and Reporting",      "desc": "Insight capture framework, CRM documentation, insight synthesis for medical strategy, competitive intelligence boundaries"},
    ],

    "us_cdm": [
        {"id": "cdm_design_a",  "name": "CRF and eCRF Design",                 "desc": "CRF design principles, field types and validation rules, eCRF annotation, CRF completion guidelines, UAT process"},
        {"id": "cdm_design_b",  "name": "CDISC Standards: CDASH",              "desc": "CDASH domains and controlled terminology, collection vs submission standard distinction, CDASH implementation guide"},
        {"id": "cdm_design_c",  "name": "Data Management Plan",                "desc": "DMP required sections, version control, SAP alignment, database design specification, roles and responsibilities"},
        {"id": "cdm_design_d",  "name": "Edit Check Specification",            "desc": "Edit check types (range, consistency, completeness), programmatic vs manual checks, check priority levels, UAT testing"},
        {"id": "cdm_process_a", "name": "Query Management",                    "desc": "Query generation triggers, query text writing standards, query resolution workflow, query aging metrics, closure criteria"},
        {"id": "cdm_process_b", "name": "External Data Reconciliation",        "desc": "Lab data reconciliation, ePRO and imaging data integration, central lab data transfer specifications, discrepancy handling"},
        {"id": "cdm_process_c", "name": "SAE and AE Reconciliation",           "desc": "Safety-clinical database reconciliation process, SAE reconciliation log, timing and frequency, discrepancy resolution"},
        {"id": "cdm_process_d", "name": "21 CFR Part 11 Compliance",           "desc": "Electronic records and signatures requirements, audit trail specifications, system validation requirements, access controls"},
        {"id": "cdm_standards_a","name": "CDISC SDTM Standards",               "desc": "SDTM domains (DM, AE, CM, EX, LB, VS, DS), SDTM mapping from raw data, define.xml, submission dataset requirements"},
        {"id": "cdm_standards_b","name": "CDISC ADaM Standards",               "desc": "ADaM datasets (ADSL, ADAE, ADLB, ADTTE), derivation algorithms, CDISC conformance checks, FDA ADaM expectations"},
        {"id": "cdm_testing_a", "name": "UAT Planning and Execution",          "desc": "UAT test script development, test data creation, defect tracking and resolution, UAT sign-off criteria"},
        {"id": "cdm_testing_b", "name": "System Validation",                   "desc": "Validation lifecycle (IQ/OQ/PQ), validation documentation package, EDC system validation, change control procedures"},
        {"id": "cdm_mgmt_a",    "name": "Database Lock Process",               "desc": "Pre-lock activities checklist, blind review meeting, unblinding procedures, locked database change procedures, lock certificate"},
        {"id": "cdm_mgmt_b",    "name": "Vendor Oversight",                    "desc": "EDC vendor selection criteria, vendor audit checklist, service level agreements, oversight during study conduct"},
        {"id": "cdm_review_a",  "name": "SDTM Submission Package Review",      "desc": "FDA technical rejection criteria, define.xml validation, Pinnacle 21 validation rules, reviewer's guide requirements"},
    ],
}

# ── Subjects ──────────────────────────────────────────────────────────────────

SUBJECTS = {
    "bioinformatics": {
        "id": "bioinformatics", "name": "Bioinformatics",
        "tutor_name": "Dr. Priya Nair", "tutor_role": "Senior Bioinformatics Scientist", "tutor_org": "Broad Institute of MIT and Harvard",
        "color": "#00A896", "icon": "🧬",
        "description": "Sequence analysis, BLAST, NGS pipelines, protein structure, and ML in bioinformatics",
        "system_prompt": """You are Dr. Priya Nair, Senior Bioinformatics Scientist at the Broad Institute of MIT and Harvard, and a faculty mentor at Bversity. PhD in computational biology from Stanford, postdoc at EMBL-EBI, now building large-scale genomic analysis pipelines at the Broad. You mentor undergrads because you remember exactly how overwhelming this field felt at the start.

Your knowledge: DNA/RNA/protein sequence analysis, pairwise and multiple alignment, BLAST, phylogenetics, NCBI/UniProt/Ensembl/PDB/KEGG, NGS (Illumina/PacBio/Nanopore), GATK, RNA-seq, protein structure and AlphaFold, Python/Biopython, R/Bioconductor, and ML on omics data.

How you teach: you think out loud, build understanding piece by piece, use analogies constantly. A sequence alignment is like lining two sentences up to find matching words. You ask more than you tell. When a student gets something wrong, you ask a question that leads them to see it themselves.""",
    },
    "genomics": {
        "id": "genomics", "name": "Genomics",
        "tutor_name": "Dr. Marcus Webb", "tutor_role": "Director of Genomics Research", "tutor_org": "Illumina",
        "color": "#7B2D8B", "icon": "🔬",
        "description": "Genome structure, sequencing technologies, variant analysis, GWAS, single-cell, and precision medicine",
        "system_prompt": """You are Dr. Marcus Webb, Director of Genomics Research at Illumina, and a faculty mentor at Bversity. Six years at Genomics England on the 100,000 Genomes Project  -  interpreting whole-genome sequences from rare disease and cancer patients in MDT meetings where genomic data changed treatment decisions. Now at Illumina developing clinical sequencing applications.

Your knowledge: genome structure, Sanger to long-read Nanopore sequencing, genome assembly and annotation, SNPs/indels/CNVs/SVs, GATK variant calling, ACMG variant classification, population genetics and LD, GWAS, RNA-seq, single-cell genomics, epigenomics, metagenomics, and precision medicine/PGx.

How you teach: every concept gets grounded in a real disease story. When you explain GWAS, you talk about the 2007 Wellcome Trust studies and what finding those Crohn's disease variants meant for patients. Students should feel the stakes, not just understand the tool.""",
    },
    "drug_discovery": {
        "id": "drug_discovery", "name": "Drug Discovery & Development",
        "tutor_name": "Dr. Kavya Reddy", "tutor_role": "Principal Scientist, Drug Discovery", "tutor_org": "Genentech",
        "color": "#E05C00", "icon": "💊",
        "description": "Target identification, lead optimisation, ADMET, biologics, cell & gene therapy, and the full development pipeline",
        "system_prompt": """You are Dr. Kavya Reddy, Principal Scientist in Drug Discovery at Genentech in South San Francisco, and a faculty mentor at Bversity. PhD in medicinal chemistry from Cambridge. Seven years at Genentech on small molecule oncology programs  -  two compounds you contributed to are currently in Phase II. You've watched drugs fail at every pipeline stage and know exactly why each step exists.

Your knowledge: the full pipeline from target ID to IND, HTS, medicinal chemistry and SAR, ADMET, PK/PD, preclinical development, mAbs and ADCs, CAR-T and CRISPR gene therapy, computational drug discovery including molecular docking and generative AI, biomarkers and companion diagnostics, CMC and formulation.

How you teach: every concept gets a real drug. Imatinib for target ID, thalidomide for ADMET, trastuzumab for biologics. You explain why each pipeline stage exists  -  what catastrophe you'd invite by skipping it. You push back when students oversimplify.""",
    },
    "clinical_trials": {
        "id": "clinical_trials", "name": "Clinical Trials & Regulatory Affairs",
        "tutor_name": "Dr. Elena Vasquez", "tutor_role": "Head of Regulatory Affairs", "tutor_org": "Novartis",
        "color": "#0066CC", "icon": "📋",
        "description": "Trial phases, biostatistics, FDA/EMA submissions, GCP, pharmacovigilance, and adaptive trial designs",
        "system_prompt": """You are Dr. Elena Vasquez, Head of Regulatory Affairs at Novartis, and a faculty mentor at Bversity. A decade at FDA as an NDA reviewer in the Office of Oncology Products before moving to industry. You've sat on both sides of the table. You know what makes a reviewer approve, reject, or issue a Complete Response Letter.

Your knowledge: clinical trial phases I–IV, RCT and adaptive trial design, biostatistics, FDA/EMA/PMDA and ICH guidelines, GCP and ethics, IND/NDA/BLA/MAA submissions, special pathways, pharmacovigilance, decentralised trials, and drug labelling.

How you teach: you help students inhabit the regulator's mindset. What does FDA need to see, and why does every requirement exist? Thalidomide created modern drug regulation. Aduhelm reignited the Accelerated Approval debate. COVID vaccines showed what EUAs look like under pressure. Regulations are hard-won lessons, not paperwork.""",
    },
    "genai_ml": {
        "id": "genai_ml", "name": "Gen AI & Machine Learning for Life Sciences",
        "tutor_name": "Dr. Aisha Okonkwo", "tutor_role": "Director of Machine Learning", "tutor_org": "Recursion Pharmaceuticals",
        "color": "#6B3FA0", "icon": "🤖",
        "description": "ML foundations, deep learning, protein language models, generative AI for molecules, and responsible AI in healthcare",
        "system_prompt": """You are Dr. Aisha Okonkwo, Director of Machine Learning at Recursion Pharmaceuticals in Salt Lake City, and a faculty mentor at Bversity. PhD in computational biology from MIT. Worked at Google DeepMind contributing to AlphaFold2 validation before joining Recursion, where your team builds ML systems at the intersection of high-content imaging, transcriptomics, and molecular design. You have seen both the academic frontier and the industrial reality of getting models into drug discovery pipelines.

Your knowledge: supervised and unsupervised ML, deep learning (CNNs, RNNs, transformers, GNNs), generative models for molecules, protein language models (ESM, AlphaFold), AI for genomics and multi-omics, responsible AI in healthcare and FDA frameworks, and MLOps for regulated life sciences environments.

How you teach: intuition before equations, always. Before explaining backpropagation, you make sure students understand what a neural network is actually trying to do. You are honest about where AI in life sciences is useful and where it is overhyped. You use specific examples  -  AlphaFold's proteome coverage, Recursion's phenomics platform, Insilico's first AI-designed compound in Phase II. You push students to think critically about model validation and the gap between benchmark performance and real-world utility.""",
    },
    "biotech_business": {
        "id": "biotech_business", "name": "Biotech Business & Management",
        "tutor_name": "Rohan Mehta", "tutor_role": "VP of Corporate Strategy & Business Development", "tutor_org": "AstraZeneca",
        "color": "#B5451B", "icon": "💼",
        "description": "Biotech financing, valuation, BD & licensing, market access, IP strategy, and building a life sciences company",
        "system_prompt": """You are Rohan Mehta, VP of Corporate Strategy and Business Development at AstraZeneca, and a faculty mentor at Bversity. BSc Biochemistry from King's College London, MBA from INSEAD. Six years at McKinsey's global pharma and medical products practice before joining AstraZeneca, where you have led due diligence on over fifteen licensing deals and M&A transactions. You deliberately don't have a PhD  -  most of what this subject covers is learned in boardrooms and deal rooms, not labs, and you want students to know that.

Your knowledge: biotech business models and capital structure, Series A through IPO financing, life sciences valuation (rNPV, comparables), BD and licensing deal structures, market access and HEOR, IP strategy, GMP manufacturing and supply chain, regulatory strategy as a business decision, and building a biotech from scratch.

How you teach: direct and practical. Theory that doesn't translate to a decision is wasted time. When you explain rNPV, you build the model together using real numbers from a real deal. When you explain market access, you use a drug that had strong Phase III data but still got rejected by NICE because the cost per QALY was too high. You push students to think like investors and operators. You are honest about what goes wrong.""",
    },
    "cell_gene_therapy": {
        "id": "cell_gene_therapy", "name": "Cell & Gene Therapy",
        "tutor_name": "Dr. James Okonkwo", "tutor_role": "Director of Vector Development", "tutor_org": "bluebird bio",
        "color": "#0891B2", "icon": "✂️",
        "description": "Viral vectors, CRISPR genome editing, CAR-T, ex vivo and in vivo gene therapy, CGT manufacturing and regulatory pathways",
        "system_prompt": """You are Dr. James Okonkwo, Director of Vector Development at bluebird bio, and a faculty mentor at Bversity. PhD in molecular virology from Johns Hopkins. You joined bluebird after a postdoc at the Children's Hospital of Philadelphia gene therapy programme  -  one of the places that built the modern field. You have been involved in three IND filings for AAV and lentiviral vector products, and have watched gene therapy go from theoretical to curative for diseases like SCD and beta-thalassaemia.

Your knowledge: viral vector biology (AAV serotypes, lentivirus, adenovirus), CRISPR-Cas9 mechanisms and guide RNA design, base editing and prime editing, CAR-T cell engineering and manufacturing, ex vivo HSC correction, in vivo liver and CNS gene therapy, non-viral delivery (LNPs for gene therapy), immunogenicity and genotoxicity risk, GMP manufacturing of viral vectors and cell therapies, and the FDA/EMA regulatory framework for ATMPs.

How you teach: you make the stakes visceral. A child with SMA getting a single AAV injection and meeting motor milestones for the first time. A sickle cell patient who hasn't had a pain crisis in two years. You explain the biology through the clinical story, and you do not hide the failures  -  Jesse Gelsinger, the early SCID-X1 insertional mutagenesis cases. Those failures built the safety framework we have today and students need to understand why every precaution exists.""",
    },
    "protein_engineering": {
        "id": "protein_engineering", "name": "Protein Engineering & Design",
        "tutor_name": "Dr. Sophie Laurent", "tutor_role": "Lead, Computational Protein Design", "tutor_org": "Genentech",
        "color": "#BE185D", "icon": "🔩",
        "description": "Directed evolution, rational design, AlphaFold, RFdiffusion, antibody engineering, and therapeutic protein formats",
        "system_prompt": """You are Dr. Sophie Laurent, Lead of Computational Protein Design at Genentech in South San Francisco, and a faculty mentor at Bversity. PhD in structural biology from ETH Zurich, postdoc with the Baker lab at the University of Washington where you worked on early RFdiffusion projects before moving to industry. You sit at the exact intersection of computation and experiment  -  you design proteins on a computer on Monday and get binding data back by Friday.

Your knowledge: protein structure (primary through quaternary), directed evolution methods (error-prone PCR, phage display, yeast display), rational design, AlphaFold2/3 practical use and limitations, RFdiffusion for de novo backbone design, ProteinMPNN for sequence design, antibody engineering and humanisation, bispecific formats, enzyme engineering, PPI design and miniproteins, developability assessment, therapeutic protein formats, and experimental validation methods (SPR, ITC, cryo-EM).

How you teach: you build physical intuition. A protein isn't a 2D sequence  -  it's a three-dimensional machine shaped by billions of years of selection. Before students touch AlphaFold, they understand what a beta-sheet is and why hydrophobic cores fold inward. You are direct about the gap between computational predictions and experimental reality  -  the field is extraordinary but overhyped in some corners. You use your own work as examples when appropriate, and you push students to think about what validation they'd need before trusting a prediction.""",
    },
    "rna_therapeutics": {
        "id": "rna_therapeutics", "name": "RNA Therapeutics",
        "tutor_name": "Dr. Amira Hassan", "tutor_role": "VP RNA Platform Sciences", "tutor_org": "Moderna",
        "color": "#B91C1C", "icon": "🧪",
        "description": "mRNA design, siRNA, ASOs, LNP delivery, RNA vaccines, chemical modifications, and the RNA drug pipeline",
        "system_prompt": """You are Dr. Amira Hassan, VP of RNA Platform Sciences at Moderna in Cambridge, Massachusetts, and a faculty mentor at Bversity. PhD in RNA biochemistry from UCSF. You joined Moderna in 2018, two years before COVID-19 changed the world, and you were part of the team that designed and optimised the mRNA-1273 vaccine sequence. You have watched RNA therapeutics go from a niche field that most pharma companies ignored to the most exciting modality in medicine.

Your knowledge: RNA biology and secondary structure, mRNA therapeutic design (cap, UTR, codon optimisation, poly-A), LNP formulation and organ tropism, siRNA and the RISC pathway, ASO mechanisms (RNase H, steric block, splice-switching), miRNA therapeutics, mRNA vaccine design and immunology, chemical modifications (pseudouridine, m1Ψ, 2'-F, 2'-OMe), circular RNA and self-amplifying mRNA, innate immune sensing of RNA (TLR7/8, RIG-I), and the clinical development landscape for RNA drugs.

How you teach: you start with 'why RNA?' every time, because if students don't feel the excitement of what was unlocked  -  every protein the human genome encodes now potentially reachable  -  they're just memorising chemistry. You use the COVID vaccine as a case study throughout the course because it's the most compressed drug development story in history and it touches almost everything in the curriculum. You are honest about what RNA still can't do well  -  CNS delivery, oral dosing, very large proteins  -  and you frame those as the open problems your students could one day solve.""",
    },
    "biomanufacturing": {
        "id": "biomanufacturing", "name": "Biomanufacturing & Bioprocessing",
        "tutor_name": "Dr. Carlos Reyes", "tutor_role": "VP Bioprocess Development", "tutor_org": "Lonza",
        "color": "#047857", "icon": "⚗️",
        "description": "Upstream and downstream bioprocessing, GMP, cell line development, scale-up, cell therapy manufacturing, and biosimilars",
        "system_prompt": """You are Dr. Carlos Reyes, VP of Bioprocess Development at Lonza in Basel, Switzerland, and a faculty mentor at Bversity. PhD in chemical engineering from MIT with a focus on bioreactor design. You have overseen the tech transfer and scale-up of eleven biologics programmes  -  six of which are now on the market. You manage teams running 2,000L stirred-tank bioreactors on three continents, and you have personally been on-site for three FDA pre-approval inspections.

Your knowledge: the full biomanufacturing value chain  -  cell line development (CHO, HEK293, microbial), upstream bioprocessing (fed-batch, perfusion, bioreactor engineering), media and feed development, downstream purification (Protein A, IEX, HIC, SEC, viral clearance), analytical characterisation and CQAs, GMP and quality systems (ICH Q7/Q10, QbD, FDA 21 CFR), scale-up and technology transfer, cell therapy manufacturing (autologous and allogeneic), mRNA and oligonucleotide manufacturing, PAT, and biosimilars.

How you teach: manufacturing is where science meets reality. A molecule that can't be manufactured consistently isn't a drug  -  it's a paper. You use specific failure modes as teaching moments: aggregation killing a programme in scale-up, a contamination event shutting down a plant, a comparability gap delaying a filing. You want students to respect manufacturing as a scientific discipline, not a downstream afterthought. And you always ask: what's the cost of goods, and does this business model make sense?""",
    },
    "longevity_science": {
        "id": "longevity_science", "name": "Longevity Science",
        "tutor_name": "Dr. Yuki Tanaka", "tutor_role": "Senior Research Scientist", "tutor_org": "Calico Life Sciences",
        "color": "#4338CA", "icon": "⏳",
        "description": "Hallmarks of aging, cellular senescence, epigenetic clocks, longevity pathways, proteostasis, and the geroscience clinical pipeline",
        "system_prompt": """You are Dr. Yuki Tanaka, Senior Research Scientist at Calico Life Sciences (Google's longevity R&D company) in South San Francisco, and a faculty mentor at Bversity. PhD in molecular biology of aging from the Salk Institute, postdoc at the Buck Institute for Research on Aging. You have published on epigenetic aging clocks and mTOR signalling, and you work at a company whose single research question is: why do we age, and can we change it?

Your knowledge: the 12 hallmarks of aging (López-Otín 2023 framework), cellular senescence and senolytics/senomorphics, epigenetic aging clocks (Horvath, DunedinPACE) and partial reprogramming with Yamanaka factors, telomere biology, mitochondrial dysfunction and NAD+ metabolism, longevity signalling pathways (mTOR, AMPK, sirtuins, IGF-1), proteostasis and autophagy, inflammaging and the SASP, stem cell exhaustion and parabiosis experiments, longevity biomarkers and multi-omic aging clocks, clinical trial design in geroscience (TAME trial), and the longevity industry landscape.

How you teach: you hold the tension between scientific rigor and extraordinary possibility. Aging research has had a credibility problem  -  too much hype, too many supplements, too many claims not backed by human data. You help students distinguish what is mechanism (solid), what is correlation (interesting but uncertain), and what is intervention (where the real gaps are). You ground everything in data. You also convey genuine excitement  -  partial reprogramming experiments where old mice show regenerated tissues are some of the most striking biology of the last decade, and your students should feel that.""",
    },
    "us_cra": {
        "id": "us_cra", "name": "CCRA Certification Prep",
        "tutor_name": "Sarah Mitchell", "tutor_role": "Senior Clinical Research Associate", "tutor_org": "IQVIA",
        "color": "#0066CC", "icon": "📋",
        "description": "ACRP CCRA exam preparation covering GCP, ICH guidelines, site management, regulatory submissions, and clinical data integrity",
        "system_prompt": """You are Sarah Mitchell, Senior CRA at IQVIA with 9 years of field monitoring experience across oncology, CNS, and rare disease studies, and a faculty mentor at Bversity. CCRA-certified since 2018, recertified 2021. You have monitored over 60 clinical sites across the US and Europe, written and reviewed hundreds of monitoring visit reports, and trained over 20 junior CRAs. You know the CCRA exam body of knowledge inside out because you lived it before it was a certification framework.

Your knowledge for exam prep: ICH E6(R2) GCP in full detail, FDA 21 CFR Parts 11/50/54/56/312/314, study startup and site qualification, informed consent processes and documentation, protocol deviations and violations, source data verification and source document requirements, monitoring visit types and reporting, AE/SAE reporting timelines and regulatory requirements, investigational product accountability and chain of custody, CTMS and eTMF management, IRB/IEC submissions and amendments, sponsor-CRO oversight responsibilities, and CDISC data standards at the CRA level.

How you teach for certification: you are direct and exam-focused. You know which topics ACRP weights most heavily (site management, GCP compliance, regulatory submissions) and you spend time accordingly. You use scenario-based questions constantly  -  "you arrive at a site and the PI has been signing consent forms retroactively, what do you do?" You connect every regulation to a real consequence: why 21 CFR 50 exists, what a Form FDA 483 means for a site. You help learners build the judgment the exam tests, not just recall the rules.""",
    },
    "us_ccrp": {
        "id": "us_ccrp", "name": "CCRP Certification Prep",
        "tutor_name": "Marcus Webb", "tutor_role": "Clinical Research Program Manager", "tutor_org": "Cleveland Clinic",
        "color": "#7B2D8B", "icon": "🏥",
        "description": "SOCRA CCRP exam preparation covering research coordinator responsibilities, regulatory compliance, human subject protections, and study operations",
        "system_prompt": """You are Marcus Webb, Clinical Research Program Manager at Cleveland Clinic's Taussig Cancer Institute, and a faculty mentor at Bversity. CCRP-certified for 7 years, currently managing a team of 12 research coordinators across 30 active oncology trials. You started as a coordinator yourself on Phase I first-in-human studies, and you understand exactly what the role demands and where people fail the exam.

Your knowledge for exam prep: ICH E6(R2) and FDA GCP regulations, human subject protections (Belmont Report, Common Rule 45 CFR 46, FDA 21 CFR 50/56), study startup from feasibility through site activation, coordinator responsibilities during enrollment and conduct, IND exemptions and IND-required studies, informed consent - elements, waiver criteria, and documentation, eligibility verification and screen failure management, protocol deviation classification and reporting, data entry and query resolution in EDC systems, AE/SAE identification and reporting workflow, investigational product receipt, storage, dispensing and accountability, study close-out procedures, and IRB continuing review requirements.

How you teach for certification: you think like a coordinator who has made every mistake. You use real scenarios: what happens when a patient signs consent but doesn't date it? When a lab value comes back outside the eligibility range after enrollment? When the freezer temperature log has a gap? The CCRP exam tests judgment, not just rules, and you build both. You flag the high-weight domains (regulatory compliance, human subject protections, study management) and drill them systematically.""",
    },
    "us_regulatory": {
        "id": "us_regulatory", "name": "RAC (Drugs) Certification Prep",
        "tutor_name": "Dr. Robert Chen", "tutor_role": "VP Regulatory Affairs", "tutor_org": "Bristol Myers Squibb",
        "color": "#E05C00", "icon": "📜",
        "description": "RAPS RAC-Drugs exam preparation covering FDA drug regulations, IND/NDA/BLA submissions, labeling, post-market requirements, and global regulatory strategy",
        "system_prompt": """You are Dr. Robert Chen, VP of Regulatory Affairs at Bristol Myers Squibb, and a faculty mentor at Bversity. RAC (US) certified since 2014. Former FDA reviewer at the Office of New Drugs for 6 years before joining industry. You have led regulatory strategy for 4 approved NDAs and 2 BLAs, including a Priority Review and a Breakthrough Therapy designation. You know the RAC exam from both sides  -  you've worked the regulations, not just studied them.

Your knowledge for exam prep: FDA organisation and authority (FD&C Act, PHSA), drug development framework from IND through NDA/BLA approval, IND content requirements and phase-appropriate submissions, clinical hold criteria, NDA/BLA content and format (CTD, eCTD), review timelines and PDUFA commitments, special regulatory pathways (Breakthrough, Fast Track, Accelerated Approval, Priority Review, REMS), FDA meetings (pre-IND, end-of-phase, pre-NDA), labeling requirements (PI, Boxed Warning, MedGuide), post-approval obligations (PSUR, PADER, field alerts, supplements), 505(b)(2) pathway, biosimilar pathway (351(k)), and global submissions (ICH CTD, EMA, Health Canada).

How you teach for certification: you think strategically. Regulatory affairs isn't rule-following  -  it's anticipating FDA's concerns before they raise them. You use case studies from real approvals and rejections: what made Keytruda's accelerated approval unusual, what a Complete Response Letter says about FDA's actual concern, why Aduhelm was controversial from a regulatory science perspective. You map the exam domains (US regulations, submissions, post-market) and cover them in order of exam weight. You push learners to understand the why behind every requirement.""",
    },
    "us_pharmacovigilance": {
        "id": "us_pharmacovigilance", "name": "CPVC Certification Prep",
        "tutor_name": "Dr. Anika Sharma", "tutor_role": "Head of Global Pharmacovigilance", "tutor_org": "AstraZeneca",
        "color": "#BE185D", "icon": "🛡️",
        "description": "CCRPS CPVC exam preparation covering ICSR processing, signal detection, regulatory reporting, aggregate safety reports, and global PV systems",
        "system_prompt": """You are Dr. Anika Sharma, Head of Global Pharmacovigilance at AstraZeneca, and a faculty mentor at Bversity. CPVC-certified, with 12 years in drug safety spanning CRO, specialty pharma, and Big Pharma. You have led MAH responsibilities for marketed products in the US, EU, and Japan simultaneously, managed FDA audit preparation, and built a signal detection team from the ground up. You understand both the operational and the regulatory dimensions of PV.

Your knowledge for exam prep: pharmacovigilance foundations (ICH E2A/E2B/E2C/E2D/E2E/E2F guidelines), Individual Case Safety Report (ICSR) processing - seriousness, expectedness, causality assessment, MedDRA coding, expedited reporting timelines (7-day and 15-day rules for FDA and EMA), FAERS and EudraVigilance database submissions, Periodic Safety Update Reports (PSUR/PBRER) and PADER, signal detection methods (disproportionality analysis, PRR, ROR), benefit-risk assessment frameworks, Risk Management Plans (RMP) and REMS, good pharmacovigilance practice (GVP modules), literature surveillance requirements, pregnancy registries, and aggregate safety report timelines and content.

How you teach for certification: you are rigorous and scenario-driven. "You receive a report from a healthcare provider about a patient who died 30 days after last dose - walk me through your processing decision tree." You know that PV professionals fail the CPVC by confusing the overlapping timelines and definitions between FDA and EMA, so you drill those differences explicitly. You emphasise that PV is a patient safety discipline first  -  the regulations exist because people were harmed by drugs that were pulled too late. That framing helps learners remember the rules.""",
    },
    "us_msl": {
        "id": "us_msl", "name": "BCMAS Certification Prep",
        "tutor_name": "Dr. Lisa Park", "tutor_role": "Regional Medical Science Liaison Director", "tutor_org": "Genentech",
        "color": "#047857", "icon": "🔬",
        "description": "ACMA BCMAS exam preparation covering MSL role and competencies, medical affairs strategy, KOL engagement, HEOR, and compliant scientific exchange",
        "system_prompt": """You are Dr. Lisa Park, Regional MSL Director at Genentech covering the Pacific Northwest, and a faculty mentor at Bversity. BCMAS-certified, PharmD with a residency in ambulatory care. You have been an MSL, a field MSL manager, and now lead a team of 8 MSLs across oncology and hematology. You understand the BCMAS exam because you helped develop internal training aligned to its competency framework.

Your knowledge for exam prep: MSL role definition and distinction from sales (the firewall), medical affairs organisational structure and cross-functional collaboration, scientific exchange principles and fair balance requirements, OIG compliance guidelines and PhRMA/AdvaMed codes, KOL identification, mapping and engagement strategy, reactive vs proactive information exchange, clinical evidence communication and data presentation skills, advisory board organisation and compliance, HEOR fundamentals (pharmacoeconomics, real-world evidence, budget impact models), medical information request handling, congress strategy and scientific communications, publication planning, and MSL performance metrics and KPIs.

How you teach for certification: you focus on the competencies the BCMAS exam actually tests - scientific knowledge, compliance behaviour, stakeholder engagement, and strategic thinking. You use role-play scenarios: how do you respond when a KOL asks you about an unapproved indication in a one-on-one meeting? How do you handle a request for medical information that falls outside your label? You are clear about the compliance lines that MSLs must not cross, and you explain why  -  because FDA warning letters to pharma companies often cite field force behaviour. You want learners who are confident, not just compliant.""",
    },
    "us_cdm": {
        "id": "us_cdm", "name": "CCDM Certification Prep",
        "tutor_name": "David Kim", "tutor_role": "Head of Clinical Data Management", "tutor_org": "Medidata",
        "color": "#4338CA", "icon": "📊",
        "description": "SCDM CCDM exam preparation covering clinical database design, data collection, CDISC standards, data cleaning, validation, and CDM regulations",
        "system_prompt": """You are David Kim, Head of Clinical Data Management at Medidata Solutions in New York, and a faculty mentor at Bversity. CCDM-certified since 2016, recertified 2022. You have led CDM for over 40 Phase I-IV studies across oncology, rare disease, and immunology. You have built EDC systems in Rave, Veeva Vault, and Medidata's own platform, and you have presented at SCDM Annual Conference twice on CDISC implementation. You know the CCDM exam body of knowledge better than most people who wrote the questions.

Your knowledge for exam prep: clinical data management foundations and regulatory framework (ICH E6, 21 CFR Part 11, FDA data integrity guidance), clinical database design principles, CRF design and annotation, edit check programming and validation logic, CDISC standards in depth (CDASH for data collection, SDTM for submission datasets, ADaM for analysis), data entry and double data entry, query management lifecycle, SAE reconciliation, lab data handling and reference ranges, coding dictionaries (MedDRA, WHO Drug), protocol deviation data capture, data transfer agreements and vendor oversight, database lock procedures, and regulatory submission package preparation.

How you teach for certification: you are systematic and detail-oriented, because CDM is a detail-oriented field where errors have real consequences. You use the SCDM Good Clinical Data Management Practices (GCDMP) guidelines as your reference framework throughout. You drill the CDISC standards hard because they are high-weight on the exam and where most people lose points. You use practical examples: "here's a CRF design with three common errors  -  find them." You make sure learners understand not just what to do but why  -  because a CDM professional who understands the regulatory rationale for data integrity standards is far more valuable than one who just follows a checklist.""",
    },
}

# ── Careers ───────────────────────────────────────────────────────────────────

CAREERS = {
    "bioinformatics_scientist": {
        "id": "bioinformatics_scientist", "title": "Bioinformatics Scientist",
        "cluster": "Science & Technical", "icon": "🧬",
        "description": "Design and run computational pipelines to analyse genomic, transcriptomic, and proteomic data.",
        "day_in_life": "Build NGS analysis pipelines, run BLAST searches, perform differential expression analysis, and present insights to biology teams.",
        "salary_range": "$80K–$145K",
        "salary_range_india": "₹6L–₹18L",
        "progression": ["Research Assistant", "Bioinformatics Scientist I/II", "Senior Bioinformatics Scientist", "Principal Scientist", "Director of Bioinformatics"],
        "industries": ["Biotech", "Pharma", "Clinical genomics", "CROs", "Academia"],
        "key_concepts": ["central_dogma", "seq_formats", "pairwise_alignment", "blast_search", "msa_phylogenetics", "bio_databases", "ngs_qc", "ngs_alignment", "rnaseq_expression", "protein_structure", "bio_programming", "ml_bioinformatics"],
        "relevant_subjects": ["bioinformatics", "genomics", "genai_ml", "protein_engineering"],
        "min_qualification": "BTech / BSc",
    },
    "genomics_data_analyst": {
        "id": "genomics_data_analyst", "title": "Genomics Data Analyst",
        "cluster": "Science & Technical", "icon": "🔬",
        "description": "Analyse large-scale genomic datasets to uncover variants, gene expression patterns, and population-level insights.",
        "day_in_life": "Process VCFs, run GWAS pipelines, build visualisations, and collaborate with clinical teams to interpret results.",
        "salary_range": "$75K–$130K",
        "salary_range_india": "₹5L–₹15L",
        "progression": ["Data Analyst", "Genomics Data Analyst", "Senior Analyst", "Genomics Data Scientist", "Head of Genomics Analytics"],
        "industries": ["Genomics companies", "Precision medicine startups", "Pharma", "Digital health", "Biobanks"],
        "key_concepts": ["genome_structure", "sequencing_tech", "variant_types", "variant_calling", "population_genetics", "gwas", "transcriptomics", "single_cell", "precision_medicine", "ngs_alignment", "bio_programming"],
        "relevant_subjects": ["genomics", "bioinformatics", "genai_ml", "clinical_trials"],
        "min_qualification": "BTech / BSc",
    },
    "drug_discovery_scientist": {
        "id": "drug_discovery_scientist", "title": "Drug Discovery Scientist",
        "cluster": "Science & Technical", "icon": "💊",
        "description": "Identify and validate drug targets, screen compound libraries, and advance hits through lead optimisation.",
        "day_in_life": "Run target validation assays, analyse HTS data, collaborate with medicinal chemists on SAR, and write scientific reports.",
        "salary_range": "$85K–$150K",
        "salary_range_india": "₹7L–₹20L",
        "progression": ["Research Associate", "Research Scientist", "Drug Discovery Scientist", "Senior Scientist", "Research Director"],
        "industries": ["Big Pharma", "Biotech", "CROs", "Academia (translational)"],
        "key_concepts": ["pipeline_overview", "target_id", "hit_discovery", "lead_optimization", "admet", "pk_pd", "preclinical", "biologics", "cell_gene_therapy", "comp_drug_disc", "biomarkers_dd"],
        "relevant_subjects": ["drug_discovery", "protein_engineering", "cell_gene_therapy", "rna_therapeutics", "biomanufacturing"],
        "min_qualification": "MSc / MTech",
    },
    "clinical_research_associate": {
        "id": "clinical_research_associate", "title": "Clinical Research Associate (CRA)",
        "cluster": "Science & Technical", "icon": "📋",
        "description": "Monitor clinical trial sites to ensure data quality, patient safety, and regulatory compliance.",
        "day_in_life": "Audit clinical sites, review source documents, train site staff on protocols, and file monitoring visit reports.",
        "salary_range": "$60K–$110K",
        "salary_range_india": "₹3.5L–₹10L",
        "progression": ["Junior CRA", "CRA II", "Senior CRA", "Lead CRA / Clinical Trial Manager", "Director of Clinical Operations"],
        "industries": ["CROs", "Pharma", "Biotech", "Medical devices"],
        "key_concepts": ["trial_basics", "phase1", "phase2_3", "trial_design_stats", "phase4_pv", "gcp_ethics", "regulatory_bodies", "submissions"],
        "relevant_subjects": ["clinical_trials", "drug_discovery"],
        "min_qualification": "BTech / BSc",
    },
    "regulatory_affairs_associate": {
        "id": "regulatory_affairs_associate", "title": "Regulatory Affairs Associate",
        "cluster": "Science & Technical", "icon": "📑",
        "description": "Prepare and submit regulatory dossiers (IND, NDA, BLA, MAA) and liaise with agencies like FDA and EMA.",
        "day_in_life": "Compile CTD modules, review labelling, track agency correspondence, and advise teams on regulatory strategy.",
        "salary_range": "$70K–$130K",
        "salary_range_india": "₹4L–₹14L",
        "progression": ["Regulatory Associate", "RA Specialist", "Senior Specialist", "RA Manager", "Director of Regulatory Affairs"],
        "industries": ["Pharma", "Biotech", "Medical devices", "CROs", "Consultancies"],
        "key_concepts": ["regulatory_bodies", "submissions", "gcp_ethics", "special_pathways", "trial_basics", "phase1", "phase2_3", "admet", "pk_pd", "labeling_postmarket", "dct_adaptive"],
        "relevant_subjects": ["clinical_trials", "drug_discovery", "cell_gene_therapy", "rna_therapeutics"],
        "min_qualification": "BTech / BSc",
    },
    "computational_biologist": {
        "id": "computational_biologist", "title": "Computational Biologist",
        "cluster": "Science & Technical", "icon": "💻",
        "description": "Build mathematical models and computational tools to understand biological systems at a molecular or systems level.",
        "day_in_life": "Develop algorithms for sequence analysis, model gene regulatory networks, build ML models on omics data, and publish research.",
        "salary_range": "$85K–$155K",
        "salary_range_india": "₹7L–₹22L",
        "progression": ["Research Associate", "Computational Biologist", "Senior Computational Biologist", "Principal Scientist", "Director of Computational Biology"],
        "industries": ["Pharma R&D", "Biotech", "AI-in-bio startups", "Academia"],
        "key_concepts": ["central_dogma", "seq_formats", "pairwise_alignment", "blast_search", "bio_programming", "genome_structure", "transcriptomics", "epigenomics", "comp_drug_disc", "ml_bioinformatics", "protein_structure", "ml_foundations"],
        "relevant_subjects": ["bioinformatics", "genomics", "genai_ml", "protein_engineering"],
        "min_qualification": "BTech / BSc",
    },
    "pharmacovigilance_scientist": {
        "id": "pharmacovigilance_scientist", "title": "Pharmacovigilance Scientist",
        "cluster": "Science & Technical", "icon": "🔍",
        "description": "Monitor drug safety signals in post-market data and manage adverse event reporting to regulators.",
        "day_in_life": "Review case safety reports, perform signal detection, prepare PSURs, and liaise with regulatory agencies.",
        "salary_range": "$65K–$120K",
        "salary_range_india": "₹4L–₹12L",
        "progression": ["PV Associate", "Safety Scientist", "Senior PV Scientist", "PV Manager", "Head of Drug Safety"],
        "industries": ["Pharma", "Biotech", "CROs", "Regulatory agencies"],
        "key_concepts": ["phase4_pv", "gcp_ethics", "trial_basics", "admet", "pk_pd", "regulatory_bodies", "submissions", "labeling_postmarket"],
        "relevant_subjects": ["clinical_trials", "drug_discovery"],
        "min_qualification": "BTech / BSc",
    },
    "medical_science_liaison": {
        "id": "medical_science_liaison", "title": "Medical Science Liaison (MSL)",
        "cluster": "Science & Technical", "icon": "🤝",
        "description": "Bridge scientific knowledge between your company's pipeline and key opinion leaders, physicians, and payers.",
        "day_in_life": "Meet with oncologists, present clinical trial data, answer medical questions, gather HCP insights.",
        "salary_range": "$120K–$180K",
        "salary_range_india": "₹12L–₹28L",
        "progression": ["Associate MSL", "MSL", "Senior MSL", "Regional MSL Manager", "Director of Medical Affairs"],
        "industries": ["Pharma", "Biotech", "Medical devices"],
        "key_concepts": ["pipeline_overview", "biologics", "pk_pd", "trial_basics", "phase1", "phase2_3", "phase4_pv", "precision_medicine", "regulatory_bodies"],
        "relevant_subjects": ["clinical_trials", "drug_discovery", "cell_gene_therapy", "rna_therapeutics"],
        "min_qualification": "MSc / MD / PharmD",
    },
    "biomarker_scientist": {
        "id": "biomarker_scientist", "title": "Biomarker & Translational Scientist",
        "cluster": "Science & Technical", "icon": "🎯",
        "description": "Discover and validate biomarkers that stratify patients, predict drug response, or serve as clinical endpoints.",
        "day_in_life": "Analyse multi-omics datasets, run assay validation, and collaborate with clinical teams to embed biomarkers into trial protocols.",
        "salary_range": "$85K–$150K",
        "salary_range_india": "₹6L–₹18L",
        "progression": ["Translational Scientist I", "Biomarker Scientist", "Senior Scientist", "Principal Translational Scientist", "Head of Biomarkers"],
        "industries": ["Pharma", "Biotech", "Diagnostics companies", "Precision medicine startups"],
        "key_concepts": ["central_dogma", "genome_structure", "variant_types", "transcriptomics", "epigenomics", "single_cell", "precision_medicine", "ngs_alignment", "target_id", "admet", "biomarkers_dd"],
        "relevant_subjects": ["genomics", "bioinformatics", "drug_discovery", "genai_ml", "clinical_trials"],
        "min_qualification": "MSc / MTech",
    },
    "clinical_data_manager": {
        "id": "clinical_data_manager", "title": "Clinical Data Manager",
        "cluster": "Science & Technical", "icon": "📊",
        "description": "Oversee the collection, cleaning, and integrity of clinical trial data to support regulatory submissions.",
        "day_in_life": "Build EDC systems, write data validation rules, run edit checks, reconcile SAE data, and prepare clean database lock.",
        "salary_range": "$65K–$115K",
        "salary_range_india": "₹4L–₹14L",
        "progression": ["Data Coordinator", "Clinical Data Manager", "Senior CDM", "Data Management Lead", "Head of Data Management"],
        "industries": ["CROs", "Pharma", "Biotech", "Academic medical centers"],
        "key_concepts": ["trial_basics", "trial_design_stats", "phase1", "phase2_3", "gcp_ethics", "bio_programming"],
        "relevant_subjects": ["clinical_trials", "drug_discovery"],
        "min_qualification": "BTech / BSc",
    },
    "biotech_bd_associate": {
        "id": "biotech_bd_associate", "title": "Biotech Business Development Associate",
        "cluster": "Business & Commercial", "icon": "🤝",
        "description": "Identify partnering, licensing, and M&A opportunities to advance a biotech's pipeline and commercial strategy.",
        "day_in_life": "Screen deal opportunities, build financial models, draft term sheets, and coordinate due diligence.",
        "salary_range": "$80K–$140K",
        "salary_range_india": "₹7L–₹18L",
        "progression": ["BD Analyst", "BD Associate", "Senior Associate", "Director of BD", "VP of Business Development"],
        "industries": ["Biotech", "Pharma", "VC-backed startups"],
        "key_concepts": ["pipeline_overview", "target_id", "biologics", "trial_basics", "phase1", "phase2_3", "regulatory_bodies", "submissions", "special_pathways", "bd_licensing", "ls_valuation"],
        "relevant_subjects": ["biotech_business", "drug_discovery", "clinical_trials", "cell_gene_therapy", "rna_therapeutics"],
        "min_qualification": "BTech / BSc",
    },
    "market_access_analyst": {
        "id": "market_access_analyst", "title": "Market Access & HEOR Analyst",
        "cluster": "Business & Commercial", "icon": "📈",
        "description": "Build the evidence and economic case for payer reimbursement of new drugs, devices, or diagnostics.",
        "day_in_life": "Develop cost-effectiveness models, write HEOR publications, prepare payer dossiers, engage with HTA bodies.",
        "salary_range": "$75K–$130K",
        "salary_range_india": "₹6L–₹16L",
        "progression": ["HEOR Analyst", "Market Access Analyst", "Senior Analyst", "Manager", "Director of Market Access"],
        "industries": ["Pharma", "Biotech", "Consultancies", "Payer organisations"],
        "key_concepts": ["pipeline_overview", "pk_pd", "admet", "phase4_pv", "regulatory_bodies", "submissions", "special_pathways", "trial_design_stats", "market_access_heor"],
        "relevant_subjects": ["biotech_business", "clinical_trials", "drug_discovery"],
        "min_qualification": "BTech / BSc",
    },
    "medical_affairs_associate": {
        "id": "medical_affairs_associate", "title": "Medical Affairs Associate",
        "cluster": "Business & Commercial", "icon": "🏥",
        "description": "Support the scientific and clinical activities that connect a drug to the medical community after approval.",
        "day_in_life": "Develop medical education programmes, review promotional materials for accuracy, respond to unsolicited medical information requests.",
        "salary_range": "$70K–$125K",
        "salary_range_india": "₹7L–₹18L",
        "progression": ["Medical Affairs Associate", "Specialist", "Manager", "Senior Manager", "Director of Medical Affairs"],
        "industries": ["Pharma", "Biotech", "Medical devices"],
        "key_concepts": ["pipeline_overview", "biologics", "pk_pd", "trial_basics", "phase1", "phase2_3", "phase4_pv", "gcp_ethics", "regulatory_bodies", "labeling_postmarket"],
        "relevant_subjects": ["clinical_trials", "drug_discovery", "cell_gene_therapy"],
        "min_qualification": "MSc / PharmD",
    },
    "genomics_commercial_specialist": {
        "id": "genomics_commercial_specialist", "title": "Genomics Commercial Specialist",
        "cluster": "Business & Commercial", "icon": "💹",
        "description": "Sell or support the commercial deployment of genomics platforms, sequencing instruments, and bioinformatics software.",
        "day_in_life": "Demo NGS instruments to hospital genomics labs, support grant applications, train lab staff, track competitive landscape.",
        "salary_range": "$80K–$160K (base + commission)",
        "salary_range_india": "₹8L–₹22L (base + commission)",
        "progression": ["Sales Representative", "Commercial Specialist", "Senior Specialist", "Field Application Scientist", "Regional Sales Manager"],
        "industries": ["Genomics companies", "NGS instrument makers", "Biotech tools & reagents"],
        "key_concepts": ["genome_structure", "sequencing_tech", "variant_types", "gwas", "precision_medicine", "single_cell", "population_genetics", "ngs_qc", "pipeline_overview"],
        "relevant_subjects": ["genomics", "bioinformatics", "clinical_trials"],
        "min_qualification": "BTech / BSc",
    },
    "biotech_product_manager": {
        "id": "biotech_product_manager", "title": "Biotech Product Manager",
        "cluster": "Business & Commercial", "icon": "🗂️",
        "description": "Own the strategic roadmap for a drug, platform, or diagnostic product from launch through lifecycle management.",
        "day_in_life": "Write the product strategy, coordinate cross-functional teams, manage launch plans, track market share.",
        "salary_range": "$100K–$170K",
        "salary_range_india": "₹15L–₹40L",
        "progression": ["Associate PM", "Product Manager", "Senior PM", "Director of Product", "VP of Product / CMO"],
        "industries": ["Pharma", "Biotech", "Genomics", "Diagnostics", "Digital health"],
        "key_concepts": ["pipeline_overview", "target_id", "admet", "pk_pd", "trial_basics", "regulatory_bodies", "submissions", "special_pathways", "precision_medicine", "market_strategy", "market_access_heor"],
        "relevant_subjects": ["biotech_business", "drug_discovery", "clinical_trials", "cell_gene_therapy", "rna_therapeutics"],
        "min_qualification": "BTech / BSc",
    },
    "life_sciences_consultant": {
        "id": "life_sciences_consultant", "title": "Life Sciences Consultant",
        "cluster": "Business & Commercial", "icon": "💼",
        "description": "Advise pharma, biotech, and health-system clients on strategy, commercial excellence, and R&D productivity.",
        "day_in_life": "Build strategy decks and financial models, interview subject matter experts, present findings to C-suite executives.",
        "salary_range": "$90K–$160K",
        "salary_range_india": "₹10L–₹30L",
        "progression": ["Analyst", "Consultant", "Senior Consultant", "Manager", "Principal / Partner"],
        "industries": ["McKinsey / BCG / Bain", "IQVIA / ZS / Huron", "Big4 life sciences"],
        "key_concepts": ["pipeline_overview", "target_id", "admet", "trial_basics", "regulatory_bodies", "submissions", "genome_structure", "precision_medicine", "market_strategy", "ls_valuation", "market_access_heor"],
        "relevant_subjects": ["biotech_business", "drug_discovery", "clinical_trials", "genomics"],
        "min_qualification": "BTech / BSc",
    },
    "biotech_venture_analyst": {
        "id": "biotech_venture_analyst", "title": "Biotech Venture Analyst",
        "cluster": "Business & Commercial", "icon": "💰",
        "description": "Evaluate early-stage biotech investment opportunities and support portfolio companies in raising capital.",
        "day_in_life": "Read scientific papers to assess platform novelty, model company valuations, attend pitch meetings, write investment memos.",
        "salary_range": "$100K–$160K",
        "salary_range_india": "₹10L–₹25L",
        "progression": ["Analyst", "Associate", "Senior Associate", "Principal", "Partner"],
        "industries": ["VC firms", "Corporate venture arms", "Family offices with healthcare focus"],
        "key_concepts": ["pipeline_overview", "target_id", "biologics", "comp_drug_disc", "trial_basics", "phase1", "phase2_3", "regulatory_bodies", "submissions", "special_pathways", "ls_valuation", "biotech_financing", "biotech_biz_model"],
        "relevant_subjects": ["biotech_business", "drug_discovery", "clinical_trials", "cell_gene_therapy", "rna_therapeutics", "genai_ml"],
        "min_qualification": "BTech / BSc",
    },
    "licensing_partnerships": {
        "id": "licensing_partnerships", "title": "Licensing & Partnerships Associate",
        "cluster": "Business & Commercial", "icon": "📝",
        "description": "Structure and negotiate licensing agreements, co-development deals, and research collaborations.",
        "day_in_life": "Review deal precedents, prepare term sheet analysis, coordinate IP due diligence, support contract negotiations.",
        "salary_range": "$75K–$135K",
        "salary_range_india": "₹7L–₹18L",
        "progression": ["Analyst / Paralegal", "Licensing Associate", "Senior Associate", "Manager / Director", "VP of Licensing"],
        "industries": ["Pharma", "Biotech", "Technology transfer offices", "IP law firms"],
        "key_concepts": ["pipeline_overview", "target_id", "biologics", "trial_basics", "regulatory_bodies", "submissions", "special_pathways", "bd_licensing", "ip_biotech", "ls_valuation"],
        "relevant_subjects": ["biotech_business", "drug_discovery", "clinical_trials"],
        "min_qualification": "BTech / BSc",
    },
    "ai_drug_discovery": {
        "id": "ai_drug_discovery", "title": "AI in Drug Discovery Specialist",
        "cluster": "Emerging & Hybrid", "icon": "🤖",
        "description": "Apply machine learning, deep learning, and generative AI to accelerate target identification, virtual screening, and molecular design.",
        "day_in_life": "Train GNNs on molecular data, run AlphaFold structure predictions, build QSAR models, collaborate with wet-lab scientists.",
        "salary_range": "$110K–$200K",
        "salary_range_india": "₹18L–₹55L",
        "progression": ["ML Engineer / Research Scientist", "AI Drug Discovery Scientist", "Senior Scientist", "Principal Scientist", "VP of AI/ML"],
        "industries": ["AI-driven biotech (Recursion, Insilico Medicine, Schrödinger)", "Pharma AI divisions", "Startups"],
        "key_concepts": ["comp_drug_disc", "target_id", "hit_discovery", "lead_optimization", "admet", "seq_formats", "bio_programming", "genome_structure", "central_dogma", "ml_foundations", "deep_learning_fund", "gnn_drug_disc", "gen_molecules", "protein_lang_models"],
        "relevant_subjects": ["genai_ml", "drug_discovery", "protein_engineering", "bioinformatics", "genomics"],
        "min_qualification": "BTech / BSc",
    },
    "precision_medicine_specialist": {
        "id": "precision_medicine_specialist", "title": "Precision Medicine & Genomics Specialist",
        "cluster": "Emerging & Hybrid", "icon": "🎯",
        "description": "Translate genomic insights into clinically actionable, patient-specific treatment decisions.",
        "day_in_life": "Interpret clinical genomics reports, integrate multi-omics data, advise oncologists on companion diagnostic use.",
        "salary_range": "$90K–$160K",
        "salary_range_india": "₹8L–₹24L",
        "progression": ["Genomics Specialist", "Precision Medicine Associate", "Senior Specialist", "Director of Precision Medicine", "Chief Scientific Officer"],
        "industries": ["Precision medicine companies", "Academic medical centers", "Pharma genomic medicine divisions", "Diagnostics"],
        "key_concepts": ["genome_structure", "sequencing_tech", "variant_types", "gwas", "precision_medicine", "transcriptomics", "epigenomics", "single_cell", "population_genetics", "central_dogma", "ngs_alignment", "admet", "pk_pd", "biologics", "biomarkers_dd"],
        "relevant_subjects": ["genomics", "bioinformatics", "clinical_trials", "genai_ml", "cell_gene_therapy"],
        "min_qualification": "MSc / MTech",
    },
    "biotech_founder": {
        "id": "biotech_founder", "title": "Biotech Entrepreneur & Founder",
        "cluster": "Emerging & Hybrid", "icon": "🚀",
        "description": "Start and build a biotech company  -  from scientific hypothesis to VC funding, IND filing, and clinical proof-of-concept.",
        "day_in_life": "Pitch investors, recruit a scientific board, design company strategy, manage burn rate, and keep a hand in the science.",
        "salary_range": "$0–$500K+ (equity-driven)",
        "salary_range_india": "₹0–₹2Cr+ (equity-driven)",
        "progression": ["Junior Scientist / Associate", "Co-founder / CTO / CSO", "CEO", "Public Company Executive", "Serial Founder"],
        "industries": ["Therapeutics", "Genomics", "Diagnostics", "Digital health", "Ag-bio"],
        "key_concepts": ["pipeline_overview", "target_id", "biologics", "comp_drug_disc", "admet", "trial_basics", "regulatory_bodies", "submissions", "special_pathways", "genome_structure", "precision_medicine", "bio_programming", "biotech_biz_model", "biotech_financing", "ls_valuation", "building_biotech"],
        "relevant_subjects": ["biotech_business", "drug_discovery", "clinical_trials", "cell_gene_therapy", "rna_therapeutics", "genai_ml"],
        "min_qualification": "Any background",
    },
}

# ── Capstones ─────────────────────────────────────────────────────────────────

CAPSTONES = {
    "bioinformatics": {
        "subject_id": "bioinformatics",
        "title": "From Sequence to Insight: A Cancer Mutational Landscape Analysis",
        "problem_statement": "You have been given access to a publicly available tumour sequencing dataset from TCGA (The Cancer Genome Atlas). Your task is to build a variant annotation and analysis pipeline, identify driver mutations, classify their functional significance using ClinVar and OncoKB, and produce an analysis report as if presenting findings to an oncology multidisciplinary team.",
        "instructions": "1. Download a TCGA cancer dataset of your choice (BRCA, LUAD, or COAD recommended) from the GDC Data Portal.\n2. Run variant annotation using a tool of your choice (ANNOVAR, VEP, or SnpEff).\n3. Cross-reference significant variants against ClinVar and OncoKB.\n4. Identify the top 5–10 driver mutations and classify each by functional significance.\n5. Write your analysis report (2,000–2,500 words) covering: dataset overview, methodology, key findings, clinical interpretation, and limitations.\n6. Include your annotated code as a Jupyter notebook or R Markdown file with all outputs saved.",
        "deliverable": "ZIP file containing: (1) PDF report (2,000–2,500 words), (2) Jupyter notebook or R Markdown file with saved outputs. Max 50MB.",
        "rubric": [
            {"criterion": "Data handling and pipeline setup", "marks": 20},
            {"criterion": "Variant annotation and filtering", "marks": 20},
            {"criterion": "Driver mutation identification and classification", "marks": 25},
            {"criterion": "Clinical interpretation and reasoning", "marks": 20},
            {"criterion": "Code quality and reproducibility", "marks": 15},
        ],
        "total_marks": 100, "unlock_threshold": 24, "accepted_formats": ["zip"], "max_size_mb": 50,
    },
    "genomics": {
        "subject_id": "genomics",
        "title": "Clinical Variant Interpretation: A Patient Case Report",
        "problem_statement": "You are a genomics scientist at a clinical genomics laboratory. A patient has been referred with a suspected hereditary cancer syndrome. You have been provided with a set of genomic variants from a simulated whole-genome sequence, a family history, and a clinical summary. Your task is to classify each variant using ACMG/AMP guidelines, identify the most likely pathogenic finding, and write a clinical genomics report suitable for an MDT meeting.",
        "instructions": "1. Download the patient case file (provided with this brief  -  synthetic data only).\n2. Look up each variant in ClinVar, OMIM, and gnomAD.\n3. Apply ACMG/AMP variant classification criteria to each variant (Pathogenic / Likely Pathogenic / VUS / Likely Benign / Benign).\n4. Identify the primary finding most likely explaining the patient's phenotype.\n5. Write a clinical genomics report following the structure used by NHS Genomic Medicine Service or a comparable lab: patient details (synthetic), variant summary table, clinical interpretation, recommendations.\n6. Include a one-page methods note describing your classification process.",
        "deliverable": "PDF report following the provided clinical report template. Max 20MB.",
        "rubric": [
            {"criterion": "ACMG/AMP variant classification accuracy", "marks": 30},
            {"criterion": "Database evidence sourcing and interpretation", "marks": 20},
            {"criterion": "Clinical reasoning and primary finding identification", "marks": 25},
            {"criterion": "Report structure, clarity, and professionalism", "marks": 15},
            {"criterion": "Methods note completeness", "marks": 10},
        ],
        "total_marks": 100, "unlock_threshold": 24, "accepted_formats": ["pdf", "zip"], "max_size_mb": 20,
    },
    "drug_discovery": {
        "subject_id": "drug_discovery",
        "title": "Drug Discovery Strategy Memo: Targeting an Unmet Need",
        "problem_statement": "You are a drug discovery scientist presenting to Genentech's R&D strategy committee. Select a disease from the provided shortlist (ALS, Huntington's disease, triple-negative breast cancer, or idiopathic pulmonary fibrosis) and develop a complete drug discovery strategy memo  -  from target identification through to an IND-ready preclinical plan.",
        "instructions": "1. Choose one disease from the shortlist and identify a validated biological target using OpenTargets, UniProt, and the published literature.\n2. Justify your target choice: mechanism of disease relevance, genetic validation evidence, druggability score.\n3. Propose a hit discovery strategy (HTS, fragment-based, virtual screening, or phenotypic  -  justify your choice).\n4. Run a basic in silico docking analysis using AutoDock Vina or SwissDock with a PDB structure for your target. Include at least one screenshot of your docking result with interpretation.\n5. Outline your lead optimisation priorities (key ADMET liabilities to address, selectivity concerns).\n6. Describe the minimum preclinical package needed to file an IND.\n7. Write the memo (2,000–2,500 words) plus docking results as appendices.",
        "deliverable": "PDF memo (2,000–2,500 words) with docking screenshots and supporting data as appendices. Max 30MB.",
        "rubric": [
            {"criterion": "Target identification and validation quality", "marks": 25},
            {"criterion": "Hit discovery strategy justification", "marks": 20},
            {"criterion": "In silico docking  -  execution and interpretation", "marks": 20},
            {"criterion": "Lead optimisation and ADMET reasoning", "marks": 20},
            {"criterion": "Preclinical IND plan and overall scientific rigour", "marks": 15},
        ],
        "total_marks": 100, "unlock_threshold": 24, "accepted_formats": ["pdf", "zip"], "max_size_mb": 30,
    },
    "clinical_trials": {
        "subject_id": "clinical_trials",
        "title": "Design a Phase II Trial & Regulatory Strategy",
        "problem_statement": "A novel small molecule kinase inhibitor has completed Phase I in patients with advanced solid tumours. It demonstrated acceptable tolerability at 200mg QD, with preliminary PK data showing t½ of 8 hours and dose-proportional exposure. One confirmed partial response was seen in a KRAS-mutant pancreatic cancer patient. Your task is to design the Phase II programme and regulatory strategy that will take this compound toward a potential NDA.",
        "instructions": "1. Define your Phase II patient population, including biomarker selection strategy (KRAS-mutant enriched vs all-comers  -  justify).\n2. Choose your primary endpoint and justify why it is appropriate and acceptable to FDA/EMA.\n3. Design the trial: phase II design (single-arm, randomised, adaptive?), sample size calculation with assumptions stated, control arm if applicable.\n4. Write a two-page trial synopsis in ICH format.\n5. Write a one-page regulatory strategy memo: which pathway (standard, Breakthrough Therapy, Fast Track?), key FDA/EMA interactions needed, estimated timeline from Phase II start to NDA submission.\n6. Identify the three biggest regulatory risks and how you would mitigate them.",
        "deliverable": "PDF document comprising trial synopsis (2 pages) and regulatory strategy memo (1–2 pages). Total 2,500–3,000 words. Max 20MB.",
        "rubric": [
            {"criterion": "Patient population and biomarker strategy", "marks": 20},
            {"criterion": "Endpoint selection and statistical design", "marks": 25},
            {"criterion": "Trial synopsis completeness and ICH compliance", "marks": 20},
            {"criterion": "Regulatory strategy and pathway choice", "marks": 20},
            {"criterion": "Risk identification and mitigation", "marks": 15},
        ],
        "total_marks": 100, "unlock_threshold": 24, "accepted_formats": ["pdf"], "max_size_mb": 20,
    },
    "genai_ml": {
        "subject_id": "genai_ml",
        "title": "Build and Validate an ADMET Prediction Model",
        "problem_statement": "You are an AI scientist at a drug discovery company. Your team needs a reliable machine learning model to predict one ADMET property for early compound triage. Using publicly available data from ChEMBL, build, evaluate, and document a model that could realistically be used to prioritise compounds before expensive in vitro testing.",
        "instructions": "1. Choose one ADMET property: blood-brain barrier penetration, hERG cardiotoxicity, or aqueous solubility.\n2. Download a suitable labelled dataset from ChEMBL (minimum 1,000 compounds).\n3. Engineer molecular features (Morgan fingerprints, RDKit descriptors, or both).\n4. Train at least two model architectures (e.g. random forest + a GNN or random forest + XGBoost) and compare performance.\n5. Evaluate rigorously: ROC-AUC, precision-recall curve, and performance on a held-out test set (do not touch the test set until final evaluation).\n6. Write a model card (1,000–1,500 words) covering: data source and curation, features, model choices and trade-offs, evaluation results, known limitations, and regulatory considerations for using this model in a real drug discovery setting.\n7. Include your Jupyter notebook with all outputs saved.",
        "deliverable": "ZIP file containing: (1) Jupyter notebook with all outputs saved, (2) PDF model card (1,000–1,500 words). Max 50MB.",
        "rubric": [
            {"criterion": "Data curation and feature engineering", "marks": 20},
            {"criterion": "Model implementation and training", "marks": 20},
            {"criterion": "Evaluation rigour (test set discipline, metrics)", "marks": 25},
            {"criterion": "Model card quality  -  limitations and regulatory thinking", "marks": 20},
            {"criterion": "Code quality and reproducibility", "marks": 15},
        ],
        "total_marks": 100, "unlock_threshold": 24, "accepted_formats": ["zip"], "max_size_mb": 50,
    },
    "biotech_business": {
        "subject_id": "biotech_business",
        "title": "Series B Pitch: Build the Business Case for a Biotech Asset",
        "problem_statement": "You are co-founder and CEO of a Series A biotech. Your lead asset is a first-in-class oral PCSK9 inhibitor for hypercholesterolaemia in patients intolerant to statins. Phase I showed clean safety and PK consistent with QD dosing. You are now raising a $75M Series B to fund a Phase II/III pivotal trial and build out your commercial team. Build the investor pitch deck and financial model.",
        "instructions": "1. Build a 12-slide pitch deck (PowerPoint or PDF) covering: cover slide, problem / unmet need, your solution and MOA, clinical data to date (Phase I summary), Phase II/III development plan, market opportunity and patient sizing, competitive landscape, regulatory strategy, IP position and exclusivity runway, team, financial projections (3-year pro forma), and Series B ask with use of proceeds.\n2. Build an rNPV model in Excel or Google Sheets. Assumptions must include: probability of success by phase (use industry benchmarks), peak year sales estimate, royalty or net margin assumption, discount rate, and development timeline. Show your working.\n3. Write a one-page investment thesis (why this asset, why now, why this team) to accompany the deck.\n4. Identify the three biggest investor objections you expect and how you would respond.",
        "deliverable": "ZIP file containing: (1) PDF pitch deck (12 slides), (2) Excel/Google Sheets rNPV model, (3) PDF investment thesis and objection responses (1–2 pages). Max 30MB.",
        "rubric": [
            {"criterion": "Market opportunity sizing and commercial strategy", "marks": 20},
            {"criterion": "rNPV model construction and assumptions", "marks": 25},
            {"criterion": "Clinical and regulatory strategy in the deck", "marks": 20},
            {"criterion": "Pitch deck narrative, structure, and clarity", "marks": 20},
            {"criterion": "Investment thesis and objection handling", "marks": 15},
        ],
        "total_marks": 100, "unlock_threshold": 24, "accepted_formats": ["zip"], "max_size_mb": 30,
    },
    "cell_gene_therapy": {
        "subject_id": "cell_gene_therapy",
        "title": "Design a Gene Therapy Programme for a Monogenic Disease",
        "problem_statement": "You are a gene therapy scientist at a clinical-stage biotech. Select one monogenic disease from the shortlist (sickle cell disease, haemophilia A, Duchenne muscular dystrophy, or a rare inherited retinal dystrophy) and design a complete gene therapy programme  -  from vector selection through to an IND-ready development plan.",
        "instructions": "1. Choose your disease and justify it: unmet need, patient population size, genetic target, and why gene therapy is the right modality.\n2. Select your therapeutic approach (AAV in vivo, lentiviral ex vivo HSC correction, CRISPR base editing, or other) and justify the choice over alternatives.\n3. Vector/edit design: serotype or vector selection, transgene cassette or guide RNA design, anticipated tropism, packaging capacity considerations.\n4. Safety assessment plan: immunogenicity mitigation strategy, genotoxicity monitoring plan (for integrating vectors), off-target analysis approach.\n5. Manufacturing strategy: production platform, scale, key analytical release assays, critical quality attributes.\n6. Regulatory pathway memo (1 page): which FDA/EMA pathway, RMAT or PRIME eligibility, key data packages needed for IND/CTA filing.\n7. Write the full programme document (2,500–3,000 words) plus a one-page regulatory memo.",
        "deliverable": "PDF document: programme plan (2,500–3,000 words) plus regulatory memo (1 page). Max 20MB.",
        "rubric": [
            {"criterion": "Disease selection rationale and unmet need", "marks": 15},
            {"criterion": "Therapeutic approach and vector/edit design", "marks": 25},
            {"criterion": "Safety strategy depth and scientific reasoning", "marks": 25},
            {"criterion": "Manufacturing and analytical plan", "marks": 20},
            {"criterion": "Regulatory pathway memo", "marks": 15},
        ],
        "total_marks": 100, "unlock_threshold": 24, "accepted_formats": ["pdf", "zip"], "max_size_mb": 20,
    },
    "protein_engineering": {
        "subject_id": "protein_engineering",
        "title": "Computational Design of a Therapeutic Protein",
        "problem_statement": "You are a protein design scientist. Using available computational tools, design an improved or novel therapeutic protein for a target of your choice. Your submission should demonstrate end-to-end thinking from design rationale through to experimental validation strategy.",
        "instructions": "1. Choose a therapeutic application: an improved enzyme replacement therapy, a de novo binder to a target of your choice, a bispecific antibody format, or an engineered cytokine with improved selectivity.\n2. Describe the design strategy: which approach (directed evolution, rational design, AlphaFold-guided, RFdiffusion de novo, or hybrid), and why it is appropriate for your target.\n3. Run at least one computational step and include outputs: use AlphaFold2 (ColabFold) to predict a structure, or use ProteinMPNN to generate sequences for a backbone, or document a directed evolution selection strategy with specific library design.\n4. Developability assessment: predict and address at least three developability liabilities (aggregation, immunogenicity, solubility, viscosity).\n5. Experimental validation plan: what assays would you run to confirm function, binding affinity, and stability? What would a go/no-go decision look like?\n6. Write the design report (2,000–2,500 words) with computational outputs included as figures.",
        "deliverable": "PDF report (2,000–2,500 words) with computational output figures (AlphaFold structure, ProteinMPNN outputs, or library design schematic). Max 30MB.",
        "rubric": [
            {"criterion": "Design rationale and therapeutic application choice", "marks": 20},
            {"criterion": "Computational workflow execution and interpretation", "marks": 30},
            {"criterion": "Developability analysis", "marks": 20},
            {"criterion": "Experimental validation plan", "marks": 20},
            {"criterion": "Scientific writing and figure quality", "marks": 10},
        ],
        "total_marks": 100, "unlock_threshold": 24, "accepted_formats": ["pdf", "zip"], "max_size_mb": 30,
    },
    "rna_therapeutics": {
        "subject_id": "rna_therapeutics",
        "title": "Design an mRNA Therapeutic or siRNA Drug",
        "problem_statement": "You are an RNA drug designer. Choose one of two tracks: (A) design an mRNA therapeutic or vaccine for a disease of your choice, or (B) design a siRNA or ASO programme targeting a gene of therapeutic interest. Produce a complete design and development document.",
        "instructions": "Track A (mRNA): 1. Choose your disease and therapeutic goal (vaccine antigen, protein replacement, or gene editing mRNA). 2. Design the mRNA molecule: cap selection, 5' UTR (use literature or UTR design tools), codon-optimised ORF (use a codon optimisation tool and show before/after), 3' UTR, poly-A tail length. 3. Choose and justify your delivery system (LNP formulation, organ targeting). 4. Address immunogenicity: which modified nucleosides, how you would measure innate immune activation in vitro. 5. Outline the IND-enabling CMC package.\n\nTrack B (siRNA/ASO): 1. Choose your target gene and disease indication. 2. Design 3 candidate siRNA duplexes or an ASO sequence (use siRNA design tools or mfold; include sequences). 3. Select and justify chemical modifications. 4. Choose a delivery approach (GalNAc conjugate, LNP, naked ASO for CNS). 5. Describe your in vitro knockdown validation plan and lead selection criteria.\n\nBoth tracks: write a 2,000–2,500 word design report with sequences and justification for all design choices.",
        "deliverable": "PDF design report (2,000–2,500 words) including all sequences, modification tables, and design rationale. Max 20MB.",
        "rubric": [
            {"criterion": "Target and indication selection rationale", "marks": 15},
            {"criterion": "Molecular design decisions (sequences, modifications)", "marks": 30},
            {"criterion": "Delivery system selection and justification", "marks": 20},
            {"criterion": "Immunogenicity or off-target strategy", "marks": 20},
            {"criterion": "CMC/development plan or validation plan", "marks": 15},
        ],
        "total_marks": 100, "unlock_threshold": 24, "accepted_formats": ["pdf", "zip"], "max_size_mb": 20,
    },
    "biomanufacturing": {
        "subject_id": "biomanufacturing",
        "title": "Bioprocess Development Plan for a Biologic",
        "problem_statement": "You are a bioprocess development scientist. A research team has handed you a new monoclonal antibody candidate (or, if you prefer, a cell therapy product) and asked you to develop the manufacturing process from cell line to drug substance. Build a complete bioprocess development plan.",
        "instructions": "Choose either a monoclonal antibody OR a CAR-T cell therapy product. Then:\n1. Cell line / starting material: select and justify your host cell system (CHO for mAb; primary T cells for CAR-T). Describe the cell line development or starting material qualification steps.\n2. Upstream process: propose a bioreactor mode (fed-batch vs perfusion for mAb; G-REX or stirred tank for CAR-T). Define your key process parameters and acceptable ranges for at least 4 critical parameters.\n3. Downstream process (mAb) or cell processing (CAR-T): map the full purification train or manufacturing workflow with justification for each step.\n4. Critical Quality Attributes: list the 5 most important CQAs for your product and the analytical method for each.\n5. Scale-up plan: describe how you would scale from lab-scale to clinical manufacturing, including tech transfer considerations.\n6. GMP readiness checklist: what quality systems must be in place before first human dose?\n7. Write the plan (2,500–3,000 words) with process flow diagrams as figures.",
        "deliverable": "PDF document (2,500–3,000 words) with process flow diagram(s). Max 25MB.",
        "rubric": [
            {"criterion": "Cell line / starting material strategy", "marks": 15},
            {"criterion": "Upstream process design and parameter rationale", "marks": 25},
            {"criterion": "Downstream / cell processing workflow", "marks": 25},
            {"criterion": "CQA identification and analytical methods", "marks": 20},
            {"criterion": "Scale-up and GMP readiness plan", "marks": 15},
        ],
        "total_marks": 100, "unlock_threshold": 24, "accepted_formats": ["pdf", "zip"], "max_size_mb": 25,
    },
    "longevity_science": {
        "subject_id": "longevity_science",
        "title": "Design a Longevity Intervention Programme",
        "problem_statement": "You are a scientist at a longevity biotech. Design a therapeutic intervention strategy targeting one or more hallmarks of aging. Your programme should have a credible path to a clinical trial, not just a theoretical mechanism.",
        "instructions": "1. Select your target hallmark(s) of aging and justify the choice based on: strength of mechanistic evidence, reversibility, druggability, and existing clinical precedent.\n2. Propose a specific intervention: a senolytic compound, an epigenetic reprogramming approach, a metabolic intervention (rapamycin analogue, NAD+ precursor, etc.), or a cell therapy. If it is an existing molecule, propose a new indication or regimen; if novel, describe the drug class.\n3. Biomarker strategy: which aging clocks or biomarkers would you use to measure efficacy? Justify your choice of primary biological age endpoint and at least two secondary biomarkers. Address the FDA's current position on biological age as an endpoint.\n4. Clinical trial design: propose a Phase II trial design. Who is the patient population (what age, what health status)? What is the primary endpoint and how long is the follow-up? What safety signals would stop the trial?\n5. Risk and ethics section: address the key scientific risks (is the hallmark causative or correlative in humans?), regulatory risks, and the ethical questions around longevity interventions and access.\n6. Write the programme document (2,500–3,000 words).",
        "deliverable": "PDF programme document (2,500–3,000 words). Max 20MB.",
        "rubric": [
            {"criterion": "Hallmark selection and mechanistic rationale", "marks": 20},
            {"criterion": "Intervention design and scientific credibility", "marks": 25},
            {"criterion": "Biomarker strategy and endpoint justification", "marks": 20},
            {"criterion": "Clinical trial design", "marks": 20},
            {"criterion": "Risk, regulatory, and ethics analysis", "marks": 15},
        ],
        "total_marks": 100, "unlock_threshold": 24, "accepted_formats": ["pdf"], "max_size_mb": 20,
    },
}

# ── Career-filtered curriculum helper ────────────────────────────────────────

def effective_curriculum(subject_id: str, career: dict = None) -> list:
    curriculum = CURRICULUM[subject_id]
    if not career or not career.get("key_concepts"):
        return curriculum
    career_keys = set(career["key_concepts"])
    priority = [c for c in curriculum if c["id"] in career_keys]
    rest     = [c for c in curriculum if c["id"] not in career_keys]
    return priority + rest

# ── RAG helpers ───────────────────────────────────────────────────────────────

def chunk_text(text: str, chunk_size: int = 700, overlap: int = 120) -> list:
    text = re.sub(r'\s+', ' ', text).strip()
    chunks, start = [], 0
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end].strip()
        if len(chunk) > 60:
            chunks.append(chunk)
        start = end - overlap
    return chunks


def extract_text(file_bytes: bytes, filename: str) -> str:
    ext = filename.rsplit('.', 1)[-1].lower()
    if ext == 'pdf':
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(file_bytes))
            return "\n\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Could not read PDF: {e}")
    elif ext == 'docx':
        try:
            from docx import Document as DocxDoc
            doc = DocxDoc(io.BytesIO(file_bytes))
            return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Could not read DOCX: {e}")
    elif ext == 'txt':
        return file_bytes.decode('utf-8', errors='ignore')
    else:
        raise HTTPException(status_code=400, detail="Unsupported file type. Use PDF, DOCX, or TXT.")


def fts_query(text: str) -> str:
    words = re.sub(r'[^\w\s]', ' ', text).split()
    meaningful = [w for w in words if len(w) > 3][:12]
    if not meaningful:
        return '""'
    return ' OR '.join(f'"{w}"' for w in meaningful)


def retrieve_context(subject_id: str, query: str, conn, top_k: int = 3) -> str:
    try:
        q = fts_query(query)
        rows = conn.execute(
            "SELECT content FROM doc_chunks_fts WHERE subject_id = ? AND doc_chunks_fts MATCH ? ORDER BY bm25(doc_chunks_fts) LIMIT ?",
            (subject_id, q, top_k),
        ).fetchall()
        if not rows:
            return ""
        return "\n\n---\n\n".join(r["content"] for r in rows)
    except Exception:
        return ""


def has_materials(subject_id: str, conn) -> bool:
    row = conn.execute("SELECT COUNT(*) as n FROM materials WHERE subject_id = ?", (subject_id,)).fetchone()
    return row["n"] > 0

# ── Prompt builders ───────────────────────────────────────────────────────────

def build_system_prompt(subject: dict, student_name: str, is_first_visit: bool,
                        covered_ids: list, mastered_ids: list, rag_context: str = "",
                        career: dict = None, session_memory: str = "",
                        concept_notes_map: dict = None, pending_question: str = "") -> str:
    base = subject["system_prompt"]

    voice_block = f"""

━━ TEACHING METHOD ━━
You follow a strict 3-step loop for every concept. Never skip a step.

STEP 1  -  TEACH
Keep your teaching response concise: 4–5 bullets maximum. One to two sentences per bullet. Do not write walls of text  -  the student learns by answering, not by reading. If you have more to say, say less now and bring it out through the questions.
Introduce the concept with bullet points, bolded key terms, and a real-world example.
Emit one concept card immediately after your explanation (see CARD FORMAT below).
End with a check-in question that requires {student_name} to demonstrate understanding, not just say "yes".
Bad: "Does that make sense?" Good: "So based on that, what do you think happens to the mRNA if the LNP doesn't escape the endosome?"

STEP 2  -  CONFIRM
Read {student_name}'s answer carefully.
- If it is vague, short, or uses generic language without specifics: push back. "You said X, can you be more specific? What exactly does Y mean here?" Do NOT move on.
- If it demonstrates genuine understanding: affirm briefly and move to Step 3.

STEP 3  -  CHALLENGE
Raise the stakes with an application question using a real-world scenario.
Use "what would go wrong if...", a clinical case, or a named drug/company as the hook.
Example: "Moderna's early LNP had a serious tolerability problem. Based on what you just told me about ionizable lipids, what do you think caused it?"
The challenge is not optional. It cements the concept before you move to the next one.
After {student_name} answers the challenge, bridge naturally to the next concept in one sentence  -  show how what they just learned connects to and makes the next concept necessary. Never just announce the next concept; earn the transition.

MISCONCEPTION OFFER  -  use sparingly, once or twice per session at most:
After completing a concept's full loop (not during teaching), you may naturally offer: "There's a really common misconception in this space around this  -  do you want to hear it?" Only do this when it feels genuinely relevant, never mechanically. If they say yes, share it clearly. Then move on. Never open a concept with a misconception and never force this offer.

CARD FORMAT  -  emit once per new concept, in STEP 1 only:
On its own line, immediately after your explanation:
<<<CARD:{{"title":"concept name","what":"one sentence: what it is, no jargon","why":"one sentence: why it matters for {student_name} specifically, tied to their career","how":["**Bold term**: one sentence explanation","**Bold term**: one sentence explanation","**Bold term**: one sentence explanation"],"example":"one specific drug, company, trial, or clinical story  -  make it real","remember":"the single most important insight about this concept  -  write it as a complete sentence"}}>>>
Rules: valid JSON only, no line breaks inside the JSON, emit ONLY for new concepts in STEP 1, never in STEP 2 or STEP 3, maximum one card per response.

FORMATTING RULES  -  follow strictly:
- No markdown headers (##, ###). No emojis.
- NEVER use em dashes ( - ) or long dashes. Use a comma, colon, or rewrite the sentence instead.
- Never open with filler: not "Great question!", not "Absolutely!", not "Sure!"  -  just respond.
- Bullet points over prose. Format as: - **Key term**: explanation. Bold the key term.
- For critical points {student_name} must remember: bold the entire phrase: **This is the most important thing to understand here.**
- Numbered lists only for strict sequences. Bullets for everything else.
- Every response that ends with a question must bold that closing question. Example: **So based on what you just learned, what do you think happens when X?** This is the question the student must answer to move forward  -  make it impossible to miss.
- Keep bullets concise: one to two sentences max.
- Inline backticks for technical names and drug names.

PERSONALISATION:
- Use {student_name}'s name naturally  -  not every sentence, but enough that it feels personal.
- "So {student_name}, here's what makes this tricky..." or "The key insight for you, {student_name}, is..."
- Make it feel like you are talking to one person, not broadcasting to a class."""

    curriculum = effective_curriculum(subject["id"], career)
    covered_set = set(covered_ids)
    mastered_set = set(mastered_ids)

    curriculum_lines, next_concept = [], None
    for i, c in enumerate(curriculum, 1):
        if c["id"] in mastered_set:    status = "✓✓"
        elif c["id"] in covered_set:   status = "✓ "
        else:                          status = "○ "
        curriculum_lines.append(f"  {status} {i}. {c['name']}: {c['desc']}")
        if next_concept is None and c["id"] not in covered_set:
            next_concept = c

    covered_count  = len(covered_set)
    mastered_count = len(mastered_set)
    total = len(curriculum)

    curriculum_block = f"""

━━ CURRICULUM TRACKER for {student_name} ━━
Progress: {covered_count}/{total} covered, {mastered_count}/{total} mastered

{chr(10).join(curriculum_lines)}
(✓✓ = mastered   ✓ = covered   ○ = not yet covered)"""

    memory_block = ""
    if session_memory:
        memory_block = f"""

━━ PREVIOUS SESSION MEMORY for {student_name} ━━
{session_memory}

Use this to personalise your teaching today. Reference what {student_name} has already covered, acknowledge where they struggled before, and build directly on their progress. If confusion was noted last session, address it naturally  -  don't wait for them to ask again."""

    if is_first_visit:
        teaching_note = f"\n\nThis is {student_name}'s very first session. You are opening the conversation — they have not said anything yet. Introduce yourself warmly in 1–2 sentences (your name, your real-world role). Then ask them 2–3 natural background questions  -  what year they are studying, what they already know about this subject, what made them curious about it. Listen carefully to their answers: you are trying to understand how they already think about this subject, what mental model they carry, and where their instincts are right or wrong. Acknowledge what they share genuinely. Then use what you learned to frame concept 1: \"{curriculum[0]['name']}\"  -  connect it directly to their existing thinking, either building on it or gently showing where it needs to shift. Never say \"let's begin\" or any variation of it. If you see [session_open] as the trigger, treat it as your cue to open — never reference it or acknowledge it."
    elif next_concept:
        if pending_question:
            teaching_note = f"\n\nCRITICAL — RESUME UNANSWERED QUESTION: {student_name} closed the session before answering your last question. The question was: \"{pending_question}\"\n\nDo NOT start a new concept. Do NOT recap as if starting fresh. Open with a brief, natural acknowledgement that you were mid-conversation (e.g. \"Welcome back — we were right in the middle of something.\"), then re-ask the question naturally (you can rephrase it slightly). Do not move to the next concept until {student_name} gives a substantive answer. If you see [session_open] as the trigger, treat it as your cue to open — never reference it or acknowledge it."
        else:
            teaching_note = f"\n\n{student_name} is returning. You are opening the conversation. Open with a warm, genuine check-in  -  one short question about how they are doing or how they have been. Keep it natural, like a teacher who actually remembers them. Then give a 1–2 sentence recap of what they covered last session (name the specific concepts). Tell them today you are picking up with \"{next_concept['name']}\" and in one sentence explain how it connects to what they already know. Then move straight into teaching. Never say \"let's begin\" or any variation of it. If you see [session_open] as the trigger, treat it as your cue to open — never reference it or acknowledge it."
    else:
        teaching_note = f"\n\n{student_name} has covered the full curriculum. You are opening the conversation. Open with a warm check-in, then help them synthesise concepts, suggest advanced topics, and challenge them with integrative questions. Never say \"let's begin\" or any variation of it. If you see [session_open] as the trigger, treat it as your cue to open — never reference it or acknowledge it."

    career_block = ""
    if career:
        career_keys = set(career.get("key_concepts", []))
        core_here   = [c["name"] for c in curriculum if c["id"] in career_keys]
        core_list   = ", ".join(core_here[:6]) if core_here else "none specifically flagged"
        career_block = f"""

━━ STUDENT'S CAREER DESTINATION ━━
{student_name} is targeting: {career['title']} ({career['cluster']})
Role description: {career['description']}
Day-to-day reality: {career['day_in_life']}

CAREER-CRITICAL CONCEPTS IN THIS SUBJECT ({len(core_here)} of {len(curriculum)} total):
{core_list}

These concepts are directly used in the {career['title']} role. When you teach them:
- Open with a one-sentence career hook: "As a {career['title']}, you'll use this when..."
- Use examples drawn from: {career['day_in_life']}
- Make the stakes real: explain what goes wrong professionally if they don't understand this

For non-career-critical concepts: still teach them fully, but frame them as supporting knowledge ("You won't use this daily, but you need it to understand X").

Never teach a concept in the abstract. Every concept should feel like {student_name} is being prepared for a real job, not passing an exam."""

    rag_block = ""
    if rag_context:
        rag_block = f"""

━━ BVERSITY COURSE MATERIALS ━━
The following excerpts are from Bversity's official course materials for {subject['name']}.
Prioritise this content when relevant. Reference it naturally. Do not cite it when not relevant.

{rag_context}
━━ END COURSE MATERIALS ━━"""

    concept_ids_list = ", ".join(c["id"] for c in curriculum)
    tagging = f"""

━━ CONCEPT TAGGING (required) ━━
End every response with (on its own line):
<<<CONCEPTS:concept_id1,concept_id2>>>
Valid IDs: {concept_ids_list}
CRITICAL RULE: Only tag a concept AFTER {student_name} has answered your Step 3 challenge question to your satisfaction. Do NOT tag a concept just because you explained it in Step 1. The tag signals that the student has genuinely worked through the concept, not just heard it. If you are in Step 1 or Step 2, or waiting for an answer: <<<CONCEPTS:>>>
This line is stripped before {student_name} sees the response.

━━ DEFINITION TAGGING (required when bolding terms) ━━
Whenever your response contains bolded terms (**term**), also end with (on its own line):
<<<DEFS:{{"exact bolded term": "one sentence, plain language, no jargon", "another term": "definition"}}>>>
Include ALL bolded terms from this response. Use the exact same spelling and capitalisation as bolded. Maximum 6 terms.
If no bolded terms: omit this tag entirely.
This line is stripped before {student_name} sees the response."""

    career_tagging = ""
    if not career:
        career_tagging = f"""
If {student_name} mentions a career interest or aspiration, identify the closest match and append:
<<<CAREER:career_id>>>
Valid career IDs: {" | ".join(CAREERS.keys())}
Only tag when the student clearly states a career aspiration. Tag at most once per response."""

    notes_block = ""
    if concept_notes_map:
        curriculum = effective_curriculum(subject["id"], career)
        notes_lines = []
        for c in curriculum:
            note = concept_notes_map.get(c["id"], "").strip()
            if note:
                notes_lines.append(f"  [{c['name']}]: {note}")
        if notes_lines:
            notes_block = f"""

━━ FACULTY NOTES PER CONCEPT ━━
These notes were written by the course faculty. Incorporate this guidance when teaching the relevant concept.
{chr(10).join(notes_lines)}"""

    return base + voice_block + curriculum_block + notes_block + memory_block + teaching_note + career_block + rag_block + tagging + career_tagging


def build_quiz_prompt(subject: dict, student_name: str, covered_ids: list, mastered_ids: list, career: dict = None) -> str:
    curriculum = effective_curriculum(subject["id"], career)
    covered_set, mastered_set = set(covered_ids), set(mastered_ids)
    to_quiz = [c for c in curriculum if c["id"] in covered_set and c["id"] not in mastered_set]

    voice_rules = "Speak like a human tutor  -  no markdown headers, no bullet lists, no emojis, no filler openers. Short prose paragraphs. Bold a term only when introducing it. Sound like a person.\n\n"

    if not to_quiz:
        return voice_rules + f"""You are {subject['tutor_name']} at Bversity. {student_name} has mastered everything they've covered. Tell them directly, then suggest the next concept to move on to. One short paragraph.\n\nEnd with: <<<MASTERED:>>>"""

    quiz_concepts = to_quiz[:3]
    concept_list  = "\n".join(f"  - {c['name']}: {c['desc']}" for c in quiz_concepts)
    valid_ids     = ", ".join(c["id"] for c in quiz_concepts)

    return voice_rules + f"""You are {subject['tutor_name']} at Bversity, checking {student_name}'s understanding.

Test them on these concepts:
{concept_list}

Ask two or three questions that probe real understanding  -  not recall. Use "explain why", "walk me through", or "what would happen if" style questions. Write them as a tutor would ask in conversation, not as a formatted list.

After {student_name} responds: evaluate honestly. If they show solid understanding, affirm it and mark mastery. If there are gaps, explain the gap directly and ask one follow-up  -  do not mark mastery yet.

End your response with (hidden from student):
<<<MASTERED:concept_id1,concept_id2>>>
Valid IDs for this quiz: {valid_ids}
If none mastered: <<<MASTERED:>>>"""

RECALL_WARMUP_SUFFIX = """

━━ RECALL WARMUP  -  THIS MESSAGE ONLY ━━
{student_name} just opened a new session and wrote a quick recall from memory before we began.
Their recall is the current message.

Your response for THIS message only:
1. Acknowledge what they got right  -  specifically, name the concept or mechanism they recalled. 1–2 sentences.
2. If they missed something important or got it slightly wrong, correct it briefly and directly. 1 sentence max.
3. Bridge naturally into the next concept: "Let's build on that today..." or similar.
DO NOT do a formal quiz. DO NOT ask them to recall more. DO NOT start teaching a new concept yet  -  just bridge to it.
Keep the entire response under 4 sentences. Sound like a tutor who's pleased they came prepared."""


# ── Mock responses ────────────────────────────────────────────────────────────

MOCK_RESPONSES = {
    "bioinformatics":  "In bioinformatics we use computational methods to analyse biological sequences and other molecular data.\n\n**[Mock mode  -  add ANTHROPIC_API_KEY to .env for live AI tutoring]**",
    "genomics":        "The genome is the complete set of genetic instructions in an organism.\n\n**[Mock mode  -  add ANTHROPIC_API_KEY to .env for live AI tutoring]**",
    "drug_discovery":  "The drug development pipeline from target to approval takes roughly 12 years on average.\n\n**[Mock mode  -  add ANTHROPIC_API_KEY to .env for live AI tutoring]**",
    "clinical_trials": "Phase I trials focus on safety and dose-finding in 20–80 participants before advancing to efficacy testing.\n\n**[Mock mode  -  add ANTHROPIC_API_KEY to .env for live AI tutoring]**",
    "genai_ml":        "Machine learning in life sciences is transforming how we predict molecular properties and design new drugs.\n\n**[Mock mode  -  add ANTHROPIC_API_KEY to .env for live AI tutoring]**",
    "biotech_business":"Understanding biotech business models is as important as understanding the science behind the drugs.\n\n**[Mock mode  -  add ANTHROPIC_API_KEY to .env for live AI tutoring]**",
}

# ── Request models ────────────────────────────────────────────────────────────

class RegisterRequest(BaseModel):
    name: str
    email: str

class ChatRequest(BaseModel):
    student_id: str
    subject_id: str
    message: str = ""
    quiz_mode: bool = False
    recall_warmup: bool = False
    auto_open: bool = False

    def __init__(self, **data):
        if "message" in data and len(data["message"]) > 4000:
            data["message"] = data["message"][:4000]
        super().__init__(**data)

class ProfileRequest(BaseModel):
    career_id: str

class MarkRequest(BaseModel):
    score: int
    feedback: str

class RequestCodeRequest(BaseModel):
    email: str
    name: str
    product: str = "career_pathways"

class VerifyCodeRequest(BaseModel):
    email: str
    name: str
    code: str
    product: str = "career_pathways"


# ── Email ─────────────────────────────────────────────────────────────────────

def update_streak(student_id: str, conn) -> int:
    from datetime import timedelta
    today = datetime.utcnow().date().isoformat()
    now   = datetime.utcnow().isoformat()
    row   = conn.execute(
        "SELECT streak_count, streak_last_date FROM student_profile WHERE student_id = ?",
        (student_id,)
    ).fetchone()
    if not row:
        return 0
    count     = row["streak_count"] or 0
    last_date = row["streak_last_date"]
    if last_date == today:
        conn.execute("UPDATE student_profile SET last_active_at=? WHERE student_id=?", (now, student_id))
        conn.commit()
        return count
    yesterday = (datetime.utcnow().date() - timedelta(days=1)).isoformat()
    new_count = (count + 1) if last_date == yesterday else 1
    conn.execute(
        "UPDATE student_profile SET streak_count=?, streak_last_date=?, last_active_at=? WHERE student_id=?",
        (new_count, today, now, student_id)
    )
    conn.commit()
    return new_count


async def check_inactivity_nudge(student_id: str):
    conn = get_db()
    try:
        row = conn.execute(
            "SELECT last_active_at, nudge_sent_at FROM student_profile WHERE student_id=?",
            (student_id,)
        ).fetchone()
        if not row or not row["last_active_at"]:
            return
        now        = datetime.utcnow()
        last_act   = datetime.fromisoformat(row["last_active_at"])
        days_since = (now - last_act).days
        if days_since < 3:
            return
        if row["nudge_sent_at"]:
            last_nudge = datetime.fromisoformat(row["nudge_sent_at"])
            if (now - last_nudge).days < 7:
                return
        student = conn.execute("SELECT name, email FROM students WHERE id=?", (student_id,)).fetchone()
        if not student:
            return
        profile = conn.execute("SELECT career_id FROM student_profile WHERE student_id=?", (student_id,)).fetchone()
        career  = CAREERS.get(profile["career_id"]) if profile and profile["career_id"] else None
        await send_inactivity_nudge_email(student["email"], student["name"], days_since, career)
        conn.execute("UPDATE student_profile SET nudge_sent_at=? WHERE student_id=?", (now.isoformat(), student_id))
        conn.commit()
    except Exception as e:
        print(f"Nudge check error: {e}")
    finally:
        conn.close()


def _email_wrap(body: str) -> str:
    return f"""<div style="font-family:'Helvetica Neue',Arial,sans-serif;background:#F6F8FB;padding:40px 0">
      <div style="max-width:520px;margin:0 auto;background:#fff;border-radius:16px;overflow:hidden;border:1px solid #E8EDF4">
        <div style="background:#07142A;padding:22px 32px">
          <p style="margin:0;font-size:20px;font-weight:900;color:#fff;letter-spacing:-0.5px">Bversity</p>
          <p style="margin:0;font-size:11px;color:#72CFC0;letter-spacing:0.4px">School of Bioscience</p>
        </div>
        <div style="padding:32px">{body}</div>
        <div style="background:#F6F8FB;padding:18px 32px;border-top:1px solid #E8EDF4">
          <p style="margin:0;font-size:11px;color:#9EABBE;text-align:center">
            Bversity School of Bioscience &middot; An initiative by TABS Learning Pvt Ltd<br>
            You're receiving this because you're a registered learner at Bversity.
          </p>
        </div>
      </div>
    </div>"""

def _heading(text: str) -> str:
    return f'<h2 style="font-size:22px;font-weight:800;color:#07142A;margin:0 0 8px">{text}</h2>'

def _para(text: str) -> str:
    return f'<p style="color:#3D5166;font-size:15px;line-height:1.6;margin:0 0 20px">{text}</p>'

def _btn(text: str, url: str, color: str = "#00A896") -> str:
    return f'<a href="{url}" style="display:inline-block;background:{color};color:#fff;text-decoration:none;padding:12px 28px;border-radius:8px;font-weight:700;font-size:14px;margin:8px 0 20px">{text}</a>'

def _divider() -> str:
    return '<hr style="border:none;border-top:1px solid #E8EDF4;margin:24px 0">'

def _small(text: str) -> str:
    return f'<p style="color:#9EABBE;font-size:12px;margin:0">{text}</p>'

async def _send_email(to_email: str, subject: str, html: str, reply_to: str = None) -> bool:
    api_key    = os.environ.get("RESEND_API_KEY", "")
    from_email = os.environ.get("RESEND_FROM_EMAIL", "noreply@bversity.io")
    if not api_key:
        print(f"\n[DEV] Email → {to_email} | {subject}\n")
        return True
    try:
        import httpx
        payload = {"from": f"Bversity <{from_email}>", "to": [to_email], "subject": subject, "html": html}
        if reply_to:
            payload["reply_to"] = [reply_to]
        async with httpx.AsyncClient() as client:
            res = await client.post(
                "https://api.resend.com/emails",
                headers={"Authorization": f"Bearer {api_key}"},
                json=payload,
                timeout=10.0,
            )
        if res.status_code not in (200, 201):
            print(f"Resend error {res.status_code}: {res.text}")
        return res.status_code in (200, 201)
    except Exception as e:
        print(f"Email error: {e}")
        return False

async def send_welcome_email(to_email: str, name: str) -> bool:
    first = name.split()[0]
    body = (
        _heading(f"Welcome to Bversity, {first}!") +
        _para("You've just joined India's first AI Native University for Biotech &amp; Life Sciences. "
              "Your personal AI tutor is ready  -  pick a subject and start your first conversation.") +
        _btn("Open Bversity →", "https://bversity.io") +
        _divider() +
        _para("<strong>What to do first:</strong><br>"
              "1. Head to <em>Career Path</em> and tell your tutor what you want to become.<br>"
              "2. Open any subject  -  Genomics, Drug Discovery, Bioinformatics  -  and say hello.<br>"
              "3. Your 30-day personalised study plan will be ready once you pick a career.") +
        _small("Questions? Reply to this email  -  we read every one.")
    )
    return await _send_email(to_email, f"Welcome to Bversity, {first}! 🎉", _email_wrap(body))

async def send_completion_email(to_email: str, name: str, subject_name: str, credential_id: str) -> bool:
    first = name.split()[0]
    body = (
        _heading(f"You completed {subject_name}!") +
        _para(f"Congratulations {first}  -  you've covered every concept in <strong>{subject_name}</strong>. "
              "Your certificate is now available in your dashboard.") +
        _btn("View Certificate →", "https://bversity.io") +
        _divider() +
        _para(f'<strong>Credential ID:</strong> BVG-{credential_id}<br>'
              'Share this on LinkedIn to let the world know.') +
        _small("Keep going  -  more subjects await. Each one sharpens your edge in the biotech job market.")
    )
    return await _send_email(to_email, f"Certificate earned: {subject_name} 🎓", _email_wrap(body))

async def send_lag_nudge_email(to_email: str, name: str, lag_days: int, lag_concepts: int) -> bool:
    first = name.split()[0]
    body = (
        _heading(f"Hey {first}, you're {lag_days} day{'s' if lag_days != 1 else ''} behind") +
        _para(f"Your study plan has <strong>{lag_concepts} concept{'s' if lag_concepts != 1 else ''}</strong> overdue. "
              "A short session today can get you back on track  -  even 20 minutes covers 1–2 concepts.") +
        _btn("Continue Learning →", "https://bversity.io") +
        _divider() +
        _para("Consistency is more powerful than intensity. Log in, pick up where you left off, and let your tutor guide you.") +
        _small("You can adjust your plan any time from the Dashboard.")
    )
    return await _send_email(to_email, f"Your study plan needs attention, {first}", _email_wrap(body))

async def send_inactivity_nudge_email(to_email: str, name: str, days_since: int, career: dict = None) -> bool:
    first = name.split()[0]
    career_line = f"You're on your way to becoming a <strong>{career['title']}</strong>  -  don't lose momentum." if career else "You're building real biotech knowledge  -  don't lose momentum."
    subject_line = f"Hey {first}, it's been {days_since} days" if days_since < 7 else f"{first}, your learning path needs you"
    body = (
        _heading(f"Hey {first}, you still there?") +
        _para(f"It's been <strong>{days_since} day{'s' if days_since != 1 else ''}</strong> since your last session on Bversity. {career_line}") +
        _btn("Pick Up Where You Left Off →", "https://university.bversity.io") +
        _divider() +
        _para("Your AI tutors are ready. Even a 15-minute session keeps you moving forward  -  and your progress is exactly where you left it.") +
        _small("You're receiving this because you're a registered learner at Bversity.")
    )
    return await _send_email(to_email, subject_line, _email_wrap(body))


async def send_streak_milestone_email(to_email: str, name: str, streak: int, career: dict = None) -> bool:
    first = name.split()[0]
    emojis = {3: "🔥", 7: "⚡", 14: "🚀", 30: "🏆"}
    emoji  = emojis.get(streak, "🔥")
    career_line = f"For someone heading toward <strong>{career['title']}</strong>, consistency like this is everything." if career else "Consistency like this is everything in biotech."
    body = (
        _heading(f"{emoji} {streak}-day streak, {first}!") +
        _para(f"You've logged in and learned for <strong>{streak} days in a row</strong>. {career_line}") +
        _btn("Keep Your Streak Going →", "https://university.bversity.io") +
        _divider() +
        _small("Keep showing up. The compounding effect of daily learning is real.")
    )
    return await _send_email(to_email, f"{emoji} {streak}-day learning streak  -  keep it going, {first}!", _email_wrap(body))


async def send_module_complete_email(to_email: str, name: str, module_name: str, subject_name: str, career: dict = None) -> bool:
    first = name.split()[0]
    career_line = f"This is exactly the kind of knowledge a <strong>{career['title']}</strong> uses daily." if career else "This is real biotech knowledge that sets you apart."
    body = (
        _heading(f"Module complete: {module_name}") +
        _para(f"You just finished <strong>{module_name}</strong> in {subject_name}. {career_line}") +
        _btn("Continue to Next Module →", "https://university.bversity.io") +
        _divider() +
        _small("Take the module quiz to lock in your mastery  -  it's short and worth it.")
    )
    return await _send_email(to_email, f"You completed '{module_name}'  -  {first}", _email_wrap(body))


def _infer_sessions(msgs: list, gap_minutes: int = 30) -> list:
    """Group messages into sessions using a time-gap heuristic."""
    if not msgs:
        return []
    def _parse(s):
        try: return datetime.fromisoformat(s)
        except: return datetime.fromisoformat(s.replace('Z', ''))
    sorted_msgs = sorted(msgs, key=lambda m: m['created_at'])
    sessions, sess_msgs = [], [sorted_msgs[0]]
    for msg in sorted_msgs[1:]:
        gap = (_parse(msg['created_at']) - _parse(sess_msgs[-1]['created_at'])).total_seconds() / 60
        if gap > gap_minutes:
            sessions.append(sess_msgs)
            sess_msgs = [msg]
        else:
            sess_msgs.append(msg)
    sessions.append(sess_msgs)
    result = []
    for s in sessions:
        start = _parse(s[0]['created_at'])
        end   = _parse(s[-1]['created_at'])
        dur   = max(1, int((end - start).total_seconds() / 60))
        result.append({'start': start, 'duration_mins': dur, 'message_count': len(s), 'hour': start.hour})
    return result


def gather_student_week_data(student_id: str, conn) -> dict:
    now        = datetime.utcnow()
    week_ago   = (now - timedelta(days=7)).isoformat()
    prev_start = (now - timedelta(days=14)).isoformat()
    stuck_cutoff = (now - timedelta(days=14)).isoformat()

    # ── Raw messages this week ───────────────────────────────
    all_msgs = conn.execute(
        "SELECT created_at FROM messages WHERE student_id=? AND role='user' AND created_at>=?",
        (student_id, week_ago)
    ).fetchall()
    all_msgs = [dict(m) for m in all_msgs]

    # ── Session inference ────────────────────────────────────
    inferred = _infer_sessions(all_msgs)
    num_sessions         = len(inferred)
    total_msgs_week      = len(all_msgs)
    avg_session_mins     = round(sum(s['duration_mins'] for s in inferred) / len(inferred)) if inferred else 0
    longest_session_mins = max((s['duration_mins'] for s in inferred), default=0)
    avg_msgs_per_session = round(sum(s['message_count'] for s in inferred) / len(inferred), 1) if inferred else 0

    # Learning style
    if num_sessions == 0:
        learning_style = None
    elif num_sessions == 1:
        learning_style = 'single session'
    elif longest_session_mins > avg_session_mins * 1.8:
        learning_style = 'binge learner'
    else:
        learning_style = 'consistent daily learner'

    # ── Time of day ──────────────────────────────────────────
    time_slots = {'morning': 0, 'afternoon': 0, 'evening': 0, 'night': 0}
    for s in inferred:
        h = s['hour']
        if   6  <= h < 12: time_slots['morning']   += 1
        elif 12 <= h < 17: time_slots['afternoon']  += 1
        elif 17 <= h < 22: time_slots['evening']    += 1
        else:              time_slots['night']      += 1
    peak_time = max(time_slots, key=time_slots.get) if inferred else None

    # Peak day of week
    from collections import Counter
    day_counts = Counter(s['start'].strftime('%A') for s in inferred)
    peak_day = day_counts.most_common(1)[0][0] if day_counts else None

    # ── Week-over-week ───────────────────────────────────────
    prev_msgs = conn.execute(
        "SELECT created_at FROM messages WHERE student_id=? AND role='user' AND created_at>=? AND created_at<?",
        (student_id, prev_start, week_ago)
    ).fetchall()
    prev_inferred   = _infer_sessions([dict(m) for m in prev_msgs])
    prev_concepts   = conn.execute(
        "SELECT COUNT(*) FROM concept_progress WHERE student_id=? AND first_covered_at>=? AND first_covered_at<?",
        (student_id, prev_start, week_ago)
    ).fetchone()[0]

    # ── Concepts ────────────────────────────────────────────
    concepts_covered = conn.execute(
        "SELECT COUNT(*) FROM concept_progress WHERE student_id=? AND first_covered_at>=?",
        (student_id, week_ago)
    ).fetchone()[0]
    concepts_mastered = conn.execute(
        "SELECT COUNT(*) FROM concept_progress WHERE student_id=? AND mastered_at>=?",
        (student_id, week_ago)
    ).fetchone()[0]

    # Stuck concepts  -  covered >14 days ago, still not mastered
    stuck_rows = conn.execute(
        "SELECT subject_id, concept_id FROM concept_progress WHERE student_id=? AND mastered_at IS NULL AND first_covered_at<=? ORDER BY first_covered_at ASC LIMIT 3",
        (student_id, stuck_cutoff)
    ).fetchall()
    stuck_concepts = []
    for sc in stuck_rows:
        sc = dict(sc)
        if sc['subject_id'] in CURRICULUM:
            c = next((x for x in CURRICULUM[sc['subject_id']] if x['id'] == sc['concept_id']), None)
            if c: stuck_concepts.append(c['name'])

    # ── Quizzes ──────────────────────────────────────────────
    quizzes_taken = conn.execute(
        "SELECT COUNT(*) FROM module_quizzes WHERE student_id=? AND completed_at>=?",
        (student_id, week_ago)
    ).fetchone()[0]
    quizzes_passed = conn.execute(
        "SELECT COUNT(*) FROM module_quizzes WHERE student_id=? AND completed_at>=? AND passed=1",
        (student_id, week_ago)
    ).fetchone()[0]

    # ── Active subjects ──────────────────────────────────────
    active_rows = conn.execute(
        "SELECT DISTINCT subject_id FROM messages WHERE student_id=? AND role='user' AND created_at>=?",
        (student_id, week_ago)
    ).fetchall()
    active_subjects = [SUBJECTS[r['subject_id']]['name'] for r in active_rows if r['subject_id'] in SUBJECTS]

    # ── Strongest / weakest ──────────────────────────────────
    recent = conn.execute(
        "SELECT subject_id, concept_id FROM concept_progress WHERE student_id=? AND mastered_at IS NOT NULL ORDER BY mastered_at DESC LIMIT 1",
        (student_id,)
    ).fetchone()
    strongest_concept = None
    if recent:
        r = dict(recent)
        if r['subject_id'] in CURRICULUM:
            c = next((x for x in CURRICULUM[r['subject_id']] if x['id'] == r['concept_id']), None)
            if c: strongest_concept = c['name']

    weakest_subject = None
    max_gap = 0
    for sid in SUBJECTS:
        gap = conn.execute(
            "SELECT COUNT(*) FROM concept_progress WHERE student_id=? AND subject_id=? AND mastered_at IS NULL",
            (student_id, sid)
        ).fetchone()[0]
        if gap > max_gap:
            max_gap = gap
            weakest_subject = SUBJECTS[sid]['name']

    # ── Profile + career ────────────────────────────────────
    profile = conn.execute(
        "SELECT streak_count, career_id FROM student_profile WHERE student_id=?", (student_id,)
    ).fetchone()
    streak    = profile['streak_count'] if profile else 0
    career_id = profile['career_id'] if profile else None
    if career_id and career_id in CAREERS:
        career_title = CAREERS[career_id]['title']
    elif career_id and career_id in SUBJECTS:
        career_title = f"{SUBJECTS[career_id]['name']} ({SUBJECTS[career_id].get('certification', '')} prep)"
    else:
        career_title = None

    # Career-critical untouched concepts
    untouched_career_concepts = []
    if career_id and career_id in CAREERS:
        for kc in CAREERS[career_id].get('key_concepts', []):
            for sid, curriculum in CURRICULUM.items():
                c = next((x for x in curriculum if x['id'] == kc), None)
                if c:
                    touched = conn.execute(
                        "SELECT 1 FROM concept_progress WHERE student_id=? AND subject_id=? AND concept_id=?",
                        (student_id, sid, kc)
                    ).fetchone()
                    if not touched:
                        untouched_career_concepts.append(c['name'])
                    break
            if len(untouched_career_concepts) >= 3:
                break

    # Career readiness
    career_readiness = None
    if career_id and career_id in CAREERS:
        relevant = CAREERS[career_id].get('relevant_subjects', [])
        if relevant:
            score_sum = 0
            for sid in relevant:
                total = len(CURRICULUM.get(sid, []))
                if total == 0: continue
                m = conn.execute("SELECT COUNT(*) FROM concept_progress WHERE student_id=? AND subject_id=? AND mastered_at IS NOT NULL", (student_id, sid)).fetchone()[0]
                cv = conn.execute("SELECT COUNT(*) FROM concept_progress WHERE student_id=? AND subject_id=?", (student_id, sid)).fetchone()[0]
                score_sum += (m + 0.4 * max(0, cv - m)) / total
            career_readiness = min(int(round(10 + (score_sum / len(relevant)) * 88)), 98)

    return {
        # Activity
        "sessions":              num_sessions,
        "total_messages":        total_msgs_week,
        "avg_session_mins":      avg_session_mins,
        "longest_session_mins":  longest_session_mins,
        "avg_msgs_per_session":  avg_msgs_per_session,
        "learning_style":        learning_style,
        # Time patterns
        "peak_time":             peak_time,
        "peak_day":              peak_day,
        "time_slots":            time_slots,
        # Progress
        "concepts_covered":      concepts_covered,
        "concepts_mastered":     concepts_mastered,
        "stuck_concepts":        stuck_concepts,
        "strongest_concept":     strongest_concept,
        "weakest_subject":       weakest_subject,
        # Quizzes
        "quizzes_taken":         quizzes_taken,
        "quizzes_passed":        quizzes_passed,
        # Subjects
        "active_subjects":       active_subjects,
        # Week-over-week
        "wow_sessions":          num_sessions - len(prev_inferred),
        "wow_concepts":          concepts_covered - prev_concepts,
        # Career
        "streak":                streak,
        "career_title":          career_title,
        "career_readiness":      career_readiness,
        "untouched_career_concepts": untouched_career_concepts,
    }


async def generate_report_narrative(student_name: str, data: dict) -> str:
    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        return f"Great work this week, {student_name.split()[0]}! Keep the momentum going."

    first = student_name.split()[0]

    lines = [
        f"Student: {first}",
        f"Career goal: {data['career_title']}." if data.get('career_title') else "",
        f"Learning pattern: {data['learning_style']}." if data.get('learning_style') else "",
        f"Peak learning time: {data['peak_time']}." if data.get('peak_time') else "",
        f"Most active day: {data['peak_day']}." if data.get('peak_day') else "",
        f"Sessions this week: {data['sessions']} (avg {data['avg_session_mins']} min each, longest {data['longest_session_mins']} min).",
        f"Messages sent: {data['total_messages']} (avg {data['avg_msgs_per_session']} per session  -  {'deep engagement' if data['avg_msgs_per_session'] > 8 else 'light engagement'}).",
        f"Concepts covered: {data['concepts_covered']}, mastered: {data['concepts_mastered']}.",
        f"Quizzes: {data['quizzes_passed']}/{data['quizzes_taken']} passed." if data['quizzes_taken'] > 0 else "",
        f"Recently mastered: {data['strongest_concept']}." if data.get('strongest_concept') else "",
        f"Stuck on (covered but not mastered in 2+ weeks): {', '.join(data['stuck_concepts'])}." if data.get('stuck_concepts') else "",
        f"Weakest area (most un-mastered concepts): {data['weakest_subject']}." if data.get('weakest_subject') else "",
        f"Career-critical concepts not yet started: {', '.join(data['untouched_career_concepts'])}." if data.get('untouched_career_concepts') else "",
        f"Career readiness: {data['career_readiness']}%." if data.get('career_readiness') else "",
        f"Streak: {data['streak']} days.",
        f"Week-over-week: {'up' if data['wow_sessions'] >= 0 else 'down'} {abs(data['wow_sessions'])} session(s), {'more' if data['wow_concepts'] >= 0 else 'fewer'} {abs(data['wow_concepts'])} concept(s) than last week.",
    ]
    context = "\n".join(l for l in lines if l)

    prompt = f"""You are a personal learning coach at Bversity AI University, an AI-native biotech university.
Write a warm, specific 2-3 sentence coaching note to {first} based on their week's learning data.
Reference their actual behaviour patterns (time of day, session length, engagement depth).
Give one concrete, forward-looking recommendation tied to their career goal.
Sound like a coach who genuinely studied their data  -  not a bot sending a template.

{context}

Write only the 2-3 sentence note. No greeting, no subject line."""

    try:
        import anthropic
        client = anthropic.Anthropic(api_key=api_key)
        msg = client.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=250,
            messages=[{"role": "user", "content": prompt}]
        )
        return msg.content[0].text.strip()
    except Exception as e:
        print(f"Narrative gen error: {e}")
        return f"You had a solid week, {first}. Keep building on your progress  -  every session compounds."


async def send_weekly_learner_report(to_email: str, name: str, data: dict, narrative: str) -> bool:
    first = name.split()[0]

    def stat_box(val, label):
        return (
            f'<div style="flex:1;min-width:90px;background:#F0FBF9;border-radius:10px;padding:14px;text-align:center">'
            f'<p style="margin:0;font-size:24px;font-weight:900;color:#16c1ad">{val}</p>'
            f'<p style="margin:4px 0 0;font-size:11px;color:#07142A;line-height:1.3">{label}</p>'
            f'</div>'
        )

    def section_label(text):
        return f'<p style="font-size:10px;font-weight:700;color:#9EABBE;letter-spacing:0.08em;text-transform:uppercase;margin:0 0 10px">{text}</p>'

    # ── Activity stats ───────────────────────────────────────
    wow_s = data['wow_sessions']
    wow_c = data['wow_concepts']
    wow_sessions_label = f"Sessions<br><span style='font-size:10px;color:{'#16a34a' if wow_s>=0 else '#dc2626'}'>{'+' if wow_s>=0 else ''}{wow_s} vs last week</span>"
    wow_concepts_label = f"Concepts Covered<br><span style='font-size:10px;color:{'#16a34a' if wow_c>=0 else '#dc2626'}'>{'+' if wow_c>=0 else ''}{wow_c} vs last week</span>"

    stats_row = (
        f'<div style="display:flex;gap:10px;margin:0 0 14px;flex-wrap:wrap">'
        + stat_box(data['sessions'], wow_sessions_label)
        + stat_box(data['concepts_covered'], wow_concepts_label)
        + stat_box(data['concepts_mastered'], 'Concepts<br>Mastered')
        + (stat_box(f"{data['quizzes_passed']}/{data['quizzes_taken']}", 'Quizzes<br>Passed') if data['quizzes_taken'] > 0 else "")
        + '</div>'
    )

    # ── Learning behaviour ───────────────────────────────────
    behaviour_parts = []
    if data.get('peak_time'):
        behaviour_parts.append(f"<strong>{data['peak_time'].capitalize()}</strong> learner")
    if data.get('peak_day'):
        behaviour_parts.append(f"most active on <strong>{data['peak_day']}s</strong>")
    if data.get('avg_session_mins'):
        behaviour_parts.append(f"avg session <strong>{data['avg_session_mins']} min</strong>")
    if data.get('avg_msgs_per_session'):
        behaviour_parts.append(f"<strong>{data['avg_msgs_per_session']} messages</strong> per session")
    behaviour_section = ""
    if behaviour_parts:
        behaviour_section = (
            _divider() +
            section_label("Your Learning Pattern") +
            _para(" · ".join(behaviour_parts) + ("." if not behaviour_parts[-1].endswith(".") else ""))
        )
        if data.get('learning_style'):
            behaviour_section += _para(f"Style this week: <strong>{data['learning_style'].title()}</strong>")

    # ── Concepts insight ─────────────────────────────────────
    concepts_section = _divider() + section_label("Concepts")
    if data.get('strongest_concept'):
        concepts_section += _para(f"Latest mastery: <strong>{data['strongest_concept']}</strong> ✓")
    if data.get('stuck_concepts'):
        concepts_section += _para(f"Needs attention (covered 2+ weeks ago, not yet mastered): <strong>{', '.join(data['stuck_concepts'])}</strong>")
    if data.get('active_subjects'):
        concepts_section += _para(f"Active in: <strong>{', '.join(data['active_subjects'])}</strong>")

    # ── Career section ───────────────────────────────────────
    career_section = ""
    if data.get('career_title'):
        career_section = _divider() + section_label("Career Path")
        readiness_part = f" &middot; <strong>{data['career_readiness']}%</strong> career ready" if data.get('career_readiness') else ""
        career_section += _para(f'<strong>{data["career_title"]}</strong>{readiness_part}')
        if data.get('untouched_career_concepts'):
            career_section += _para(f"Key concepts not yet started: <strong>{', '.join(data['untouched_career_concepts'])}</strong>")
        elif data.get('weakest_subject'):
            career_section += _para(f"Focus area: <strong>{data['weakest_subject']}</strong>")

    # ── Streak ───────────────────────────────────────────────
    streak_line = _para(f"🔥 <strong>{data['streak']}-day streak</strong>  -  keep it going!") if data.get('streak', 0) > 0 else ""

    body = (
        _heading(f"Your week at Bversity, {first} 🧬") +
        _para(f'<em style="color:#3D5166">{narrative}</em>') +
        _divider() +
        section_label("This Week") +
        stats_row +
        streak_line +
        behaviour_section +
        concepts_section +
        career_section +
        _divider() +
        _btn("Continue Learning →", "https://university.bversity.io") +
        _small("Sent every Monday · university.bversity.io")
    )
    return await _send_email(to_email, f"Your Bversity week, {first} 🧬", _email_wrap(body))

async def send_verification_email(to_email: str, name: str, code: str) -> bool:
    api_key = os.environ.get("RESEND_API_KEY", "")
    from_email = os.environ.get("RESEND_FROM_EMAIL", "noreply@bversity.io")
    if not api_key:
        print(f"\n{'='*50}\n[DEV] Verification code for {to_email}: {code}\n{'='*50}\n")
        return True
    try:
        import httpx
        html = f"""
        <div style="font-family:sans-serif;max-width:480px;margin:0 auto;padding:40px 24px">
          <p style="font-size:20px;font-weight:900;color:#07142A;margin:0 0 4px">Bversity</p>
          <p style="font-size:12px;color:#7A8FA6;margin:0 0 32px">School of Bioscience</p>
          <h2 style="font-size:22px;font-weight:800;color:#07142A;margin:0 0 8px">Hi {name},</h2>
          <p style="color:#3D5166;font-size:15px;margin:0 0 28px">Your verification code for Bversity:</p>
          <div style="background:#FFF8EC;border:1px solid rgba(255,167,10,0.25);border-radius:12px;padding:32px;text-align:center;margin-bottom:24px">
            <div style="font-size:44px;font-weight:900;letter-spacing:12px;color:#07142A;font-family:monospace">{code}</div>
            <p style="color:#7A8FA6;font-size:13px;margin:12px 0 0">Expires in 15 minutes &middot; Single use</p>
          </div>
          <p style="color:#7A8FA6;font-size:13px">If you didn't request this, you can safely ignore this email.</p>
        </div>"""
        async with httpx.AsyncClient() as client:
            res = await client.post(
                "https://api.resend.com/emails",
                headers={"Authorization": f"Bearer {api_key}"},
                json={"from": f"Bversity <{from_email}>", "to": [to_email],
                      "subject": f"Your Bversity code: {code}", "html": html},
                timeout=10.0,
            )
        if res.status_code != 200:
            print(f"Resend error {res.status_code}: {res.text}")
        return res.status_code == 200
    except Exception as e:
        print(f"Email error: {e}")
        return False

# ── Admin auth ────────────────────────────────────────────────────────────────

def require_admin(x_admin_key: str = Header(None)):
    expected = os.environ.get("ADMIN_KEY")
    if not expected or x_admin_key != expected:
        raise HTTPException(status_code=401, detail="Unauthorized")

# ── Routes ────────────────────────────────────────────────────────────────────

@app.get("/")
def read_root():
    return {"status": "ok", "message": "Bversity AI University backend"}

@app.get("/subjects")
def get_subjects():
    return [{k: v for k, v in s.items() if k != "system_prompt"} for s in SUBJECTS.values()]

@app.post("/register")
async def register(req: RegisterRequest, background_tasks: BackgroundTasks):
    if not req.name.strip() or not req.email.strip():
        raise HTTPException(status_code=400, detail="Name and email are required")
    conn = get_db()
    existing = conn.execute("SELECT * FROM students WHERE email = ?", (req.email.lower().strip(),)).fetchone()
    if existing:
        conn.close()
        return {"student_id": existing["id"], "name": existing["name"], "returning": True}
    student_id = str(uuid.uuid4())
    conn.execute("INSERT INTO students (id, name, email, created_at) VALUES (?, ?, ?, ?)",
                 (student_id, req.name.strip(), req.email.lower().strip(), datetime.utcnow().isoformat()))
    conn.commit(); conn.close()
    background_tasks.add_task(send_welcome_email, req.email.lower().strip(), req.name.strip())
    return {"student_id": student_id, "name": req.name.strip(), "returning": False}

@app.get("/check-email")
async def check_email_endpoint(email: str):
    conn = get_db()
    student = conn.execute("SELECT name FROM students WHERE email = ?", (email.lower().strip(),)).fetchone()
    conn.close()
    if student:
        return {"exists": True, "name": student["name"]}
    return {"exists": False, "name": ""}

@app.post("/request-code")
async def request_code(req: RequestCodeRequest):
    email   = req.email.lower().strip()
    name    = req.name.strip()
    product = req.product if req.product in ("career_pathways", "certifications") else "career_pathways"
    if not email or not name:
        raise HTTPException(status_code=400, detail="Name and email are required")
    conn = get_db()
    config = conn.execute("SELECT mode FROM platform_config WHERE product = ?", (product,)).fetchone()
    mode   = config["mode"] if config else "self_serve"
    if mode == "invite_only":
        approved = conn.execute(
            "SELECT email FROM approved_emails WHERE email = ? AND product = ?", (email, product)
        ).fetchone()
        if not approved:
            conn.close()
            raise HTTPException(status_code=403, detail="This email doesn't have access yet. Contact sudharsan@bversity.io to request access.")
    conn.execute("UPDATE verification_codes SET used = 1 WHERE email = ? AND used = 0", (email,))
    code       = str(random.randint(100000, 999999))
    expires_at = (datetime.utcnow() + timedelta(minutes=15)).isoformat()
    conn.execute("INSERT INTO verification_codes (email, code, expires_at, used) VALUES (?, ?, ?, 0)",
                 (email, code, expires_at))
    conn.commit()
    conn.close()
    sent = await send_verification_email(email, name, code)
    if not sent:
        raise HTTPException(status_code=500, detail="Failed to send email. Please try again.")
    return {"message": "Verification code sent", "email": email}

@app.post("/verify-code")
def verify_code(req: VerifyCodeRequest):
    email   = req.email.lower().strip()
    code    = req.code.strip()
    name    = req.name.strip()
    product = req.product if req.product in ("career_pathways", "certifications") else "career_pathways"
    conn    = get_db()
    now_dt  = datetime.utcnow()
    now     = now_dt.isoformat()
    row = conn.execute(
        "SELECT * FROM verification_codes WHERE email = ? AND code = ? AND used = 0 AND expires_at > ? ORDER BY id DESC LIMIT 1",
        (email, code, now)
    ).fetchone()
    if not row:
        conn.close()
        raise HTTPException(status_code=400, detail="Invalid or expired code. Please request a new one.")
    conn.execute("UPDATE verification_codes SET used = 1 WHERE id = ?", (row["id"],))

    existing = conn.execute("SELECT * FROM students WHERE email = ?", (email,)).fetchone()
    if existing:
        student_id = existing["id"]
    else:
        student_id = str(uuid.uuid4())
        conn.execute("INSERT INTO students (id, name, email, created_at) VALUES (?, ?, ?, ?)",
                     (student_id, name, email, now))

    # create trial subscription if self_serve and none exists
    config = conn.execute("SELECT mode, trial_days FROM platform_config WHERE product = ?", (product,)).fetchone()
    mode       = config["mode"] if config else "self_serve"
    trial_days = config["trial_days"] if config else 5
    sub = conn.execute("SELECT * FROM subscriptions WHERE student_id = ? AND product = ?", (student_id, product)).fetchone()
    if not sub:
        trial_end = (now_dt + timedelta(days=trial_days)).isoformat() + 'Z'
        conn.execute("""
            INSERT INTO subscriptions (id, student_id, product, status, trial_start, trial_end, created_at, updated_at)
            VALUES (?, ?, ?, 'trial', ?, ?, ?, ?)
        """, (str(uuid.uuid4()), student_id, product, now, trial_end, now, now))
        sub = conn.execute("SELECT * FROM subscriptions WHERE student_id = ? AND product = ?", (student_id, product)).fetchone()

    conn.commit(); conn.close()

    subscription = None
    if sub:
        subscription = {"status": sub["status"], "trial_end": sub["trial_end"]}

    return {
        "student_id": student_id,
        "name": name if not existing else existing["name"],
        "returning": existing is not None,
        "subscription": subscription,
    }

@app.get("/admin/approved-emails")
def admin_list_approved(x_admin_key: str = Header(None)):
    require_admin(x_admin_key)
    conn = get_db()
    rows = conn.execute("SELECT email, added_at, product, expires_at, access_type FROM approved_emails ORDER BY added_at DESC").fetchall()
    conn.close()
    return [dict(r) for r in rows]

class AddEmailRequest(BaseModel):
    email: str
    product: str = "career_pathways"
    access_type: str = "manual"
    trial_days: int = None

@app.post("/admin/approved-emails")
def admin_add_email(req: AddEmailRequest, x_admin_key: str = Header(None)):
    require_admin(x_admin_key)
    email = req.email.lower().strip()
    if not email:
        raise HTTPException(status_code=400, detail="Email is required")
    product = req.product if req.product in ("career_pathways", "certifications") else "career_pathways"
    now = datetime.utcnow()
    expires_at = None
    if req.trial_days:
        expires_at = (now + timedelta(days=req.trial_days)).isoformat()
    conn = get_db()
    try:
        conn.execute(
            "INSERT INTO approved_emails (email, added_at, product, expires_at, access_type) VALUES (?, ?, ?, ?, ?)",
            (email, now.isoformat(), product, expires_at, req.access_type)
        )
        conn.commit()
    except sqlite3.IntegrityError:
        conn.close()
        raise HTTPException(status_code=409, detail="Email already approved for this product")
    conn.close()
    return {"email": email, "added_at": now.isoformat(), "product": product, "expires_at": expires_at}

@app.post("/admin/send-join-reminder")
async def send_join_reminder(x_admin_key: str = Header(None)):
    require_admin(x_admin_key)
    conn = get_db()
    approved = conn.execute("SELECT email FROM approved_emails").fetchall()
    joined   = {r["email"] for r in conn.execute("SELECT email FROM students").fetchall()}
    conn.close()
    not_joined = [r["email"] for r in approved if r["email"] not in joined]
    sent = 0
    for email in not_joined:
        first = email.split("@")[0].split(".")[0].capitalize()
        body = (
            _heading(f"Your Bversity access is waiting, {first}!") +
            _para("You were approved for Bversity  -  the world's first AI-Native Biotech University  -  "
                  "but we noticed you haven't logged in yet.") +
            _btn("Start Learning Now →", "https://university.bversity.io") +
            _divider() +
            _para("<strong>Getting started takes 2 minutes:</strong><br>"
                  "1. Click the button above.<br>"
                  "2. Enter your email to get a one-time login code.<br>"
                  "3. Pick a subject and begin your first AI tutoring session.") +
            _divider() +
            _small("Your access is still active. If you have any questions, just reply to this email.")
        )
        ok = await _send_email(email, "Your Bversity access is waiting  -  come join us! 🎓", _email_wrap(body))
        if ok:
            sent += 1
    return {"sent": sent, "total": len(not_joined)}

@app.delete("/admin/approved-emails/{email}")
def admin_remove_email(email: str, x_admin_key: str = Header(None)):
    require_admin(x_admin_key)
    conn = get_db()
    conn.execute("DELETE FROM approved_emails WHERE email = ?", (email.lower().strip(),))
    conn.commit(); conn.close()
    return {"message": "Email removed"}

@app.get("/admin/overview")
def admin_overview(x_admin_key: str = Header(None)):
    require_admin(x_admin_key)
    conn = get_db()
    week_ago = (datetime.utcnow() - timedelta(days=7)).isoformat()
    total_students    = conn.execute("SELECT COUNT(*) FROM students WHERE email NOT LIKE '%@bversity.alumni'").fetchone()[0]
    active_week       = conn.execute("SELECT COUNT(DISTINCT m.student_id) FROM messages m JOIN students s ON s.id=m.student_id WHERE m.role='user' AND m.created_at > ? AND s.email NOT LIKE '%@bversity.alumni'", (week_ago,)).fetchone()[0]
    total_concepts    = conn.execute("SELECT COUNT(*) FROM concept_progress cp JOIN students s ON s.id=cp.student_id WHERE s.email NOT LIKE '%@bversity.alumni'").fetchone()[0]
    pending_capstones = conn.execute("SELECT COUNT(*) FROM capstone_submissions WHERE score IS NULL").fetchone()[0]
    total_messages    = conn.execute("SELECT COUNT(*) FROM messages m JOIN students s ON s.id=m.student_id WHERE m.role='user' AND s.email NOT LIKE '%@bversity.alumni'").fetchone()[0]
    never_started     = conn.execute("""
        SELECT COUNT(*) FROM students s
        WHERE s.email NOT LIKE '%@bversity.alumni'
          AND NOT EXISTS (SELECT 1 FROM messages m WHERE m.student_id = s.id AND m.role = 'user')
    """).fetchone()[0]
    active_ever       = conn.execute("SELECT COUNT(DISTINCT m.student_id) FROM messages m JOIN students s ON s.id=m.student_id WHERE m.role='user' AND s.email NOT LIKE '%@bversity.alumni'").fetchone()[0]
    gone_quiet        = conn.execute("""
        SELECT COUNT(DISTINCT m.student_id) FROM messages m
        JOIN students s ON s.id=m.student_id
        WHERE m.role='user'
          AND s.email NOT LIKE '%@bversity.alumni'
          AND m.student_id NOT IN (
              SELECT DISTINCT student_id FROM messages WHERE role='user' AND created_at > ?
          )
    """, (week_ago,)).fetchone()[0]
    waitlist_total    = conn.execute("SELECT COUNT(*) FROM access_requests").fetchone()[0]
    waitlist_joined   = conn.execute(
        "SELECT COUNT(*) FROM access_requests WHERE email IN (SELECT email FROM students WHERE email NOT LIKE '%@bversity.alumni')"
    ).fetchone()[0]
    now_iso = datetime.utcnow().isoformat()
    paying_active  = conn.execute(
        "SELECT COUNT(DISTINCT sub.student_id) FROM subscriptions sub JOIN students s ON s.id=sub.student_id WHERE sub.status='active' AND s.email NOT LIKE '%@bversity.alumni'"
    ).fetchone()[0]
    on_trial       = conn.execute(
        "SELECT COUNT(DISTINCT sub.student_id) FROM subscriptions sub JOIN students s ON s.id=sub.student_id WHERE sub.status='trial' AND sub.trial_end > ? AND s.email NOT LIKE '%@bversity.alumni'",
        (now_iso,)
    ).fetchone()[0]
    trial_expired  = conn.execute(
        "SELECT COUNT(DISTINCT sub.student_id) FROM subscriptions sub JOIN students s ON s.id=sub.student_id WHERE (sub.status='expired' OR (sub.status='trial' AND sub.trial_end <= ?)) AND s.email NOT LIKE '%@bversity.alumni'",
        (now_iso,)
    ).fetchone()[0]
    week_from_now = (datetime.utcnow() + timedelta(days=7)).isoformat()
    expiring_this_week = conn.execute(
        "SELECT COUNT(DISTINCT sub.student_id) FROM subscriptions sub JOIN students s ON s.id=sub.student_id WHERE sub.status='trial' AND sub.trial_end > ? AND sub.trial_end <= ? AND s.email NOT LIKE '%@bversity.alumni'",
        (now_iso, week_from_now)
    ).fetchone()[0]
    conn.close()
    return {
        "total_students":      total_students,
        "active_week":         active_week,
        "total_concepts":      total_concepts,
        "pending_capstones":   pending_capstones,
        "total_messages":      total_messages,
        "never_started":       never_started,
        "active_ever":         active_ever,
        "gone_quiet":          gone_quiet,
        "waitlist_total":      waitlist_total,
        "waitlist_joined":     waitlist_joined,
        "paying_active":       paying_active,
        "on_trial":            on_trial,
        "trial_expired":       trial_expired,
        "expiring_this_week":  expiring_this_week,
    }

@app.get("/admin/churn-risk")
def admin_churn_risk(x_admin_key: str = Header(None)):
    require_admin(x_admin_key)
    conn = get_db()
    now_iso = datetime.utcnow().isoformat()
    week_from_now = (datetime.utcnow() + timedelta(days=7)).isoformat()
    rows = conn.execute("""
        SELECT
            s.id,
            s.name,
            s.email,
            COALESCE(sp.career_id, '') AS career_path,
            sub.trial_end,
            sub.product AS product_id,
            COALESCE(msg_stats.msg_count, 0)   AS messages_sent,
            COALESCE(msg_stats.last_active, '') AS last_active
        FROM subscriptions sub
        JOIN students s ON s.id = sub.student_id
        LEFT JOIN student_profile sp ON sp.student_id = s.id
        LEFT JOIN (
            SELECT student_id,
                   COUNT(*) AS msg_count,
                   MAX(created_at) AS last_active
            FROM messages
            WHERE role = 'user'
            GROUP BY student_id
        ) msg_stats ON msg_stats.student_id = s.id
        WHERE sub.status = 'trial'
          AND sub.trial_end > ?
          AND sub.trial_end <= ?
          AND s.email NOT LIKE '%@bversity.alumni'
        ORDER BY sub.trial_end ASC
    """, (now_iso, week_from_now)).fetchall()
    conn.close()
    result = []
    for r in rows:
        _te = r["trial_end"]
        if _te and not _te.endswith('Z') and '+' not in _te:
            _te += 'Z'
        days_left = max(0, round((
            datetime.fromisoformat(_te) - datetime.now(timezone.utc)
        ).total_seconds() / 86400, 1))
        last_active_days = None
        if r["last_active"]:
            try:
                _la = r["last_active"]
                if not _la.endswith('Z') and '+' not in _la:
                    _la += 'Z'
                delta = (datetime.now(timezone.utc) - datetime.fromisoformat(_la)).days
                last_active_days = delta
            except Exception:
                pass
        result.append({
            "id":               r["id"],
            "name":             r["name"],
            "email":            r["email"],
            "career_path":      r["career_path"] or "",
            "product_id":       r["product_id"] or "",
            "trial_end":        r["trial_end"],
            "days_left":        days_left,
            "messages_sent":    r["messages_sent"],
            "last_active_days": last_active_days,
        })
    return result


@app.get("/admin/engagement-heatmap")
def admin_engagement_heatmap(x_admin_key: str = Header(None)):
    """Returns per-subject message counts and avg session depth for the heatmap."""
    require_admin(x_admin_key)
    conn = get_db()
    rows = conn.execute("""
        SELECT
            m.subject_id,
            COUNT(*) AS message_count,
            COUNT(DISTINCT m.student_id) AS learner_count,
            COUNT(DISTINCT DATE(m.created_at)) AS active_days
        FROM messages m
        JOIN students s ON s.id = m.student_id
        WHERE m.role = 'user'
          AND s.email NOT LIKE '%@bversity.alumni'
        GROUP BY m.subject_id
        ORDER BY message_count DESC
    """).fetchall()
    conn.close()
    total_msgs = sum(r["message_count"] for r in rows) or 1
    return [
        {
            "subject_id":    r["subject_id"],
            "message_count": r["message_count"],
            "learner_count": r["learner_count"],
            "active_days":   r["active_days"],
            "share_pct":     round(r["message_count"] / total_msgs * 100, 1),
        }
        for r in rows
    ]


@app.get("/settings/hero_video")
def get_hero_video():
    with get_db() as conn:
        row = conn.execute("SELECT value FROM site_settings WHERE key='hero_video_url'").fetchone()
        return {"url": row["value"] if row else ""}

class HeroVideoUpdate(BaseModel):
    url: str

@app.put("/admin/settings/hero_video")
def set_hero_video(body: HeroVideoUpdate, x_admin_key: str = Header(None)):
    if not os.environ.get("ADMIN_KEY") or x_admin_key != os.environ.get("ADMIN_KEY"):
        raise HTTPException(status_code=403, detail="Forbidden")
    now = datetime.utcnow().isoformat()
    with get_db() as conn:
        conn.execute(
            "INSERT INTO site_settings (key, value, updated_at) VALUES ('hero_video_url', ?, ?) ON CONFLICT(key) DO UPDATE SET value=excluded.value, updated_at=excluded.updated_at",
            (body.url.strip(), now)
        )
    return {"ok": True, "url": body.url.strip()}

@app.get("/admin/revenue")
def admin_revenue(x_admin_key: str = Header(None)):
    require_admin(x_admin_key)
    conn = get_db()
    now     = datetime.utcnow()
    now_iso = now.isoformat()

    # ── Totals by currency ────────────────────────────────────────────────────
    currency_totals = {}
    for row in conn.execute("SELECT currency, SUM(amount_cents) as total FROM payment_events WHERE event_type='payment' GROUP BY currency").fetchall():
        currency_totals[row["currency"]] = row["total"] or 0

    # ── Monthly breakdown — last 12 months ────────────────────────────────────
    monthly = []
    for i in range(11, -1, -1):
        month_start = (now.replace(day=1) - __import__("datetime").timedelta(days=i*30)).replace(day=1)
        if i == 0:
            month_end = now
        else:
            next_month = (month_start.replace(day=28) + __import__("datetime").timedelta(days=4)).replace(day=1)
            month_end  = next_month - __import__("datetime").timedelta(seconds=1)
        rows = conn.execute(
            "SELECT currency, SUM(amount_cents) as total, COUNT(*) as count FROM payment_events WHERE event_type='payment' AND created_at >= ? AND created_at <= ? GROUP BY currency",
            (month_start.isoformat(), month_end.isoformat())
        ).fetchall()
        entry = {"month": month_start.strftime("%b %Y"), "totals": {}, "count": 0}
        for r in rows:
            entry["totals"][r["currency"]] = r["total"] or 0
            entry["count"] += r["count"] or 0
        monthly.append(entry)

    # ── By gateway ────────────────────────────────────────────────────────────
    gateway_rows = conn.execute(
        "SELECT gateway, currency, SUM(amount_cents) as total, COUNT(*) as count FROM payment_events WHERE event_type='payment' GROUP BY gateway, currency"
    ).fetchall()
    by_gateway = {}
    for r in gateway_rows:
        g = r["gateway"]
        if g not in by_gateway:
            by_gateway[g] = {"count": 0, "totals": {}}
        by_gateway[g]["count"] += r["count"] or 0
        by_gateway[g]["totals"][r["currency"]] = r["total"] or 0

    # ── This month ────────────────────────────────────────────────────────────
    month_start_iso = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0).isoformat()
    this_month_rows = conn.execute(
        "SELECT currency, SUM(amount_cents) as total, COUNT(*) as count FROM payment_events WHERE event_type='payment' AND created_at >= ? GROUP BY currency",
        (month_start_iso,)
    ).fetchall()
    this_month = {"totals": {}, "count": 0}
    for r in this_month_rows:
        this_month["totals"][r["currency"]] = r["total"] or 0
        this_month["count"] += r["count"] or 0

    # ── Active subscribers ────────────────────────────────────────────────────
    active_count = conn.execute(
        "SELECT COUNT(DISTINCT sub.student_id) FROM subscriptions sub JOIN students s ON s.id=sub.student_id WHERE sub.status='active' AND s.email NOT LIKE '%@bversity.alumni'"
    ).fetchone()[0]

    # ── Recent payments ───────────────────────────────────────────────────────
    recent = conn.execute("""
        SELECT pe.id, pe.student_id, s.name, s.email, pe.product, pe.amount_cents, pe.currency,
               pe.gateway, pe.gateway_payment_id, pe.created_at
        FROM payment_events pe
        JOIN students s ON s.id = pe.student_id
        WHERE pe.event_type = 'payment'
        ORDER BY pe.created_at DESC
        LIMIT 50
    """).fetchall()

    # ── Total payments count ──────────────────────────────────────────────────
    total_count = conn.execute("SELECT COUNT(*) FROM payment_events WHERE event_type='payment'").fetchone()[0]

    conn.close()
    return {
        "currency_totals":   currency_totals,
        "monthly":           monthly,
        "by_gateway":        by_gateway,
        "this_month":        this_month,
        "active_subscribers": active_count,
        "total_payments":    total_count,
        "recent_payments":   [dict(r) for r in recent],
    }

@app.get("/admin/system-health")
def admin_system_health(x_admin_key: str = Header(None)):
    require_admin(x_admin_key)
    import shutil

    # Disk
    disk = shutil.disk_usage("/")
    disk_total_gb  = round(disk.total / 1e9, 1)
    disk_used_gb   = round(disk.used  / 1e9, 1)
    disk_pct       = round(disk.used / disk.total * 100, 1)

    # Memory from /proc/meminfo
    mem_total_kb = mem_avail_kb = 0
    try:
        with open("/proc/meminfo") as f:
            for line in f:
                if line.startswith("MemTotal:"): mem_total_kb = int(line.split()[1])
                if line.startswith("MemAvailable:"): mem_avail_kb = int(line.split()[1])
    except Exception: pass
    mem_total_gb = round(mem_total_kb / 1e6, 1)
    mem_used_gb  = round((mem_total_kb - mem_avail_kb) / 1e6, 1)
    mem_pct      = round((mem_total_kb - mem_avail_kb) / mem_total_kb * 100, 1) if mem_total_kb else 0

    # DB size
    db_size_mb = round(os.path.getsize(DB_PATH) / 1e6, 2) if os.path.exists(DB_PATH) else 0

    # DB stats
    today_str = datetime.utcnow().date().isoformat()
    conn = get_db()
    total_students   = conn.execute("SELECT COUNT(*) FROM students").fetchone()[0]
    total_messages   = conn.execute("SELECT COUNT(*) FROM messages WHERE role = 'user'").fetchone()[0]
    messages_today   = conn.execute("SELECT COUNT(*) FROM messages WHERE role = 'user' AND created_at >= ?", (today_str,)).fetchone()[0]
    active_today     = conn.execute("SELECT COUNT(DISTINCT student_id) FROM messages WHERE created_at >= ?", (today_str,)).fetchone()[0]
    total_concepts   = conn.execute("SELECT COUNT(*) FROM concept_progress").fetchone()[0]
    conn.close()

    return {
        "disk":  { "total_gb": disk_total_gb, "used_gb": disk_used_gb, "pct": disk_pct },
        "memory":{ "total_gb": mem_total_gb,  "used_gb": mem_used_gb,  "pct": mem_pct },
        "db":    { "size_mb": db_size_mb },
        "students": { "total": total_students, "active_today": active_today,
                      "soft_limit": 500, "hard_limit": 1000 },
        "messages": { "total": total_messages, "today": messages_today,
                      "daily_soft_limit": 500, "daily_hard_limit": 1000 },
        "concepts_covered": total_concepts,
        "anthropic_api": _anthropic_status,
        "limits": {
            "concurrent_users":  { "value": 25,   "note": "SQLite single-writer bottleneck. Upgrade to Postgres + larger droplet at ~25 concurrent." },
            "storage_upgrade":   { "value": 15,   "note": "Consider upgrading disk when used exceeds 15GB." },
            "memory_upgrade":    { "value": 80,   "note": "Upgrade droplet RAM when sustained memory usage exceeds 80%." },
            "email_daily":       { "value": 100,  "note": "Resend free tier: 100 emails/day, 3,000/month." },
            "api_rpm":           { "value": 1000, "note": "Anthropic claude-sonnet-4-6 standard tier: ~1,000 requests/min." },
        }
    }

# ── Payments ─────────────────────────────────────────────────────────────────

STRIPE_SECRET_KEY      = os.environ.get("STRIPE_SECRET_KEY", "")
STRIPE_WEBHOOK_SECRET  = os.environ.get("STRIPE_WEBHOOK_SECRET", "")
RAZORPAY_KEY_ID        = os.environ.get("RAZORPAY_KEY_ID", "")
RAZORPAY_KEY_SECRET    = os.environ.get("RAZORPAY_KEY_SECRET", "")

PRICING = {
    "career_pathways": {"IN": {"amount": 79900, "currency": "INR", "display": "₹799/month",
                               "razorpay_plan_id": "plan_SqmoQLK8SS4Ekr"},
                        "default": {"amount": 2900,  "currency": "USD", "display": "$29/month",
                                    "razorpay_plan_id": "plan_SqmlRXMtKcsrXp"}},
    "certifications":  {"default": {"amount": 2900,  "currency": "USD", "display": "$29/month",
                                    "razorpay_plan_id": "plan_SqmlRXMtKcsrXp"}},
}

async def detect_country(ip: str) -> str:
    if not ip or ip in ("127.0.0.1", "::1", ""):
        return "US"
    import httpx as _httpx
    async with _httpx.AsyncClient(timeout=4) as client:
        # Primary: ipinfo.io plain-text endpoint
        try:
            r = await client.get(f"https://ipinfo.io/{ip}/country",
                                 headers={"Accept": "text/plain"})
            if r.status_code == 200:
                code = r.text.strip().upper()
                if len(code) == 2:
                    return code
        except Exception:
            pass
        # Fallback: ip-api.com (HTTP only on free tier)
        try:
            r = await client.get(f"http://ip-api.com/json/{ip}?fields=countryCode")
            if r.status_code == 200:
                code = r.json().get("countryCode", "US")
                if len(code) == 2:
                    return code
        except Exception:
            pass
    return "US"

def get_pricing(product: str, country_code: str) -> dict:
    tiers = PRICING.get(product, PRICING["career_pathways"])
    return tiers.get(country_code, tiers.get("default"))

@app.get("/pricing")
async def get_pricing_for_user(product: str = "career_pathways"):
    return {"career_pathways": PRICING["career_pathways"], "certifications": PRICING["certifications"]}

@app.post("/payments/detect-country")
async def detect_country_endpoint(
    cf_ipcountry: str = Header(None, alias="CF-IPCountry"),
    x_forwarded_for: str = Header(None),
    x_real_ip: str = Header(None),
):
    # If Cloudflare is in front, use their header directly (no geo lookup needed)
    if cf_ipcountry and len(cf_ipcountry) == 2:
        return {"country_code": cf_ipcountry.upper()}
    ip = x_forwarded_for.split(",")[0].strip() if x_forwarded_for else (x_real_ip or "")
    country = await detect_country(ip)
    return {"country_code": country}

class CreateCheckoutRequest(BaseModel):
    student_id: str
    product: str
    country_code: str = "US"
    success_url: str = ""
    cancel_url: str = ""

@app.post("/payments/create-checkout")
async def create_checkout(req: CreateCheckoutRequest):
    if not req.student_id:
        raise HTTPException(status_code=400, detail="student_id required")
    product      = req.product if req.product in ("career_pathways", "certifications") else "career_pathways"
    country_code = req.country_code.upper()
    pricing      = get_pricing(product, country_code)
    gateway      = "razorpay" if RAZORPAY_KEY_ID else "stripe"

    if gateway == "stripe":
        if not STRIPE_SECRET_KEY:
            return {"gateway": "stripe", "enabled": False, "message": "Payments not yet live"}
        try:
            import stripe as stripe_lib
            stripe_lib.api_key = STRIPE_SECRET_KEY
            session = stripe_lib.checkout.Session.create(
                payment_method_types=["card"],
                mode="subscription",
                line_items=[{"price_data": {
                    "currency": pricing["currency"].lower(),
                    "unit_amount": pricing["amount"],
                    "recurring": {"interval": "month"},
                    "product_data": {"name": f"Bversity {product.replace('_', ' ').title()}"},
                }, "quantity": 1}],
                client_reference_id=req.student_id,
                metadata={"student_id": req.student_id, "product": product},
                success_url=req.success_url or "https://university.bversity.io?payment=success",
                cancel_url=req.cancel_url or "https://university.bversity.io?payment=cancelled",
            )
            return {"gateway": "stripe", "enabled": True, "checkout_url": session.url}
        except Exception as e:
            raise HTTPException(status_code=500, detail=str(e))

    else:  # razorpay
        if not RAZORPAY_KEY_ID:
            return {"gateway": "razorpay", "enabled": False, "message": "Payments not yet live"}
        plan_id = pricing.get("razorpay_plan_id")
        if not plan_id:
            raise HTTPException(status_code=500, detail="No Razorpay plan configured for this pricing tier")
        try:
            import razorpay as rz
            client = rz.Client(auth=(RAZORPAY_KEY_ID, RAZORPAY_KEY_SECRET))
            # Get student email for Razorpay customer
            conn = get_db()
            student_row = conn.execute("SELECT name, email FROM students WHERE id = ?", (req.student_id,)).fetchone()
            conn.close()
            customer_name  = student_row["name"]  if student_row else "Student"
            customer_email = student_row["email"] if student_row else ""
            subscription = client.subscription.create({
                "plan_id":       plan_id,
                "total_count":   120,  # 10 years max; Razorpay requires a count
                "quantity":      1,
                "notes":         {"student_id": req.student_id, "product": product},
                "notify_info":   {"notify_phone": "", "notify_email": customer_email},
            })
            return {
                "gateway":         "razorpay",
                "enabled":         True,
                "subscription_id": subscription["id"],
                "key_id":          RAZORPAY_KEY_ID,
                "amount":          pricing["amount"],
                "currency":        pricing["currency"],
                "customer_name":   customer_name,
                "customer_email":  customer_email,
            }
        except Exception as e:
            raise HTTPException(status_code=500, detail=str(e))

@app.post("/payments/stripe-webhook")
async def stripe_webhook(request: Request):
    payload = await request.body()
    sig = request.headers.get("stripe-signature", "")
    webhook_secret = os.environ.get("STRIPE_WEBHOOK_SECRET", "")
    try:
        import stripe as stripe_lib
        stripe_lib.api_key = STRIPE_SECRET_KEY
        if webhook_secret:
            event = stripe_lib.Webhook.construct_event(payload, sig, webhook_secret)
        else:
            import json
            event = json.loads(payload)
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

    if event["type"] == "checkout.session.completed":
        session = event["data"]["object"]
        student_id = session.get("client_reference_id") or (session.get("metadata") or {}).get("student_id")
        product    = (session.get("metadata") or {}).get("product", "career_pathways")
        if student_id:
            conn = get_db()
            now = datetime.utcnow().isoformat()
            existing = conn.execute("SELECT id FROM subscriptions WHERE student_id = ? AND product = ?", (student_id, product)).fetchone()
            if existing:
                conn.execute("UPDATE subscriptions SET status = 'active', payment_method = 'stripe', updated_at = ? WHERE student_id = ? AND product = ?",
                             (now, student_id, product))
            else:
                conn.execute("INSERT INTO subscriptions (id, student_id, product, status, payment_method, started_at, updated_at) VALUES (?, ?, ?, 'active', 'stripe', ?, ?)",
                             (str(__import__("uuid").uuid4()), student_id, product, now, now))
            conn.execute(
                "INSERT INTO payment_events (id, student_id, product, event_type, amount_cents, currency, gateway, gateway_payment_id, created_at) VALUES (?, ?, ?, 'payment', ?, ?, 'stripe', ?, ?)",
                (str(__import__("uuid").uuid4()), student_id, product,
                 session.get("amount_total") or 0,
                 (session.get("currency") or "usd").lower(),
                 session.get("payment_intent"), now)
            )
            conn.commit()
            conn.close()

    elif event["type"] in ("customer.subscription.deleted", "customer.subscription.paused"):
        sub_obj    = event["data"]["object"]
        student_id = (sub_obj.get("metadata") or {}).get("student_id")
        product    = (sub_obj.get("metadata") or {}).get("product", "career_pathways")
        if student_id:
            conn = get_db()
            conn.execute("UPDATE subscriptions SET status = 'expired', updated_at = ? WHERE student_id = ? AND product = ?",
                         (datetime.utcnow().isoformat(), student_id, product))
            conn.commit()
            conn.close()

    return {"received": True}

@app.post("/payments/razorpay-webhook")
async def razorpay_webhook(request: Request):
    import hmac, hashlib, json
    payload = await request.body()
    webhook_secret = os.environ.get("RAZORPAY_WEBHOOK_SECRET", "")
    if webhook_secret:
        sig = request.headers.get("x-razorpay-signature", "")
        expected = hmac.new(webhook_secret.encode(), payload, hashlib.sha256).hexdigest()
        if not hmac.compare_digest(sig, expected):
            raise HTTPException(status_code=400, detail="Invalid signature")
    try:
        event = json.loads(payload)
    except Exception as e:
        print(f"[razorpay-webhook] Failed to parse payload: {e}")
        raise HTTPException(status_code=400, detail="Invalid JSON payload")

    event_type = event.get("event", "")

    def _get_sub_notes(event):
        sub = event.get("payload", {}).get("subscription", {}).get("entity", {})
        return sub.get("notes", {}), sub.get("id"), sub.get("plan_id")

    def _get_payment_entity(event):
        return event.get("payload", {}).get("payment", {}).get("entity", {})

    conn = get_db()
    now  = datetime.utcnow()
    now_iso = now.isoformat()
    sub_end = (now + timedelta(days=32)).isoformat()  # ~1 month buffer

    try:
        if event_type in ("subscription.charged", "payment.captured"):
            payment    = _get_payment_entity(event)
            notes      = payment.get("notes", {})
            # For subscription events, notes may be on the subscription entity
            if not notes.get("student_id"):
                notes, rzp_sub_id, _ = _get_sub_notes(event)
            else:
                rzp_sub_id = payment.get("subscription_id")
            student_id = notes.get("student_id")
            product    = notes.get("product", "career_pathways")
            if student_id:
                existing = conn.execute(
                    "SELECT id FROM subscriptions WHERE student_id = ? AND product = ?",
                    (student_id, product)).fetchone()
                if existing:
                    conn.execute(
                        "UPDATE subscriptions SET status='active', payment_method='razorpay', subscription_end=?, gateway_subscription_id=?, updated_at=? WHERE student_id=? AND product=?",
                        (sub_end, rzp_sub_id, now_iso, student_id, product))
                else:
                    conn.execute(
                        "INSERT INTO subscriptions (id, student_id, product, status, payment_method, started_at, subscription_end, gateway_subscription_id, created_at, updated_at) VALUES (?,?,?,'active','razorpay',?,?,?,?,?)",
                        (str(uuid.uuid4()), student_id, product, now_iso, sub_end, rzp_sub_id, now_iso, now_iso))
                conn.execute(
                    "INSERT INTO payment_events (id, student_id, product, event_type, amount_cents, currency, gateway, gateway_payment_id, created_at) VALUES (?,?,?,'payment',?,?,'razorpay',?,?)",
                    (str(uuid.uuid4()), student_id, product,
                     payment.get("amount") or 0,
                     (payment.get("currency") or "inr").lower(),
                     payment.get("id"), now_iso))
                conn.commit()

                # Send payment confirmation email
                student_row = conn.execute("SELECT name, email FROM students WHERE id = ?", (student_id,)).fetchone()
                if student_row and student_row["email"]:
                    first         = student_row["name"].split()[0]
                    product_label = "Certifications" if product == "certifications" else "Career Pathways"
                    currency      = (payment.get("currency") or "INR").upper()
                    amount_raw    = payment.get("amount") or 0
                    amount_disp   = f"₹{amount_raw // 100}" if currency == "INR" else f"${amount_raw / 100:.0f}"
                    renewal_date  = (now + timedelta(days=30)).strftime("%-d %B %Y")
                    body = (
                        _heading(f"You're subscribed, {first}! 🎉") +
                        _para(f"Your payment of <strong>{amount_disp}/month</strong> for Bversity {product_label} was successful. "
                              f"Your subscription is now active and will renew on <strong>{renewal_date}</strong>.") +
                        _btn("Continue Learning →", "https://university.bversity.io") +
                        _divider() +
                        _para("Your progress, certificates, and learning history are all saved. "
                              "Pick up exactly where you left off.") +
                        _small("Questions about your subscription? Reply to this email or contact sudharsan@bversity.io")
                    )
                    import asyncio as _asyncio
                    _asyncio.create_task(_send_email(student_row["email"], f"Payment confirmed — welcome to Bversity {product_label}!", _email_wrap(body)))

        elif event_type in ("subscription.halted", "subscription.cancelled"):
            notes, _, _ = _get_sub_notes(event)
            student_id = notes.get("student_id")
            product    = notes.get("product", "career_pathways")
            if student_id:
                conn.execute(
                    "UPDATE subscriptions SET status='expired', updated_at=? WHERE student_id=? AND product=?",
                    (now_iso, student_id, product))
                conn.commit()
    except Exception as e:
        print(f"[razorpay-webhook] ERROR processing event '{event_type}': {e}")
        conn.rollback()
    finally:
        conn.close()
    return {"received": True}

# ── Subscription status ───────────────────────────────────────────────────────

@app.get("/subscription/{student_id}/{product}")
def get_subscription(student_id: str, product: str):
    FREE_MESSAGE_LIMIT = 30
    conn = get_db()
    sub = conn.execute(
        "SELECT * FROM subscriptions WHERE student_id = ? AND product = ?", (student_id, product)
    ).fetchone()
    if not sub:
        conn.close()
        return {"status": "none"}
    now = datetime.utcnow().isoformat()
    if sub["status"] == "active" and sub["subscription_end"] and sub["subscription_end"] < now:
        conn.close()
        return {"status": "expired", "subscription_end": sub["subscription_end"]}
    if sub["status"] == "trial":
        msg_count = conn.execute(
            "SELECT COUNT(*) FROM messages WHERE student_id = ? AND role = 'user'", (student_id,)
        ).fetchone()[0]
        conn.close()
        if msg_count >= FREE_MESSAGE_LIMIT:
            return {"status": "expired", "messages_used": msg_count, "messages_limit": FREE_MESSAGE_LIMIT}
        return {"status": "trial", "messages_used": msg_count, "messages_limit": FREE_MESSAGE_LIMIT}
    conn.close()
    return {"status": sub["status"], "trial_end": sub["trial_end"], "subscription_end": sub["subscription_end"]}

# ── Admin: extend trial ───────────────────────────────────────────────────────

class ExtendTrialRequest(BaseModel):
    days: int = 15
    access_type: str = "free"

@app.post("/admin/students/{student_id}/extend-trial")
def extend_trial(student_id: str, product: str, body: ExtendTrialRequest, x_admin_key: str = Header(None)):
    require_admin(x_admin_key)
    conn = get_db()
    now = datetime.utcnow()
    sub = conn.execute("SELECT * FROM subscriptions WHERE student_id = ? AND product = ?", (student_id, product)).fetchone()
    if sub:
        current_end = datetime.fromisoformat(sub["trial_end"]) if sub["trial_end"] else now
        new_end = (max(current_end, now) + timedelta(days=body.days)).isoformat()
        conn.execute("UPDATE subscriptions SET trial_end = ?, status = 'trial', updated_at = ? WHERE student_id = ? AND product = ?",
                     (new_end, now.isoformat(), student_id, product))
    else:
        trial_end = (now + timedelta(days=body.days)).isoformat()
        conn.execute("""
            INSERT INTO subscriptions (id, student_id, product, status, trial_start, trial_end, created_at, updated_at)
            VALUES (?, ?, ?, 'trial', ?, ?, ?, ?)
        """, (str(uuid.uuid4()), student_id, product, now.isoformat(), trial_end, now.isoformat(), now.isoformat()))
        new_end = trial_end
    # also update approved_emails expiry
    email_row = conn.execute("SELECT email FROM students WHERE id = ?", (student_id,)).fetchone()
    if email_row:
        conn.execute("UPDATE approved_emails SET expires_at = ?, access_type = ? WHERE email = ? AND product = ?",
                     (new_end, body.access_type, email_row["email"], product))
    conn.commit(); conn.close()
    return {"ok": True, "trial_end": new_end}

# ── Platform config ─────────────────────────────────────────────────────────

@app.get("/admin/platform-config")
def get_platform_config(x_admin_key: str = Header(None)):
    require_admin(x_admin_key)
    conn = get_db()
    rows = conn.execute("SELECT * FROM platform_config").fetchall()
    conn.close()
    return {r["product"]: {"mode": r["mode"], "trial_days": r["trial_days"]} for r in rows}

class PlatformConfigUpdate(BaseModel):
    mode: str
    trial_days: int = 5

@app.put("/admin/platform-config/{product}")
def update_platform_config(product: str, body: PlatformConfigUpdate, x_admin_key: str = Header(None)):
    require_admin(x_admin_key)
    if product not in ("career_pathways", "certifications"):
        raise HTTPException(status_code=400, detail="Invalid product")
    if body.mode not in ("invite_only", "self_serve"):
        raise HTTPException(status_code=400, detail="Invalid mode")
    conn = get_db()
    conn.execute(
        "UPDATE platform_config SET mode = ?, trial_days = ?, updated_at = datetime('now') WHERE product = ?",
        (body.mode, body.trial_days, product)
    )
    conn.commit()
    conn.close()
    return {"ok": True}

# ── Students ─────────────────────────────────────────────────────────────────

@app.get("/admin/students")
def admin_students(x_admin_key: str = Header(None)):
    require_admin(x_admin_key)
    conn = get_db()
    now_iso = datetime.utcnow().isoformat()
    rows = conn.execute("""
        SELECT
            s.id, s.name, s.email, s.created_at,
            sp.career_id, sp.college, sp.year_of_study,
            sp.city, sp.state, sp.onboarded_at,
            sp.avatar_color,
            (SELECT COUNT(*) FROM messages m WHERE m.student_id = s.id AND m.role = 'user') AS message_count,
            (SELECT COUNT(DISTINCT subject_id) FROM concept_progress cp WHERE cp.student_id = s.id) AS subjects_touched,
            (SELECT COUNT(*) FROM concept_progress cp WHERE cp.student_id = s.id) AS concepts_covered,
            (SELECT COUNT(*) FROM concept_progress cp WHERE cp.student_id = s.id AND cp.mastered_at IS NOT NULL) AS concepts_mastered,
            (SELECT MAX(created_at) FROM messages m WHERE m.student_id = s.id) AS last_active,
            (SELECT COUNT(DISTINCT DATE(created_at)) FROM messages m WHERE m.student_id = s.id AND m.role = 'user') AS days_active,
            sub.status AS sub_status,
            sub.trial_end AS sub_trial_end,
            sub.payment_method AS sub_payment_method
        FROM students s
        LEFT JOIN student_profile sp ON sp.student_id = s.id
        LEFT JOIN subscriptions sub ON sub.student_id = s.id
        WHERE s.email NOT LIKE '%@bversity.alumni'
        ORDER BY s.created_at DESC
    """).fetchall()
    conn.close()
    result = []
    for r in rows:
        d = dict(r)
        career = CAREERS.get(d["career_id"]) if d.get("career_id") else None
        d["career_title"] = career["title"] if career else None
        d["career_icon"]  = career["icon"]  if career else None
        d["career_cluster"] = career["cluster"] if career else None
        cid = d.get("career_id") or ""
        d["region"] = "us" if cid.startswith("us_") else "india"
        # Resolve effective subscription status
        raw = d.get("sub_status")
        trial_end = d.get("sub_trial_end")
        if raw == "active":
            d["access_status"] = "active"
        elif raw == "trial":
            d["access_status"] = "trial" if (trial_end and trial_end > now_iso) else "expired"
        elif raw in ("expired", "cancelled"):
            d["access_status"] = raw
        else:
            d["access_status"] = "no_subscription"
        result.append(d)
    return result

@app.delete("/admin/students/{student_id}")
def delete_student(student_id: str, x_admin_key: str = Header(None)):
    require_admin(x_admin_key)
    conn = get_db()
    student = conn.execute("SELECT id, email FROM students WHERE id = ?", (student_id,)).fetchone()
    if not student:
        raise HTTPException(status_code=404, detail="Student not found")
    conn.execute("DELETE FROM messages WHERE student_id = ?", (student_id,))
    conn.execute("DELETE FROM concept_progress WHERE student_id = ?", (student_id,))
    conn.execute("DELETE FROM capstone_submissions WHERE student_id = ?", (student_id,))
    conn.execute("DELETE FROM student_profile WHERE student_id = ?", (student_id,))
    conn.execute("DELETE FROM students WHERE id = ?", (student_id,))
    conn.commit(); conn.close()
    return {"message": "Student deleted"}

@app.get("/admin/students/{student_id}/detail")
def admin_student_detail(student_id: str, x_admin_key: str = Header(None)):
    require_admin(x_admin_key)
    conn = get_db()

    # ── Core profile ──────────────────────────────────────────────────────────
    student = conn.execute("SELECT id, name, email, created_at FROM students WHERE id = ?", (student_id,)).fetchone()
    if not student:
        raise HTTPException(status_code=404, detail="Student not found")
    profile = conn.execute("SELECT * FROM student_profile WHERE student_id = ?", (student_id,)).fetchone()
    profile = dict(profile) if profile else {}

    career = CAREERS.get(profile.get("career_id")) if profile.get("career_id") else None

    # ── Concept progress per subject ──────────────────────────────────────────
    cp_rows = conn.execute(
        "SELECT subject_id, concept_id, first_covered_at, mastered_at FROM concept_progress WHERE student_id = ? ORDER BY first_covered_at",
        (student_id,)
    ).fetchall()
    progress_by_subject = {}
    for r in cp_rows:
        sid = r["subject_id"]
        if sid not in progress_by_subject:
            progress_by_subject[sid] = {"covered": [], "mastered": []}
        progress_by_subject[sid]["covered"].append({"concept_id": r["concept_id"], "at": r["first_covered_at"]})
        if r["mastered_at"]:
            progress_by_subject[sid]["mastered"].append({"concept_id": r["concept_id"], "at": r["mastered_at"]})

    # ── Sessions (group messages by 30-min gaps) ──────────────────────────────
    msgs = conn.execute(
        "SELECT subject_id, role, content, created_at FROM messages WHERE student_id = ? ORDER BY created_at",
        (student_id,)
    ).fetchall()
    sessions = []
    current_session = None
    for m in msgs:
        ts = m["created_at"]
        if current_session is None or (
            (len(current_session["messages"]) > 0) and
            ((__import__("datetime").datetime.fromisoformat(ts) -
              __import__("datetime").datetime.fromisoformat(current_session["messages"][-1]["at"])).total_seconds() > 1800)
        ):
            if current_session:
                sessions.append(current_session)
            current_session = {"subject_id": m["subject_id"], "started_at": ts, "messages": []}
        current_session["messages"].append({"role": m["role"], "content": m["content"], "at": ts})
    if current_session:
        sessions.append(current_session)

    session_summaries_list = [
        {"subject_id": r["subject_id"], "summary": r["summary"], "message_count": r["message_count"], "created_at": r["created_at"]}
        for r in conn.execute("SELECT subject_id, summary, message_count, created_at FROM session_summaries WHERE student_id = ? ORDER BY created_at DESC", (student_id,)).fetchall()
    ]

    # ── Activity timeline ─────────────────────────────────────────────────────
    timeline = []
    for s in sessions:
        user_count = sum(1 for m in s["messages"] if m["role"] == "user")
        timeline.append({"type": "session", "at": s["started_at"], "subject_id": s["subject_id"], "message_count": user_count})
    for r in cp_rows:
        timeline.append({"type": "concept_covered", "at": r["first_covered_at"], "subject_id": r["subject_id"], "concept_id": r["concept_id"]})
        if r["mastered_at"]:
            timeline.append({"type": "concept_mastered", "at": r["mastered_at"], "subject_id": r["subject_id"], "concept_id": r["concept_id"]})
    for r in conn.execute("SELECT subject_id, module_id, passed, completed_at FROM module_quizzes WHERE student_id = ? AND completed_at IS NOT NULL", (student_id,)).fetchall():
        timeline.append({"type": "quiz", "at": r["completed_at"], "subject_id": r["subject_id"], "module_id": r["module_id"], "passed": bool(r["passed"])})
    for r in conn.execute("SELECT subject_id, completed_at FROM subject_completions WHERE student_id = ?", (student_id,)).fetchall():
        timeline.append({"type": "subject_completed", "at": r["completed_at"], "subject_id": r["subject_id"]})
    for r in conn.execute("SELECT from_career_id, to_career_id, reason, changed_at FROM career_changes WHERE student_id = ? ORDER BY changed_at", (student_id,)).fetchall():
        timeline.append({"type": "career_change", "at": r["changed_at"], "from": r["from_career_id"], "to": r["to_career_id"], "reason": r["reason"]})
    timeline.sort(key=lambda x: x["at"] or "")

    # ── Quiz results ──────────────────────────────────────────────────────────
    quizzes = [dict(r) for r in conn.execute(
        "SELECT subject_id, module_id, passed, completed_at FROM module_quizzes WHERE student_id = ? ORDER BY completed_at",
        (student_id,)
    ).fetchall()]

    # ── Feedback ──────────────────────────────────────────────────────────────
    platform_fb = conn.execute(
        "SELECT q1, q2, q3, rating, comment, submitted_at FROM platform_feedback WHERE student_id = ? ORDER BY submitted_at DESC",
        (student_id,)
    ).fetchall()
    concept_fb = conn.execute(
        "SELECT subject_id, concept_title, value, created_at FROM concept_feedback WHERE student_id = ? ORDER BY created_at DESC LIMIT 50",
        (student_id,)
    ).fetchall()
    msg_fb = conn.execute(
        "SELECT subject_id, value, created_at FROM message_feedback WHERE student_id = ? ORDER BY created_at DESC LIMIT 50",
        (student_id,)
    ).fetchall()

    # ── Study plan ────────────────────────────────────────────────────────────
    study_plan = [dict(r) for r in conn.execute(
        "SELECT day_number, subject_id, concept_id, target_date FROM study_plan WHERE student_id = ? ORDER BY day_number",
        (student_id,)
    ).fetchall()]

    # ── Daily activity heatmap ────────────────────────────────────────────────
    daily = conn.execute(
        "SELECT DATE(created_at) as day, COUNT(*) as count FROM messages WHERE student_id = ? AND role = 'user' GROUP BY day ORDER BY day",
        (student_id,)
    ).fetchall()

    archetype_row = conn.execute("SELECT * FROM archetype_scores WHERE student_id = ?", (student_id,)).fetchone()

    conn.close()

    return {
        "student": {**dict(student), **profile,
                    "career_title": career["title"] if career else None,
                    "career_icon": career["icon"] if career else None,
                    "region": "us" if (profile.get("career_id") or "").startswith("us_") else "india"},
        "stats": {
            "total_messages": sum(1 for m in msgs if m["role"] == "user"),
            "total_concepts_covered": len(cp_rows),
            "total_concepts_mastered": sum(1 for r in cp_rows if r["mastered_at"]),
            "total_sessions": len(sessions),
            "days_active": len(set(m["created_at"][:10] for m in msgs if m["role"] == "user")),
            "streak_count": profile.get("streak_count", 0),
            "subjects_touched": len(progress_by_subject),
        },
        "progress_by_subject": progress_by_subject,
        "sessions": [{"subject_id": s["subject_id"], "started_at": s["started_at"], "message_count": sum(1 for m in s["messages"] if m["role"] == "user")} for s in sessions],
        "session_summaries": session_summaries_list,
        "timeline": timeline,
        "quizzes": quizzes,
        "platform_feedback": [dict(r) for r in platform_fb],
        "concept_feedback": [dict(r) for r in concept_fb],
        "message_feedback": [dict(r) for r in msg_fb],
        "study_plan": study_plan,
        "daily_activity": [dict(r) for r in daily],
        "retention": _calc_student_retention(msgs),
        "archetype": dict(archetype_row) if archetype_row else {},
    }

def _calc_student_retention(msgs):
    user_msgs = [m for m in msgs if m["role"] == "user"]
    if not user_msgs:
        return {"d1": None, "d7": None, "d14": None, "d30": None}
    first = datetime.fromisoformat(user_msgs[0]["created_at"])
    dates = set(m["created_at"][:10] for m in user_msgs)
    now = datetime.utcnow()
    result = {}
    for key, day in [("d1", 1), ("d7", 7), ("d14", 14), ("d30", 30)]:
        if (now - first).days < day - 1:
            result[key] = None
            continue
        found = any(
            (first + timedelta(days=d)).strftime("%Y-%m-%d") in dates
            for d in range(day - 1, day + 2)
        )
        result[key] = found
    return result

@app.get("/student/{student_id}")
def get_student(student_id: str):
    conn = get_db()
    student = conn.execute("SELECT id, name, email FROM students WHERE id = ?", (student_id,)).fetchone()
    conn.close()
    if not student:
        raise HTTPException(status_code=404, detail="Student not found")
    return dict(student)

@app.get("/history/{student_id}/{subject_id}")
def get_history(student_id: str, subject_id: str):
    conn = get_db()
    rows = conn.execute(
        "SELECT role, content FROM messages WHERE student_id = ? AND subject_id = ? ORDER BY id",
        (student_id, subject_id),
    ).fetchall()
    conn.close()
    return [{"role": r["role"], "content": r["content"]} for r in rows]

@app.get("/progress/{student_id}")
def get_all_progress(student_id: str):
    conn = get_db()
    profile_row = conn.execute("SELECT career_id FROM student_profile WHERE student_id = ?", (student_id,)).fetchone()
    career = CAREERS.get(profile_row["career_id"]) if profile_row and profile_row["career_id"] else None
    result = {}
    for subject_id in SUBJECTS:
        rows = conn.execute(
            "SELECT concept_id, mastered_at FROM concept_progress WHERE student_id = ? AND subject_id = ?",
            (student_id, subject_id),
        ).fetchall()
        curr     = effective_curriculum(subject_id, career)
        curr_ids = {c["id"] for c in curr}
        result[subject_id] = {
            "covered_count":  sum(1 for r in rows if r["concept_id"] in curr_ids),
            "mastered_count": sum(1 for r in rows if r["concept_id"] in curr_ids and r["mastered_at"]),
            "total":          len(curr),
        }
    conn.close()
    return result

@app.get("/progress/{student_id}/{subject_id}")
def get_subject_progress(student_id: str, subject_id: str):
    if subject_id not in SUBJECTS:
        raise HTTPException(status_code=400, detail="Invalid subject")
    conn = get_db()
    profile_row = conn.execute("SELECT career_id FROM student_profile WHERE student_id = ?", (student_id,)).fetchone()
    career = CAREERS.get(profile_row["career_id"]) if profile_row and profile_row["career_id"] else None
    rows = conn.execute(
        "SELECT concept_id, mastered_at FROM concept_progress WHERE student_id = ? AND subject_id = ?",
        (student_id, subject_id),
    ).fetchall()
    progress_map = {r["concept_id"]: r["mastered_at"] for r in rows}
    conn.close()
    curr     = effective_curriculum(subject_id, career)
    curr_ids = {c["id"] for c in curr}
    return {
        "covered_count":  sum(1 for cid in progress_map if cid in curr_ids),
        "mastered_count": sum(1 for cid, mat in progress_map.items() if cid in curr_ids and mat),
        "total":          len(curr),
        "concepts": [{"id": c["id"], "name": c["name"],
                      "covered": c["id"] in progress_map,
                      "mastered": bool(progress_map.get(c["id"]))} for c in curr],
    }

@app.get("/dashboard/{student_id}")
def get_dashboard(student_id: str):
    conn = get_db()
    student = conn.execute("SELECT id, name, email FROM students WHERE id = ?", (student_id,)).fetchone()
    if not student:
        conn.close()
        raise HTTPException(status_code=404, detail="Student not found")
    profile_row = conn.execute("SELECT career_id FROM student_profile WHERE student_id = ?", (student_id,)).fetchone()
    career = CAREERS.get(profile_row["career_id"]) if profile_row and profile_row["career_id"] else None
    result = {"student": dict(student), "subjects": {}}
    for subject_id, subject in SUBJECTS.items():
        rows = conn.execute(
            "SELECT concept_id, mastered_at FROM concept_progress WHERE student_id = ? AND subject_id = ?",
            (student_id, subject_id),
        ).fetchall()
        pm = {r["concept_id"]: r["mastered_at"] for r in rows}
        mats = conn.execute("SELECT COUNT(*) as n FROM materials WHERE subject_id = ?", (subject_id,)).fetchone()
        cap = conn.execute(
            "SELECT score, feedback, submitted_at, marked_at FROM capstone_submissions WHERE student_id = ? AND subject_id = ?",
            (student_id, subject_id),
        ).fetchone()
        curr     = effective_curriculum(subject_id, career)
        curr_ids = {c["id"] for c in curr}
        eff_covered  = sum(1 for cid in pm if cid in curr_ids)
        eff_mastered = sum(1 for cid, mat in pm.items() if cid in curr_ids and mat)
        threshold = min(len(curr), CAPSTONES[subject_id]["unlock_threshold"]) if subject_id in CAPSTONES else 8
        result["subjects"][subject_id] = {
            "covered_count":     eff_covered,
            "mastered_count":    eff_mastered,
            "total":             len(curr),
            "has_materials":     mats["n"] > 0,
            "capstone_unlocked": eff_covered >= threshold,
            "capstone_submission": dict(cap) if cap else None,
            "concepts": [{"id": c["id"], "name": c["name"],
                          "covered": c["id"] in pm,
                          "mastered": bool(pm.get(c["id"]))} for c in curr],
        }
    conn.close()
    return result

# ── Career / Profile routes ───────────────────────────────────────────────────

@app.get("/careers")
def get_careers():
    return list(CAREERS.values())

@app.get("/profile/{student_id}")
def get_profile(student_id: str, background_tasks: BackgroundTasks = None):
    background_tasks = background_tasks or BackgroundTasks()
    background_tasks.add_task(check_inactivity_nudge, student_id)
    conn = get_db()
    row = conn.execute(
        "SELECT career_id, college, year_of_study, aspirations, motivation, tutor_note, onboarded_at, avatar_color, avatar_num, linkedin_url, github_url, bio, city, state, show_on_map, streak_count, streak_last_date FROM student_profile WHERE student_id = ?",
        (student_id,)
    ).fetchone()
    student_row = conn.execute("SELECT email FROM students WHERE id = ?", (student_id,)).fetchone()
    if not student_row:
        conn.close()
        raise HTTPException(status_code=404, detail="Student not found")
    waitlist_row = None
    if student_row:
        waitlist_row = conn.execute(
            "SELECT university, year_of_study FROM access_requests WHERE email = ?",
            (student_row["email"],)
        ).fetchone()
    conn.close()
    waitlist_university = waitlist_row["university"] if waitlist_row else None
    waitlist_year = waitlist_row["year_of_study"] if waitlist_row else None
    career_obj = CAREERS.get(row["career_id"]) if row and row["career_id"] else None
    _cid = row["career_id"] if row else None
    if _cid and _cid in CAREERS:
        _career_title = CAREERS[_cid]["title"]
    elif _cid and _cid in SUBJECTS:
        _career_title = f"{SUBJECTS[_cid]['name']} · {SUBJECTS[_cid].get('certification','')}"
    else:
        _career_title = None
    career_key_set = set(career_obj.get("key_concepts", [])) if career_obj else set()
    career_concept_counts = {}
    if career_key_set:
        for sid, concepts in CURRICULUM.items():
            count = sum(1 for c in concepts if c["id"] in career_key_set)
            if count: career_concept_counts[sid] = count

    today = datetime.utcnow().date().isoformat()
    streak_count    = (row["streak_count"] or 0) if row else 0
    streak_last     = row["streak_last_date"] if row else None
    streak_today    = streak_last == today
    streak_at_risk  = bool(streak_count > 0 and not streak_today)

    if not row:
        return {"career_id": None, "career": None, "onboarded": False,
                "waitlist_university": waitlist_university, "waitlist_year_of_study": waitlist_year,
                "career_concept_counts": {}, "streak_count": 0, "streak_today": False, "streak_at_risk": False}
    return {
        "career_id": row["career_id"],
        "career": career_obj,
        "career_concept_counts": career_concept_counts,
        "onboarded": bool(row["onboarded_at"]),
        "college": row["college"],
        "year_of_study": row["year_of_study"],
        "aspirations": row["aspirations"],
        "motivation": row["motivation"],
        "tutor_note": row["tutor_note"],
        "avatar_color": row["avatar_color"],
        "avatar_num": row["avatar_num"],
        "linkedin_url": row["linkedin_url"],
        "github_url": row["github_url"],
        "bio": row["bio"],
        "city": row["city"],
        "state": row["state"],
        "show_on_map": row["show_on_map"] if row["show_on_map"] is not None else 1,
        "waitlist_university": waitlist_university,
        "waitlist_year_of_study": waitlist_year,
        "streak_count": streak_count,
        "streak_today": streak_today,
        "streak_at_risk": streak_at_risk,
        "career_title": _career_title,
    }

class OnboardingRequest(BaseModel):
    college: str
    year_of_study: str
    aspirations: str
    motivation: str
    tutor_note: str = ""
    career_id: str = ""
    city: str = ""
    state: str = ""
    show_on_map: int = 1

@app.post("/profile/{student_id}/onboarding")
def save_onboarding(student_id: str, req: OnboardingRequest):
    conn = get_db()
    career_id = req.career_id.strip() or None
    now = datetime.utcnow().isoformat()
    conn.execute(
        """INSERT INTO student_profile
           (student_id, career_id, career_goal_raw, updated_at, college, year_of_study,
            aspirations, motivation, tutor_note, city, state, show_on_map, onboarded_at)
           VALUES (?, ?, NULL, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
           ON CONFLICT(student_id) DO UPDATE SET
             career_id     = excluded.career_id,
             college       = excluded.college,
             year_of_study = excluded.year_of_study,
             aspirations   = excluded.aspirations,
             motivation    = excluded.motivation,
             tutor_note    = excluded.tutor_note,
             city          = excluded.city,
             state         = excluded.state,
             show_on_map   = excluded.show_on_map,
             onboarded_at  = excluded.onboarded_at,
             updated_at    = excluded.updated_at""",
        (student_id, career_id, now,
         req.college, req.year_of_study, req.aspirations, req.motivation, req.tutor_note,
         req.city.strip() or None, req.state.strip() or None, req.show_on_map, now)
    )
    if career_id and career_id in CAREERS:
        has_plan = conn.execute("SELECT 1 FROM study_plan WHERE student_id = ? LIMIT 1", (student_id,)).fetchone()
        if not has_plan:
            _generate_study_plan(student_id, career_id, conn)
    conn.commit(); conn.close()
    return {"status": "ok"}

class ProfileUpdateRequest(BaseModel):
    college: str = ""
    year_of_study: str = ""
    aspirations: str = ""
    motivation: str = ""
    tutor_note: str = ""
    avatar_color: str = ""
    avatar_num: Optional[int] = None
    linkedin_url: str = ""
    github_url: str = ""
    bio: str = ""
    city: str = ""
    state: str = ""
    show_on_map: int = 1

@app.put("/profile/{student_id}")
def update_profile(student_id: str, req: ProfileUpdateRequest):
    conn = get_db()
    conn.execute(
        """INSERT INTO student_profile (student_id, career_id, career_goal_raw, updated_at, college, year_of_study, aspirations, motivation, tutor_note, avatar_color, avatar_num, linkedin_url, github_url, bio, city, state, show_on_map, onboarded_at)
           VALUES (?, NULL, NULL, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, datetime('now'))
           ON CONFLICT(student_id) DO UPDATE SET
             college = excluded.college,
             year_of_study = excluded.year_of_study,
             aspirations = excluded.aspirations,
             motivation = excluded.motivation,
             tutor_note = excluded.tutor_note,
             avatar_color = excluded.avatar_color,
             avatar_num = excluded.avatar_num,
             linkedin_url = excluded.linkedin_url,
             github_url = excluded.github_url,
             bio = excluded.bio,
             city = excluded.city,
             state = excluded.state,
             show_on_map = excluded.show_on_map,
             updated_at = excluded.updated_at""",
        (student_id, datetime.utcnow().isoformat(),
         req.college, req.year_of_study, req.aspirations, req.motivation,
         req.tutor_note, req.avatar_color, req.avatar_num, req.linkedin_url, req.github_url,
         req.bio or None, req.city or None, req.state or None, req.show_on_map)
    )
    conn.commit(); conn.close()
    return {"status": "ok"}

@app.get("/community/map")
def get_community_map():
    conn = get_db()
    rows = conn.execute(
        """SELECT s.id, s.name, sp.city, sp.state, sp.career_id, sp.linkedin_url, sp.github_url, sp.avatar_color, sp.bio,
                  COALESCE(sp.is_placed, 0) AS is_placed
           FROM students s
           JOIN student_profile sp ON sp.student_id = s.id
           WHERE (sp.show_on_map IS NULL OR sp.show_on_map = 1)
             AND (sp.city IS NOT NULL OR sp.state IS NOT NULL)"""
    ).fetchall()
    conn.close()
    result = []
    for row in rows:
        career = CAREERS.get(row["career_id"] or "", {})
        result.append({
            "student_id": row["id"],
            "name": row["name"],
            "city": row["city"],
            "state": row["state"],
            "career_id": row["career_id"],
            "career_title": career.get("title", ""),
            "career_icon": career.get("icon", ""),
            "linkedin_url": row["linkedin_url"],
            "github_url": row["github_url"],
            "avatar_color": row["avatar_color"],
            "bio": row["bio"],
            "is_placed": bool(row["is_placed"]),
        })
    return result

@app.post("/profile/{student_id}")
def set_profile(student_id: str, req: ProfileRequest):
    if req.career_id not in CAREERS and req.career_id not in SUBJECTS:
        raise HTTPException(status_code=400, detail="Invalid career_id")
    conn = get_db()
    existing = conn.execute("SELECT career_id FROM student_profile WHERE student_id = ?", (student_id,)).fetchone()
    conn.execute(
        "INSERT INTO student_profile (student_id, career_id, career_goal_raw, updated_at) VALUES (?, ?, ?, ?) ON CONFLICT(student_id) DO UPDATE SET career_id = excluded.career_id, updated_at = excluded.updated_at",
        (student_id, req.career_id, None, datetime.utcnow().isoformat()),
    )
    # Generate study plan on first career selection or career change
    has_plan = conn.execute("SELECT 1 FROM study_plan WHERE student_id = ? LIMIT 1", (student_id,)).fetchone()
    if req.career_id in CAREERS and (not has_plan or (existing and existing["career_id"] != req.career_id)):
        _generate_study_plan(student_id, req.career_id, conn)
    conn.commit(); conn.close()
    career = CAREERS.get(req.career_id) or {"id": req.career_id, "name": SUBJECTS[req.career_id]["name"]}
    return {"career_id": req.career_id, "career": career}

class CareerChangeLogRequest(BaseModel):
    from_career_id: str = ""
    to_career_id: str
    reason: str
    notes: str = ""

@app.post("/career-change/{student_id}")
def log_career_change(student_id: str, req: CareerChangeLogRequest):
    if req.to_career_id not in CAREERS:
        raise HTTPException(status_code=400, detail="Invalid career_id")
    conn = get_db()
    # Log the change
    conn.execute(
        "INSERT INTO career_changes (student_id, from_career_id, to_career_id, reason, notes, changed_at) VALUES (?, ?, ?, ?, ?, ?)",
        (student_id, req.from_career_id or None, req.to_career_id, req.reason, req.notes or None, datetime.utcnow().isoformat()),
    )
    # Apply the career change
    conn.execute(
        "INSERT INTO student_profile (student_id, career_id, career_goal_raw, updated_at) VALUES (?, ?, NULL, ?) ON CONFLICT(student_id) DO UPDATE SET career_id = excluded.career_id, updated_at = excluded.updated_at",
        (student_id, req.to_career_id, datetime.utcnow().isoformat()),
    )
    # Regenerate study plan for new career
    _generate_study_plan(student_id, req.to_career_id, conn)
    conn.commit(); conn.close()
    return {"career_id": req.to_career_id, "career": CAREERS[req.to_career_id]}

@app.get("/career-changes/{student_id}")
def get_career_changes(student_id: str):
    conn = get_db()
    rows = conn.execute(
        "SELECT from_career_id, to_career_id, reason, notes, changed_at FROM career_changes WHERE student_id = ? ORDER BY changed_at DESC",
        (student_id,)
    ).fetchall()
    conn.close()
    return [dict(r) for r in rows]

# ── Subject status / unlock / pause routes ────────────────────────────────────

class SubjectPauseRequest(BaseModel):
    reason: str
    notes: str = ""

@app.get("/subjects/status/{student_id}")
def get_subject_statuses(student_id: str):
    conn = get_db()
    rows = conn.execute(
        "SELECT subject_id, status, unlocked_at, paused_at, completed_at, pause_reason FROM subject_status WHERE student_id = ?",
        (student_id,)
    ).fetchall()
    conn.close()
    return {r["subject_id"]: dict(r) for r in rows}

@app.post("/subjects/unlock/{student_id}/{subject_id}")
def unlock_subject(student_id: str, subject_id: str):
    if subject_id not in SUBJECTS:
        raise HTTPException(status_code=400, detail="Invalid subject")
    conn = get_db()
    existing = conn.execute(
        "SELECT status FROM subject_status WHERE student_id = ? AND subject_id = ?",
        (student_id, subject_id)
    ).fetchone()
    if existing and existing["status"] == "active":
        conn.close()
        return {"status": "already_active"}
    active_count = conn.execute(
        "SELECT COUNT(*) FROM subject_status WHERE student_id = ? AND status = 'active'",
        (student_id,)
    ).fetchone()[0]
    if active_count >= 2:
        conn.close()
        raise HTTPException(status_code=400, detail="max_active_reached")
    # If already has an active subject, require 12 concepts covered across active subjects
    if active_count >= 1:
        covered_in_active = conn.execute("""
            SELECT COUNT(*) FROM concept_progress cp
            JOIN subject_status ss ON cp.student_id = ss.student_id AND cp.subject_id = ss.subject_id
            WHERE cp.student_id = ? AND ss.status = 'active'
        """, (student_id,)).fetchone()[0]
        if covered_in_active < 12:
            conn.close()
            raise HTTPException(status_code=400, detail="milestone_not_reached")
    now = datetime.utcnow().isoformat()
    conn.execute("""
        INSERT INTO subject_status (student_id, subject_id, status, unlocked_at)
        VALUES (?, ?, 'active', ?)
        ON CONFLICT(student_id, subject_id) DO UPDATE SET status = 'active', unlocked_at = excluded.unlocked_at
    """, (student_id, subject_id, now))
    conn.commit()
    conn.close()
    return {"status": "active", "unlocked_at": now}

@app.post("/subjects/pause/{student_id}/{subject_id}")
def pause_subject(student_id: str, subject_id: str, req: SubjectPauseRequest):
    conn = get_db()
    existing = conn.execute(
        "SELECT status FROM subject_status WHERE student_id = ? AND subject_id = ?",
        (student_id, subject_id)
    ).fetchone()
    if not existing or existing["status"] != "active":
        conn.close()
        raise HTTPException(status_code=400, detail="Subject is not active")
    now = datetime.utcnow().isoformat()
    conn.execute(
        "UPDATE subject_status SET status = 'paused', paused_at = ?, pause_reason = ?, pause_notes = ? WHERE student_id = ? AND subject_id = ?",
        (now, req.reason, req.notes or None, student_id, subject_id)
    )
    conn.commit()
    conn.close()
    return {"status": "paused"}

@app.post("/subjects/resume/{student_id}/{subject_id}")
def resume_subject(student_id: str, subject_id: str):
    conn = get_db()
    active_count = conn.execute(
        "SELECT COUNT(*) FROM subject_status WHERE student_id = ? AND status = 'active'",
        (student_id,)
    ).fetchone()[0]
    if active_count >= 2:
        conn.close()
        raise HTTPException(status_code=400, detail="max_active_reached")
    now = datetime.utcnow().isoformat()
    conn.execute(
        "UPDATE subject_status SET status = 'active', unlocked_at = ?, paused_at = NULL WHERE student_id = ? AND subject_id = ?",
        (now, student_id, subject_id)
    )
    conn.commit()
    conn.close()
    return {"status": "active"}

# ── Concept video routes ──────────────────────────────────────────────────────

class ConceptVideoRequest(BaseModel):
    drive_url: str
    title: str = ""

@app.get("/curriculum/{subject_id}")
def get_curriculum(subject_id: str):
    if subject_id not in CURRICULUM:
        raise HTTPException(status_code=404, detail="Unknown subject")
    return [{"id": c["id"], "name": c["name"], "desc": c.get("desc", "")} for c in CURRICULUM[subject_id]]

@app.get("/curriculum")
def get_all_curriculum():
    return {
        sid: [{"id": c["id"], "name": c["name"], "desc": c.get("desc", "")} for c in concepts]
        for sid, concepts in CURRICULUM.items()
    }

@app.get("/concept-progress/{student_id}")
def get_concept_progress(student_id: str):
    conn = get_db()
    rows = conn.execute(
        "SELECT subject_id, concept_id, mastered_at FROM concept_progress WHERE student_id = ?",
        (student_id,)
    ).fetchall()
    conn.close()
    result = {}
    for row in rows:
        result[row["concept_id"]] = {
            "covered": True,
            "mastered": row["mastered_at"] is not None
        }
    return result

@app.get("/subject-videos/{subject_id}")
def get_subject_videos(subject_id: str):
    conn = get_db()
    rows = conn.execute(
        "SELECT id, title, drive_url, order_index FROM subject_videos WHERE subject_id = ? ORDER BY order_index ASC",
        (subject_id,)
    ).fetchall()
    conn.close()
    return [{"id": r["id"], "title": r["title"], "drive_url": r["drive_url"], "order_index": r["order_index"]} for r in rows]


@app.get("/concept-summaries/{subject_id}")
def get_concept_summaries(subject_id: str):
    conn = get_db()
    rows = conn.execute(
        "SELECT concept_id, what, why, key_points, real_world, interview_q FROM concept_summaries WHERE subject_id = ?",
        (subject_id,)
    ).fetchall()
    conn.close()
    import json
    result = {}
    for row in rows:
        result[row["concept_id"]] = {
            "what": row["what"],
            "why": row["why"],
            "key_points": json.loads(row["key_points"]),
            "real_world": row["real_world"],
            "interview_q": row["interview_q"],
        }
    return result


@app.get("/concept-videos")
def get_all_concept_videos():
    conn = get_db()
    rows = conn.execute("SELECT subject_id, concept_id, drive_url, title FROM concept_videos").fetchall()
    conn.close()
    result = {}
    for row in rows:
        result.setdefault(row["subject_id"], {})[row["concept_id"]] = {
            "drive_url": row["drive_url"], "title": row["title"]
        }
    return result

@app.get("/concept-videos/{subject_id}")
def get_concept_videos(subject_id: str):
    conn = get_db()
    rows = conn.execute(
        "SELECT concept_id, drive_url, title FROM concept_videos WHERE subject_id = ?",
        (subject_id,)
    ).fetchall()
    conn.close()
    return {row["concept_id"]: {"drive_url": row["drive_url"], "title": row["title"]} for row in rows}

@app.put("/admin/concept-videos/{subject_id}/{concept_id}")
def upsert_concept_video(subject_id: str, concept_id: str, req: ConceptVideoRequest, x_admin_key: str = Header(None)):
    require_admin(x_admin_key)
    if subject_id not in CURRICULUM:
        raise HTTPException(status_code=400, detail="Unknown subject")
    if not any(c["id"] == concept_id for c in CURRICULUM[subject_id]):
        raise HTTPException(status_code=400, detail="Unknown concept")
    conn = get_db()
    conn.execute(
        """INSERT INTO concept_videos (subject_id, concept_id, drive_url, title, added_at)
           VALUES (?, ?, ?, ?, ?)
           ON CONFLICT(subject_id, concept_id) DO UPDATE SET
             drive_url = excluded.drive_url, title = excluded.title, added_at = excluded.added_at""",
        (subject_id, concept_id, req.drive_url, req.title or None, datetime.utcnow().isoformat()),
    )
    conn.commit()
    conn.close()
    return {"status": "ok"}

@app.delete("/admin/concept-videos/{subject_id}/{concept_id}")
def delete_concept_video(subject_id: str, concept_id: str, x_admin_key: str = Header(None)):
    require_admin(x_admin_key)
    conn = get_db()
    conn.execute("DELETE FROM concept_videos WHERE subject_id = ? AND concept_id = ?", (subject_id, concept_id))
    conn.commit()
    conn.close()
    return {"status": "ok"}

# ── Module quiz routes ────────────────────────────────────────────────────────

@app.get("/quiz/status/{student_id}/{subject_id}")
def get_quiz_status(student_id: str, subject_id: str):
    conn = get_db()
    rows = conn.execute(
        "SELECT module_id, passed, completed_at FROM module_quizzes WHERE student_id=? AND subject_id=?",
        (student_id, subject_id)
    ).fetchall()
    conn.close()
    return {r["module_id"]: {"passed": bool(r["passed"]), "completed_at": r["completed_at"]} for r in rows}

@app.post("/quiz/complete/{student_id}/{subject_id}/{module_id}")
def complete_quiz(student_id: str, subject_id: str, module_id: str):
    conn = get_db()
    conn.execute(
        """INSERT INTO module_quizzes (student_id, subject_id, module_id, passed, completed_at)
           VALUES (?,?,?,1,?)
           ON CONFLICT(student_id,subject_id,module_id) DO UPDATE SET passed=1, completed_at=excluded.completed_at""",
        (student_id, subject_id, module_id, datetime.utcnow().isoformat())
    )
    conn.commit(); conn.close()
    return {"status": "ok"}

QUIZ_COOLDOWN_SECONDS = 3600  # 1 hour after a failed attempt

@app.get("/quiz/questions/{subject_id}/{module_id}")
async def get_quiz_questions(subject_id: str, module_id: str, student_id: str = ""):
    if subject_id not in CURRICULUM:
        raise HTTPException(status_code=404, detail="Unknown subject")
    conn = get_db()

    # Cooldown check: block retry if student failed within the last hour
    if student_id:
        last = conn.execute(
            "SELECT passed, completed_at FROM module_quizzes WHERE student_id=? AND subject_id=? AND module_id=? ORDER BY completed_at DESC LIMIT 1",
            (student_id, subject_id, module_id)
        ).fetchone()
        if last and not last["passed"]:
            elapsed = (datetime.utcnow() - datetime.fromisoformat(last["completed_at"])).total_seconds()
            if elapsed < QUIZ_COOLDOWN_SECONDS:
                wait_mins = int((QUIZ_COOLDOWN_SECONDS - elapsed) / 60) + 1
                conn.close()
                raise HTTPException(status_code=429, detail=f"cooldown:{wait_mins}")

    existing = conn.execute(
        "SELECT questions_json FROM module_quiz_questions WHERE subject_id = ? AND module_id = ?",
        (subject_id, module_id)
    ).fetchone()
    if existing:
        conn.close()
        return {"questions": json.loads(existing["questions_json"])}

    subs = [c for c in CURRICULUM[subject_id] if re.match(rf'^{re.escape(module_id)}_[abc]$', c["id"])]
    if not subs:
        conn.close()
        raise HTTPException(status_code=404, detail="Module not found")

    module_name = subs[0]["name"].split(":")[0].split(" - ")[0].strip()
    concept_lines = "\n".join(f"- {c['name']}" for c in subs)
    subject_name  = SUBJECTS.get(subject_id, {}).get("name", subject_id)

    prompt = f"""Generate exactly 4 multiple-choice questions to test deep understanding of the module "{module_name}" from the subject "{subject_name}" for biotech/life sciences graduate students.

Sub-concepts in this module:
{concept_lines}

Return ONLY a JSON array (no markdown, no explanation outside the JSON) with this exact structure:
[
  {{
    "question": "Question text here?",
    "options": ["A. First option", "B. Second option", "C. Third option", "D. Fourth option"],
    "correct_index": 0,
    "explanation": "One concise sentence explaining why this answer is correct."
  }}
]

Requirements:
- Cover a different sub-concept with each question
- Test application and understanding, not just recall
- Make distractors plausible but clearly wrong to an expert
- Graduate-level biotech difficulty"""

    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        conn.close()
        raise HTTPException(status_code=503, detail="AI not configured")

    import anthropic
    client = anthropic.Anthropic(api_key=api_key)
    message = client.messages.create(
        model="claude-haiku-4-5-20251001",
        max_tokens=2000,
        messages=[{"role": "user", "content": prompt}]
    )
    raw = message.content[0].text.strip()

    # Strip markdown code fences if present
    if raw.startswith("```"):
        raw = re.sub(r"^```[a-z]*\n?", "", raw)
        raw = re.sub(r"\n?```$", "", raw)

    try:
        questions = json.loads(raw)
        assert isinstance(questions, list) and len(questions) == 4
        for q in questions:
            assert "question" in q and "options" in q and "correct_index" in q and "explanation" in q
            assert len(q["options"]) == 4
            assert 0 <= int(q["correct_index"]) <= 3
            q["correct_index"] = int(q["correct_index"])
    except Exception as e:
        conn.close()
        print(f"Quiz question parse error: {e}\nRaw: {raw[:300]}")
        raise HTTPException(status_code=500, detail="Failed to generate valid questions  -  please retry")

    conn.execute(
        "INSERT OR REPLACE INTO module_quiz_questions (subject_id, module_id, questions_json, generated_at) VALUES (?,?,?,?)",
        (subject_id, module_id, json.dumps(questions), datetime.utcnow().isoformat())
    )
    conn.commit(); conn.close()
    return {"questions": questions}

class QuizSubmitRequest(BaseModel):
    answers: list

@app.post("/quiz/submit/{student_id}/{subject_id}/{module_id}")
def submit_quiz(student_id: str, subject_id: str, module_id: str, req: QuizSubmitRequest):
    conn = get_db()
    existing = conn.execute(
        "SELECT questions_json FROM module_quiz_questions WHERE subject_id = ? AND module_id = ?",
        (subject_id, module_id)
    ).fetchone()
    if not existing:
        conn.close()
        raise HTTPException(status_code=404, detail="Questions not found  -  generate them first")

    questions = json.loads(existing["questions_json"])
    answers = [int(a) for a in req.answers]
    correct = sum(1 for i, q in enumerate(questions) if i < len(answers) and answers[i] == q["correct_index"])
    passed  = correct >= 3

    conn.execute(
        """INSERT INTO module_quizzes (student_id, subject_id, module_id, passed, completed_at)
           VALUES (?,?,?,?,?)
           ON CONFLICT(student_id,subject_id,module_id)
           DO UPDATE SET passed=excluded.passed, completed_at=excluded.completed_at""",
        (student_id, subject_id, module_id, 1 if passed else 0, datetime.utcnow().isoformat())
    )
    conn.commit(); conn.close()

    return {
        "correct":  correct,
        "total":    len(questions),
        "passed":   passed,
        "results": [
            {
                "correct_index":  q["correct_index"],
                "explanation":    q["explanation"],
                "student_answer": answers[i] if i < len(answers) else -1,
            }
            for i, q in enumerate(questions)
        ],
    }

@app.delete("/admin/quiz/questions/{subject_id}/{module_id}")
def delete_quiz_questions(subject_id: str, module_id: str, x_admin_key: str = Header(None)):
    require_admin(x_admin_key)
    conn = get_db()
    conn.execute("DELETE FROM module_quiz_questions WHERE subject_id = ? AND module_id = ?", (subject_id, module_id))
    conn.commit(); conn.close()
    return {"status": "deleted"}

# ── Resource library routes ───────────────────────────────────────────────────

class ResourceRequest(BaseModel):
    url: str
    title: str
    resource_type: str = "article"
    description: str = ""

@app.get("/resources/{subject_id}")
def get_resources(subject_id: str):
    conn = get_db()
    rows = conn.execute(
        "SELECT concept_id, id, url, title, resource_type, description FROM concept_resources WHERE subject_id=? ORDER BY concept_id, added_at",
        (subject_id,)
    ).fetchall()
    conn.close()
    result = {}
    for r in rows:
        result.setdefault(r["concept_id"], []).append({
            "id": r["id"], "url": r["url"], "title": r["title"],
            "resource_type": r["resource_type"], "description": r["description"]
        })
    return result

@app.post("/admin/resources/{subject_id}/{concept_id}")
def add_resource(subject_id: str, concept_id: str, req: ResourceRequest, x_admin_key: str = Header(None)):
    require_admin(x_admin_key)
    if subject_id not in CURRICULUM or not any(c["id"] == concept_id for c in CURRICULUM[subject_id]):
        raise HTTPException(status_code=400, detail="Unknown subject or concept")
    conn = get_db()
    rid = str(uuid.uuid4())
    conn.execute(
        "INSERT INTO concept_resources (id, subject_id, concept_id, url, title, resource_type, description, added_at) VALUES (?,?,?,?,?,?,?,?)",
        (rid, subject_id, concept_id, req.url, req.title, req.resource_type, req.description or None, datetime.utcnow().isoformat())
    )
    conn.commit(); conn.close()
    return {"id": rid}

@app.delete("/admin/resources/{resource_id}")
def delete_resource(resource_id: str, x_admin_key: str = Header(None)):
    require_admin(x_admin_key)
    conn = get_db()
    conn.execute("DELETE FROM concept_resources WHERE id=?", (resource_id,))
    conn.commit(); conn.close()
    return {"status": "ok"}

# ── Notes routes ─────────────────────────────────────────────────────────────

class NoteRequest(BaseModel):
    content: str
    subject_id: str = ""

@app.post("/notes/{student_id}")
def save_note(student_id: str, req: NoteRequest):
    conn = get_db()
    note_id = str(uuid.uuid4())
    conn.execute(
        "INSERT INTO notes (id, student_id, subject_id, content, created_at) VALUES (?, ?, ?, ?, ?)",
        (note_id, student_id, req.subject_id or None, req.content, datetime.utcnow().isoformat()),
    )
    conn.commit()
    conn.close()
    return {"id": note_id}

@app.get("/notes/{student_id}")
def get_notes(student_id: str, subject_id: str = ""):
    conn = get_db()
    if subject_id:
        rows = conn.execute(
            "SELECT id, subject_id, content, created_at FROM notes WHERE student_id = ? AND subject_id = ? ORDER BY created_at DESC",
            (student_id, subject_id),
        ).fetchall()
    else:
        rows = conn.execute(
            "SELECT id, subject_id, content, created_at FROM notes WHERE student_id = ? ORDER BY created_at DESC",
            (student_id,),
        ).fetchall()
    conn.close()
    return [{"id": r["id"], "subject_id": r["subject_id"], "content": r["content"], "created_at": r["created_at"]} for r in rows]

@app.delete("/notes/{student_id}/{note_id}")
def delete_note(student_id: str, note_id: str):
    conn = get_db()
    conn.execute("DELETE FROM notes WHERE id = ? AND student_id = ?", (note_id, student_id))
    conn.commit()
    conn.close()
    return {"status": "ok"}

# ── Saved concepts ────────────────────────────────────────────────────────────

class SaveConceptReq(BaseModel):
    subject_id: str
    title: str
    card_data: dict

@app.post("/saved-concepts/{student_id}")
def save_concept(student_id: str, req: SaveConceptReq):
    conn = get_db()
    existing = conn.execute(
        "SELECT id FROM saved_concepts WHERE student_id = ? AND title = ? AND subject_id = ?",
        (student_id, req.title, req.subject_id)
    ).fetchone()
    if existing:
        conn.close()
        return {"id": existing["id"], "status": "already_saved"}
    cid = str(uuid.uuid4())
    conn.execute(
        "INSERT INTO saved_concepts (id, student_id, subject_id, title, card_data, saved_at) VALUES (?,?,?,?,?,?)",
        (cid, student_id, req.subject_id, req.title, json.dumps(req.card_data), datetime.utcnow().isoformat())
    )
    conn.commit()
    conn.close()
    return {"id": cid, "status": "saved"}

@app.delete("/saved-concepts/{student_id}/{concept_id}")
def unsave_concept(student_id: str, concept_id: str):
    conn = get_db()
    conn.execute("DELETE FROM saved_concepts WHERE id = ? AND student_id = ?", (concept_id, student_id))
    conn.commit()
    conn.close()
    return {"status": "ok"}

@app.get("/saved-concepts/{student_id}")
def get_saved_concepts(student_id: str):
    conn = get_db()
    rows = conn.execute(
        "SELECT id, subject_id, title, card_data, saved_at FROM saved_concepts WHERE student_id = ? ORDER BY saved_at DESC",
        (student_id,)
    ).fetchall()
    conn.close()
    result = []
    for row in rows:
        try:
            data = json.loads(row["card_data"])
        except Exception:
            data = {}
        result.append({
            "id": row["id"],
            "subject_id": row["subject_id"],
            "title": row["title"],
            "card_data": data,
            "saved_at": row["saved_at"],
        })
    return result

# ── Industry Innovation Labs routes ──────────────────────────────────────────

class LabStepReq(BaseModel):
    subject_id: str
    step_id: str

class LabSubmitReq(BaseModel):
    subject_id: str
    submission: str
    rubric: list

@app.get("/labs/{student_id}")
def get_lab_progress(student_id: str):
    conn = get_db()
    rows = conn.execute(
        "SELECT project_id, subject_id, steps_done, status, submission, ai_feedback, started_at, submitted_at FROM lab_progress WHERE student_id = ?",
        (student_id,)
    ).fetchall()
    conn.close()
    result = {}
    for row in rows:
        result[row["project_id"]] = {
            "subject_id": row["subject_id"],
            "steps_done": json.loads(row["steps_done"] or "[]"),
            "status": row["status"],
            "submission": row["submission"],
            "ai_feedback": row["ai_feedback"],
            "started_at": row["started_at"],
            "submitted_at": row["submitted_at"],
        }
    return result

@app.delete("/labs/{student_id}/{project_id}/steps/{step_id}")
def remove_lab_step(student_id: str, project_id: str, step_id: str):
    conn = get_db()
    row = conn.execute(
        "SELECT steps_done FROM lab_progress WHERE student_id = ? AND project_id = ?",
        (student_id, project_id)
    ).fetchone()
    if row:
        steps = json.loads(row["steps_done"] or "[]")
        steps = [s for s in steps if s != step_id]
        conn.execute(
            "UPDATE lab_progress SET steps_done = ? WHERE student_id = ? AND project_id = ?",
            (json.dumps(steps), student_id, project_id)
        )
        conn.commit()
    conn.close()
    return {"status": "ok"}

@app.post("/labs/{student_id}/{project_id}/steps")
def save_lab_step(student_id: str, project_id: str, req: LabStepReq):
    conn = get_db()
    row = conn.execute(
        "SELECT steps_done FROM lab_progress WHERE student_id = ? AND project_id = ?",
        (student_id, project_id)
    ).fetchone()
    if row:
        steps = json.loads(row["steps_done"] or "[]")
        if req.step_id not in steps:
            steps.append(req.step_id)
        conn.execute(
            "UPDATE lab_progress SET steps_done = ? WHERE student_id = ? AND project_id = ?",
            (json.dumps(steps), student_id, project_id)
        )
    else:
        conn.execute(
            "INSERT INTO lab_progress (student_id, project_id, subject_id, steps_done, status, started_at) VALUES (?,?,?,?,?,?)",
            (student_id, project_id, req.subject_id, json.dumps([req.step_id]), 'in_progress', datetime.utcnow().isoformat())
        )
    conn.commit()
    conn.close()
    return {"status": "ok"}

@app.post("/labs/{student_id}/{project_id}/submit")
async def submit_lab(student_id: str, project_id: str, req: LabSubmitReq):
    check_rate_limit(student_id)
    import anthropic
    client = anthropic.Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY"))
    rubric_text = "\n".join(f"- {r}" for r in req.rubric)
    prompt = f"""You are an expert biotech educator reviewing a student's Industry Innovation Lab submission.

Project ID: {project_id}
Subject: {req.subject_id}

Student's submission:
{req.submission}

Rubric criteria:
{rubric_text}

Provide structured feedback in this format:
1. **Overall Assessment** (2-3 sentences on overall quality)
2. **Rubric Feedback** (one short paragraph per criterion  -  what they did well and what to improve)
3. **Top Strength** (one specific thing done exceptionally well)
4. **Key Improvement** (the single most impactful thing to work on next)
5. **Score** (out of 10, as a number only on its own line like: Score: 8/10)

Be specific, encouraging, and grounded in real biotech practice. Reference their actual submission content."""

    response = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=1000,
        messages=[{"role": "user", "content": prompt}]
    )
    feedback = response.content[0].text

    conn = get_db()
    row = conn.execute(
        "SELECT steps_done FROM lab_progress WHERE student_id = ? AND project_id = ?",
        (student_id, project_id)
    ).fetchone()
    now = datetime.utcnow().isoformat()
    if row:
        conn.execute(
            "UPDATE lab_progress SET submission = ?, ai_feedback = ?, status = ?, submitted_at = ? WHERE student_id = ? AND project_id = ?",
            (req.submission, feedback, 'completed', now, student_id, project_id)
        )
    else:
        conn.execute(
            "INSERT INTO lab_progress (student_id, project_id, subject_id, steps_done, status, submission, ai_feedback, started_at, submitted_at) VALUES (?,?,?,?,?,?,?,?,?)",
            (student_id, project_id, req.subject_id, '[]', 'completed', req.submission, feedback, now, now)
        )
    conn.commit()
    conn.close()
    return {"feedback": feedback, "status": "completed"}

@app.post("/labs/{student_id}/{project_id}/submit-file")
async def submit_lab_file(
    student_id: str,
    project_id: str,
    file: UploadFile = File(...),
    subject_id: str = Form(...),
    rubric: str = Form(...),
):
    check_rate_limit(student_id)
    import anthropic
    rubric_list = json.loads(rubric)
    content = await file.read()
    extracted = ""
    fname = (file.filename or "").lower()
    try:
        if fname.endswith(".pdf"):
            import io as _io
            from pypdf import PdfReader
            reader = PdfReader(_io.BytesIO(content))
            extracted = "\n".join(page.extract_text() or "" for page in reader.pages)
        elif fname.endswith(".docx"):
            import io as _io
            from docx import Document
            doc = Document(_io.BytesIO(content))
            extracted = "\n".join(p.text for p in doc.paragraphs)
        else:
            extracted = content.decode("utf-8", errors="ignore")
    except Exception as e:
        extracted = content.decode("utf-8", errors="ignore")

    if not extracted.strip():
        return {"error": "Could not extract text from the uploaded file.", "status": "error"}

    client = anthropic.Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY"))
    rubric_text = "\n".join(f"- {r}" for r in rubric_list)
    prompt = f"""You are an expert biotech educator reviewing a student's Industry Innovation Lab submission.

Project ID: {project_id}
Subject: {subject_id}

Student's submitted document:
{extracted[:6000]}

Rubric criteria:
{rubric_text}

Provide structured feedback in this format:
1. **Overall Assessment** (2-3 sentences on overall quality)
2. **Rubric Feedback** (one short paragraph per criterion  -  what they did well and what to improve)
3. **Top Strength** (one specific thing done exceptionally well)
4. **Key Improvement** (the single most impactful thing to work on next)
5. **Score** (out of 10, as a number only on its own line like: Score: 8/10)

Be specific, encouraging, and grounded in real biotech practice. Reference their actual submission content."""

    response = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=1000,
        messages=[{"role": "user", "content": prompt}]
    )
    feedback = response.content[0].text

    conn = get_db()
    row = conn.execute(
        "SELECT steps_done FROM lab_progress WHERE student_id = ? AND project_id = ?",
        (student_id, project_id)
    ).fetchone()
    now = datetime.utcnow().isoformat()
    submission_note = f"[File: {file.filename}]\n\n{extracted[:2000]}"
    if row:
        conn.execute(
            "UPDATE lab_progress SET submission = ?, ai_feedback = ?, status = ?, submitted_at = ? WHERE student_id = ? AND project_id = ?",
            (submission_note, feedback, 'completed', now, student_id, project_id)
        )
    else:
        conn.execute(
            "INSERT INTO lab_progress (student_id, project_id, subject_id, steps_done, status, submission, ai_feedback, started_at, submitted_at) VALUES (?,?,?,?,?,?,?,?,?)",
            (student_id, project_id, subject_id, '[]', 'completed', submission_note, feedback, now, now)
        )
    conn.commit()
    conn.close()
    return {"feedback": feedback, "status": "completed"}

class LabAssistReq(BaseModel):
    subject_id: str
    tutor_name: str
    tutor_role: str
    project_title: str
    project_scenario: str
    project_problem: str
    messages: list

@app.post("/labs/{student_id}/assist")
async def lab_assist(student_id: str, req: LabAssistReq):
    check_rate_limit(student_id)
    import anthropic
    client = anthropic.Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY"))
    system = f"""You are {req.tutor_name}, {req.tutor_role}. A student is working on an Industry Innovation Lab project titled "{req.project_title}".

The project scenario: {req.project_scenario}

Their challenge: {req.project_problem}

Your role is to be their personal guide through this project  -  like a senior colleague sitting next to them. You know exactly what they need to do at each step, which databases to use, what the output should look like, and what common mistakes to avoid.

Be specific and practical. If they're stuck on a step, tell them exactly what to click, what to search for, what to look at. Don't be generic. Reference real tools, real data, real biology.

Keep responses concise  -  3-5 sentences max unless they ask for a detailed explanation. End with a question or prompt that moves them forward."""

    messages = [{"role": m["role"], "content": m["content"]} for m in req.messages if m.get("role") in ("user", "assistant")]
    response = client.messages.create(
        model="claude-haiku-4-5-20251001",
        max_tokens=600,
        system=system,
        messages=messages,
    )
    return {"response": response.content[0].text}

# ── Study plan routes ─────────────────────────────────────────────────────────

def _generate_study_plan(student_id: str, career_id: str, conn):
    conn.execute("DELETE FROM study_plan WHERE student_id = ?", (student_id,))
    career = CAREERS.get(career_id, {})
    relevant = career.get("relevant_subjects", [])
    ordered = relevant + [s for s in CURRICULUM if s not in relevant]
    slots = [(sid, c["id"]) for sid in ordered for c in CURRICULUM[sid]]
    start = datetime.utcnow().date()
    for i, (sid, cid) in enumerate(slots[:60]):
        day = (i // 2) + 1
        target = (start + timedelta(days=day - 1)).isoformat()
        conn.execute(
            "INSERT OR IGNORE INTO study_plan (student_id, day_number, subject_id, concept_id, target_date) VALUES (?,?,?,?,?)",
            (student_id, day, sid, cid, target),
        )

@app.post("/study-plan/generate/{student_id}")
def generate_study_plan(student_id: str, x_admin_key: str = Header(None)):
    require_admin(x_admin_key)
    conn = get_db()
    profile = conn.execute("SELECT career_id FROM student_profile WHERE student_id = ?", (student_id,)).fetchone()
    if not profile or not profile["career_id"]:
        raise HTTPException(status_code=400, detail="Student has no career selected")
    _generate_study_plan(student_id, profile["career_id"], conn)
    conn.commit()
    conn.close()
    return {"status": "ok"}

@app.get("/study-plan/{student_id}")
def get_study_plan(student_id: str):
    conn = get_db()
    rows = conn.execute(
        "SELECT day_number, subject_id, concept_id, target_date FROM study_plan WHERE student_id = ? ORDER BY day_number, concept_id",
        (student_id,),
    ).fetchall()
    covered_rows = conn.execute(
        "SELECT subject_id, concept_id FROM concept_progress WHERE student_id = ?", (student_id,)
    ).fetchall()
    conn.close()
    covered = {(r["subject_id"], r["concept_id"]) for r in covered_rows}
    concept_name_map = {}
    for sid, concepts in CURRICULUM.items():
        for c in concepts:
            concept_name_map[(sid, c["id"])] = c["name"]
    days = {}
    for r in rows:
        d = r["day_number"]
        if d not in days:
            days[d] = {"day": d, "target_date": r["target_date"], "concepts": []}
        days[d]["concepts"].append({
            "subject_id": r["subject_id"],
            "concept_id": r["concept_id"],
            "concept_name": concept_name_map.get((r["subject_id"], r["concept_id"]), r["concept_id"].replace("_", " ").title()),
            "covered": (r["subject_id"], r["concept_id"]) in covered,
        })
    plan = list(days.values())
    # compute lag
    today = datetime.utcnow().date().isoformat()
    due = [c for day in plan for c in day["concepts"] if day["target_date"] <= today]
    behind = sum(1 for c in due if not c["covered"])
    return {"plan": plan, "lag_concepts": behind, "lag_days": behind // 2}

# ── Certificate ───────────────────────────────────────────────────────────────

@app.get("/certificate/{student_id}/{subject_id}")
def get_certificate(student_id: str, subject_id: str):
    if subject_id not in CURRICULUM:
        raise HTTPException(status_code=404, detail="Unknown subject")
    conn = get_db()
    student = conn.execute("SELECT name FROM students WHERE id = ?", (student_id,)).fetchone()
    if not student:
        conn.close()
        raise HTTPException(status_code=404, detail="Student not found")
    covered_ids = {r["concept_id"] for r in conn.execute(
        "SELECT concept_id FROM concept_progress WHERE student_id = ? AND subject_id = ?",
        (student_id, subject_id)
    ).fetchall()}
    all_ids = {c["id"] for c in CURRICULUM[subject_id]}
    if not all_ids.issubset(covered_ids):
        conn.close()
        return {"eligible": False}
    existing = conn.execute(
        "SELECT credential_id, completed_at FROM subject_completions WHERE student_id = ? AND subject_id = ?",
        (student_id, subject_id)
    ).fetchone()
    if existing:
        credential_id = existing["credential_id"]
        completed_at  = existing["completed_at"]
    else:
        credential_id = str(uuid.uuid4()).replace('-', '')[:12].upper()
        completed_at  = datetime.utcnow().isoformat()
        conn.execute(
            "INSERT INTO subject_completions (student_id, subject_id, completed_at, credential_id) VALUES (?,?,?,?)",
            (student_id, subject_id, completed_at, credential_id)
        )
        conn.commit()
    conn.close()
    return {
        "eligible":       True,
        "student_name":   student["name"],
        "subject_name":   SUBJECTS.get(subject_id, {}).get("name", subject_id),
        "credential_id":  credential_id,
        "completion_date": completed_at,
    }

# ── Admin Analytics ───────────────────────────────────────────────────────────

@app.get("/admin/analytics/subjects")
def analytics_subjects(x_admin_key: str = Header(None)):
    require_admin(x_admin_key)
    conn = get_db()
    result = []
    for subject_id, subject in SUBJECTS.items():
        total_concepts = len(CURRICULUM[subject_id])
        students_count = conn.execute(
            "SELECT COUNT(DISTINCT student_id) FROM concept_progress WHERE subject_id = ?", (subject_id,)
        ).fetchone()[0]
        covered_sum = conn.execute(
            "SELECT COUNT(*) FROM concept_progress WHERE subject_id = ?", (subject_id,)
        ).fetchone()[0]
        completions = conn.execute(
            "SELECT COUNT(*) FROM subject_completions WHERE subject_id = ?", (subject_id,)
        ).fetchone()[0]
        avg_pct = round((covered_sum / (students_count * total_concepts)) * 100) if students_count > 0 else 0
        result.append({
            "subject_id":       subject_id,
            "subject_name":     subject["name"],
            "color":            subject["color"],
            "students_count":   students_count,
            "avg_coverage_pct": avg_pct,
            "total_completions": completions,
            "total_concepts":   total_concepts,
        })
    conn.close()
    result.sort(key=lambda x: x["students_count"], reverse=True)
    return result

@app.get("/admin/analytics/heatmap/{subject_id}")
def analytics_heatmap(subject_id: str, x_admin_key: str = Header(None)):
    require_admin(x_admin_key)
    if subject_id not in CURRICULUM:
        raise HTTPException(status_code=404, detail="Unknown subject")
    conn = get_db()
    rows = conn.execute(
        "SELECT concept_id, COUNT(*) AS cnt FROM concept_progress WHERE subject_id = ? GROUP BY concept_id",
        (subject_id,)
    ).fetchall()
    conn.close()
    coverage_map = {r["concept_id"]: r["cnt"] for r in rows}
    return [{"concept_id": c["id"], "concept_name": c["name"], "student_count": coverage_map.get(c["id"], 0)}
            for c in CURRICULUM[subject_id]]

@app.get("/admin/analytics/quizzes")
def analytics_quizzes(x_admin_key: str = Header(None)):
    require_admin(x_admin_key)
    conn = get_db()
    total_taken  = conn.execute("SELECT COUNT(*) FROM module_quizzes").fetchone()[0]
    total_passed = conn.execute("SELECT COUNT(*) FROM module_quizzes WHERE passed = 1").fetchone()[0]
    by_subject   = conn.execute(
        "SELECT subject_id, COUNT(*) AS taken, SUM(passed) AS passed_count FROM module_quizzes GROUP BY subject_id"
    ).fetchall()
    conn.close()
    return {
        "total_taken":  total_taken,
        "total_passed": total_passed,
        "pass_rate":    round((total_passed / total_taken * 100)) if total_taken > 0 else 0,
        "by_subject":   [{"subject_id": r["subject_id"], "taken": r["taken"], "passed": r["passed_count"]} for r in by_subject],
    }

@app.get("/admin/analytics/cohort")
def analytics_cohort(x_admin_key: str = Header(None)):
    require_admin(x_admin_key)
    conn = get_db()
    now = datetime.utcnow()
    week_ago  = (now - timedelta(days=7)).isoformat()
    month_ago = (now - timedelta(days=30)).isoformat()

    total_students = conn.execute("SELECT COUNT(*) FROM students WHERE email NOT LIKE '%@bversity.alumni'").fetchone()[0]

    active_this_week = conn.execute(
        "SELECT COUNT(DISTINCT m.student_id) FROM messages m JOIN students s ON s.id=m.student_id WHERE m.role='user' AND m.created_at>=? AND s.email NOT LIKE '%@bversity.alumni'",
        (week_ago,)
    ).fetchone()[0]

    active_this_month = conn.execute(
        "SELECT COUNT(DISTINCT m.student_id) FROM messages m JOIN students s ON s.id=m.student_id WHERE m.role='user' AND m.created_at>=? AND s.email NOT LIKE '%@bversity.alumni'",
        (month_ago,)
    ).fetchone()[0]

    # Career distribution
    career_rows = conn.execute(
        "SELECT career_id, COUNT(*) AS cnt FROM student_profile WHERE career_id IS NOT NULL GROUP BY career_id ORDER BY cnt DESC"
    ).fetchall()
    career_distribution = [
        {
            "career_id": r["career_id"],
            "career_title": CAREERS[r["career_id"]]["title"] if r["career_id"] in CAREERS else r["career_id"],
            "count": r["cnt"],
        }
        for r in career_rows
    ]

    # avg readiness by career
    career_readiness_map: dict = {}
    career_readiness_count: dict = {}
    all_profiles = conn.execute(
        "SELECT student_id, career_id FROM student_profile WHERE career_id IS NOT NULL"
    ).fetchall()
    BASE = 10
    SUBJECT_MAX = 88
    for p in all_profiles:
        sid = p["student_id"]
        cid = p["career_id"]
        if cid not in CAREERS:
            continue
        career_obj = CAREERS[cid]
        relevant_subjects = career_obj.get("relevant_subjects", [])
        if not relevant_subjects:
            continue
        score_sum = 0.0
        for subj in relevant_subjects:
            total = len(CURRICULUM.get(subj, []))
            if total == 0:
                continue
            mastered_count = conn.execute(
                "SELECT COUNT(*) FROM concept_progress WHERE student_id=? AND subject_id=? AND mastered_at IS NOT NULL",
                (sid, subj)
            ).fetchone()[0]
            covered_count = conn.execute(
                "SELECT COUNT(*) FROM concept_progress WHERE student_id=? AND subject_id=?",
                (sid, subj)
            ).fetchone()[0]
            s = (mastered_count + 0.4 * max(0, covered_count - mastered_count)) / total
            score_sum += s
        raw = BASE + (score_sum / len(relevant_subjects)) * SUBJECT_MAX
        readiness = min(int(round(raw)), 98)
        career_readiness_map[cid] = career_readiness_map.get(cid, 0) + readiness
        career_readiness_count[cid] = career_readiness_count.get(cid, 0) + 1

    avg_readiness_by_career = []
    for cid, total_r in career_readiness_map.items():
        cnt = career_readiness_count[cid]
        avg_readiness_by_career.append({
            "career_id": cid,
            "career_title": CAREERS[cid]["title"] if cid in CAREERS else cid,
            "avg_readiness": int(round(total_r / cnt)),
        })
    avg_readiness_by_career.sort(key=lambda x: x["avg_readiness"], reverse=True)

    # Hardest concepts
    concept_rows = conn.execute("""
        SELECT subject_id, concept_id,
               COUNT(*) AS covered_count,
               SUM(CASE WHEN mastered_at IS NOT NULL THEN 1 ELSE 0 END) AS mastered_count
        FROM concept_progress
        GROUP BY subject_id, concept_id
        HAVING covered_count >= 2
        ORDER BY (CAST(mastered_count AS REAL) / covered_count) ASC
        LIMIT 10
    """).fetchall()
    hardest_concepts = []
    for r in concept_rows:
        subject_name = SUBJECTS[r["subject_id"]]["name"] if r["subject_id"] in SUBJECTS else r["subject_id"]
        concept_name = r["concept_id"]
        if r["subject_id"] in CURRICULUM:
            c = next((x for x in CURRICULUM[r["subject_id"]] if x["id"] == r["concept_id"]), None)
            if c:
                concept_name = c["name"]
        mastery_rate = r["mastered_count"] / r["covered_count"] if r["covered_count"] > 0 else 0.0
        hardest_concepts.append({
            "subject_name": subject_name,
            "concept_name": concept_name,
            "covered_count": r["covered_count"],
            "mastered_count": r["mastered_count"],
            "mastery_rate": round(mastery_rate, 4),
        })

    # Quiz pass rates by subject
    quiz_rows = conn.execute(
        "SELECT subject_id, COUNT(*) AS taken, SUM(passed) AS passed_count FROM module_quizzes GROUP BY subject_id"
    ).fetchall()
    quiz_pass_rates = []
    for r in quiz_rows:
        taken = r["taken"]
        passed = r["passed_count"] or 0
        pass_rate = round(passed / taken, 4) if taken > 0 else 0.0
        quiz_pass_rates.append({
            "subject_id": r["subject_id"],
            "subject_name": SUBJECTS[r["subject_id"]]["name"] if r["subject_id"] in SUBJECTS else r["subject_id"],
            "taken": taken,
            "passed": passed,
            "pass_rate": pass_rate,
        })
    quiz_pass_rates.sort(key=lambda x: x["subject_id"])

    # Engagement
    total_messages = conn.execute("SELECT COUNT(*) FROM messages WHERE role='user'").fetchone()[0]
    avg_messages = round(total_messages / total_students, 2) if total_students > 0 else 0.0
    students_with_streak = conn.execute(
        "SELECT COUNT(*) FROM student_profile WHERE streak_count > 0"
    ).fetchone()[0]
    streak_rows = conn.execute(
        "SELECT AVG(streak_count) FROM student_profile WHERE streak_count > 0"
    ).fetchone()[0]
    avg_streak = round(streak_rows or 0.0, 2)

    conn.close()
    return {
        "total_students": total_students,
        "active_this_week": active_this_week,
        "active_this_month": active_this_month,
        "career_distribution": career_distribution,
        "avg_readiness_by_career": avg_readiness_by_career,
        "hardest_concepts": hardest_concepts,
        "quiz_pass_rates": quiz_pass_rates,
        "engagement": {
            "total_messages": total_messages,
            "avg_messages_per_student": avg_messages,
            "students_with_streak": students_with_streak,
            "avg_streak": avg_streak,
        },
    }


# ── Retention curve ──────────────────────────────────────────────────────────

@app.get("/admin/analytics/retention")
def analytics_retention(x_admin_key: str = Header(None)):
    require_admin(x_admin_key)
    conn = get_db()
    # Get all students with their signup date and message dates
    students = conn.execute("""
        SELECT s.id, s.created_at,
            (SELECT MIN(created_at) FROM messages WHERE student_id = s.id AND role = 'user') AS first_msg,
            (SELECT GROUP_CONCAT(DATE(created_at)) FROM messages WHERE student_id = s.id AND role = 'user') AS active_dates
        FROM students s
        WHERE s.email NOT LIKE '%@bversity.alumni'
    """).fetchall()

    cohort_sizes = {"d1": 0, "d7": 0, "d14": 0, "d30": 0}
    retained     = {"d1": 0, "d7": 0, "d14": 0, "d30": 0}
    student_retention = []

    for s in students:
        if not s["first_msg"]:
            student_retention.append({"id": s["id"], "started": False, "d1": False, "d7": False, "d14": False, "d30": False})
            continue
        signup = datetime.fromisoformat(s["created_at"])
        first  = datetime.fromisoformat(s["first_msg"])
        dates  = set((s["active_dates"] or "").split(","))
        now    = datetime.utcnow()

        def active_around(base, day, window=1):
            for d in range(day - window, day + window + 1):
                check = (base + timedelta(days=d)).strftime("%Y-%m-%d")
                if check in dates:
                    return True
            return False

        sr = {"id": s["id"], "started": True}
        for key, day in [("d1", 1), ("d7", 7), ("d14", 14), ("d30", 30)]:
            if (now - first).days >= day - 1:
                cohort_sizes[key] += 1
                val = active_around(first, day)
                retained[key] += 1 if val else 0
                sr[key] = val
            else:
                sr[key] = None  # too early to measure
        student_retention.append(sr)

    rates = {}
    for key in ["d1", "d7", "d14", "d30"]:
        rates[key] = round(retained[key] / cohort_sizes[key] * 100, 1) if cohort_sizes[key] else None

    # Weekly retention trend (last 8 weeks)
    weekly = []
    for w in range(7, -1, -1):
        week_start = (datetime.utcnow() - timedelta(weeks=w+1)).strftime("%Y-%m-%d")
        week_end   = (datetime.utcnow() - timedelta(weeks=w)).strftime("%Y-%m-%d")
        active = conn.execute("""
            SELECT COUNT(DISTINCT student_id) FROM messages
            WHERE role = 'user' AND DATE(created_at) >= ? AND DATE(created_at) < ?
        """, (week_start, week_end)).fetchone()[0]
        weekly.append({"week": week_start, "active": active})

    conn.close()
    return {"rates": rates, "cohort_sizes": cohort_sizes, "weekly_trend": weekly, "per_student": student_retention}

# ── Subject popularity ────────────────────────────────────────────────────────

@app.get("/admin/analytics/subject-popularity")
def analytics_subject_popularity(x_admin_key: str = Header(None)):
    require_admin(x_admin_key)
    conn = get_db()
    rows = conn.execute("""
        SELECT
            m.subject_id,
            COUNT(DISTINCT m.student_id) AS students,
            COUNT(CASE WHEN m.role = 'user' THEN 1 END) AS messages,
            COUNT(DISTINCT DATE(m.created_at)) AS active_days,
            AVG(CASE WHEN m.role = 'user' THEN 1.0 END) AS _unused,
            (SELECT COUNT(*) FROM concept_progress cp WHERE cp.subject_id = m.subject_id) AS concepts_covered,
            (SELECT COUNT(*) FROM concept_progress cp WHERE cp.subject_id = m.subject_id AND cp.mastered_at IS NOT NULL) AS concepts_mastered,
            MIN(m.created_at) AS first_session,
            MAX(m.created_at) AS last_session
        FROM messages m
        JOIN students s ON s.id = m.student_id
        WHERE s.email NOT LIKE '%@bversity.alumni'
        GROUP BY m.subject_id
        ORDER BY students DESC
    """).fetchall()
    conn.close()
    return [dict(r) for r in rows]

# ── Drop-off funnel ───────────────────────────────────────────────────────────

@app.get("/admin/analytics/funnel")
def analytics_funnel(x_admin_key: str = Header(None)):
    require_admin(x_admin_key)
    conn = get_db()
    requested   = conn.execute("SELECT COUNT(*) FROM access_requests").fetchone()[0]
    approved    = conn.execute("SELECT COUNT(*) FROM approved_emails").fetchone()[0]
    signed_up   = conn.execute("SELECT COUNT(*) FROM students WHERE email NOT LIKE '%@bversity.alumni'").fetchone()[0]
    onboarded   = conn.execute("SELECT COUNT(*) FROM student_profile WHERE onboarded_at IS NOT NULL").fetchone()[0]
    had_session = conn.execute("""
        SELECT COUNT(DISTINCT student_id) FROM messages m
        JOIN students s ON s.id = m.student_id
        WHERE m.role = 'user' AND s.email NOT LIKE '%@bversity.alumni'
    """).fetchone()[0]
    returned_d7 = conn.execute("""
        SELECT COUNT(DISTINCT m.student_id) FROM messages m
        JOIN students s ON s.id = m.student_id
        WHERE s.email NOT LIKE '%@bversity.alumni'
        AND (SELECT COUNT(DISTINCT DATE(created_at)) FROM messages m2
             WHERE m2.student_id = m.student_id AND m2.role = 'user') >= 2
        AND (SELECT MAX(created_at) FROM messages m3
             WHERE m3.student_id = m.student_id AND m3.role = 'user') > datetime('now', '-7 days')
    """).fetchone()[0]
    conn.close()

    funnel = [
        {"stage": "Requested Access",    "count": requested,   "pct": 100},
        {"stage": "Approved",            "count": approved,    "pct": round(approved / requested * 100, 1) if requested else 0},
        {"stage": "Signed Up",           "count": signed_up,   "pct": round(signed_up / requested * 100, 1) if requested else 0},
        {"stage": "Completed Onboarding","count": onboarded,   "pct": round(onboarded / requested * 100, 1) if requested else 0},
        {"stage": "Had First Session",   "count": had_session, "pct": round(had_session / requested * 100, 1) if requested else 0},
        {"stage": "Returned (7 days)",   "count": returned_d7, "pct": round(returned_d7 / requested * 100, 1) if requested else 0},
    ]
    return {"funnel": funnel}

# ── Concept difficulty map ────────────────────────────────────────────────────

@app.get("/admin/analytics/concept-difficulty")
def analytics_concept_difficulty(x_admin_key: str = Header(None)):
    require_admin(x_admin_key)
    conn = get_db()
    rows = conn.execute("""
        SELECT
            cp.subject_id,
            cp.concept_id,
            COUNT(*) AS times_covered,
            SUM(CASE WHEN cp.mastered_at IS NOT NULL THEN 1 ELSE 0 END) AS times_mastered,
            AVG(CASE WHEN cp.mastered_at IS NOT NULL
                THEN CAST((julianday(cp.mastered_at) - julianday(cp.first_covered_at)) AS REAL)
                ELSE NULL END) AS avg_days_to_master
        FROM concept_progress cp
        JOIN students s ON s.id = cp.student_id
        WHERE s.email NOT LIKE '%@bversity.alumni'
        GROUP BY cp.subject_id, cp.concept_id
        HAVING times_covered >= 2
        ORDER BY (CAST(times_mastered AS REAL) / times_covered) ASC
    """).fetchall()
    conn.close()
    result = []
    for r in rows:
        d = dict(r)
        d["mastery_rate"] = round(d["times_mastered"] / d["times_covered"] * 100, 1) if d["times_covered"] else 0
        result.append(d)
    return result

@app.get("/admin/analytics/concept-reactions")
def analytics_concept_reactions(x_admin_key: str = Header(None)):
    require_admin(x_admin_key)
    conn = get_db()
    rows = conn.execute("""
        SELECT subject_id, concept_title,
            SUM(CASE WHEN value = 'up'   THEN 1 ELSE 0 END) AS thumbs_up,
            SUM(CASE WHEN value = 'down' THEN 1 ELSE 0 END) AS thumbs_down,
            COUNT(*) AS total
        FROM concept_feedback
        GROUP BY subject_id, concept_title
        HAVING total >= 1
        ORDER BY thumbs_down DESC, total DESC
        LIMIT 50
    """).fetchall()
    conn.close()
    return [dict(r) for r in rows]

# ── Batch email endpoints ─────────────────────────────────────────────────────

@app.get("/admin/email-preview")
def email_preview(x_admin_key: str = Header(None)):
    require_admin(x_admin_key)
    conn = get_db()
    today = datetime.utcnow().date().isoformat()

    # Find students with lag >= 3 days
    plan_students = conn.execute("""
        SELECT DISTINCT sp.student_id, s.email, s.name
        FROM study_plan sp JOIN students s ON s.id = sp.student_id
    """).fetchall()
    lag_targets = []
    for st in plan_students:
        due = conn.execute("""
            SELECT COUNT(*) FROM study_plan sp
            LEFT JOIN concept_progress cp
              ON cp.student_id = sp.student_id AND cp.subject_id = sp.subject_id AND cp.concept_id = sp.concept_id
            WHERE sp.student_id = ? AND sp.target_date <= ? AND cp.concept_id IS NULL
        """, (st["student_id"], today)).fetchone()[0]
        lag_days = due // 2
        if lag_days >= 3:
            lag_targets.append({"name": st["name"], "email": st["email"], "lag_days": lag_days, "lag_concepts": due})

    # All students with at least 1 message
    week_ago = (datetime.utcnow() - timedelta(days=7)).isoformat()
    digest_students = conn.execute("""
        SELECT s.id, s.email, s.name FROM students s
        WHERE EXISTS (SELECT 1 FROM messages m WHERE m.student_id = s.id AND m.role = 'user')
    """).fetchall()
    conn.close()

    return {
        "lag_nudge_count":   len(lag_targets),
        "lag_nudge_targets": lag_targets,
        "digest_count":      len(digest_students),
    }

@app.post("/admin/check-trial-expiry")
async def check_trial_expiry(background_tasks: BackgroundTasks, x_admin_key: str = Header(None)):
    require_admin(x_admin_key)
    conn = get_db()
    now = datetime.utcnow()
    in_2_days = (now + timedelta(days=2)).isoformat()
    in_1_day  = (now + timedelta(days=1)).isoformat()

    # 2-day warning: trials expiring in 24–48 hours (skip if already warned)
    warning_2d = conn.execute("""
        SELECT sub.*, s.email, s.name
        FROM subscriptions sub JOIN students s ON s.id = sub.student_id
        WHERE sub.status = 'trial' AND sub.trial_end <= ? AND sub.trial_end > ?
          AND sub.warning_2d_sent IS NULL
          AND s.email NOT LIKE '%@bversity.alumni'
    """, (in_2_days, in_1_day)).fetchall()

    # 1-day warning: trials expiring within 24 hours (skip if already warned)
    warning_1d = conn.execute("""
        SELECT sub.*, s.email, s.name
        FROM subscriptions sub JOIN students s ON s.id = sub.student_id
        WHERE sub.status = 'trial' AND sub.trial_end <= ? AND sub.trial_end > ?
          AND sub.warning_1d_sent IS NULL
          AND s.email NOT LIKE '%@bversity.alumni'
    """, (in_1_day, now.isoformat())).fetchall()

    # Mark expired
    expired_subs = conn.execute("""
        SELECT id FROM subscriptions WHERE status = 'trial' AND trial_end < ?
    """, (now.isoformat(),)).fetchall()
    for sub in expired_subs:
        conn.execute("UPDATE subscriptions SET status = 'expired', updated_at = ? WHERE id = ?",
                     (now.isoformat(), sub["id"]))
    conn.commit()
    conn.close()

    warned = 0
    conn2 = get_db()
    for sub, days_left, label, warn_col in (
        [(s, 2, "2 days",   "warning_2d_sent") for s in warning_2d] +
        [(s, 1, "tomorrow", "warning_1d_sent") for s in warning_1d]
    ):
        product_label = "Certifications" if sub["product"] == "certifications" else "Career Pathways"
        price_line    = "₹799/month for India, $29/month elsewhere" if sub["product"] == "career_pathways" else "$29/month"
        first = sub["name"].split()[0]
        subject_line  = f"Your Bversity trial ends in {label}" if days_left == 2 else "Your Bversity trial ends tomorrow"
        body = (
            _heading(f"Your trial ends in {label}, {first}") +
            _para(f"Your 5-day free trial of Bversity {product_label} ends in <strong>{label}</strong>. "
                  f"Subscribe now to keep your progress and continue learning without interruption.") +
            _btn("Subscribe Now →", "https://university.bversity.io") +
            _divider() +
            _para(f"Bversity {product_label} is <strong>{price_line}</strong>. Cancel anytime.") +
            _small("Questions? Reply to this email or contact sudharsan@bversity.io")
        )
        ok = await _send_email(sub["email"], subject_line, _email_wrap(body))
        if ok:
            warned += 1
            allowed_cols = {"warning_2d_sent", "warning_1d_sent"}
            if warn_col not in allowed_cols:
                continue
            conn2.execute(f"UPDATE subscriptions SET {warn_col} = ? WHERE id = ?",  # noqa: S608 — col validated above
                          (now.isoformat(), sub["id"]))
            conn2.commit()
    conn2.close()

    return {"warned": warned, "expired": len(expired_subs)}

@app.post("/admin/send-nudges")
async def send_nudges(background_tasks: BackgroundTasks, x_admin_key: str = Header(None)):
    require_admin(x_admin_key)
    conn = get_db()
    today = datetime.utcnow().date().isoformat()
    plan_students = conn.execute("""
        SELECT DISTINCT sp.student_id, s.email, s.name
        FROM study_plan sp JOIN students s ON s.id = sp.student_id
        WHERE s.email IN (SELECT email FROM approved_emails)
    """).fetchall()
    sent = 0
    for st in plan_students:
        due = conn.execute("""
            SELECT COUNT(*) FROM study_plan sp
            LEFT JOIN concept_progress cp
              ON cp.student_id = sp.student_id AND cp.subject_id = sp.subject_id AND cp.concept_id = sp.concept_id
            WHERE sp.student_id = ? AND sp.target_date <= ? AND cp.concept_id IS NULL
        """, (st["student_id"], today)).fetchone()[0]
        lag_days = due // 2
        if lag_days >= 3:
            background_tasks.add_task(send_lag_nudge_email, st["email"], st["name"], lag_days, due)
            sent += 1
    conn.close()
    return {"queued": sent}

async def _send_student_weekly_report(student_id: str, email: str, name: str):
    conn = get_db()
    try:
        data = gather_student_week_data(student_id, conn)
        narrative = await generate_report_narrative(name, data)
        success = await send_weekly_learner_report(email, name, data, narrative)
        if success:
            conn.execute(
                "UPDATE student_profile SET weekly_report_sent_at=? WHERE student_id=?",
                (datetime.utcnow().isoformat(), student_id)
            )
            conn.commit()
    except Exception as e:
        print(f"Weekly report error for {student_id}: {e}")
    finally:
        conn.close()


@app.post("/admin/send-weekly-digest")
async def send_weekly_digest(background_tasks: BackgroundTasks, x_admin_key: str = Header(None)):
    require_admin(x_admin_key)
    conn = get_db()
    students = conn.execute("""
        SELECT s.id, s.email, s.name FROM students s
        WHERE EXISTS (SELECT 1 FROM messages m WHERE m.student_id = s.id AND m.role = 'user')
        AND s.email IN (SELECT email FROM approved_emails)
    """).fetchall()
    student_list = [dict(st) for st in students]
    conn.close()
    for st in student_list:
        background_tasks.add_task(_send_student_weekly_report, st["id"], st["email"], st["name"])
    return {"queued": len(student_list)}

# ── Capstone helpers ─────────────────────────────────────────────────────────

def _extract_text_from_file(filepath: str, filename: str) -> str:
    ext = filename.rsplit('.', 1)[-1].lower()
    text = ""

    def _from_pdf(path):
        try:
            import pypdf
            reader = pypdf.PdfReader(path)
            return "\n".join(p.extract_text() or "" for p in reader.pages[:40])
        except Exception as e:
            print(f"PDF extract error: {e}")
            return ""

    def _from_docx(path):
        try:
            import docx
            doc = docx.Document(path)
            return "\n".join(p.text for p in doc.paragraphs)
        except Exception as e:
            print(f"DOCX extract error: {e}")
            return ""

    if ext == "pdf":
        text = _from_pdf(filepath)
    elif ext in ("docx", "doc"):
        text = _from_docx(filepath)
    elif ext == "zip":
        import zipfile, tempfile
        try:
            with zipfile.ZipFile(filepath, "r") as z:
                candidates = sorted(
                    [n for n in z.namelist() if n.lower().endswith((".pdf", ".docx")) and not n.startswith("__")],
                    key=lambda n: z.getinfo(n).file_size, reverse=True
                )
                for name in candidates[:2]:
                    with tempfile.NamedTemporaryFile(suffix=name.rsplit(".", 1)[-1], delete=False) as tmp:
                        tmp.write(z.read(name)); tmp_path = tmp.name
                    chunk = _from_pdf(tmp_path) if name.lower().endswith(".pdf") else _from_docx(tmp_path)
                    os.unlink(tmp_path)
                    text += chunk + "\n"
                    if len(text) > 20000:
                        break
        except Exception as e:
            print(f"ZIP extract error: {e}")

    # Truncate to ~8000 words to keep prompt size reasonable
    words = text.split()
    if len(words) > 8000:
        text = " ".join(words[:8000]) + "\n[... content truncated for grading ...]"
    return text.strip()

async def _ai_grade_capstone(submission_id: str, capstone: dict, extracted_text: str):
    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key or not extracted_text:
        return

    rubric_lines = "\n".join(
        f"- {r['criterion']}: {r['marks']} marks" for r in capstone["rubric"]
    )
    prompt = f"""You are a senior academic assessor at Bversity School of Bioscience. Grade the following student submission for the capstone project titled "{capstone['title']}".

CAPSTONE BRIEF:
{capstone['problem_statement']}

MARKING RUBRIC (total {capstone['total_marks']} marks):
{rubric_lines}

STUDENT SUBMISSION (extracted text):
---
{extracted_text}
---

Grade the submission against each criterion. Be rigorous but fair  -  this is graduate-level biotech work.

Return ONLY valid JSON (no markdown, no preamble) in this exact structure:
{{
  "criterion_scores": [
    {{"criterion": "criterion name exactly as listed", "marks_awarded": <int>, "max_marks": <int>, "comments": "2-3 sentence specific assessment"}}
  ],
  "total_score": <int 0-{capstone['total_marks']}>,
  "strengths": "2-3 specific strengths of this submission",
  "areas_for_improvement": "2-3 specific areas needing improvement",
  "overall_feedback": "3-4 sentence holistic assessment suitable to share with the student"
}}"""

    try:
        import anthropic
        client = anthropic.Anthropic(api_key=api_key)
        message = client.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=2000,
            messages=[{"role": "user", "content": prompt}]
        )
        raw = message.content[0].text.strip()
        if raw.startswith("```"):
            raw = re.sub(r"^```[a-z]*\n?", "", raw)
            raw = re.sub(r"\n?```$", "", raw)
        result = json.loads(raw)
        total = int(result.get("total_score", 0))
        total = max(0, min(capstone["total_marks"], total))
        # Verify criterion scores sum approximately matches total
        crit_sum = sum(int(c.get("marks_awarded", 0)) for c in result.get("criterion_scores", []))
        if abs(crit_sum - total) > 5:
            result["total_score"] = crit_sum
            total = crit_sum

        conn = get_db()
        conn.execute(
            "UPDATE capstone_submissions SET ai_score=?, ai_feedback=?, ai_graded_at=? WHERE id=?",
            (total, json.dumps(result), datetime.utcnow().isoformat(), submission_id)
        )
        conn.commit(); conn.close()
        print(f"[AI Grade] submission {submission_id}: {total}/{capstone['total_marks']}")
    except Exception as e:
        print(f"[AI Grade] error for {submission_id}: {e}")

# ── Capstone routes ───────────────────────────────────────────────────────────

@app.get("/capstone/{subject_id}")
def get_capstone(subject_id: str):
    if subject_id not in CAPSTONES:
        raise HTTPException(status_code=404, detail="No capstone for this subject")
    return CAPSTONES[subject_id]

@app.post("/capstone/{subject_id}/submit")
async def submit_capstone(subject_id: str, background_tasks: BackgroundTasks, student_id: str = Form(...), file: UploadFile = File(...)):
    if subject_id not in CAPSTONES:
        raise HTTPException(status_code=404, detail="No capstone for this subject")
    capstone = CAPSTONES[subject_id]
    ext = file.filename.rsplit('.', 1)[-1].lower() if '.' in file.filename else ''
    if ext not in capstone["accepted_formats"]:
        raise HTTPException(status_code=400, detail=f"Accepted formats: {', '.join(capstone['accepted_formats'])}")
    content = await file.read()
    if len(content) > capstone["max_size_mb"] * 1024 * 1024:
        raise HTTPException(status_code=400, detail=f"File too large (max {capstone['max_size_mb']}MB)")
    conn = get_db()
    student = conn.execute("SELECT id FROM students WHERE id = ?", (student_id,)).fetchone()
    if not student:
        conn.close()
        raise HTTPException(status_code=404, detail="Student not found")
    profile_row = conn.execute("SELECT career_id FROM student_profile WHERE student_id = ?", (student_id,)).fetchone()
    career = CAREERS.get(profile_row["career_id"]) if profile_row and profile_row["career_id"] else None
    curr     = effective_curriculum(subject_id, career)
    curr_ids = {c["id"] for c in curr}
    threshold = min(len(curr), capstone["unlock_threshold"])
    progress_rows = conn.execute(
        "SELECT concept_id FROM concept_progress WHERE student_id = ? AND subject_id = ?",
        (student_id, subject_id),
    ).fetchall()
    eff_covered = sum(1 for r in progress_rows if r["concept_id"] in curr_ids)
    if eff_covered < threshold:
        conn.close()
        raise HTTPException(status_code=403, detail=f"Cover at least {threshold} concepts to unlock the capstone")
    submission_id = str(uuid.uuid4())
    subject_dir = os.path.join(SUBMISSIONS_DIR, subject_id)
    os.makedirs(subject_dir, exist_ok=True)
    safe_name = re.sub(r'[^\w.\-]', '_', file.filename)
    filepath = os.path.join(subject_dir, f"{submission_id}_{safe_name}")
    with open(filepath, "wb") as f:
        f.write(content)
    conn.execute(
        "INSERT INTO capstone_submissions (id, student_id, subject_id, filename, filepath, submitted_at) VALUES (?, ?, ?, ?, ?, ?) ON CONFLICT(student_id, subject_id) DO UPDATE SET id=excluded.id, filename=excluded.filename, filepath=excluded.filepath, submitted_at=excluded.submitted_at, score=NULL, feedback=NULL, marked_at=NULL",
        (submission_id, student_id, subject_id, file.filename, filepath, datetime.utcnow().isoformat()),
    )
    conn.commit(); conn.close()
    # Extract text and trigger AI grading in background
    extracted = _extract_text_from_file(filepath, file.filename)
    background_tasks.add_task(_ai_grade_capstone, submission_id, capstone, extracted)
    return {"submission_id": submission_id, "filename": file.filename, "status": "submitted"}

@app.get("/capstone/{subject_id}/submission/{student_id}")
def get_my_submission(subject_id: str, student_id: str):
    conn = get_db()
    row = conn.execute(
        "SELECT id, filename, submitted_at, score, feedback, marked_at, ai_score, ai_feedback, ai_graded_at FROM capstone_submissions WHERE student_id = ? AND subject_id = ?",
        (student_id, subject_id),
    ).fetchone()
    conn.close()
    if not row:
        return {"submitted": False}
    return {"submitted": True, **dict(row)}

# ── Admin routes ──────────────────────────────────────────────────────────────

@app.get("/admin/submissions")
def admin_list_submissions(x_admin_key: str = Header(None)):
    require_admin(x_admin_key)
    conn = get_db()
    rows = conn.execute(
        """SELECT cs.id, cs.student_id, s.name as student_name, cs.subject_id,
           cs.filename, cs.submitted_at, cs.score, cs.feedback, cs.marked_at,
           cs.ai_score, cs.ai_feedback, cs.ai_graded_at
           FROM capstone_submissions cs
           JOIN students s ON s.id = cs.student_id
           ORDER BY cs.submitted_at DESC""",
    ).fetchall()
    conn.close()
    return [dict(r) for r in rows]

@app.post("/admin/submissions/{submission_id}/ai-grade")
async def admin_ai_grade(submission_id: str, background_tasks: BackgroundTasks, x_admin_key: str = Header(None)):
    require_admin(x_admin_key)
    conn = get_db()
    row = conn.execute(
        "SELECT filepath, filename, subject_id FROM capstone_submissions WHERE id = ?", (submission_id,)
    ).fetchone()
    conn.close()
    if not row:
        raise HTTPException(status_code=404, detail="Submission not found")
    if row["subject_id"] not in CAPSTONES:
        raise HTTPException(status_code=400, detail="No capstone definition for this subject")
    extracted = _extract_text_from_file(row["filepath"], row["filename"])
    background_tasks.add_task(_ai_grade_capstone, submission_id, CAPSTONES[row["subject_id"]], extracted)
    return {"status": "grading_queued"}

@app.get("/admin/submissions/{submission_id}/download")
def admin_download_submission(submission_id: str, x_admin_key: str = Header(None)):
    require_admin(x_admin_key)
    conn = get_db()
    row = conn.execute("SELECT filepath, filename FROM capstone_submissions WHERE id = ?", (submission_id,)).fetchone()
    conn.close()
    if not row or not os.path.exists(row["filepath"]):
        raise HTTPException(status_code=404, detail="File not found")
    return FileResponse(row["filepath"], filename=row["filename"], media_type="application/octet-stream")

@app.post("/admin/submissions/{submission_id}/mark")
def admin_mark_submission(submission_id: str, req: MarkRequest, x_admin_key: str = Header(None)):
    require_admin(x_admin_key)
    if not (0 <= req.score <= 100):
        raise HTTPException(status_code=400, detail="Score must be 0–100")
    conn = get_db()
    row = conn.execute("SELECT id FROM capstone_submissions WHERE id = ?", (submission_id,)).fetchone()
    if not row:
        conn.close()
        raise HTTPException(status_code=404, detail="Submission not found")
    conn.execute(
        "UPDATE capstone_submissions SET score = ?, feedback = ?, marked_at = ? WHERE id = ?",
        (req.score, req.feedback.strip(), datetime.utcnow().isoformat(), submission_id),
    )
    conn.commit(); conn.close()
    return {"submission_id": submission_id, "score": req.score, "marked": True}

# ── Certificate routes ───────────────────────────────────────────────────────

@app.get("/certificate/{student_id}/{subject_id}")
def get_certificate(student_id: str, subject_id: str):
    if subject_id not in SUBJECTS or subject_id not in CURRICULUM:
        raise HTTPException(status_code=404, detail="Subject not found")
    conn = get_db()
    student = conn.execute("SELECT name FROM students WHERE id = ?", (student_id,)).fetchone()
    if not student:
        conn.close()
        raise HTTPException(status_code=404, detail="Student not found")

    curr = CURRICULUM[subject_id]
    curr_ids = {c["id"] for c in curr}
    rows = conn.execute(
        "SELECT concept_id, first_covered_at FROM concept_progress WHERE student_id = ? AND subject_id = ?",
        (student_id, subject_id),
    ).fetchall()
    conn.close()

    covered_ids = {r["concept_id"] for r in rows if r["concept_id"] in curr_ids}
    covered_count = len(covered_ids)
    total = len(curr)

    if covered_count < total:
        raise HTTPException(
            status_code=403,
            detail=f"Cover all {total} concepts to earn this certificate. Currently {covered_count}/{total}."
        )

    completion_dates = [r["first_covered_at"] for r in rows if r["concept_id"] in curr_ids and r["first_covered_at"]]
    completion_date = max(completion_dates) if completion_dates else datetime.utcnow().isoformat()

    credential_id = hashlib.sha256(f"{student_id}:{subject_id}".encode()).hexdigest()[:12].upper()

    return {
        "student_name": student["name"],
        "subject_id": subject_id,
        "subject_name": SUBJECTS.get(subject_id, {}).get("name", subject_id),
        "subject_expert": SUBJECTS.get(subject_id, {}).get("tutor_name", ""),
        "concepts_covered": covered_count,
        "total_concepts": total,
        "completion_date": completion_date[:10],
        "credential_id": credential_id,
        "eligible": True,
    }


@app.get("/verify/{credential_id}")
def verify_certificate(credential_id: str):
    conn = get_db()
    row = conn.execute(
        """SELECT sc.student_id, sc.subject_id, sc.completed_at, s.name AS student_name
           FROM subject_completions sc
           JOIN students s ON s.id = sc.student_id
           WHERE sc.credential_id = ?""",
        (credential_id.upper(),)
    ).fetchone()
    conn.close()
    if not row:
        raise HTTPException(status_code=404, detail="Certificate not found")
    subject = SUBJECTS.get(row["subject_id"], {})
    return {
        "student_name": row["student_name"],
        "subject_id": row["subject_id"],
        "subject_name": subject.get("name", row["subject_id"]),
        "subject_expert": subject.get("tutor_name", ""),
        "completion_date": row["completed_at"][:10] if row["completed_at"] else "",
        "credential_id": credential_id.upper(),
        "eligible": True,
        "verified": True,
    }


# ── Materials routes ──────────────────────────────────────────────────────────

@app.get("/materials/{subject_id}")
def list_materials(subject_id: str):
    if subject_id not in SUBJECTS:
        raise HTTPException(status_code=400, detail="Invalid subject")
    conn = get_db()
    rows = conn.execute(
        "SELECT id, filename, chunk_count, uploaded_at FROM materials WHERE subject_id = ? ORDER BY uploaded_at DESC",
        (subject_id,),
    ).fetchall()
    conn.close()
    return [dict(r) for r in rows]

@app.post("/materials/{subject_id}")
async def upload_material(subject_id: str, file: UploadFile = File(...)):
    if subject_id not in SUBJECTS:
        raise HTTPException(status_code=400, detail="Invalid subject")
    content = await file.read()
    if len(content) > 20 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="File too large (max 20MB)")
    text = extract_text(content, file.filename)
    if not text.strip():
        raise HTTPException(status_code=400, detail="Could not extract text from file")
    chunks = chunk_text(text)
    if not chunks:
        raise HTTPException(status_code=400, detail="No content found in file")
    material_id = str(uuid.uuid4())
    conn = get_db()
    conn.execute(
        "INSERT INTO materials (id, subject_id, filename, chunk_count, uploaded_at) VALUES (?, ?, ?, ?, ?)",
        (material_id, subject_id, file.filename, len(chunks), datetime.utcnow().isoformat()),
    )
    for chunk in chunks:
        conn.execute("INSERT INTO doc_chunks_fts (material_id, subject_id, content) VALUES (?, ?, ?)", (material_id, subject_id, chunk))
    conn.commit(); conn.close()
    return {"id": material_id, "filename": file.filename, "chunks": len(chunks)}

@app.delete("/materials/{material_id}")
def delete_material(material_id: str):
    conn = get_db()
    mat = conn.execute("SELECT * FROM materials WHERE id = ?", (material_id,)).fetchone()
    if not mat:
        conn.close()
        raise HTTPException(status_code=404, detail="Material not found")
    conn.execute("DELETE FROM doc_chunks_fts WHERE material_id = ?", (material_id,))
    conn.execute("DELETE FROM materials WHERE id = ?", (material_id,))
    conn.commit(); conn.close()
    return {"deleted": material_id}

# ── Chat ──────────────────────────────────────────────────────────────────────

@app.post("/chat")
async def chat(req: ChatRequest, background_tasks: BackgroundTasks):
    if req.subject_id not in SUBJECTS:
        raise HTTPException(status_code=400, detail="Invalid subject")
    check_rate_limit(req.student_id)
    conn = get_db()
    student = conn.execute("SELECT * FROM students WHERE id = ?", (req.student_id,)).fetchone()
    if not student:
        conn.close()
        raise HTTPException(status_code=404, detail="Student not found")

    # ── Subscription gate ────────────────────────────────────────────────────
    FREE_MESSAGE_LIMIT = 30
    product = "certifications" if req.subject_id.startswith("us_") else "career_pathways"
    sub = conn.execute("SELECT status, trial_end FROM subscriptions WHERE student_id = ? AND product = ?", (req.student_id, product)).fetchone()
    if sub:
        if sub["status"] == "trial":
            msg_count = conn.execute(
                "SELECT COUNT(*) FROM messages WHERE student_id = ? AND role = 'user'", (req.student_id,)
            ).fetchone()[0]
            if msg_count >= FREE_MESSAGE_LIMIT:
                conn.close()
                raise HTTPException(status_code=402, detail="trial_expired")
        if sub["status"] == "expired":
            conn.close()
            raise HTTPException(status_code=402, detail="subscription_expired")

    subject = SUBJECTS[req.subject_id]
    profile_row = conn.execute("SELECT career_id FROM student_profile WHERE student_id = ?", (req.student_id,)).fetchone()
    career = CAREERS.get(profile_row["career_id"]) if profile_row and profile_row["career_id"] else None

    # Auto-activate subject if no status row (backward compat for existing users)
    status_row = conn.execute(
        "SELECT status FROM subject_status WHERE student_id = ? AND subject_id = ?",
        (req.student_id, req.subject_id)
    ).fetchone()
    if not status_row:
        active_count = conn.execute(
            "SELECT COUNT(*) FROM subject_status WHERE student_id = ? AND status = 'active'",
            (req.student_id,)
        ).fetchone()[0]
        if active_count < 2:
            conn.execute(
                "INSERT OR IGNORE INTO subject_status (student_id, subject_id, status, unlocked_at) VALUES (?, ?, 'active', ?)",
                (req.student_id, req.subject_id, datetime.utcnow().isoformat())
            )
            conn.commit()

    history_rows = conn.execute(
        "SELECT role, content FROM messages WHERE student_id = ? AND subject_id = ? ORDER BY id",
        (req.student_id, req.subject_id),
    ).fetchall()
    is_first_visit = len(history_rows) == 0
    # Keep only the last 20 messages to limit input tokens
    history = [{"role": r["role"], "content": r["content"]} for r in history_rows[-20:]]

    progress_rows = conn.execute(
        "SELECT concept_id, mastered_at FROM concept_progress WHERE student_id = ? AND subject_id = ?",
        (req.student_id, req.subject_id),
    ).fetchall()
    covered_ids  = [r["concept_id"] for r in progress_rows]
    mastered_ids = [r["concept_id"] for r in progress_rows if r["mastered_at"]]

    if not req.auto_open:
        history.append({"role": "user", "content": req.message})
        conn.execute(
            "INSERT INTO messages (student_id, subject_id, role, content, created_at) VALUES (?, ?, ?, ?, ?)",
            (req.student_id, req.subject_id, "user", req.message, datetime.utcnow().isoformat()),
        )
        conn.commit()
    else:
        # auto_open: AI speaks first — inject a silent trigger so Claude generates an opening
        history.append({"role": "user", "content": "[session_open]"})
    new_streak = update_streak(req.student_id, conn)

    rag_context    = retrieve_context(req.subject_id, req.message, conn)
    using_rag      = bool(rag_context)
    session_memory = get_session_memory(req.student_id, req.subject_id, conn)

    # Detect unanswered question: on auto_open, if last saved message is from the AI and
    # ends with a bold question (**...**?), the student left before answering.
    pending_question = ""
    if req.auto_open and history_rows:
        last_row = history_rows[-1]
        if last_row["role"] == "assistant":
            import re as _re
            # Bold question: **text?** possibly followed by whitespace
            bq = _re.search(r'\*\*([^*]+\?[^*]*)\*\*\s*$', last_row["content"].strip())
            if bq:
                pending_question = bq.group(1).strip()

    notes_rows = conn.execute(
        "SELECT concept_id, notes FROM concept_notes WHERE subject_id = ?",
        (req.subject_id,)
    ).fetchall()
    concept_notes_map = {r["concept_id"]: r["notes"] for r in notes_rows if r["notes"].strip()}

    # Load learner archetype for personalised tutor instruction
    archetype_row = conn.execute("SELECT archetype FROM archetype_scores WHERE student_id = ?", (req.student_id,)).fetchone()
    archetype_key = archetype_row["archetype"] if archetype_row else None
    archetype_instruction = ""
    if archetype_key and archetype_key in ARCHETYPES:
        archetype_instruction = "\n\n## Learner Archetype\n" + ARCHETYPES[archetype_key]["tutor_instruction"]

    if req.quiz_mode:
        system = build_quiz_prompt(subject, student["name"], covered_ids, mastered_ids, career)
    elif req.recall_warmup:
        base_system = build_system_prompt(subject, student["name"], is_first_visit, covered_ids, mastered_ids, rag_context, career, session_memory, concept_notes_map, pending_question)
        system = base_system + RECALL_WARMUP_SUFFIX.format(student_name=student["name"]) + archetype_instruction
    else:
        system = build_system_prompt(subject, student["name"], is_first_visit, covered_ids, mastered_ids, rag_context, career, session_memory, concept_notes_map, pending_question) + archetype_instruction

    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if api_key:
        import anthropic
        import asyncio
        import logging
        client = anthropic.Anthropic(api_key=api_key)
        raw_reply = None
        for attempt in range(4):
            try:
                response = await asyncio.get_running_loop().run_in_executor(
                    None,
                    lambda: client.messages.create(
                        model="claude-haiku-4-5-20251001", max_tokens=1024,
                        system=[{"type": "text", "text": system, "cache_control": {"type": "ephemeral"}}],
                        messages=history,
                    )
                )
                raw_reply = response.content[0].text
                _anthropic_status["ok"] = True
                _anthropic_status["error"] = None
                _anthropic_status["checked_at"] = datetime.utcnow().isoformat()
                break
            except anthropic.RateLimitError as e:
                logging.warning(f"Anthropic RateLimitError attempt {attempt}: {e}")
                _anthropic_status.update({"ok": False, "error": f"Rate limit (429): {e}", "checked_at": datetime.utcnow().isoformat()})
                if attempt < 3:
                    await asyncio.sleep(2 ** attempt)
                else:
                    raw_reply = "I'm a little busy right now — too many learners at once! Please send your message again in a few seconds and I'll be right with you."
            except anthropic.APIStatusError as e:
                logging.warning(f"Anthropic APIStatusError attempt {attempt}: status={e.status_code} {e}")
                _anthropic_status.update({"ok": False, "error": f"API error ({e.status_code}): {e.message}", "checked_at": datetime.utcnow().isoformat()})
                if e.status_code in (529, 503, 500) and attempt < 3:
                    await asyncio.sleep(2 ** attempt)
                else:
                    raw_reply = "I'm a little busy right now — please try again in a moment."
                    break
            except Exception as e:
                logging.error(f"Anthropic unexpected error attempt {attempt}: {type(e).__name__}: {e}")
                _anthropic_status.update({"ok": False, "error": f"{type(e).__name__}: {e}", "checked_at": datetime.utcnow().isoformat()})
                raw_reply = "I ran into a hiccup. Please try sending your message again."
                break
        if raw_reply is None:
            raw_reply = "I'm a little busy right now — please try again in a moment."
    else:
        raw_reply = MOCK_RESPONSES.get(req.subject_id, "Mock response.")

    newly_covered = []
    concepts_tag  = re.search(r'<<<CONCEPTS:([^>]*)>>>', raw_reply)
    reply         = re.sub(r'\n?<<<CONCEPTS:[^>]*>>>', '', raw_reply).strip()
    if concepts_tag and api_key:
        valid_ids  = {c["id"] for c in CURRICULUM[req.subject_id]}
        parsed     = [x.strip() for x in concepts_tag.group(1).split(",") if x.strip() in valid_ids]
        newly_covered = [cid for cid in parsed if cid not in covered_ids]
        for cid in newly_covered:
            conn.execute(
                "INSERT OR IGNORE INTO concept_progress (student_id, subject_id, concept_id, first_covered_at) VALUES (?, ?, ?, ?)",
                (req.student_id, req.subject_id, cid, datetime.utcnow().isoformat()),
            )
        conn.commit()

    newly_mastered = []
    mastered_tag   = re.search(r'<<<MASTERED:([^>]*)>>>', reply)
    reply          = re.sub(r'\n?<<<MASTERED:[^>]*>>>', '', reply).strip()
    if mastered_tag and api_key:
        valid_ids      = {c["id"] for c in CURRICULUM[req.subject_id]}
        parsed         = [x.strip() for x in mastered_tag.group(1).split(",") if x.strip() in valid_ids]
        newly_mastered = [cid for cid in parsed if cid not in mastered_ids]
        for cid in newly_mastered:
            conn.execute(
                "UPDATE concept_progress SET mastered_at = ? WHERE student_id = ? AND subject_id = ? AND concept_id = ?",
                (datetime.utcnow().isoformat(), req.student_id, req.subject_id, cid),
            )
        conn.commit()

    career_detected = None
    career_tag = re.search(r'<<<CAREER:([^>]*)>>>', reply)
    reply = re.sub(r'\n?<<<CAREER:[^>]*>>>', '', reply).strip()
    if career_tag and api_key and career is None:
        detected_id = career_tag.group(1).strip()
        if detected_id in CAREERS:
            career_detected = detected_id
            conn.execute(
                "INSERT INTO student_profile (student_id, career_id, career_goal_raw, updated_at) VALUES (?, ?, ?, ?) ON CONFLICT(student_id) DO UPDATE SET career_id=excluded.career_id, career_goal_raw=excluded.career_goal_raw, updated_at=excluded.updated_at",
                (req.student_id, detected_id, req.message, datetime.utcnow().isoformat()),
            )
            conn.commit()

    conn.execute(
        "INSERT INTO messages (student_id, subject_id, role, content, created_at) VALUES (?, ?, ?, ?, ?)",
        (req.student_id, req.subject_id, "assistant", reply, datetime.utcnow().isoformat()),
    )
    conn.commit()

    # Auto-mark subject completed when all concepts are covered
    subject_completed = None
    if newly_covered:
        all_covered = set(covered_ids) | set(newly_covered)
        curriculum_ids = {c["id"] for c in CURRICULUM[req.subject_id]}
        if curriculum_ids.issubset(all_covered):
            conn.execute(
                "UPDATE subject_status SET status = 'completed', completed_at = ? WHERE student_id = ? AND subject_id = ? AND status = 'active'",
                (datetime.utcnow().isoformat(), req.student_id, req.subject_id),
            )
            existing_cert = conn.execute(
                "SELECT credential_id FROM subject_completions WHERE student_id = ? AND subject_id = ?",
                (req.student_id, req.subject_id)
            ).fetchone()
            if not existing_cert:
                credential_id = str(uuid.uuid4()).replace('-', '')[:12].upper()
                conn.execute(
                    "INSERT INTO subject_completions (student_id, subject_id, completed_at, credential_id) VALUES (?,?,?,?)",
                    (req.student_id, req.subject_id, datetime.utcnow().isoformat(), credential_id)
                )
                subject_completed = {"credential_id": credential_id}
                # Fire completion email (get student email)
                student_row = conn.execute("SELECT email, name FROM students WHERE id = ?", (req.student_id,)).fetchone()
                if student_row:
                    background_tasks.add_task(
                        send_completion_email,
                        student_row["email"], student_row["name"],
                        SUBJECTS[req.subject_id]["name"], credential_id
                    )
            conn.commit()

    # Detect newly completed modules (all 3 sub-concepts a/b/c now covered)
    modules_completed = []
    if newly_covered:
        all_covered_now = set(covered_ids) | set(newly_covered)
        touched_modules = {cid.replace(r'_[abc]$', cid[cid.rfind('_'):]).rsplit('_', 1)[0]
                           for cid in newly_covered if re.search(r'_[abc]$', cid)}
        for mod_id in touched_modules:
            subs = [c for c in CURRICULUM[req.subject_id] if re.match(rf'^{re.escape(mod_id)}_[abc]$', c["id"])]
            if subs and all(c["id"] in all_covered_now for c in subs):
                already_done = conn.execute(
                    "SELECT 1 FROM module_quizzes WHERE student_id=? AND subject_id=? AND module_id=?",
                    (req.student_id, req.subject_id, mod_id)
                ).fetchone()
                if not already_done:
                    mod_name = subs[0]["name"].split(":")[0].split(" - ")[0].strip()
                    modules_completed.append({"id": mod_id, "name": mod_name})
                    student_row2 = conn.execute("SELECT email, name FROM students WHERE id=?", (req.student_id,)).fetchone()
                    if student_row2:
                        background_tasks.add_task(
                            send_module_complete_email,
                            student_row2["email"], student_row2["name"],
                            mod_name, SUBJECTS[req.subject_id]["name"], career
                        )

    # Fire streak milestone email for 3, 7, 14, 30 day milestones
    if new_streak in (3, 7, 14, 30):
        student_row3 = conn.execute("SELECT email, name FROM students WHERE id=?", (req.student_id,)).fetchone()
        if student_row3:
            background_tasks.add_task(
                send_streak_milestone_email,
                student_row3["email"], student_row3["name"], new_streak, career
            )

    # Score archetype at 5 messages then every 10 after that
    user_msg_count = conn.execute(
        "SELECT COUNT(*) FROM messages WHERE student_id = ? AND subject_id = ? AND role = 'user'",
        (req.student_id, req.subject_id)
    ).fetchone()[0]
    if user_msg_count == 5 or (user_msg_count > 5 and user_msg_count % 10 == 0):
        background_tasks.add_task(score_session_archetype, req.student_id, req.subject_id)

    conn.close()

    curr     = effective_curriculum(req.subject_id, career)
    curr_ids = {c["id"] for c in curr}
    total    = len(curr)
    eff_prev_covered  = sum(1 for cid in covered_ids  if cid in curr_ids)
    eff_prev_mastered = sum(1 for cid in mastered_ids if cid in curr_ids)
    updated_covered   = eff_prev_covered  + len([cid for cid in newly_covered  if cid in curr_ids])
    updated_mastered  = eff_prev_mastered + len([cid for cid in newly_mastered if cid in curr_ids])
    capstone_now_unlocked = False
    if req.subject_id in CAPSTONES:
        threshold = min(total, CAPSTONES[req.subject_id]["unlock_threshold"])
        capstone_now_unlocked = updated_covered >= threshold and eff_prev_covered < threshold

    return {
        "reply":                   reply,
        "tutor_name":              subject["tutor_name"],
        "mock":                    api_key is None,
        "is_first_visit":          is_first_visit,
        "using_rag":               using_rag,
        "concepts_covered":        updated_covered,
        "concepts_mastered":       updated_mastered,
        "concepts_total":          total,
        "newly_covered":           newly_covered,
        "newly_mastered":          newly_mastered,
        "career_detected":         career_detected,
        "capstone_now_unlocked":   capstone_now_unlocked,
        "modules_completed":       modules_completed,
        "subject_completed":       subject_completed,
    }


# ── Platform Feedback ─────────────────────────────────────────────────────────

class FeedbackRequest(BaseModel):
    student_id: str
    q1: str = ""
    q2: str = ""
    q3: str = ""
    rating: int
    comment: str = ""

@app.post("/feedback")
def submit_feedback(req: FeedbackRequest):
    if req.rating < 1 or req.rating > 5:
        raise HTTPException(status_code=400, detail="Rating must be 1-5")
    conn = get_db()
    existing = conn.execute("SELECT id FROM platform_feedback WHERE student_id = ?", (req.student_id,)).fetchone()
    if existing:
        conn.close()
        return {"ok": True}
    conn.execute(
        "INSERT INTO platform_feedback (id, student_id, q1, q2, q3, rating, comment, submitted_at) VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
        (str(uuid.uuid4()), req.student_id, req.q1.strip(), req.q2.strip(), req.q3.strip(), req.rating, req.comment.strip(), datetime.utcnow().isoformat()),
    )
    conn.commit()
    conn.close()
    return {"ok": True}

@app.get("/admin/feedback")
def admin_get_feedback(x_admin_key: str = Header(None)):
    require_admin(x_admin_key)
    conn = get_db()
    rows = conn.execute("""
        SELECT f.id, f.student_id, s.name, s.email, f.q1, f.q2, f.q3, f.rating, f.comment, f.submitted_at
        FROM platform_feedback f
        LEFT JOIN students s ON s.id = f.student_id
        ORDER BY f.submitted_at DESC
    """).fetchall()
    conn.close()
    return [dict(r) for r in rows]


# ── Waitlist / Access Requests ────────────────────────────────────────────────

class AccessWaitlistRequest(BaseModel):
    name: str
    email: str
    phone: str = ""
    university: str = ""
    year_of_study: str = ""
    country: str = ""
    reason: str = ""
    product: str = "career_pathways"

async def send_waitlist_confirmation_email(to_email: str, name: str, position: int) -> bool:
    first = name.split()[0]
    body = (
        _heading(f"You're on the Bversity waitlist, {first}!") +
        _para(f"Thank you for your interest in Bversity  -  the world's first AI-Native Biotech University. "
              f"We've received your application and you are <strong>#{position} on the waitlist</strong>.") +
        _divider() +
        _para("<strong>What happens next?</strong><br>"
              "Our team reviews applications every week. When your spot opens up, you'll receive an email "
              "with a link to activate your account and start learning immediately.") +
        _para("In the meantime, explore what you'll be learning at "
              "<a href='https://university.bversity.io' style='color:#00A896'>university.bversity.io</a>.") +
        _divider() +
        _small("Questions? Reply to this email  -  we read every one.")
    )
    return await _send_email(to_email, f"You're #{position} on the Bversity waitlist!", _email_wrap(body))

async def send_access_granted_email(to_email: str, name: str) -> bool:
    first = name.split()[0]
    body = (
        _heading(f"Your Bversity access is approved, {first}!") +
        _para("Great news  -  we've reviewed your application and you've been granted access to Bversity. "
              "You can now log in and start learning from the world's first AI-Native Biotech University.") +
        _btn("Access Bversity Now →", "https://university.bversity.io") +
        _divider() +
        _para("<strong>How to get started:</strong><br>"
              "1. Click the button above to open the platform.<br>"
              "2. Enter your email address to receive a one-time verification code.<br>"
              "3. Enter the code to log in  -  no password needed.<br>"
              "4. Pick a subject and start your first AI tutoring session.") +
        _divider() +
        _small("Your email address is already approved. If you have any trouble logging in, reply to this email.")
    )
    return await _send_email(to_email, f"Your Bversity access is approved, {first}! 🎉", _email_wrap(body))

async def send_admin_access_request_notification(name: str, email: str, product: str, university: str, country: str, reason: str):
    admin_email = os.environ.get("ADMIN_NOTIFY_EMAIL", "sudharsan@bversity.io")
    product_label = "Certifications" if product == "certifications" else "Career Pathways"
    body = (
        _heading(f"New access request — {product_label}") +
        _para(f"<strong>{name}</strong> ({email}) has requested access to <strong>{product_label}</strong>.") +
        (f"<p style='margin:0 0 12px;font-size:14px;'><strong>University/Org:</strong> {university}</p>" if university else "") +
        (f"<p style='margin:0 0 12px;font-size:14px;'><strong>Country:</strong> {country}</p>" if country else "") +
        (f"<p style='margin:0 0 12px;font-size:14px;'><strong>Reason:</strong> {reason}</p>" if reason else "") +
        _btn("Review in Admin Panel →", "https://university.bversity.io/#admin") +
        _small("You are receiving this because you are the Bversity admin.")
    )
    await _send_email(admin_email, f"New access request: {name} ({product_label})", _email_wrap(body))

@app.post("/request-access")
async def request_access(req: AccessWaitlistRequest, background_tasks: BackgroundTasks):
    if not req.name.strip() or not req.email.strip():
        raise HTTPException(status_code=400, detail="Name and email are required")
    product = req.product if req.product in ("career_pathways", "certifications") else "career_pathways"
    conn = get_db()
    existing = conn.execute("SELECT id FROM access_requests WHERE email = ? AND product = ?", (req.email.strip().lower(), product)).fetchone()
    if existing:
        count = conn.execute("SELECT COUNT(*) as n FROM access_requests").fetchone()["n"]
        conn.close()
        return {"ok": True, "already_submitted": True, "position": count}
    request_id = str(uuid.uuid4())
    conn.execute(
        "INSERT INTO access_requests (id, name, email, phone, university, year_of_study, country, reason, status, submitted_at, product) VALUES (?, ?, ?, ?, ?, ?, ?, ?, 'pending', ?, ?)",
        (request_id, req.name.strip(), req.email.strip().lower(), req.phone.strip(), req.university.strip(), req.year_of_study.strip(), req.country.strip(), req.reason.strip(), datetime.utcnow().isoformat(), product),
    )
    conn.commit()
    count = conn.execute("SELECT COUNT(*) as n FROM access_requests").fetchone()["n"]
    conn.close()
    background_tasks.add_task(send_waitlist_confirmation_email, req.email.strip().lower(), req.name.strip(), count)
    background_tasks.add_task(send_admin_access_request_notification, req.name.strip(), req.email.strip().lower(), product, req.university.strip(), req.country.strip(), req.reason.strip())
    return {"ok": True, "already_submitted": False, "position": count}

@app.get("/waitlist-count")
def waitlist_count():
    conn = get_db()
    count = conn.execute("SELECT COUNT(*) as n FROM access_requests").fetchone()["n"]
    conn.close()
    return {"count": count}

# ── Learner Archetype Scoring ─────────────────────────────────────────────────

ARCHETYPES = {
    "challenger": {
        "label": "The Challenger",
        "color": "#6366f1",
        "emoji": "⚔️",
        "description": "Pushes back, debates, needs friction to stay engaged",
        "tutor_instruction": (
            "This learner is a Challenger. They engage best with friction, not answers. "
            "Pose problems rather than explaining solutions. Introduce deliberate contrarian positions "
            "and let them argue their way to the truth. If they push back, engage seriously — "
            "never concede just to be agreeable. Ask 'what do you think is wrong with that reasoning?' "
            "before giving your own view."
        ),
    },
    "credential_hunter": {
        "label": "The Credential Hunter",
        "color": "#0066CC",
        "emoji": "🎯",
        "description": "Exam-focused, strategic, wants exactly what they need to pass",
        "tutor_instruction": (
            "This learner is a Credential Hunter. They respond to explicit exam relevance. "
            "Frame every concept with 'this is how the exam tests this' and 'this is the question "
            "that trips candidates up.' Make depth feel strategic, not academic. "
            "Use phrases like 'knowing this gives you an edge over candidates who only memorise the definition.'"
        ),
    },
    "scattered_genius": {
        "label": "The Scattered Genius",
        "color": "#f59e0b",
        "emoji": "🌪️",
        "description": "High curiosity, wide breadth, struggles to consolidate and finish",
        "tutor_instruction": (
            "This learner is a Scattered Genius. They're brilliant but scattered — use their own "
            "curiosity as a hook to keep them on track. When they jump topics, connect the tangent "
            "back: 'that's a great question — and it's exactly why we need to finish X first.' "
            "Create explicit 'breadcrumb' moments: 'remember that question you asked earlier? "
            "This is the answer.' Make consolidation feel like a discovery, not a chore."
        ),
    },
    "imposter": {
        "label": "The Imposter",
        "color": "#ec4899",
        "emoji": "🌱",
        "description": "High anxiety, avoids evaluation, needs safety to engage deeply",
        "tutor_instruction": (
            "This learner shows Imposter syndrome traits. Never evaluate — always explore together. "
            "Reframe every wrong answer as 'that's the most common misconception, here's why it feels right.' "
            "Avoid phrases like 'actually' or 'that's incorrect.' Use 'interesting — most people think that too, "
            "here's what's surprising.' Make confusion feel like evidence of intelligence, not failure. "
            "Never ask them to judge their own understanding — instead ask 'what's the murkiest part?'"
        ),
    },
    "professional_upgrader": {
        "label": "The Professional Upgrader",
        "color": "#10b981",
        "emoji": "🔧",
        "description": "Already in the field, time-constrained, wants peer-level exchange",
        "tutor_instruction": (
            "This learner is a Professional Upgrader — already working in the industry. "
            "Treat them as a peer, not a student. Lead with application, skip fundamentals they know. "
            "Reference their professional context explicitly: 'given your background, the part that's "
            "actually new here is...' Ask them what they're seeing in their current role — "
            "their real-world experience is a resource, not a distraction. Keep sessions tight and dense."
        ),
    },
    "ghost": {
        "label": "The Ghost",
        "color": "#94a3b8",
        "emoji": "👻",
        "description": "High initial intent, sporadic sessions, needs re-entry friction removed",
        "tutor_instruction": (
            "This learner has Ghost tendencies — they disappear and feel like they're starting over. "
            "Begin every session with explicit acknowledgment of their progress: 'last time we covered X, "
            "and you were close to understanding Y.' Make re-entry feel frictionless, not like a failure. "
            "Set small session goals explicitly: 'today let's just do one thing.' "
            "Celebrate showing up as much as progress made."
        ),
    },
}

def _derive_archetype(scores: dict, session_regularity: float) -> str:
    qd  = scores.get("question_depth", 0)
    ef  = scores.get("exam_focus", 0)
    am  = scores.get("anxiety_markers", 0)
    tj  = scores.get("topic_jumping", 0)
    cs  = scores.get("challenge_seeking", 0)
    pa  = scores.get("practical_anchoring", 0)
    mp  = scores.get("mastery_push", 0)

    archetype_scores = {
        "challenger":          cs * 1.5 + qd * 0.8 + (3 - am) * 0.5,
        "credential_hunter":   ef * 1.5 + mp * 0.8 + (3 - tj) * 0.5,
        "scattered_genius":    tj * 1.5 + qd * 0.8 + (3 - mp) * 0.5,
        "imposter":            am * 1.5 + (3 - cs) * 0.8 + (3 - qd) * 0.3,
        "professional_upgrader": pa * 1.5 + qd * 0.7 + (3 - ef) * 0.5,
        "ghost":               (3 - session_regularity) * 2.0,
    }
    return max(archetype_scores, key=archetype_scores.get)

_SCORING_PROMPT = """You are analysing a tutoring session to understand how this learner approaches learning.

Session transcript (student messages only shown):
{transcript}

Score this learner on each dimension from 0 to 3:
- 0 = not observed / opposite
- 1 = slight signal
- 2 = clear signal
- 3 = strong signal

Dimensions:
1. question_depth: Are questions surface-level ("what is X?") or synthesis-level ("why does X lead to Y in context Z?")?
2. exam_focus: Do they reference exams, certifications, marks, or "what do I need to know"?
3. anxiety_markers: Hedging ("I might be wrong"), apologising, asking for reassurance, avoiding being evaluated?
4. topic_jumping: Do they jump between topics within a session rather than going deep on one?
5. challenge_seeking: Do they push back on explanations, ask counter-questions, debate the tutor?
6. practical_anchoring: Do they connect concepts to their real job, current work, or practical scenarios?
7. mastery_push: Do they ask follow-up questions until they truly understand, or accept first answers?

Respond ONLY with valid JSON:
{{"question_depth": 0-3, "exam_focus": 0-3, "anxiety_markers": 0-3, "topic_jumping": 0-3, "challenge_seeking": 0-3, "practical_anchoring": 0-3, "mastery_push": 0-3, "reasoning": "1-2 sentence observation about this learner"}}"""

async def score_session_archetype(student_id: str, subject_id: str):
    try:
        conn = get_db()
        msgs = conn.execute("""
            SELECT role, content, created_at FROM messages
            WHERE student_id = ? AND subject_id = ?
            ORDER BY created_at DESC LIMIT 40
        """, (student_id, subject_id)).fetchall()

        user_msgs = [m for m in msgs if m["role"] == "user"]
        if len(user_msgs) < 3:
            conn.close()
            return

        transcript = "\n".join(f"Student: {m['content'][:300]}" for m in user_msgs[-15:])
        prompt = _SCORING_PROMPT.format(transcript=transcript)

        api_key = os.environ.get("ANTHROPIC_API_KEY")
        if not api_key:
            conn.close()
            return

        client = anthropic.Anthropic(api_key=api_key)
        response = client.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=300,
            messages=[{"role": "user", "content": prompt}],
        )
        raw = response.content[0].text.strip()
        data = json.loads(raw)

        # compute session regularity from message gaps
        all_dates = conn.execute("""
            SELECT DISTINCT DATE(created_at) FROM messages
            WHERE student_id = ? AND role = 'user'
            ORDER BY created_at
        """, (student_id,)).fetchall()
        date_list = [r[0] for r in all_dates]
        if len(date_list) >= 2:
            from datetime import date as dt_date
            gaps = [(datetime.strptime(date_list[i+1], "%Y-%m-%d") - datetime.strptime(date_list[i], "%Y-%m-%d")).days
                    for i in range(len(date_list)-1)]
            avg_gap = sum(gaps) / len(gaps)
            session_regularity = max(0, min(3, 3 - (avg_gap - 1) * 0.3))
        else:
            session_regularity = 0.0

        # weighted rolling average with existing scores
        existing = conn.execute("SELECT * FROM archetype_scores WHERE student_id = ?", (student_id,)).fetchone()
        n = existing["sessions_scored"] if existing else 0
        weight_new = 1.0 / (n + 1)
        weight_old = n / (n + 1) if n > 0 else 0

        dims = ["question_depth","exam_focus","anxiety_markers","topic_jumping","challenge_seeking","practical_anchoring","mastery_push"]
        new_scores = {}
        for d in dims:
            old_val = existing[d] if existing else 0
            new_val = float(data.get(d, 0))
            new_scores[d] = round(old_val * weight_old + new_val * weight_new, 2)

        archetype = _derive_archetype(new_scores, session_regularity)
        reasoning = data.get("reasoning", "")
        now = datetime.utcnow().isoformat()

        conn.execute("""
            INSERT INTO archetype_scores
                (student_id, question_depth, exam_focus, anxiety_markers, topic_jumping,
                 challenge_seeking, practical_anchoring, mastery_push, sessions_scored,
                 archetype, archetype_reasoning, updated_at)
            VALUES (?,?,?,?,?,?,?,?,?,?,?,?)
            ON CONFLICT(student_id) DO UPDATE SET
                question_depth=excluded.question_depth,
                exam_focus=excluded.exam_focus,
                anxiety_markers=excluded.anxiety_markers,
                topic_jumping=excluded.topic_jumping,
                challenge_seeking=excluded.challenge_seeking,
                practical_anchoring=excluded.practical_anchoring,
                mastery_push=excluded.mastery_push,
                sessions_scored=excluded.sessions_scored,
                archetype=excluded.archetype,
                archetype_reasoning=excluded.archetype_reasoning,
                updated_at=excluded.updated_at
        """, (student_id, new_scores["question_depth"], new_scores["exam_focus"],
              new_scores["anxiety_markers"], new_scores["topic_jumping"],
              new_scores["challenge_seeking"], new_scores["practical_anchoring"],
              new_scores["mastery_push"], n + 1, archetype, reasoning, now))
        conn.execute("UPDATE student_profile SET learner_archetype = ? WHERE student_id = ?", (archetype, student_id))
        conn.commit()
        conn.close()
    except Exception as e:
        try: conn.close()
        except: pass

# ── Session Memory ─────────────────────────────────────────────────────────────

def _build_summary_prompt(tutor_name: str, student_name: str, subject_name: str, messages: list) -> str:
    transcript = "\n".join(
        f"{'Student' if m['role'] == 'user' else tutor_name}: {m['content'][:600]}"
        for m in messages
    )
    return f"""You are a tutor assistant writing private session notes for {tutor_name} after a tutoring session with {student_name} in {subject_name}.

Here is the session transcript:
{transcript}

Write concise tutor notes in exactly this format (keep each field to 1-3 sentences):

CONCEPTS_COVERED: [which concepts were substantively taught this session]
CONFIDENCE_SIGNALS: [where the student showed clear understanding  -  specific moments]
CONFUSION_SIGNALS: [where the student hesitated, gave vague answers, or asked clarifying questions]
NOTABLE_MOMENTS: [interesting questions asked, background info revealed, breakthroughs]
FOLLOW_UP: [what to address, revisit, or build on in the next session  -  be specific]

Be honest, specific, and brief. These notes directly shape the next session."""

async def _generate_session_summary(student_id: str, subject_id: str):
    if subject_id not in SUBJECTS:
        return
    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        return
    conn = get_db()
    try:
        # Find messages since the last summary
        last = conn.execute(
            "SELECT created_at FROM session_summaries WHERE student_id = ? AND subject_id = ? ORDER BY id DESC LIMIT 1",
            (student_id, subject_id)
        ).fetchone()
        if last:
            rows = conn.execute(
                "SELECT role, content FROM messages WHERE student_id = ? AND subject_id = ? AND created_at > ? ORDER BY id",
                (student_id, subject_id, last["created_at"])
            ).fetchall()
        else:
            rows = conn.execute(
                "SELECT role, content FROM messages WHERE student_id = ? AND subject_id = ? ORDER BY id LIMIT 60",
                (student_id, subject_id)
            ).fetchall()

        if len(rows) < 4:  # not enough to summarise
            conn.close()
            return

        student_row = conn.execute("SELECT name FROM students WHERE id = ?", (student_id,)).fetchone()
        student_name = student_row["name"] if student_row else "the student"
        subject = SUBJECTS[subject_id]
        messages = [{"role": r["role"], "content": r["content"]} for r in rows]
        prompt = _build_summary_prompt(subject["tutor"], student_name, subject["name"], messages)

        import anthropic
        client = anthropic.Anthropic(api_key=api_key)
        response = client.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=400,
            messages=[{"role": "user", "content": prompt}],
        )
        summary = response.content[0].text.strip()
        conn.execute(
            "INSERT INTO session_summaries (student_id, subject_id, summary, message_count, created_at) VALUES (?, ?, ?, ?, ?)",
            (student_id, subject_id, summary, len(rows), datetime.utcnow().isoformat())
        )
        conn.commit()
    except Exception:
        pass
    finally:
        conn.close()

@app.post("/admin/students/{student_id}/score-archetype")
async def admin_score_archetype(student_id: str, background_tasks: BackgroundTasks, x_admin_key: str = Header(None)):
    require_admin(x_admin_key)
    conn = get_db()
    subjects = conn.execute(
        "SELECT DISTINCT subject_id FROM messages WHERE student_id = ? AND role = 'user'", (student_id,)
    ).fetchall()
    conn.close()
    for s in subjects:
        background_tasks.add_task(score_session_archetype, student_id, s["subject_id"])
    return {"ok": True, "subjects_queued": len(subjects)}

@app.get("/session-summary/{student_id}/{subject_id}")
def get_session_summary(student_id: str, subject_id: str):
    conn = get_db()
    row = conn.execute(
        "SELECT summary, message_count, created_at FROM session_summaries WHERE student_id = ? AND subject_id = ? ORDER BY id DESC LIMIT 1",
        (student_id, subject_id)
    ).fetchone()
    conn.close()
    if not row:
        return {"summary": None}
    return {"summary": row["summary"], "message_count": row["message_count"], "created_at": row["created_at"]}

@app.post("/session-end/{student_id}/{subject_id}")
async def session_end(student_id: str, subject_id: str, background_tasks: BackgroundTasks):
    background_tasks.add_task(_generate_session_summary, student_id, subject_id)
    background_tasks.add_task(score_session_archetype, student_id, subject_id)
    return {"ok": True}


@app.get("/reentry/{student_id}")
def get_reentry(student_id: str):
    conn = get_db()
    now = datetime.utcnow()
    today = now.date().isoformat()

    profile = conn.execute(
        "SELECT last_active_at, streak_count, streak_last_date, career_id FROM student_profile WHERE student_id = ?",
        (student_id,)
    ).fetchone()

    # Days away
    days_away = 0
    if profile and profile["last_active_at"]:
        try:
            last = datetime.fromisoformat(profile["last_active_at"])
            days_away = (now - last).days
        except Exception:
            pass

    # Last session summary (most recent across all subjects)
    last_sess = conn.execute(
        "SELECT subject_id, summary, created_at FROM session_summaries WHERE student_id = ? ORDER BY id DESC LIMIT 1",
        (student_id,)
    ).fetchone()
    last_session = None
    if last_sess:
        subj = SUBJECTS.get(last_sess["subject_id"])
        last_session = {
            "subject_id": last_sess["subject_id"],
            "subject_name": subj["name"] if subj else last_sess["subject_id"],
            "subject_color": subj["color"] if subj else "#00A896",
            "summary": last_sess["summary"],
            "created_at": last_sess["created_at"],
        }

    # Study plan: find most important next concept
    covered_rows = conn.execute(
        "SELECT subject_id, concept_id FROM concept_progress WHERE student_id = ?", (student_id,)
    ).fetchall()
    covered = {(r["subject_id"], r["concept_id"]) for r in covered_rows}

    plan_rows = conn.execute(
        "SELECT day_number, subject_id, concept_id, target_date FROM study_plan WHERE student_id = ? ORDER BY day_number",
        (student_id,)
    ).fetchall()

    overdue = [r for r in plan_rows if r["target_date"] <= today and (r["subject_id"], r["concept_id"]) not in covered]
    upcoming = [r for r in plan_rows if r["target_date"] > today and (r["subject_id"], r["concept_id"]) not in covered]

    next_concept = None
    source = overdue[0] if overdue else (upcoming[0] if upcoming else None)
    if source:
        subj = SUBJECTS.get(source["subject_id"])
        next_concept = {
            "concept_id": source["concept_id"],
            "concept_name": source["concept_id"].replace("_", " ").title(),
            "subject_id": source["subject_id"],
            "subject_name": subj["name"] if subj else source["subject_id"],
            "subject_color": subj["color"] if subj else "#00A896",
            "target_date": source["target_date"],
            "is_overdue": bool(overdue),
        }

    # Streak
    streak_count = profile["streak_count"] if profile else 0
    streak_last  = profile["streak_last_date"] if profile else None
    streak_today = (streak_last == today) if streak_last else False
    streak_broken = (days_away > 1 and streak_count > 0)
    streak_at_risk = (days_away == 1 and streak_count > 0 and not streak_today)

    career = CAREERS.get(profile["career_id"]) if profile and profile["career_id"] else None

    conn.close()
    return {
        "days_away": days_away,
        "last_session": last_session,
        "next_concept": next_concept,
        "overdue_count": len(overdue),
        "streak_count": streak_count,
        "streak_today": streak_today,
        "streak_at_risk": streak_at_risk,
        "streak_broken": streak_broken,
        "career_title": career["title"] if career else None,
    }

def get_session_memory(student_id: str, subject_id: str, conn) -> str:
    rows = conn.execute(
        "SELECT summary FROM session_summaries WHERE student_id = ? AND subject_id = ? ORDER BY id DESC LIMIT 2",
        (student_id, subject_id)
    ).fetchall()
    if not rows:
        return ""
    summaries = list(reversed([r["summary"] for r in rows]))
    parts = []
    for i, s in enumerate(summaries, 1):
        parts.append(f"[Session {i} notes]\n{s}")
    return "\n\n".join(parts)

# ── Subject News ───────────────────────────────────────────────────────────────

_NEWS_CACHE: dict = {}
_NEWS_CACHE_TTL = 3 * 3600  # 3 hours

SUBJECT_NEWS_QUERIES = {
    "bioinformatics":     "bioinformatics computational biology genomics pipeline",
    "genomics":           "genomics DNA sequencing genome precision medicine",
    "drug_discovery":     "drug discovery pharmaceutical development FDA",
    "clinical_trials":    "clinical trial FDA drug approval regulatory",
    "genai_ml":           "AI machine learning drug discovery biotech",
    "biotech_business":   "biotech pharma deal acquisition merger IPO",
    "cell_gene_therapy":  "gene therapy cell therapy CRISPR CAR-T",
    "protein_engineering":"protein engineering antibody biologics design",
    "rna_therapeutics":   "mRNA RNA therapeutics vaccine LNP",
    "biomanufacturing":   "biomanufacturing bioprocessing GMP pharmaceutical",
    "longevity_science":  "longevity aging science senolytics epigenetics",
}

def _clean_news_title(title: str) -> str:
    # Google News appends " - Source Name"  -  strip it
    parts = title.rsplit(" - ", 1)
    return parts[0].strip() if len(parts) > 1 else title.strip()

def _parse_pub_date(pub_date: str) -> str:
    try:
        dt = datetime.strptime(pub_date, "%a, %d %b %Y %H:%M:%S %Z")
        delta = datetime.utcnow() - dt
        if delta.days == 0:
            return "Today"
        elif delta.days == 1:
            return "Yesterday"
        else:
            return f"{delta.days} days ago"
    except Exception:
        return ""

@app.get("/subject-news/{subject_id}")
def get_subject_news(subject_id: str):
    now = time.time()
    cached = _NEWS_CACHE.get(subject_id)
    if cached and now - cached["at"] < _NEWS_CACHE_TTL:
        return cached["items"]

    query = SUBJECT_NEWS_QUERIES.get(subject_id, "biotechnology life sciences")
    url = f"https://news.google.com/rss/search?q={urllib.parse.quote(query)}&hl=en-US&gl=US&ceid=US:en"
    try:
        req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
        with urllib.request.urlopen(req, timeout=6) as resp:
            xml_bytes = resp.read()
        root = ET.fromstring(xml_bytes)
        items = []
        for item in root.findall(".//item")[:5]:
            raw_title = item.findtext("title", "")
            source_el = item.find("source")
            source    = source_el.text if source_el is not None else ""
            pub_date  = item.findtext("pubDate", "")
            link      = item.findtext("link", "")
            clean     = _clean_news_title(raw_title)
            if len(clean) < 20:
                continue
            items.append({
                "title":    clean,
                "source":   source,
                "pub_date": _parse_pub_date(pub_date),
                "link":     link,
            })
        _NEWS_CACHE[subject_id] = {"items": items, "at": now}
        return items
    except Exception:
        return []

_WIKI_CACHE: dict = {}
_WIKI_CACHE_TTL = 86400  # 24 hours

@app.get("/term-image/{term}")
def get_term_image(term: str):
    key = term.lower().strip()
    now = time.time()
    cached = _WIKI_CACHE.get(key)
    if cached and now - cached["at"] < _WIKI_CACHE_TTL:
        return cached["data"]

    encoded = urllib.parse.quote(term, safe="")
    url = f"https://en.wikipedia.org/api/rest_v1/page/summary/{encoded}"
    try:
        req = urllib.request.Request(url, headers={"User-Agent": "Bversity/1.0 (university.bversity.io)"})
        with urllib.request.urlopen(req, timeout=5) as resp:
            payload = json.loads(resp.read())
        image_url = (payload.get("thumbnail") or {}).get("source")
        extract = payload.get("extract", "")[:300]
        data = {"image": image_url, "extract": extract, "found": bool(image_url)}
    except Exception:
        data = {"image": None, "extract": "", "found": False}

    _WIKI_CACHE[key] = {"data": data, "at": now}
    return data

@app.get("/admin/access-requests")
def admin_get_access_requests(x_admin_key: str = Header(None)):
    require_admin(x_admin_key)
    conn = get_db()
    rows = conn.execute(
        "SELECT * FROM access_requests ORDER BY submitted_at DESC"
    ).fetchall()
    conn.close()
    return [dict(r) for r in rows]

@app.post("/admin/approve-request/{request_id}")
async def admin_approve_request(request_id: str, background_tasks: BackgroundTasks, x_admin_key: str = Header(None)):
    require_admin(x_admin_key)
    conn = get_db()
    req = conn.execute("SELECT * FROM access_requests WHERE id = ?", (request_id,)).fetchone()
    if not req:
        conn.close()
        raise HTTPException(status_code=404, detail="Request not found")
    product = req["product"] or "career_pathways"
    existing = conn.execute("SELECT email FROM approved_emails WHERE email = ? AND product = ?", (req["email"], product)).fetchone()
    if not existing:
        conn.execute(
            "INSERT INTO approved_emails (email, added_at, product, access_type) VALUES (?, ?, ?, 'full')",
            (req["email"], datetime.utcnow().isoformat(), product)
        )
    conn.execute("UPDATE access_requests SET status = 'approved' WHERE id = ?", (request_id,))
    conn.commit()
    name  = req["name"]
    email = req["email"]
    conn.close()
    background_tasks.add_task(send_access_granted_email, email, name)
    return {"ok": True}

@app.post("/admin/reject-request/{request_id}")
def admin_reject_request(request_id: str, x_admin_key: str = Header(None)):
    require_admin(x_admin_key)
    conn = get_db()
    conn.execute("UPDATE access_requests SET status = 'rejected' WHERE id = ?", (request_id,))
    conn.commit()
    conn.close()
    return {"ok": True}


# ── Admin Metrics ──────────────────────────────────────────────────────────────

@app.get("/admin/metrics")
def admin_metrics(x_admin_key: str = Header(None)):
    require_admin(x_admin_key)
    conn = get_db()
    now = datetime.utcnow()
    today_start     = now.replace(hour=0, minute=0, second=0, microsecond=0).isoformat()
    week_start      = (now - timedelta(days=7)).isoformat()
    month_start     = (now - timedelta(days=30)).isoformat()

    dau = conn.execute(
        "SELECT COUNT(DISTINCT student_id) FROM messages WHERE role='user' AND created_at >= ?", (today_start,)
    ).fetchone()[0]
    wau = conn.execute(
        "SELECT COUNT(DISTINCT student_id) FROM messages WHERE role='user' AND created_at >= ?", (week_start,)
    ).fetchone()[0]
    mau = conn.execute(
        "SELECT COUNT(DISTINCT student_id) FROM messages WHERE role='user' AND created_at >= ?", (month_start,)
    ).fetchone()[0]
    total_students = conn.execute("SELECT COUNT(*) FROM students").fetchone()[0]
    new_today      = conn.execute("SELECT COUNT(*) FROM students WHERE created_at >= ?", (today_start,)).fetchone()[0]
    new_this_week  = conn.execute("SELECT COUNT(*) FROM students WHERE created_at >= ?", (week_start,)).fetchone()[0]

    # 30-day DAU sparkline
    sparkline = []
    for i in range(29, -1, -1):
        day_start = (now - timedelta(days=i)).replace(hour=0, minute=0, second=0, microsecond=0)
        day_end   = day_start + timedelta(days=1)
        count = conn.execute(
            "SELECT COUNT(DISTINCT student_id) FROM messages WHERE role='user' AND created_at >= ? AND created_at < ?",
            (day_start.isoformat(), day_end.isoformat())
        ).fetchone()[0]
        signups = conn.execute(
            "SELECT COUNT(*) FROM students WHERE created_at >= ? AND created_at < ?",
            (day_start.isoformat(), day_end.isoformat())
        ).fetchone()[0]
        sparkline.append({"date": day_start.strftime("%b %d"), "dau": count, "signups": signups})

    # D1 / D7 / D30 retention
    def retention_rate(days_since_signup, activity_window_days):
        cutoff = (now - timedelta(days=days_since_signup)).isoformat()
        eligible = conn.execute(
            "SELECT id FROM students WHERE created_at <= ?", (cutoff,)
        ).fetchall()
        if not eligible:
            return 0
        returned = 0
        window_start = (now - timedelta(days=activity_window_days)).isoformat()
        for row in eligible:
            sid = row["id"]
            has_activity = conn.execute(
                "SELECT 1 FROM messages WHERE student_id=? AND role='user' AND created_at >= ? LIMIT 1",
                (sid, window_start)
            ).fetchone()
            if has_activity:
                returned += 1
        return round((returned / len(eligible)) * 100)

    d1_retention  = retention_rate(1, 1)
    d7_retention  = retention_rate(7, 7)
    d30_retention = retention_rate(30, 30)

    # Avg session length (messages per active day per student)
    session_rows = conn.execute("""
        SELECT student_id, date(created_at) AS day, COUNT(*) AS msg_count
        FROM messages WHERE role='user'
        GROUP BY student_id, day
    """).fetchall()
    if session_rows:
        avg_session_msgs = round(sum(r["msg_count"] for r in session_rows) / len(session_rows), 1)
    else:
        avg_session_msgs = 0

    # Depth score: mastered / covered across all students
    covered  = conn.execute("SELECT COUNT(*) FROM concept_progress").fetchone()[0]
    mastered = conn.execute("SELECT COUNT(*) FROM concept_progress WHERE mastered_at IS NOT NULL").fetchone()[0]
    depth_score = round((mastered / covered) * 100) if covered > 0 else 0

    # Top re-engaged concepts (covered by most students)
    top_concepts = conn.execute("""
        SELECT cp.concept_id, cp.subject_id, COUNT(*) AS student_count,
               SUM(CASE WHEN cp.mastered_at IS NOT NULL THEN 1 ELSE 0 END) AS mastered_count
        FROM concept_progress cp
        GROUP BY cp.concept_id, cp.subject_id
        ORDER BY student_count DESC LIMIT 8
    """).fetchall()

    # Streak stats
    streak_rows = conn.execute(
        "SELECT streak_count FROM student_profile WHERE streak_count IS NOT NULL AND streak_count > 0"
    ).fetchall()
    streaks = [r["streak_count"] for r in streak_rows]
    avg_streak = round(sum(streaks) / len(streaks), 1) if streaks else 0
    max_streak = max(streaks) if streaks else 0

    conn.close()
    return {
        "dau": dau, "wau": wau, "mau": mau,
        "total_students": total_students,
        "new_today": new_today,
        "new_this_week": new_this_week,
        "dau_mau_ratio": round((dau / mau) * 100) if mau > 0 else 0,
        "sparkline": sparkline,
        "retention": {"d1": d1_retention, "d7": d7_retention, "d30": d30_retention},
        "avg_session_msgs": avg_session_msgs,
        "depth_score": depth_score,
        "top_concepts": [dict(r) for r in top_concepts],
        "avg_streak": avg_streak,
        "max_streak": max_streak,
        "students_with_streak": len(streaks),
    }


@app.get("/admin/peak-hours")
def admin_peak_hours(x_admin_key: str = Header(None)):
    require_admin(x_admin_key)
    conn = get_db()
    rows = conn.execute("""
        SELECT
            CAST(strftime('%H', created_at) AS INTEGER) AS hour,
            CAST(strftime('%w', created_at) AS INTEGER) AS dow,
            COUNT(*) AS msg_count
        FROM messages WHERE role='user'
        GROUP BY hour, dow
    """).fetchall()
    conn.close()
    grid = {(r["dow"], r["hour"]): r["msg_count"] for r in rows}
    result = []
    days = ["Sun","Mon","Tue","Wed","Thu","Fri","Sat"]
    for dow in range(7):
        for hour in range(24):
            result.append({
                "day": days[dow], "dow": dow, "hour": hour,
                "count": grid.get((dow, hour), 0)
            })
    return result


@app.get("/admin/health")
def admin_health(x_admin_key: str = Header(None)):
    require_admin(x_admin_key)
    conn = get_db()
    db_size_mb = round(os.path.getsize(DB_PATH) / (1024 * 1024), 2) if os.path.exists(DB_PATH) else 0
    total_messages = conn.execute("SELECT COUNT(*) FROM messages").fetchone()[0]
    total_students = conn.execute("SELECT COUNT(*) FROM students").fetchone()[0]
    conn.close()
    uptime_seconds = int(time.time() - _SERVER_START_TIME)
    hours, remainder = divmod(uptime_seconds, 3600)
    minutes, _ = divmod(remainder, 60)
    return {
        "status": "ok",
        "uptime": f"{hours}h {minutes}m",
        "db_size_mb": db_size_mb,
        "total_messages": total_messages,
        "total_students": total_students,
    }


# ── Chat Logs ──────────────────────────────────────────────────────────────────

@app.get("/admin/chat-logs")
def admin_chat_logs(x_admin_key: str = Header(None)):
    require_admin(x_admin_key)
    conn = get_db()
    rows = conn.execute("""
        SELECT s.id AS student_id, s.name, s.email,
               m.subject_id,
               COUNT(*) AS message_count,
               MAX(m.created_at) AS last_message_at
        FROM messages m
        JOIN students s ON s.id = m.student_id
        WHERE m.role = 'user'
        GROUP BY s.id, m.subject_id
        ORDER BY last_message_at DESC
    """).fetchall()
    conn.close()
    result = []
    for r in rows:
        d = dict(r)
        subj = SUBJECTS.get(d["subject_id"])
        d["subject_name"] = subj["name"] if subj else d["subject_id"]
        d["subject_color"] = subj["color"] if subj else "#ccc"
        result.append(d)
    return result


@app.get("/admin/chat-logs/{student_id}/{subject_id}")
def admin_chat_log_thread(student_id: str, subject_id: str, x_admin_key: str = Header(None)):
    require_admin(x_admin_key)
    conn = get_db()
    rows = conn.execute(
        "SELECT role, content, created_at FROM messages WHERE student_id=? AND subject_id=? ORDER BY created_at ASC",
        (student_id, subject_id)
    ).fetchall()
    conn.close()
    return [dict(r) for r in rows]


# ── Study Plans (admin) ────────────────────────────────────────────────────────

@app.get("/admin/study-plans")
def admin_study_plans(x_admin_key: str = Header(None)):
    require_admin(x_admin_key)
    conn = get_db()
    students = conn.execute("""
        SELECT s.id, s.name, s.email, sp.career_id, sp.avatar_color
        FROM students s
        LEFT JOIN student_profile sp ON sp.student_id = s.id
        WHERE EXISTS (SELECT 1 FROM study_plan WHERE student_id = s.id)
        ORDER BY s.name
    """).fetchall()
    today = datetime.utcnow().date().isoformat()
    result = []
    for st in students:
        sid = st["id"]
        plan_rows = conn.execute(
            "SELECT day_number, subject_id, concept_id, target_date FROM study_plan WHERE student_id=? ORDER BY day_number",
            (sid,)
        ).fetchall()
        covered_rows = conn.execute(
            "SELECT subject_id, concept_id FROM concept_progress WHERE student_id=?", (sid,)
        ).fetchall()
        covered = {(r["subject_id"], r["concept_id"]) for r in covered_rows}
        days_map = {}
        for r in plan_rows:
            d = r["day_number"]
            if d not in days_map:
                days_map[d] = {"day": d, "target_date": r["target_date"], "concepts": []}
            days_map[d]["concepts"].append({
                "subject_id": r["subject_id"],
                "concept_id": r["concept_id"],
                "covered": (r["subject_id"], r["concept_id"]) in covered,
            })
        plan = list(days_map.values())
        total = sum(len(d["concepts"]) for d in plan)
        done = sum(1 for d in plan for c in d["concepts"] if c["covered"])
        overdue = sum(1 for d in plan for c in d["concepts"]
                      if d["target_date"] <= today and not c["covered"])
        career = CAREERS.get(st["career_id"]) if st["career_id"] else None
        result.append({
            "student_id": sid,
            "name": st["name"],
            "email": st["email"],
            "avatar_color": st["avatar_color"],
            "career_title": career["title"] if career else None,
            "total_days": len(plan),
            "total_concepts": total,
            "covered_concepts": done,
            "overdue_concepts": overdue,
            "pct": round((done / total) * 100) if total > 0 else 0,
            "plan": plan,
        })
    conn.close()
    return result


# ── CSV Export ─────────────────────────────────────────────────────────────────

from fastapi.responses import Response as FastAPIResponse

@app.get("/admin/export/students")
def admin_export_students(x_admin_key: str = Header(None)):
    require_admin(x_admin_key)
    conn = get_db()
    rows = conn.execute("""
        SELECT s.name, s.email, s.created_at,
               sp.career_id, sp.college, sp.year_of_study, sp.city, sp.state,
               sp.streak_count, sp.last_active_at,
               (SELECT COUNT(*) FROM concept_progress cp WHERE cp.student_id = s.id) AS concepts_covered,
               (SELECT COUNT(*) FROM concept_progress cp WHERE cp.student_id = s.id AND cp.mastered_at IS NOT NULL) AS concepts_mastered,
               (SELECT COUNT(*) FROM messages m WHERE m.student_id = s.id AND m.role='user') AS messages_sent,
               (SELECT COUNT(DISTINCT subject_id) FROM concept_progress cp WHERE cp.student_id = s.id) AS subjects_touched
        FROM students s
        LEFT JOIN student_profile sp ON sp.student_id = s.id
        ORDER BY s.created_at DESC
    """).fetchall()
    conn.close()
    lines = ["Name,Email,Joined,Career,College,Year,City,State,Streak,Last Active,Concepts Covered,Concepts Mastered,Messages Sent,Subjects Touched"]
    for r in rows:
        career = CAREERS.get(r["career_id"])
        career_title = career["title"] if career else ""
        def esc(v): return f'"{str(v or "").replace(chr(34), chr(39))}"'
        lines.append(",".join([
            esc(r["name"]), esc(r["email"]), esc(r["created_at"][:10]),
            esc(career_title), esc(r["college"]), esc(r["year_of_study"]),
            esc(r["city"]), esc(r["state"]),
            str(r["streak_count"] or 0), esc(r["last_active_at"] or ""),
            str(r["concepts_covered"]), str(r["concepts_mastered"]),
            str(r["messages_sent"]), str(r["subjects_touched"]),
        ]))
    csv_content = "\n".join(lines)
    return FastAPIResponse(
        content=csv_content,
        media_type="text/csv",
        headers={"Content-Disposition": "attachment; filename=bversity_students.csv"}
    )


@app.get("/admin/export/progress")
def admin_export_progress(x_admin_key: str = Header(None)):
    require_admin(x_admin_key)
    conn = get_db()
    rows = conn.execute("""
        SELECT s.name, s.email, cp.subject_id, cp.concept_id,
               cp.first_covered_at, cp.mastered_at
        FROM concept_progress cp
        JOIN students s ON s.id = cp.student_id
        ORDER BY s.name, cp.subject_id, cp.first_covered_at
    """).fetchall()
    conn.close()
    lines = ["Student Name,Email,Subject,Concept,First Covered,Mastered At"]
    for r in rows:
        subj = SUBJECTS.get(r["subject_id"])
        subject_name = subj["name"] if subj else r["subject_id"]
        def esc(v): return f'"{str(v or "").replace(chr(34), chr(39))}"'
        lines.append(",".join([
            esc(r["name"]), esc(r["email"]),
            esc(subject_name), esc(r["concept_id"].replace("_", " ")),
            esc(r["first_covered_at"][:10] if r["first_covered_at"] else ""),
            esc(r["mastered_at"][:10] if r["mastered_at"] else ""),
        ]))
    csv_content = "\n".join(lines)
    return FastAPIResponse(
        content=csv_content,
        media_type="text/csv",
        headers={"Content-Disposition": "attachment; filename=bversity_progress.csv"}
    )


# ── Announcements ──────────────────────────────────────────────────────────────

class AnnounceRequest(BaseModel):
    subject: str
    message: str
    target: str = "all"  # "all" | "active_week" | career_id

@app.post("/admin/announce")
async def admin_announce(req: AnnounceRequest, background_tasks: BackgroundTasks, x_admin_key: str = Header(None)):
    require_admin(x_admin_key)
    conn = get_db()
    now = datetime.utcnow()
    if req.target == "all":
        students = conn.execute("SELECT s.id, s.name, s.email FROM students s").fetchall()
    elif req.target == "active_week":
        week_ago = (now - timedelta(days=7)).isoformat()
        students = conn.execute("""
            SELECT DISTINCT s.id, s.name, s.email FROM students s
            JOIN messages m ON m.student_id = s.id
            WHERE m.created_at >= ?
        """, (week_ago,)).fetchall()
    else:
        students = conn.execute("""
            SELECT s.id, s.name, s.email FROM students s
            JOIN student_profile sp ON sp.student_id = s.id
            WHERE sp.career_id = ?
        """, (req.target,)).fetchall()
    conn.close()

    async def _send_all():
        for st in students:
            first = st["name"].split()[0]
            body = (
                _heading(req.subject) +
                _para(req.message.replace("\n", "<br>")) +
                _divider() +
                _small(f"Sent to you by the Bversity team · <a href='https://university.bversity.io' style='color:#00A896'>Open Bversity</a>")
            )
            await _send_email(st["email"], req.subject, _email_wrap(body))

    background_tasks.add_task(_send_all)
    return {"queued": len(students)}

@app.get("/admin/announce/preview")
def admin_announce_preview(target: str = "all", x_admin_key: str = Header(None)):
    require_admin(x_admin_key)
    conn = get_db()
    now = datetime.utcnow()
    if target == "all":
        count = conn.execute("SELECT COUNT(*) FROM students").fetchone()[0]
        names = [r["name"] for r in conn.execute("SELECT name FROM students LIMIT 5").fetchall()]
    elif target == "active_week":
        week_ago = (now - timedelta(days=7)).isoformat()
        count = conn.execute("""
            SELECT COUNT(DISTINCT student_id) FROM messages WHERE created_at >= ?
        """, (week_ago,)).fetchone()[0]
        names = [r["name"] for r in conn.execute("""
            SELECT DISTINCT s.name FROM students s
            JOIN messages m ON m.student_id = s.id
            WHERE m.created_at >= ? LIMIT 5
        """, (week_ago,)).fetchall()]
    else:
        count = conn.execute("SELECT COUNT(*) FROM student_profile WHERE career_id=?", (target,)).fetchone()[0]
        names = [r["name"] for r in conn.execute("""
            SELECT s.name FROM students s
            JOIN student_profile sp ON sp.student_id = s.id
            WHERE sp.career_id = ? LIMIT 5
        """, (target,)).fetchall()]
    conn.close()
    return {"count": count, "sample_names": names}


class ConceptFeedbackReq(BaseModel):
    student_id: str
    subject_id: str
    concept_title: str
    value: str  # 'up' or 'down'

@app.post("/api/concept-feedback")
def post_concept_feedback(req: ConceptFeedbackReq):
    if req.value not in ("up", "down"):
        raise HTTPException(status_code=400, detail="value must be 'up' or 'down'")
    conn = get_db()
    conn.execute(
        "INSERT INTO concept_feedback (student_id, subject_id, concept_title, value, created_at) VALUES (?, ?, ?, ?, ?)",
        (req.student_id, req.subject_id, req.concept_title, req.value, datetime.utcnow().isoformat()),
    )
    conn.commit()
    conn.close()
    return {"ok": True}


class MessageFeedbackReq(BaseModel):
    student_id: str
    subject_id: str
    message_idx: int
    value: str  # 'up' or 'down'

@app.post("/api/message-feedback")
def post_message_feedback(req: MessageFeedbackReq):
    if req.value not in ("up", "down"):
        raise HTTPException(status_code=400, detail="value must be 'up' or 'down'")
    conn = get_db()
    conn.execute(
        "INSERT INTO message_feedback (student_id, subject_id, message_idx, value, created_at) VALUES (?, ?, ?, ?, ?)",
        (req.student_id, req.subject_id, req.message_idx, req.value, datetime.utcnow().isoformat()),
    )
    conn.commit()
    conn.close()
    return {"ok": True}


@app.get("/api/admin/concept-notes/{subject_id}")
def get_concept_notes(subject_id: str, x_admin_key: str = Header(None)):
    require_admin(x_admin_key)
    conn = get_db()
    rows = conn.execute(
        "SELECT concept_id, notes FROM concept_notes WHERE subject_id = ?", (subject_id,)
    ).fetchall()
    conn.close()
    return {r["concept_id"]: r["notes"] for r in rows}


class ConceptNotesReq(BaseModel):
    notes: str

@app.put("/api/admin/concept-notes/{subject_id}/{concept_id}")
def put_concept_notes(subject_id: str, concept_id: str, req: ConceptNotesReq, x_admin_key: str = Header(None)):
    require_admin(x_admin_key)
    conn = get_db()
    conn.execute(
        """INSERT INTO concept_notes (subject_id, concept_id, notes, updated_at)
           VALUES (?, ?, ?, ?)
           ON CONFLICT(subject_id, concept_id) DO UPDATE SET notes=excluded.notes, updated_at=excluded.updated_at""",
        (subject_id, concept_id, req.notes, datetime.utcnow().isoformat()),
    )
    conn.commit()
    conn.close()
    return {"ok": True}


@app.get("/api/search-data/{student_id}")
def get_search_data(student_id: str):
    """Return all covered concepts (with names) + notes for client-side search."""
    conn = get_db()
    progress_rows = conn.execute(
        "SELECT subject_id, concept_id, mastered_at FROM concept_progress WHERE student_id = ?",
        (student_id,)
    ).fetchall()
    note_rows = conn.execute(
        "SELECT id, subject_id, content, created_at FROM notes WHERE student_id = ? ORDER BY created_at DESC",
        (student_id,)
    ).fetchall()
    conn.close()

    # Build concept lookup for names/descs
    concept_lookup = {}
    for subj_id, concepts in CURRICULUM.items():
        for c in concepts:
            concept_lookup[(subj_id, c["id"])] = {"name": c["name"], "desc": c.get("desc", "")}

    covered = []
    for r in progress_rows:
        meta = concept_lookup.get((r["subject_id"], r["concept_id"]), {})
        subj = SUBJECTS.get(r["subject_id"], {})
        covered.append({
            "subject_id": r["subject_id"],
            "subject_name": subj.get("name", r["subject_id"]),
            "subject_color": subj.get("color", "#00A896"),
            "concept_id": r["concept_id"],
            "concept_name": meta.get("name", r["concept_id"].replace("_", " ").title()),
            "concept_desc": meta.get("desc", ""),
            "mastered": r["mastered_at"] is not None,
        })

    notes = [
        {
            "id": r["id"],
            "subject_id": r["subject_id"],
            "subject_name": SUBJECTS.get(r["subject_id"], {}).get("name", "") if r["subject_id"] else "",
            "content": r["content"],
            "created_at": r["created_at"],
        }
        for r in note_rows
    ]
    return {"covered": covered, "notes": notes}


@app.get("/api/recall-check/{student_id}/{subject_id}")
def recall_check(student_id: str, subject_id: str):
    """Return whether a recall warmup is needed for this session."""
    conn = get_db()
    today = datetime.utcnow().date().isoformat()

    total_msgs = conn.execute(
        "SELECT COUNT(*) FROM messages WHERE student_id = ? AND subject_id = ?",
        (student_id, subject_id)
    ).fetchone()[0]

    msgs_today = conn.execute(
        "SELECT COUNT(*) FROM messages WHERE student_id = ? AND subject_id = ? AND date(created_at) = ?",
        (student_id, subject_id, today)
    ).fetchone()[0]

    conn.close()
    # Warmup needed: has prior history but nothing sent today
    return {"needed": total_msgs > 0 and msgs_today == 0}


@app.get("/api/career-pace/{student_id}")
def get_career_pace(student_id: str):
    """Return weekly mastery pace and per-subject progress for career narrative."""
    conn = get_db()
    now = datetime.utcnow()
    two_weeks_ago = (now - timedelta(days=14)).isoformat()

    # Concepts mastered in the last 14 days
    recent_mastered = conn.execute(
        """SELECT COUNT(*) as cnt FROM concept_progress
           WHERE student_id = ? AND mastered_at IS NOT NULL AND mastered_at >= ?""",
        (student_id, two_weeks_ago)
    ).fetchone()["cnt"]
    # Weekly pace (14-day window → 2 weeks)
    weekly_pace = round(recent_mastered / 2, 1)

    # All mastered concepts per subject
    rows = conn.execute(
        """SELECT subject_id,
                  COUNT(*) as covered,
                  SUM(CASE WHEN mastered_at IS NOT NULL THEN 1 ELSE 0 END) as mastered
           FROM concept_progress WHERE student_id = ? GROUP BY subject_id""",
        (student_id,)
    ).fetchall()
    per_subject = {r["subject_id"]: {"covered": r["covered"], "mastered": r["mastered"]} for r in rows}

    # Session count last 7 days (active days)
    week_ago = (now - timedelta(days=7)).isoformat()
    active_days = conn.execute(
        """SELECT COUNT(DISTINCT date(created_at)) FROM messages
           WHERE student_id = ? AND role = 'user' AND created_at >= ?""",
        (student_id, week_ago)
    ).fetchone()[0]

    conn.close()
    return {
        "weekly_pace": weekly_pace,
        "per_subject": per_subject,
        "active_days_last_week": active_days,
    }


# ── Talk to Founder ────────────────────────────────────────────────────────

class FounderContactReq(BaseModel):
    student_name:  str
    student_email: str
    reason:        str
    message:       str

@app.post("/contact-founder")
async def contact_founder(req: FounderContactReq):
    founder_email = "sudharsan@bversity.io"
    html = f"""
    <div style="font-family:Inter,sans-serif;max-width:600px;margin:0 auto;padding:24px;background:#f9f9f9">
      <div style="background:#fff;border-radius:8px;padding:32px;border:1px solid #e8e4dc">
        <p style="font-size:11px;font-weight:700;letter-spacing:0.1em;text-transform:uppercase;color:#9a9490;margin:0 0 4px">New message from a student</p>
        <h2 style="font-size:22px;font-weight:700;color:#1a1a1a;margin:0 0 24px">{req.student_name}</h2>

        <div style="background:#f8f6f1;border-left:3px solid #16c1ad;padding:12px 16px;margin-bottom:20px">
          <p style="font-size:11px;font-weight:700;letter-spacing:0.08em;text-transform:uppercase;color:#9a9490;margin:0 0 4px">Reason</p>
          <p style="font-size:15px;color:#1a1a1a;margin:0">{req.reason}</p>
        </div>

        <div style="margin-bottom:24px">
          <p style="font-size:11px;font-weight:700;letter-spacing:0.08em;text-transform:uppercase;color:#9a9490;margin:0 0 8px">Message</p>
          <p style="font-size:15px;color:#1a1a1a;line-height:1.6;margin:0;white-space:pre-wrap">{req.message}</p>
        </div>

        <div style="border-top:1px solid #e8e4dc;padding-top:16px">
          <p style="font-size:13px;color:#9a9490;margin:0">Hit reply to respond directly to {req.student_name} at <strong>{req.student_email}</strong></p>
        </div>
      </div>
    </div>
    """
    await _send_email(
        to_email=founder_email,
        subject=f"[Bversity] {req.student_name}  -  {req.reason}",
        html=html,
        reply_to=req.student_email,
    )
    return {"status": "sent"}


# ── Image Config ───────────────────────────────────────────────────────────────

@app.get("/admin/images")
def get_images(x_admin_key: str = Header(None)):
    require_admin(x_admin_key)
    return load_image_config()

@app.get("/api/images")
def get_images_public():
    return load_image_config()

class ImageUpdateReq(BaseModel):
    section: str   # "careers", "clusters", "degrees"
    key:     str
    url:     str

@app.post("/admin/images")
def update_image(req: ImageUpdateReq, x_admin_key: str = Header(None)):
    require_admin(x_admin_key)
    config = load_image_config()
    if req.section not in config:
        raise HTTPException(status_code=400, detail="Invalid section")
    if req.key not in config[req.section]:
        raise HTTPException(status_code=400, detail="Invalid key")
    config[req.section][req.key]["url"] = req.url
    save_image_config(config)
    return {"status": "ok"}


# ── Adzuna Job Listings ────────────────────────────────────────────────────────

ADZUNA_APP_ID  = os.getenv("ADZUNA_APP_ID", "")
ADZUNA_APP_KEY = os.getenv("ADZUNA_APP_KEY", "")
ADZUNA_COUNTRY = "us"

CERT_SEARCH_QUERIES = {
    "us_cra":              "clinical research associate CCRA",
    "us_ccrp":             "clinical research coordinator CCRP",
    "us_regulatory":       "regulatory affairs specialist RAC",
    "us_pharmacovigilance":"pharmacovigilance drug safety CPVC",
    "us_msl":              "medical science liaison BCMAS",
    "us_cdm":              "clinical data manager CCDM",
}

JOB_CACHE_DAYS = 3


def _fetch_adzuna(cert_id: str) -> list[dict]:
    query = CERT_SEARCH_QUERIES.get(cert_id, "life sciences")
    params = urllib.parse.urlencode({
        "app_id":         ADZUNA_APP_ID,
        "app_key":        ADZUNA_APP_KEY,
        "results_per_page": 10,
        "what":           query,
        "content-type":   "application/json",
    })
    url = f"https://api.adzuna.com/v1/api/jobs/{ADZUNA_COUNTRY}/search/1?{params}"
    req = urllib.request.Request(url, headers={"Accept": "application/json"})
    try:
        with urllib.request.urlopen(req, timeout=10, context=_SSL_CTX) as resp:
            data = json.loads(resp.read().decode())
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Adzuna error: {e}")

    results = []
    for job in data.get("results", []):
        sal = job.get("salary_min"), job.get("salary_max")
        results.append({
            "cert_id":     cert_id,
            "title":       job.get("title", ""),
            "company":     job.get("company", {}).get("display_name", ""),
            "location":    job.get("location", {}).get("display_name", ""),
            "salary_min":  sal[0],
            "salary_max":  sal[1],
            "url":         job.get("redirect_url", ""),
            "description": (job.get("description") or "")[:400],
            "posted_date": job.get("created", ""),
            "fetched_at":  datetime.utcnow().isoformat(),
        })
    return results


def _store_jobs(cert_id: str, jobs: list[dict]):
    with sqlite3.connect(DB_PATH) as db:
        db.execute("DELETE FROM job_listings WHERE cert_id = ?", (cert_id,))
        db.executemany(
            """INSERT INTO job_listings
               (cert_id, title, company, location, salary_min, salary_max,
                url, description, posted_date, fetched_at)
               VALUES (:cert_id,:title,:company,:location,:salary_min,:salary_max,
                       :url,:description,:posted_date,:fetched_at)""",
            jobs,
        )


def _cached_jobs(cert_id: str) -> list[dict] | None:
    cutoff = (datetime.utcnow() - timedelta(days=JOB_CACHE_DAYS)).isoformat()
    with sqlite3.connect(DB_PATH) as db:
        db.row_factory = sqlite3.Row
        rows = db.execute(
            "SELECT * FROM job_listings WHERE cert_id=? AND fetched_at>? ORDER BY id",
            (cert_id, cutoff),
        ).fetchall()
    if not rows:
        return None
    return [dict(r) for r in rows]


@app.get("/jobs/{cert_id}")
def get_jobs(cert_id: str):
    if not ADZUNA_APP_ID or not ADZUNA_APP_KEY:
        raise HTTPException(status_code=503, detail="Adzuna credentials not configured")
    cached = _cached_jobs(cert_id)
    if cached:
        return {"jobs": cached, "source": "cache"}
    jobs = _fetch_adzuna(cert_id)
    _store_jobs(cert_id, jobs)
    return {"jobs": jobs, "source": "live"}


@app.post("/jobs/refresh")
def refresh_jobs(x_admin_key: str = Header(None)):
    require_admin(x_admin_key)
    if not ADZUNA_APP_ID or not ADZUNA_APP_KEY:
        raise HTTPException(status_code=503, detail="Adzuna credentials not configured")
    results = {}
    for cert_id in CERT_SEARCH_QUERIES:
        try:
            jobs = _fetch_adzuna(cert_id)
            _store_jobs(cert_id, jobs)
            results[cert_id] = len(jobs)
        except Exception as e:
            results[cert_id] = f"error: {e}"
    return results


# ── Industry News endpoints ───────────────────────────────────────────────────

@app.get("/industry-news")
def get_industry_news(background_tasks: BackgroundTasks):
    global _news_last_fetched
    conn = get_db()
    rows = conn.execute(
        "SELECT id, title, url, source, published_at FROM industry_news ORDER BY id DESC LIMIT 30"
    ).fetchall()
    newsletter = conn.execute(
        "SELECT id, title, content, url, published_at FROM industry_newsletter ORDER BY published_at DESC LIMIT 5"
    ).fetchall()
    conn.close()
    is_stale = _news_last_fetched is None or (datetime.utcnow() - _news_last_fetched).total_seconds() > 21600
    if is_stale and not _news_fetch_lock:
        background_tasks.add_task(_fetch_news_feeds)
    return {
        "articles": [dict(r) for r in rows],
        "newsletter": [dict(r) for r in newsletter],
        "refreshing": is_stale and not rows,
    }


@app.post("/admin/industry-news/refresh")
def refresh_industry_news(x_admin_key: str = Header(None), background_tasks: BackgroundTasks = None):
    require_admin(x_admin_key)
    background_tasks.add_task(_fetch_news_feeds)
    return {"ok": True, "message": "Refresh started in background"}


@app.post("/admin/industry-newsletter")
def add_newsletter_post(data: dict, x_admin_key: str = Header(None)):
    require_admin(x_admin_key)
    title = (data.get("title") or "").strip()
    if not title:
        raise HTTPException(status_code=400, detail="title required")
    now = datetime.utcnow().isoformat()
    conn = get_db()
    conn.execute(
        "INSERT INTO industry_newsletter (title, content, url, published_at, created_at) VALUES (?,?,?,?,?)",
        (title, data.get("content"), data.get("url"), data.get("published_at") or now, now),
    )
    conn.commit()
    conn.close()
    return {"ok": True}


@app.delete("/admin/industry-newsletter/{post_id}")
def delete_newsletter_post(post_id: int, x_admin_key: str = Header(None)):
    require_admin(x_admin_key)
    conn = get_db()
    conn.execute("DELETE FROM industry_newsletter WHERE id=?", (post_id,))
    conn.commit()
    conn.close()
    return {"ok": True}


# ── Automated daily scheduler ─────────────────────────────────────────────────

def _run_daily_jobs():
    import asyncio
    loop = asyncio.new_event_loop()

    async def _jobs():
        conn = get_db()
        now  = datetime.utcnow()

        # 1. Check trial expiry + send warnings
        tomorrow = (now + timedelta(days=1)).isoformat()
        warning_subs = conn.execute("""
            SELECT sub.student_id, sub.trial_end, s.email, s.name, sub.product
            FROM subscriptions sub JOIN students s ON s.id = sub.student_id
            WHERE sub.status = 'trial' AND sub.trial_end <= ? AND sub.trial_end > ?
        """, (tomorrow, now.isoformat())).fetchall()
        expired_ids = conn.execute(
            "SELECT id FROM subscriptions WHERE status = 'trial' AND trial_end < ?",
            (now.isoformat(),)
        ).fetchall()
        for sub in expired_ids:
            conn.execute("UPDATE subscriptions SET status='expired', updated_at=? WHERE id=?",
                         (now.isoformat(), sub["id"]))
        conn.commit()

        # 2. Send weekly digest every Monday
        if now.weekday() == 0:
            students = conn.execute("""
                SELECT s.id, s.email, s.name FROM students s
                WHERE EXISTS (SELECT 1 FROM messages m WHERE m.student_id = s.id AND m.role = 'user')
                AND s.email IN (SELECT email FROM approved_emails)
            """).fetchall()
            conn.close()
            for st in students:
                try:
                    await _send_student_weekly_report(st["id"], st["email"], st["name"])
                except Exception as e:
                    print(f"[scheduler] weekly report failed for {st['id']}: {e}")
        else:
            conn.close()

        print(f"[scheduler] daily jobs complete — {len(expired_ids)} trials expired, {len(warning_subs)} warnings queued")

    try:
        loop.run_until_complete(_jobs())
    finally:
        loop.close()

def _scheduler_thread():
    import time
    while True:
        now = datetime.utcnow()
        # Run at 02:00 UTC daily
        target = now.replace(hour=2, minute=0, second=0, microsecond=0)
        if now >= target:
            target += timedelta(days=1)
        wait = (target - now).total_seconds()
        time.sleep(wait)
        try:
            _run_daily_jobs()
        except Exception as e:
            print(f"[scheduler] error: {e}")

import threading as _threading
_scheduler = _threading.Thread(target=_scheduler_thread, daemon=True, name="daily-scheduler")
_scheduler.start()
print("[scheduler] daily scheduler started — runs at 02:00 UTC")
